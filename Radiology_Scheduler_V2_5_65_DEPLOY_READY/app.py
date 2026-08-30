from __future__ import annotations

import os
import json
import subprocess
import sys
import tempfile
import signal
from copy import deepcopy
import calendar
import html
import unicodedata
import smtplib
import hashlib
import re
import urllib.parse
from statistics import median
from io import BytesIO
from datetime import date, datetime, time, timezone, timedelta
from time import perf_counter
from pathlib import Path
from zoneinfo import ZoneInfo
from email.message import EmailMessage

import pandas as pd
import numpy as np
import streamlit as st
import xlsxwriter
from pypdf import PdfReader
from docx import Document
from supabase import create_client

import scheduler_engine as _scheduler_engine
from scheduler_engine import (
    Person, Slot, SolveResult, DEFAULT_PEOPLE, PERSON_COLORS, next_month, weekday_count, round_half_up,
    standard_target, make_slots, solve_schedule, attempt_swap, preview_swap, validate_schedule,
    lithuanian_public_holidays, public_holiday_days_in_month, is_public_holiday,
    serialize_result, deserialize_result, revalidate_loaded_result, calculate_targets, blocks_overlap, hard_unavailable_for_block,
    resident_hard_unavailable_for_block, absolute_unavailable_for_block,
    serialize_people_request_snapshot, people_from_request_snapshot,
    ROTATION_CATEGORIES, rotation_category, backup_required_slot, backup_best_effort_slot,
    effective_actual_assignments, calculate_live_fairness_snapshot,
    is_emergency_critical_slot, is_emergency_lower_priority_donor_slot,
    emergency_donor_source_slots, apply_emergency_critical_transfer,
    DEFAULT_RULE_PROFILE, validate_rule_profile, set_runtime_rules, get_runtime_rules, rule_value,
    FATIGUE_MAX_WORKDAYS_ROLLING7, FATIGUE_ROLLING7_HARD_CEILING_HOURS, WEEKLY_LOAD_SOFT_TARGET_HOURS,
    SWAP_ABSOLUTE_MAX_HOURS_ROLLING7, SWAP_MAX_WORKDAYS_ROLLING7, SWAP_MIN_DAILY_REST_HOURS, SWAP_MAX_HOURS_PER_DAY
)
import db
from notification_core import smtp_config as _smtp_config_core, smtp_missing as _smtp_missing_core, smtp_probe as _smtp_probe_core, send_email as _send_email_core

ENGINE_API_VERSION = str(getattr(_scheduler_engine,"ENGINE_API_VERSION","LEGACY_OR_UNKNOWN"))
APP_VERSION = "2.5.115 THEORETICAL BACKUP LAYER"
EXPECTED_ENGINE_API_VERSION = "2.5.112"
BASE = Path(__file__).parent
SENIOR_INITIALS = "SP"
RESEARCHER_INITIALS = "ŠR"
WESTON_CREDITOR_INITIALS = RESEARCHER_INITIALS  # SP generation clicks are owed to ŠR
DEFAULT_SUPABASE_URL = "https://gqdlwhjgwqmuoolybusy.supabase.co"
DEFAULT_SUPABASE_PUBLISHABLE_KEY = "sb_publishable_kHX4M55rZoHJr61S9kzdLg_tgKN-oDI"
DEFAULT_MANUAL_LT = (BASE / "manual_lt.md").read_text(encoding="utf-8")
DEFAULT_MANUAL_EN = (BASE / "manual_en.md").read_text(encoding="utf-8")
SENIOR_GUIDE_LT = (BASE / "SENIOR_USABILITY_GUIDE_LT.md").read_text(encoding="utf-8")
SENIOR_GUIDE_EN = (BASE / "SENIOR_USABILITY_GUIDE_EN.md").read_text(encoding="utf-8")

db.init_db(DEFAULT_MANUAL_LT, DEFAULT_MANUAL_EN, DEFAULT_PEOPLE)

st.set_page_config(page_title="Shift Happens", layout="wide", initial_sidebar_state="expanded")

if str(ENGINE_API_VERSION) != EXPECTED_ENGINE_API_VERSION:
    st.error(
        "APP / ENGINE VERSION MISMATCH. "
        f"App expects scheduler_engine API {EXPECTED_ENGINE_API_VERSION}, "
        f"but loaded {ENGINE_API_VERSION}. "
        "Deploy app.py AND scheduler_engine.py from the same release."
    )
    st.stop()
st.markdown("""
<style>
html, body, [class*="css"] {font-family:-apple-system,BlinkMacSystemFont,"Segoe UI",Roboto,Helvetica,Arial,sans-serif;}
.block-container {max-width:1550px;padding-top:1.6rem;}
h1 {font-weight:720!important;letter-spacing:-.025em;}
h2,h3 {font-weight:650!important;letter-spacing:-.015em;}
div[data-testid="stMarkdownContainer"] p, div[data-testid="stMarkdownContainer"] li {line-height:1.58;}
div[data-testid="stMetric"] {border:1px solid rgba(128,128,128,.22);padding:10px 14px;border-radius:10px;}
.deadline-card {border:1px solid rgba(128,128,128,.25);border-radius:12px;padding:12px 16px;margin:8px 0 18px 0;}
</style>
""", unsafe_allow_html=True)

TR = {
"LT": {
"language":"Kalba","user":"Vartotojas","profile":"Profilis","resident_profile":"Rezidento profilis","senior_profile":"Seniūnės profilis",
"resident_pin":"Asmeninis PIN","admin_pin":"Seniūnės PIN","local_resident":"Vietinis testavimo režimas: asmeniniai PIN nesukonfigūruoti.",
"local_senior":"Vietinis testavimo režimas: seniūnės profilis atrakintas tik seniūnės paskyrai.","bad_pin":"Neteisingas PIN.",
"login_title":"Prisijungimas","login":"PRISIJUNGTI","logout":"ATSIJUNGTI","signup":"SUKURTI PASKYRĄ","signup_title":"Pirma registracija","password":"Slaptažodis","password_repeat":"Pakartokite slaptažodį","auth_email":"El. paštas","auth_invalid":"Nepavyko prisijungti. Patikrinkite el. paštą ir slaptažodį.","signup_sent":"Paskyra sukurta. Jei el. pašto patvirtinimas įjungtas, patvirtinkite laišką ir tada prisijunkite.","signup_password_mismatch":"Slaptažodžiai nesutampa.","claim_title":"Susiekite paskyrą","resident_claim_tab":"Rezidentas","observer_claim_tab":"Skyriaus administratorė / stebėtoja","observer_claim_help":"Ši paskyra skirta skyriaus administracinei peržiūrai. Ji yra tik skaitymui: galima matyti paskelbtą ir aktualų grafiką, apsikeitimus, dublius, teisingumą ir auditą, bet negalima nieko keisti.","observer_invite_code":"Skyriaus stebėtojo kvietimo kodas","observer_claim":"AKTYVUOTI TIK PERŽIŪROS PRIEIGĄ","observer_read_only":"TIK PERŽIŪRA","observer_role":"Skyriaus stebėtojas","observer_portal":"Skyriaus grafiko stebėsena","observer_overview":"Apžvalga","observer_schedule":"Grafikai","observer_changes":"Pakeitimų žurnalas","observer_fairness":"Teisingumas","observer_backups":"Dubliai","observer_rules":"Taisyklės","observer_scope_note":"Ši paskyra negali generuoti, publikuoti, tvirtinti, keisti ar anuliuoti grafikų ir apsikeitimų.","observer_privacy_note":"Nerodomi privatūs rezidentų pageidavimai, privalomo negalėjimo datos, asmeninės pastabos, el. paštai ar paskyrų nustatymai.","observer_baseline_schedule":"Sistemos pradinis grafikas — paskelbimo momentu","observer_actual_schedule":"Faktinis grafikas — dabar","observer_change_count":"Pakeistų normalių pamainų","observer_normal_swaps":"Normalių pamainų apsikeitimai","observer_backup_swaps":"Dublių apsikeitimai","observer_pending_swaps":"Laukiantys","observer_approved_swaps":"Patvirtinti","observer_rejected_swaps":"Atmesti","observer_no_changes":"Po paskelbimo normalių pamainų pakeitimų nėra.","observer_from":"Buvo","observer_to":"Dabar","observer_change_log_help":"Sistemos pradinis grafikas lieka teisingumo apskaitai. Faktinis grafikas rodo realią situaciją po abipusių savanoriškų apsikeitimų.","observer_backup_status":"Būsena","observer_planned_backup":"Planuotas dublis","observer_actual_backup":"Faktinis dublis","observer_activated":"Aktyvuotas","observer_completed":"Realiai pavadavo","observer_no_schedule":"Šiam mėnesiui dar nėra paskelbto grafiko.","observer_access_ready":"Tik peržiūros skyriaus prieiga aktyvuota.","claim_help":"Pasirinkite tik savo inicialus ir įveskite jums skirtą vienkartinį beta kvietimo kodą.","invite_code":"Kvietimo kodas","claim":"SUSIETI PASKYRĄ","claim_failed":"Paskyros susieti nepavyko. Patikrinkite inicialus ir kvietimo kodą.","account_unlinked":"Prisijungėte, bet ši paskyra dar nesusieta su rezidentu.",
"app_title":"Rezidentų grafiko sistema","app_caption":"Rezidentų grafiko planavimas, pageidavimai, dubliai ir kontroliuojami pakeitimai.",
"year":"Metai","month":"Mėnuo","weekdays":"Darbo dienos","base_target":"Bazinis pamainų tikslas",
"deadline":"Pageidavimų pateikimo terminas","days_left":"Liko dienų","deadline_today":"Šiandien paskutinė diena.","deadline_passed":"Terminas praėjo prieš {n} d.",
"deadline_future":"Iki termino liko {n} d.","deadline_note":"Kito mėnesio pageidavimai pateikiami iki ankstesnio mėnesio 14 d. 00:00 (13 d. imtinai).",
"senior_dashboard":"Seniūnės skydas","preferences":"Pageidavimai","settings":"Nustatymai","generation":"Sudarymas","schedule":"Grafikas",
"summary":"Suvestinė","transparency":"Skaidrumas","credits_debts":"Kreditai","backups":"Dubliai","swaps":"Apsikeitimai","calendar":"Kalendorius","proof":"Patikra","senior_guide":"Seniūnės vadovas","rules":"Taisyklės",
"my_preferences":"Mano mėnesio pageidavimai","hard_unavailable":"Negaliu dirbti — RESIDENT HARD","hard_help":"Galite pažymėti visą dieną arba tik rytą / popietę. V2.5.107 tai yra privalomas SYSTEM generavimo apribojimas: tuo laiku jūsų skirti negalima. Šis prašymas NĖRA keičiamas į didesnį SOFT išpildymą ar gražesnį fairness. Jei neįmanoma padengti grafiko kartu išlaikant saugą, tikslų krūvį ir visus „Negaliu dirbti“, juodraštis negrąžinamas.","hard_all_day":"Visa diena","hard_morning":"Rytas (08:00–14:00)","hard_afternoon":"Popietė (14:00–20:00)","hard_partial_note":"Jei pažymite rytą arba popietę, kitu paros bloku vis tiek galite būti paskirtas į normalią pamainą arba būti dubliu.","hard_overlap":"Ta pati data negali būti kartu pažymėta kaip visa diena ir dalinis privalomas negalėjimas.",
"soft_free":"Noriu laisvos — pageidavimas","soft_help":"Pasirinkite visą dieną, rytą arba popietę. Sistema stengsis šį pageidavimą įvykdyti, jei tai neprieštarauja aukštesnio prioriteto taisyklėms.","soft_overlap":"Ta pati data negali būti kartu pažymėta kaip visa diena ir dalinis noras būti laisvam.",
"preferred":"Pageidauju dirbti — pageidavimas","preferred_help":"Pasirinkite visą dieną, rytą arba popietę. Darbo dienų pageidavimai yra SOFT. Savaitgalio pageidavimas yra registruojamas auditui, tačiau SYSTEM negali dėl jo skirti žmogui daugiau šeštadienių / sekmadienių nei leidžia ADMIN RAW water-filling.","preferred_overlap":"Ta pati data negali būti kartu pažymėta kaip visa diena ir dalinis pageidavimas dirbti.","vacation":"Atostogos — patvirtintos nedarbo dienos","vacation_help":"Pažymėkite patvirtintas atostogų dienas. Tomis dienomis sistema jūsų neskirs dirbti ar dubliuoti ir proporcingai sumažins mėnesio darbo tikslą, kad atostogos nebūtų laikomos teisingumo trūkumu.","vacation_overlap":"Ta pati diena pažymėta ir kaip atostogos, ir kaip kitas pateisinamas neatvykimas — palikite ją tik viename laukelyje.","note":"Papildomas komentaras","note_ph":"Pvz. po kelių dienų iš eilės nenorėčiau dvigubos pamainos.",
"save":"Išsaugoti","saved":"Išsaugota.","hard_conflict":"Pageidavimas dirbti kertasi su privalomu negalėjimu dirbti tuo pačiu laiku.","soft_conflict":"„Noriu laisvos“ ir „Pageidauju dirbti“ negali būti pasirinkti tam pačiam laikui.",
"all_preferences":"Visų rezidentų pageidavimai","preference_load":"Pageidavimų apimtis","review":"Peržiūrėti","normal":"Įprasta","visibility_flag":"Žyma „Peržiūrėti“ yra tik seniūnės dėmesio indikatorius, ne bauda ir ne automatinis apribojimas.",
"submitted":"Pateikta","updated":"Atnaujinta","hard_dates":"Negaliu dirbti — visa diena","hard_am_dates":"Negaliu dirbti — rytas","hard_pm_dates":"Negaliu dirbti — popietė","soft_dates":"Noriu laisvos — visa diena","soft_am_dates":"Noriu laisvos — rytas","soft_pm_dates":"Noriu laisvos — popietė","preferred_dates":"Pageidauju dirbti — visa diena","preferred_am_dates":"Pageidauju dirbti — rytas","preferred_pm_dates":"Pageidauju dirbti — popietė","comment":"Komentaras",
"settings_title":"Mano paskyros ir darbo pobūdžio nustatymai","short_term":"Trumpalaikiai mėnesio pageidavimai","legal_safety_inputs":"Darbo teisės / poilsio saugos duomenys","justified_absence":"Kitas pateisinamas neatvykimas (liga ar kita patvirtinta priežastis)","justified_absence_help":"Privaloma nedarbo data. Sistema tą dieną neskiria pamainų ir proporcingai perskaičiuoja šio mėnesio vidinį pamainų tikslą. 38 val./sav. norma čia nėra fiksuota programoje.","long_duty":"Jau žinomo ilgo budėjimo už šio grafiko ribų pradžios data","long_duty_help":"Čia žymėkite tik jau žinomą ilgą budėjimą, kurio ši sistema pati neplanuoja (pvz., kitoje darbovietėje ar kitame grafike). Jei ilgą / naktinį budėjimą paskiria pati LSMU schedulerio sistema, jo atskirai įvesti nereikia — sistema po jo automatiškai taiko 24 val. poilsio apsaugą ir neskiria įprastų LSMU pamainų.","labour_hard_summary":"Privalomas saugos sluoksnis: šiame grafike ≤12 val. per darbo dieną; ≥11 val. tarp atskirų darbo dienų; po sistemos paskirto arba ranka įvesto ilgo / naktinio budėjimo taikoma 24 val. poilsio apsauga; bent 1 visiškai laisva diena per kiekvienas slenkančias 7 d.; ≤48 šiame grafike žinomų darbo valandų per bet kurias 7 dienas. Generatorius papildomai taikosi į ~40 val./7 d. ir po dviejų iš eilės 12 val. dienų kitą dieną leidžia tik PM arba poilsį.","labour_scope_note":"Sistema gali tikrinti tik jai žinomą darbo laiką. Kitos darbovietės ar nesuvesti budėjimai turi būti įvertinti atskirai.","long_term":"Ilgalaikiai pasikartojantys pageidavimai","long_term_help":"Šie nustatymai automatiškai taikomi kiekvienam mėnesiui, kol juos pakeisite. Trumpalaikis konkretaus mėnesio SOFT pageidavimas turi pirmenybę prieš priešingą ilgalaikį SOFT pageidavimą; ilgalaikis „Negaliu dirbti“ lieka RESIDENT HARD. Laiko stulpelis (visa diena / rytas / popietė) taikomas RESIDENT HARD; savaitės dienos SOFT pageidavimai yra visos dienos.","weekday_name":"Savaitės diena","recurring_rule":"Pasikartojanti taisyklė","recurring_time":"Laikas","rec_none":"Nėra","rec_hard":"Negaliu dirbti (RESIDENT HARD)","rec_soft":"Noriu laisvos","rec_preferred":"Pageidauju dirbti","save_long_term":"IŠSAUGOTI ILGALAIKIUS PAGEIDAVIMUS","long_term_saved":"Ilgalaikiai pageidavimai išsaugoti.","email":"El. paštas","email_required":"Kiekvienoje paskyroje turi būti galiojantis el. pašto adresas.",
"shift_length_pref":"Pageidaujama darbo dienos trukmė","shift_length_help":"Ilgalaikis privatus darbo pobūdžio pasirinkimas. Sistema stengiasi formuoti darbo dienų trukmę pagal jūsų pasirinkimą, jei tai leidžia privalomos taisyklės, poilsio reikalavimai ir mėnesio darbo krūvis. Onko RO lieka atskira 9 val. pilnos dienos pamaina.","shift_length_any":"Nesvarbu","shift_length_6":"Dažniausiai 6 val.","shift_length_mixed":"Mišriai – tinka ir 6 val., ir 12 val. darbo dienos","shift_length_12":"Dažniausiai 12 val.","weekday_pref":"Darbo dienų pobūdis","weekend_pref":"Savaitgalių pobūdis","holiday_pref":"Švenčių dienos","holiday_pref_help":"Ilgalaikis pasirinkimas oficialioms Lietuvos švenčių dienoms. Sistema pirmiausia skiria šventinį darbą norintiems, po jų — neutraliems, o norinčius ilsėtis naudoja tik kai reikia. Tarp vienodai pasirinkusių žmonių šventinis darbas paskirstomas kuo tolygiau, atsižvelgiant ir į ankstesnius mėnesius.","holiday_rest":"Norėčiau ilsėtis per šventes","holiday_neutral":"Neutralu / nesvarbu","holiday_work":"Norėčiau dirbti per šventes","spread_pref":"Pamainų išdėstymas","avoid_double_shifts":"Jei įmanoma, vengti dvigubų pamainų",
"weekday_help":"−2 = santykinai mažiau darbo dienomis, 0 = nesvarbu, +2 = santykinai daugiau.","weekend_help":"−2 = mažiau savaitgalių, 0 = nesvarbu, +2 = daugiau.",
"spread_help":"−2 = labiau sutelktas grafikas, 0 = nesvarbu, +2 = labiau išsklaidytas.","notifications":"Pranešimai","notifications_on":"Gauti el. pašto priminimus apie pageidavimų pateikimo termino pabaigą","notification_default":"Pagal nutylėjimą pranešimai įjungti.",
"reminder_start":"Asmeninių priminimų pradžia (mėnesio diena)","reminder_help":"Parodo, nuo kurios mėnesio dienos pradėsite gauti asmeninius el. pašto priminimus apie savo artėjantį grafiką, jei dar nebūsite pateikę pageidavimų. Pvz.: „Liko 4 dienos iki pageidavimų pateikimo pabaigos.“ Priminimai sustoja, kai pageidavimai pateikiami arba terminas pasibaigia.","include_backups_calendar":"Rodyti dublius mano .ics kalendoriuje","backup_email_alerts":"Gauti el. laišką, kai seniūnė aktyvuoja mano dublį","phone_optional":"Telefono numeris SMS pranešimams (pasirinktinai)","sms_future":"SMS pranešimai paruošti nustatymuose, bet beta versijoje dar nesiunčiami.","backup_sms_alerts":"Gauti SMS, kai aktyvuojamas dublis","backup_activation":"Dublio aktyvavimas","activate_backup":"KVIESTI DUBLĮ DABAR","backup_activated":"Dublis aktyvuotas.","backup_email_sent":"El. pranešimas dubliui išsiųstas.","backup_email_failed":"Dublis aktyvuotas, bet el. laiško išsiųsti nepavyko.","undo_activation":"ATŠAUKTI DUBLIO AKTYVAVIMĄ","activation_undone":"Dublio aktyvavimas atšauktas.","smtp_admin_note":"Siunčiančio pašto slaptažodis yra vienas bendras sistemos secret ir jo rezidentai neįveda.","settings_saved":"Nustatymai išsaugoti.","backup_bonus":"Dublių bonusai","bonus_balance":"Sukaupti dublių bonusai","bonus_help":"Kai realiai pavaduojate kitą rezidentą, gaunate POILSIO kreditą kaip naudą būsimam mėnesiui. Pavaduotam žmogui jokia skola nesukuriama. RYTAS, POPIETĖ ir NAKTIS apskaitomi atskirai.","use_bonus":"Šį mėnesį panaudoti bonusų","bonus_target_effect":"Per vieną mėnesį galima panaudoti daugiausia 2 dieninius poilsio kreditus iš viso. RYTO ir POPIETĖS kreditai apskaitomi atskirai; NAKTIES kreditas dabartinio PGY1 dieninio targeto nemažina.","bonus_insufficient":"Pasirinkta daugiau bonusų nei turite sukaupę.",
"dashboard_title":"Seniūnės mėnesio kontrolės skydas","completion":"Pageidavimų užpildymas","missing_preferences":"Dar nepateikė","missing_email":"Nenurodė el. pašto","all_complete":"Visi pageidavimus pateikė.",
"email_ready":"El. pašto kanalo konfigūracija rasta","email_not_ready":"El. pašto kanalas dar neparuoštas. Seniūnės lange matysite vieną aiškų taisytiną punktą ir galėsite atlikti kanalo testą.",
"send_reminders":"SIŲSTI ŠIANDIENOS PRIMINIMUS","reminders_result":"Priminimų rezultatas","no_due_reminders":"Šiandien pagal nustatymus priminimų siųsti nereikia.","email_log":"El. laiškų žurnalas",
"generation_title":"Grafiko sudarymas ir paskelbimas","senior_only":"Šią funkciją gali atlikti tik seniūnė.","generate_draft":"GENERUOTI / PERKURTI JUODRAŠTĮ","solver_wait":"Sistema ieško geriausio sprendinio...",
"draft_saved":"Juodraštis sukurtas. Oficialus grafikas dar nepakeistas.","no_solution":"Pagal dabartines kietas taisykles tinkamo grafiko rasti nepavyko.","publish":"PASKELBTI IR UŽRAKINTI",
"published":"Grafikas paskelbtas ir pradinė versija užrakinta.","publication_mail":"Paskelbimo laiškai","no_draft":"Nėra juodraščio, kurį būtų galima paskelbti.","draft_outdated":"Po juodraščio sukūrimo pasikeitė pageidavimai, ilgalaikės taisyklės arba bonusų pasirinkimas. Perkurkite juodraštį prieš paskelbiant.","state":"Būsena","draft":"Juodraštis","published_state":"Paskelbtas","not_created":"Nesukurtas",
"hard_errors":"Privalomų taisyklių klaidos","fairness_score":"Teisingumo rodiklis","monthly_fairness":"Mėnesio teisingumas","cumulative_fairness":"Kaupiamasis teisingumas","fairness_hierarchy":"Grafiko vertinimo hierarchija","fairness_hierarchy_intro":"TRUE ABSOLUTE HARD → RESIDENT HARD „Negaliu dirbti“ (0 pažeidimų privaloma) → ADMIN RAW savaitgalių / šeštadienių / sekmadienių water-filling mažiausiu įmanomu spread (savaitgalio pageidavimai jo neapeina) → Dream Team SP+ŠR+GE kartu CENTRO RO bent kartą kiekvieną darbo savaitę, jei matematiškai įmanoma → SPS RO/SPS UG ir VISŲ kitų postų water-filling mažiausiu įmanomu spread → savaitinio krūvio / recovery fairness → kiti SOFT pageidavimai. Po publikavimo savanoriški swapai gali pakeisti ACTUAL balansą, bet tik po abiejų rezidentų sutikimo ir SP galutinio patvirtinimo.","hard_validity":"ABSOLUTE HARD atitiktis","hard_validity_pass":"0 ABSOLUTE HARD klaidų — tinkama","hard_validity_fail":"Yra ABSOLUTE HARD klaidų — negalima skelbti","fairness_monthly_explain":"Mėnesio teisingumas vertina tik pasirinktą mėnesį. Jis gali būti mažesnis net sąmoningai, jei šis mėnuo taiso ankstesnių mėnesių nelygybę.","fairness_cumulative_explain":"Kaupiamasis teisingumas sumuoja visus sistemoje paskelbtus ankstesnius mėnesius ir šį mėnesį. Tai pagrindinis ilgalaikio grupės lygumo rodiklis.","fairness_100_note":"100% reiškia, kad SYSTEM struktūrinis krūvis yra optimaliai subalansuotas pagal galiojančius HARD apribojimus. Savaitgalio „Pageidauju dirbti“ NEGALI nusipirkti papildomų savaitgalių: SYSTEM vertina raw šeštadienių, sekmadienių ir bendrą savaitgalių krūvį. Po publikavimo savanoriškas ACTUAL swapas gali balansą pakeisti tik po abiejų rezidentų sutikimo ir SP galutinio patvirtinimo.","fairness_formula_month":"Mėnesio formulė: 100 − 18× savaitgalių skirtumas − 7× penktadienių skirtumas − 4× dublių skirtumas − 2× darbo dienų skirtumas.","fairness_formula_cumulative":"Kaupiamojo teisingumo formulė tokia pati, bet kiekvienas skirtumas skaičiuojamas iš visų paskelbtų mėnesių sukauptų sumų.","fairness_breakdown":"Teisingumo išskaidymas","fairness_penalty":"Baudos taškai","fairness_scope":"Apimtis","fairness_metric":"Komponentas","fairness_spread":"Skirtumas (didž.−maž.)","fairness_history":"Teisingumo istorija","fairness_history_help":"Grafike mėnesio teisingumas parodo konkretaus mėnesio lygumą, o kaupiamasis teisingumas — ar sistema laikui bėgant artėja prie lygaus bendro krūvio.","fairness_ledger":"Sistemos teisingumo apskaita","actual_ledger":"Faktinio darbo apskaita","fairness_swap_neutral":"Abipusis savanoriškas apsikeitimas nekeičia teisingumo apskaitos: keičiasi faktinis darbas, bet ne algoritmo paskirstymo vertinimas.","fairness_forced_change":"Pateisinamas post-publication repair (liga, atostogos, force majeure, kritinis SPS pull-down) registruojamas ACTUAL audite, bet NEKEIČIA SYSTEM fairness / spread / postų istorijos. Savanoriški swapai taip pat fairness-neutral; keičiasi tik faktinis grafikas ir retrospektyvinis request satisfaction.","fairness_no_history":"Dar nėra pakankamai paskelbtų mėnesių teisingumo istorijai.","fairness_priority_table":"Ką reiškia hierarchija","fairness_level":"Lygis","fairness_goal":"Tikslas","fairness_interpretation":"Kaip interpretuoti","fairness_hard_goal":"ABSOLUTE HARD: 0 saugos / fizinio neįmanomumo pažeidimų","voluntary_unpopular_goal":"Vykdyti aiškiai savanoriškai pasirinktą nepopuliarų darbą, kai tai nepažeidžia HARD / saugos / coverage. Savaitgalio savanoriškas darbas gali būti virš RAW water-fill, nes fairness skaičiuoja likusį nesavanorišką krūvį.","voluntary_unpopular_explain":"Savaitgalio „Pageidauju dirbti“ yra savanoriškas nepopuliaraus darbo pasirinkimas: generatorius pirmiausia stengiasi jį įvykdyti, jei leidžia HARD / poilsis / coverage. Tokia savanoriška pamaina gali padidinti RAW šeštadienių ar sekmadienių spread, o fairness 0–1 taikomas likusiam nesavanoriškam krūviui. Po publikavimo ACTUAL swapai gali balansą keisti papildomai.","other_preferences_goal":"SOFT: griežtai SOFT-1 → SOFT-2 → SOFT-3; kiekviename range pirmiausia horizontalus water-filling, po to likęs įmanomas išpildymas","other_preferences_explain":"SOFT pageidavimai optimizuojami tik po TRUE ABSOLUTE HARD, SYSTEM HARD postų lygybės, RESIDENT HARD ir likusio workload/fatigue fairness; aukštesnis SOFT rangas užrakinamas prieš pereinant į žemesnį.","fairness_cumulative_goal":"Antrinis ilgalaikis tikslas: po gero einamojo mėnesio balanso taisyti ankstesnės SYSTEM istorijos likutinę nelygybę","fairness_monthly_goal":"SPS UG ir penktadieniai SYSTEM grafike išlieka struktūriškai water-fill'inami. Savaitgaliuose aiškiai savanoriškai pageidautos pamainos gali būti virš RAW 0–1, o likęs nesavanoriškas SPS RO / šeštadienių / sekmadienių krūvis turi 0–1 water-fill. Visos ne-Onko darbo vietos taip pat water-fill'inamos iki raw spread 0–1 prieš SOFT; platesnis postų koridorius leidžiamas tik jei siauresnis įrodytas neįmanomas.","preference_avg":"Vidutinis pageidavimų išpildymas","weekend_spread":"Savaitgalių skirtumas",
"published_schedule":"Galiojantis paskelbtas grafikas","not_published":"Šiam mėnesiui oficialus grafikas dar nepaskelbtas.","colors":"Nuolatinės žmonių spalvos","download_xlsx":"ATSISIŲSTI SPALVOTĄ GRAFIKĄ (.xlsx)","download_csv":"Atsisiųsti duomenų sąrašą (.csv)",
"summary_title":"Žmonių suvestinė","frozen_fairness":"Paskelbimo teisingumas","current_after_changes":"Dabartinė būsena po savanoriškų pakeitimų","fairness_frozen_note":"Sistemos teisingumo, postų spread ir future catch-up apskaita fiksuojama pagal paskirstymą paskelbimo momentu. Abipusiai savanoriški swapai ir pateisinami post-publication repair (liga, atostogos, force majeure, SPS pull-down) keičia ACTUAL grafiką, bet NEĮEINA į fairness / spread / istorijos. Faktinis darbas ir retrospektyvinis request satisfaction gali būti rodomi atskirai.",
"person":"Žmogus","name":"Vardas","target":"Tikslas","workload":"Krūvis","weekday_assignments":"Darbo dienų paskyrimai","weekday_days":"Atskiros darbo dienos","weekend_assignments":"Savaitgalio pamainos","saturday_assignments":"Šeštadienio pamainos","sunday_assignments":"Sekmadienio pamainos","prior_weekends":"Ankstesni savaitgaliai","cumulative_weekends":"Sukaupti savaitgaliai","fridays":"Penktadieniai","double_shifts":"12h darbo dienos (AM+PM)","max_consecutive":"Daugiausia dienų iš eilės","max_rolling7_hours":"Daugiausia val. per 7 d.","max_calendar_week_hours":"Daugiausia val. kalendorinę savaitę","free_days":"Laisvos dienos","preference_score":"Bendras prašymų išpildymas, %","planned_backups":"AUTO pavadavimai / dubliai","effective_backups":"Galiojantys pavadavimai / dubliai",
"transparency_title":"Skaidrumas","validity_heading":"1. Privalomų taisyklių patikra","validity_text":"0 klaidų reiškia, kad paskelbta bazinė versija nepažeidė nė vienos privalomos taisyklės. Tai galiojimo, o ne teisingumo procentas.",
"fairness_heading":"2. Grupės teisingumas","fairness_text":"Teisingumas skiriamas į mėnesio ir kaupiamąjį. Kaupiamasis yra pagrindinis ilgalaikis sistemos lygumo rodiklis; mėnesio teisingumas padeda suprasti konkretų mėnesį.","fair_formula":"Abiejų rodiklių formulė vienoda: 100 − 18×savaitgalių skirtumas − 7×penktadienių skirtumas − 4×dublių skirtumas − 2×darbo dienų skirtumas. Skiriasi tik apimtis: vienas mėnuo arba visų paskelbtų mėnesių suma.",
"metric_weekend":"Savaitgalių skirtumas","metric_friday":"Penktadienių skirtumas","metric_double":"Dvigubų pamainų skirtumas","metric_weekday":"Darbo dienų skirtumas",
"personal_vs_group":"Asmeninis pageidavimų išpildymas ir grupės teisingumas","balance_ratio":"Balanso santykis","ratio_help":"Balanso santykis = mažesnis procentas / didesnis procentas. 1,00 reiškia, kad abu rodikliai yra vienodo lygio; jis neparodo absoliučios kokybės.",
"baseline_personal":"Asmeninis paskelbimo momentu","current_personal":"Asmeninis dabar","not_applicable":"Netaikoma","all_resident_scores":"Visų rezidentų pageidavimų išpildymas",
"backup_title":"Dubliai / pavadavimai","backup_self_select":"Pasirink mano mėnesio dublių slotus","backup_self_select_help":"Rezervuojami privalomo dengimo slotai pagal poziciją: SPS RO bet kurią dieną / bloką, SPS UG bet kurią dieną / bloką, Centro UG 120 rytas ir Onko RO pilna 9 val. pamaina. CENTRO RO dengiama automatiškai best-effort. Galima pasirinkti kelis slotus; rezervuotas dublis blokuoja persidengiančią normalią pamainą.","backup_claim_deadline":"Dublių pasirinkimo terminas","backup_claim_saved":"Dublių slotai rezervuoti.","backup_claim_released":"Dublio pasirinkimas atšauktas.","backup_claim_taken":"Šį slotą ką tik pasirinko kitas rezidentas. Pasirinkite kitą.","backup_claim_locked":"Pasirinkimo terminas pasibaigė arba grafikas jau paskelbtas. Toliau dublio slotai keičiami per Apsikeitimai.","backup_claim_missing_penalty":"Dar nepasirinkote nė vieno dublio sloto. Sistema vis tiek AUTO paskirs dublius fairness-first water-filling principu; savitarna negali sukurti nelygaus dublių krūvio.","backup_claim_yours":"Jūsų rezervuoti dubliai","backup_claim_board":"Dublių rezervacijos","backup_claim_free":"Laisva","backup_claim_auto_queue":"Automatinio paskyrimo prioritetinė eilė","backup_claim_auto_queue_help":"Rezidentai, kurie nepasirinko nė vieno dublio sloto, gali patekti į automatinio paskyrimo eilę pirmiau, tačiau pasirinkimai yra tik tie-break. AUTO dubliai water-fill'inami pagal bendrą dublių krūvį ir niekada nepažeidžia „Negaliu dirbti“ / RESIDENT HARD ar ABSOLUTE HARD.","release_backup_claim":"ATŠAUKTI MANO PASIRINKIMĄ","backup_claim_reminder_kind":"Dublių pasirinkimo priminimas","backup_swap_title":"Dublių apsikeitimai","backup_swap_help":"Po grafiko paskelbimo galite pasiūlyti apsikeisti bet kuriuo suplanuotu privalomu dublio slotu. Apsikeitimas taikomas tik jei abu lieka tinkami naujiems slotams.","my_backup_duty":"Mano dublio vieta","their_backup_duty":"Kito rezidento dublio vieta","request_backup_swap":"SIŪLYTI DUBLIŲ APSIKEITIMĄ","backup_swap_sent":"Dublio apsikeitimo pasiūlymas išsiųstas.","backup_swap_invalid":"Šio apsikeitimo negalima atlikti, nes bent vienas rezidentas nebūtų tinkamas naujam dublio slotui.","backup_swap_accepted":"Dublių apsikeitimas patvirtintas ir pritaikytas.","backup_swap_rejected":"Dublio apsikeitimas atmestas.","backup_definition":"Privalomas vardinis dublis pagal poziciją: SPS RO bet kurią dieną / bloką, SPS UG bet kurią dieną / bloką, Centro UG 120 rytas ir Onko RO pilna 9 val. pamaina. CENTRO RO dengiama kuo plačiau pagal likusią saugią talpą; jos nepadengimas publikavimo neblokuoja. ABSOLUTE HARD ir persidengianti normali pamaina niekada neleidžiami.",
"my_backup_schedule":"Mano dublių grafikas","no_backups":"Šiam žmogui šį mėnesį dublio pareigų nėra.","covered_assignment":"Dubliuojamas žmogus ir jo grafikas","covered_person":"Dubliuojamas žmogus","covered_schedule":"Dubliuojama pamaina","planned_backup":"Planuotas dublis","actual_backup":"Faktinis dublis","effective_backup":"Galiojantis dublis","backup_note":"Pastaba",
"manage_backups":"Seniūnės dublių kontrolė","backup_coverage":"Dublių padengimas","working_person_days":"Privalomų padengti pamainų","covered_person_days":"Pamainų su vardiniu dubliu","backup_complete":"Visos privalomai dengiamos pamainos turi konkretų vardinį dublį.","backup_incomplete":"Bent viena privalomai dengiama pamaina neturi tinkamo dublio. Tokio grafiko negalima skelbti.","resync_backups":"ATNAUJINTI DUBLIUS PAGAL GALIOJANTĮ GRAFIKĄ","backup_synced":"Dubliai automatiškai perskaičiuoti pagal galiojantį grafiką.","backup_capacity_block":"Juodraščio negalima paskelbti, jei bent vienai privalomai dengiama pamainai nėra nė vieno tuo metu laisvo ir ABSOLUTE-HARD saugaus žmogaus. CENTRO RO best-effort padengimo trūkumas publikavimo neblokuoja. RESIDENT HARD / „Negaliu dirbti“ yra privalomas: SYSTEM juodraštis su bent vienu tokiu pažeidimu negrąžinamas ir negali būti publikuojamas.",
"cover_credit_type":"Automatiškai nustatoma pavadavimo rūšis","cover_6h":"RYTAS 08:00–14:00 = 6 val.","cover_12h":"RYTAS + POPIETĖ = du atskiri 6 val. įvykiai","cover_night12h":"NAKTIS 20:00–08:00 = 12 val.","cover_credit_note":"Rūšies seniūnė nepasirenka ranka: ji nustatoma pagal konkrečią realiai dubliuotą vietą. Jei žmogus realiai pavaduoja ir RYTĄ, ir POPIETĘ, registruojami du atskiri 6 val. pavadavimai. NAKTIS bus 12 val. įvykis, kai sistemoje atsiras naktinės pamainos.","actual_override":"Faktinio dublio rankinis pakeitimas","mark_backup_completed":"PAŽYMĖTI REALIAI ĮVYKDYTĄ PAVADAVIMĄ","backup_completed":"Realus pavadavimas užregistruotas. Pavaduojančiam rezidentui suteiktas poilsio kreditas; pavaduotam žmogui jokia skola nesukuriama.","undo_backup_completed":"ATŠAUKTI REALŲ PAVADAVIMĄ IR JO KREDITĄ","backup_completion_undone":"Realus pavadavimas ir pavaduojančiam suteiktas kreditas atšaukti.","completed_backup":"Įvykdyta","credit_balances":"Poilsio kreditai","bonus_units":"Kreditai","bonus_shift_value":"Galimas pamainų sumažinimas","rest_credit_bank":"Poilsio kreditų bankas","credit_type":"Kredito rūšis","credit_am":"RYTAS — 6 val.","credit_pm":"POPIETĖ — 6 val.","credit_night":"NAKTIS — 12 val.","use_credit_am":"Panaudoti RYTO poilsio kreditų","use_credit_pm":"Panaudoti POPIETĖS poilsio kreditų","credit_month_cap":"Per mėnesį galima panaudoti daugiausia 2 dieninius poilsio kreditus iš viso.","night_bank_only":"NAKTIES kreditai kol kas tik kaupiami; jie negali būti panaudoti dabartiniam dieniniam PGY1 targetui.","netting_explain":"Kreditas yra vienpusė nauda realiai pavaduojančiam rezidentui. Pavaduotam žmogui skola nesukuriama ir jo turimi poilsio kreditai dėl pavadavimo neatimami.","cover_effect_rest":"Pavaduojančiam rezidentui suteiktas naujas poilsio kreditas.","max_credit_error":"Vienam mėnesiui galima pasirinkti daugiausia 2 dieninius poilsio kreditus iš viso.","backup_record":"Dublio įrašas","record_actual":"ĮRAŠYTI FAKTINĮ DUBLĮ","actual_saved":"Faktinis dublis įrašytas. Dublių statistika naudos šį žmogų.","clear_actual":"GRĄŽINTI PLANUOTĄ DUBLĮ","actual_cleared":"Faktinis pakeitimas pašalintas; vėl galioja planuotas dublis.","no_eligible_backup":"Šiai pamainai nėra tinkamo žmogaus: kandidatas turi būti laisvas tuo pačiu laiko bloku ir nepažeisti ABSOLUTE HARD. RESIDENT HARD pirmiausia saugomas, o tik jei griežto kandidato nėra — taikomas minimalus ir kuo teisingiau paskirstytas praradimas.",
"swap_title":"Savanoriški apsikeitimai","swap_note":"Savanoriški apsikeitimai keičia faktinį grafiką tik po abiejų žmonių sutikimo ir privalomų taisyklių patikros. Tikslus mėnesio targetas ir Onko poros 0/2/4/... lieka ABSOLUTE HARD ir negali būti apeitos swapu; consecutive Onko gali būti tik aiškiai patvirtinama ACK pasekmė. Sistemos teisingumo apskaita lieka tokia, kokią paskyrė sistema prieš apsikeitimą.","repair_title":"Neplanuoti pakeitimai po publikavimo","repair_help":"Liga, atostogos ar kita pateisinama / force majeure priežastis keičia tik FAKTINĮ grafiką. Paskelbimo momento SISTEMOS teisingumo bazė ir fairness_history nekeičiami; sergančiam ar neatvykstančiam rezidentui nesukuriama teisingumo skola. V2.5.57: jei neatvykstantis žmogus dengė SPS RO / SPS UG, kritinis postas išlaikomas pirmiausia perkeliant tos pačios pamainos rezidentą iš žemesnės hierarchijos NEPRIVALOMO posto; donorinis optional postas gali likti tuščias. Tik jei saugaus donorinio perkėlimo nėra, naudojamas tame laiko bloke laisvo rezidento fallback. ABSOLUTE sauga, overlap ir mandatory coverage lieka privalomi. Pull-down ir kiti pateisinami repair NEĮEINA į SYSTEM fairness, postų spread ar future catch-up skaičiavimus.","repair_assignment":"Keičiama pamaina","repair_replacement":"Pavaduojantis rezidentas","repair_reason":"Priežasties kategorija","repair_reason_sickness":"Liga","repair_reason_leave":"Atostogos","repair_reason_approved":"Kitas pateisinamas neatvykimas","repair_reason_force":"Force majeure / nenumatytas įvykis","repair_note":"Vidinė pastaba (nebūtina)","apply_repair":"PRITAIKYTI NEPLANUOTĄ PAKEITIMĄ","repair_applied":"Pakeitimas pritaikytas FAKTINIAM grafikui. SYSTEM fairness, postų spread ir future catch-up lieka pagal publikavimo bazę; repair į juos nepridedamas. Faktinis grafikas ir request satisfaction perskaičiuoti.","repair_invalid":"Šio pakeitimo negalima taikyti dėl operacinės saugos / HARD taisyklės","repair_no_candidate":"Šiai pamainai nėra saugiai tinkamo pavaduojančio rezidento.","repair_history":"Neplanuotų pakeitimų istorija","repair_load":"Papildoma repair našta šį mėnesį","repair_load_help":"Tai tik operacinis audito skaitiklis. Jis NENAUDOJAMAS fairness, postų spread, future catch-up ar ateities kompensacijai. Kritinis pull-down reiškia tos pačios suplanuotos darbo pamainos vietos pakeitimą, o ne papildomą fairness naštą.","repair_fairness_neutral":"FAIRNESS NEUTRAL","repair_from":"Negalintis dirbti","repair_to":"Pavadavo","repair_date":"Data / pamaina","my_assignment":"Mano pamaina","their_assignment":"Kito žmogaus pamaina","request_swap":"SIŪLYTI APSIKEITIMĄ","request_sent":"Apsikeitimo pasiūlymas išsiųstas.","incoming":"Gauti pasiūlymai","accept":"PRIIMTI","reject":"ATMESTI","accepted":"Apsikeitimas patvirtintas ir pritaikytas.","accepted_pending":"Apsikeitimą patvirtino abu žmonės. Beta versijoje seniūnė atliks galutinę privalomų taisyklių patikrą ir pritaikys pakeitimą.","finalize_swap":"PRITAIKYTI PATVIRTINTĄ APSIKEITIMĄ","swap_applied":"Apsikeitimas pritaikytas, privalomos taisyklės patikrintos, dubliai perskaičiuoti.","swap_finalize_failed":"Apsikeitimo pritaikyti nepavyko, nes po galutinės patikros būtų pažeista privaloma taisyklė.","rejected":"Apsikeitimas atmestas.","hard_reject":"Apsikeitimas atmestas, nes pažeistų privalomą taisyklę.","history":"Apsikeitimų istorija","pending":"Laukiama","approved":"Patvirtinta","rejected_status":"Atmesta",
"calendar_title":"Mano grafikas kalendoriui","calendar_help":"Galite atsisiųsti vienkartinį .ics failą arba vieną kartą užsiprenumeruoti privačią kalendoriaus nuorodą. Prenumerata atnaujinama paskelbus naują grafiką ir po svarbių faktinio grafiko pakeitimų. Jei Nustatymuose įjungti dubliai, jie taip pat įtraukiami.","download_ics":"ATSISIŲSTI MANO GRAFIKĄ (.ics)","calendar_feed":"Privati kalendoriaus prenumeratos nuoroda","calendar_feed_private":"Ši nuoroda veikia kaip slaptažodis į jūsų grafiką — nesidalinkite ja. Ji turi atsitiktinį ilgą kodą ir nėra rodoma kitiems rezidentams.","calendar_google":"GOOGLE CALENDAR","calendar_apple":"APPLE CALENDAR","calendar_other":"OUTLOOK CALENDAR","calendar_google_help":"Google Calendar kompiuteryje: Add other calendars → From URL → įklijuokite žemiau esančią privačią nuorodą. Tai daroma vieną kartą.","calendar_apple_help":"Apple Calendar gali užsiprenumeruoti nuorodą tiesiogiai. Paspaudus mygtuką turėtų atsidaryti Calendar prenumeratos langas.","calendar_other_help":"Outlook gali prenumeruoti tą pačią privačią iCalendar nuorodą (Add calendar → Subscribe from web). Jei naudojate kitą programą, naudokite .ics failą arba prenumeratos nuorodą, jei ji palaikoma.",
"proof_title":"Mano grafiko patikra","proof_intro":"Vizuali patikra parodo, kas tiksliai atitiko jūsų poreikius ir kur liko neatitikimų. Čia nerodomas programinis kodas — tik galutiniai rezultatai.","matches":"ATITINKA","partial":"DALINAI","mismatch":"NEATITINKA","baseline":"Paskelbimo momentu","current":"Dabar","hard_ok":"Privalomas negalėjimas dirbti išlaikytas","hard_bad":"Privaloma taisyklė pažeista","soft_off_ok":"Norėtos laisvos dienos","preferred_ok":"Pageidautos darbo dienos","workload_ok":"Mėnesio krūvio tikslas","style_component":"Darbo pobūdžio kriterijus","missed_dates":"Neatitikusios datos","criterion":"Kriterijus","result":"Rezultatas","score":"Išpildymas","explanation":"Paaiškinimas","proof_all_good":"Pagal pateiktus duomenis privalomos taisyklės išlaikytos, o aktyvūs pageidavimai neturi ryškių neatitikimų.","proof_soft_issues":"Privalomos taisyklės išlaikytos, tačiau ne visi pageidavimai buvo įvykdyti.","proof_hard_issue":"Aptiktas privalomos taisyklės neatitikimas — grafiką būtina peržiūrėti.","no_active_preferences":"Šiai kategorijai aktyvaus pageidavimo nepateikėte.","swap_suggestion":"Jei privalomos taisyklės nepažeistos, bet pageidavimas liko neįvykdytas, galima ieškoti savanoriško sprendimo Apsikeitimų lange.",
"rules_title":"Grafiko taisyklės","read":"Skaityti","edit":"Redaguoti","save_rules":"IŠSAUGOTI TAISYKLIŲ PAKEITIMUS","rules_saved":"Taisyklės atnaujintos.","edit_senior_only":"Taisykles redaguoti gali tik seniūnė.",
"yes":"TAIP","no":"NE","reminder_kind":"Priminimas","publication_kind":"Grafiko paskelbimas","date":"Data","day":"Diena","time":"Laikas","department":"Padalinys","shift":"Pamaina","morning":"Rytas","afternoon":"Popietė","full_day":"Pilna diena","status":"Statusas","details":"Informacija","sent":"Išsiųsta","failed":"Nepavyko","skipped":"Praleista"
},
"EN": {
"language":"Language","user":"User","profile":"Profile","resident_profile":"Resident profile","senior_profile":"Senior scheduler profile","resident_pin":"Personal PIN","admin_pin":"Senior scheduler PIN","local_resident":"Local test mode: personal PINs are not configured.","local_senior":"Local test mode: senior functions are unlocked only for the senior account.","bad_pin":"Incorrect PIN.",
"login_title":"Sign in","login":"SIGN IN","logout":"SIGN OUT","signup":"CREATE ACCOUNT","signup_title":"First registration","password":"Password","password_repeat":"Repeat password","auth_email":"Email","auth_invalid":"Sign-in failed. Check your email and password.","signup_sent":"Account created. If email confirmation is enabled, confirm the email and then sign in.","signup_password_mismatch":"Passwords do not match.","claim_title":"Link this account","resident_claim_tab":"Resident","observer_claim_tab":"Department administrator / observer","observer_claim_help":"This account is for departmental oversight only. It is read-only: it can view the published and current schedules, swaps, backups, fairness and audit history, but cannot change anything.","observer_invite_code":"Department observer invite code","observer_claim":"ACTIVATE READ-ONLY ACCESS","observer_read_only":"TIK PERŽIŪRA","observer_role":"Department observer","observer_portal":"Department schedule oversight","observer_overview":"Overview","observer_schedule":"Schedules","observer_changes":"Change log","observer_fairness":"Teisingumas","observer_backups":"Backups","observer_rules":"Rules","observer_scope_note":"This account cannot generate, publish, approve, edit or undo schedules or swaps.","observer_privacy_note":"Private resident preferences, HARD-unavailable dates, personal notes, emails and account settings are not shown.","observer_baseline_schedule":"SYSTEM baseline — at publication","observer_actual_schedule":"ACTUAL schedule — now","observer_change_count":"Changed normal assignments","observer_normal_swaps":"Normal-shift swaps","observer_backup_swaps":"Backup swaps","observer_pending_swaps":"Pending","observer_approved_swaps":"Approved","observer_rejected_swaps":"Rejected","observer_no_changes":"There are no normal-assignment changes after publication.","observer_from":"From","observer_to":"Now","observer_change_log_help":"The SYSTEM baseline remains the fairness ledger. The ACTUAL schedule shows the real situation after bilateral voluntary swaps.","observer_backup_status":"Status","observer_planned_backup":"Planned backup","observer_actual_backup":"Actual backup","observer_activated":"Activated","observer_completed":"Actually covered","observer_no_schedule":"No schedule has been published for this month yet.","observer_access_ready":"Read-only department access activated.","claim_help":"Choose only your own initials and enter the one-time beta invite code provided to you.","invite_code":"Invite code","claim":"LINK ACCOUNT","claim_failed":"Could not link the account. Check the initials and invite code.","account_unlinked":"You are signed in, but this account is not linked to a resident yet.",
"app_title":"Resident scheduling system","app_caption":"Hard rules → transparency → soft preferences → controlled changes.","year":"Year","month":"Month","weekdays":"Weekdays","base_target":"Base shift target","deadline":"Preference deadline","days_left":"Days remaining","deadline_today":"Today is the deadline.","deadline_passed":"Deadline passed {n} day(s) ago.","deadline_future":"{n} day(s) remain until the deadline.","deadline_note":"Next month's preferences are due by 00:00 on the 14th of the preceding month (the 13th is the last full day).",
"senior_dashboard":"Senior dashboard","preferences":"Preferences","settings":"Settings","generation":"Generation","schedule":"Schedule","summary":"Summary","transparency":"Transparency","credits_debts":"Credits","backups":"Backups","swaps":"Swaps","calendar":"Calendar","proof":"Proof","senior_guide":"Senior guide","rules":"Rules",
"my_preferences":"My monthly preferences","hard_unavailable":"Unavailable — RESIDENT HARD","hard_help":"You may mark the whole day or only the morning / afternoon. In V2.5.107 this is a mandatory SYSTEM-generation constraint: you cannot be assigned in that blocked time. It is never traded for higher SOFT satisfaction or prettier fairness. If coverage, safety, exact workload and every Unavailable block cannot coexist, no draft is returned.","hard_all_day":"Whole day","hard_morning":"Morning (08:00–14:00)","hard_afternoon":"Afternoon (14:00–20:00)","hard_partial_note":"If only morning or afternoon is marked, you may still receive a normal shift or backup duty in the other time block.","hard_overlap":"The same date cannot be marked as both whole-day and partial required unavailability.","soft_free":"Would like time off — preference","soft_help":"Choose whole day, morning, or afternoon. The system tries to honor this unless a higher-priority rule prevents it.","soft_overlap":"The same date cannot be marked as both whole-day and partial requested time off.","preferred":"Prefer to work — preference","preferred_help":"Choose whole day, morning, or afternoon. Voluntary unpopular work is prioritized when labour-law and rest-safety rules allow it.","preferred_overlap":"The same date cannot be marked as both whole-day and partial preferred work.","vacation":"Approved vacation / leave days","vacation_help":"Select approved vacation days. The scheduler will not assign work or backup on those days and will proportionally reduce the monthly workload target so approved leave is not treated as a fairness deficit.","vacation_overlap":"The same day is entered as both vacation and another justified absence; keep it in only one field.","note":"Additional note","note_ph":"Example: I would prefer not to have a double shift after several consecutive days.","save":"Save","saved":"Saved.","hard_conflict":"A preferred-work request conflicts with required unavailability in the same time block.","soft_conflict":"Requested time off and preferred work cannot overlap in the same time block.",
"all_preferences":"All resident preferences","preference_load":"Preference volume","review":"Review","normal":"Normal","visibility_flag":"The Review flag is only a visibility prompt for the senior scheduler; it is not a penalty or automatic restriction.","submitted":"Submitted","updated":"Updated","hard_dates":"Unavailable — whole day","hard_am_dates":"Unavailable — morning","hard_pm_dates":"Unavailable — afternoon","soft_dates":"Time off — whole day","soft_am_dates":"Time off — morning","soft_pm_dates":"Time off — afternoon","preferred_dates":"Prefer work — whole day","preferred_am_dates":"Prefer work — morning","preferred_pm_dates":"Prefer work — afternoon","comment":"Comment",
"settings_title":"My account and work-style settings","short_term":"Short-term monthly preferences","legal_safety_inputs":"Labour-law / rest-safety inputs","justified_absence":"Other justified absence (sickness or another approved reason)","justified_absence_help":"Hard no-work date. The scheduler assigns no shifts that day and proportionally recalculates this month’s internal shift target. A 38h/week norm is not hard-coded.","long_duty":"Start date of a long duty (>12–24h or 24h)","long_duty_help":"Entering the start date of a real long duty blocks the entire following calendar day from normal assignments as a conservative ≥24h rest safeguard.","labour_hard_summary":"Hard safety layer: ≤12h per workday in this schedule; ≥11h between separate workdays; after a marked long/night duty the next day has no normal shifts; at least one fully free day in every rolling 7 days; ≤48 known scheduled hours in any rolling 7 days. The generator additionally targets ~40h/7d and after two consecutive 12h double days allows only PM or rest on the next day.","labour_scope_note":"The system can validate only work it knows about. Other employers or unentered duties must be assessed separately.","long_term":"Long-term recurring preferences","long_term_help":"These rules are applied automatically every month until you change them. A month-specific SOFT preference overrides an opposite recurring SOFT preference; recurring “Unavailable” remains RESIDENT HARD. The time column applies to RESIDENT HARD; recurring SOFT weekday preferences are whole-day.","weekday_name":"Weekday","recurring_rule":"Recurring rule","recurring_time":"Time","rec_none":"None","rec_hard":"Unavailable (RESIDENT HARD)","rec_soft":"Would like the day off","rec_preferred":"Prefer to work","save_long_term":"SAVE LONG-TERM PREFERENCES","long_term_saved":"Long-term preferences saved.","email":"Email","email_required":"Every account should contain a valid email address.","shift_length_pref":"Preferred workday length","shift_length_help":"Persistent private work-style preference. The system tries to shape your workday length according to this choice when mandatory rules, rest requirements, and monthly workload allow it. Onko RO remains a separate 9-hour full-day shift.","shift_length_any":"No preference","shift_length_6":"Mostly 6 hours","shift_length_mixed":"Mixed – both 6-hour and 12-hour workdays are fine","shift_length_12":"Mostly 12 hours","weekday_pref":"Weekday pattern","weekend_pref":"Weekend pattern","holiday_pref":"Public holidays","holiday_pref_help":"Long-term preference for official Lithuanian public holidays. Holiday duty is offered first to residents who prefer working holidays, then to neutral residents, while residents who prefer rest are used only when needed. Among residents with the same choice, holiday work is distributed as evenly as possible using current and prior months.","holiday_rest":"Prefer to rest on holidays","holiday_neutral":"Neutral / no preference","holiday_work":"Prefer to work on holidays","spread_pref":"Shift distribution","avoid_double_shifts":"Avoid double shifts when possible","weekday_help":"−2 = relatively fewer weekdays, 0 = neutral, +2 = relatively more.","weekend_help":"−2 = fewer weekends, 0 = neutral, +2 = more.","spread_help":"−2 = more clustered, 0 = neutral, +2 = more dispersed.","notifications":"Notifications","notifications_on":"Receive email reminders about the preference-submission deadline","notification_default":"Notifications are on by default.","reminder_start":"Start personal reminders on day of month","reminder_help":"Choose the day of the month from which you want to receive personal email reminders about your upcoming schedule if your preferences are still missing. Example: “4 days left until preference submission closes.” Reminders stop after you submit or the deadline passes.","include_backups_calendar":"Include backup duties in my .ics calendar","backup_email_alerts":"Email me when the senior activates my backup duty","phone_optional":"Phone number for SMS alerts (optional)","sms_future":"SMS preferences are prepared, but SMS delivery is not enabled in this beta yet.","backup_sms_alerts":"Send me an SMS when my backup duty is activated","backup_activation":"Backup activation","activate_backup":"CALL BACKUP NOW","backup_activated":"Backup activated.","backup_email_sent":"Backup alert email sent.","backup_email_failed":"Backup activated, but the alert email could not be sent.","undo_activation":"UNDO BACKUP ACTIVATION","activation_undone":"Backup activation undone.","smtp_admin_note":"The sending-mailbox password is one shared system secret; residents never enter it.","settings_saved":"Settings saved.","backup_bonus":"Backup bonuses","bonus_balance":"Available backup bonuses","bonus_help":"When you actually cover another resident, you earn a REST credit as a future benefit. The covered resident receives no debt. MORNING, AFTERNOON and NIGHT are tracked separately.","use_bonus":"Use bonuses this month","bonus_target_effect":"At most 2 daytime rest credits in total may be used in one month. MORNING and AFTERNOON are tracked separately; NIGHT credits cannot reduce the current PGY1 daytime target.","bonus_insufficient":"You selected more bonuses than you currently have.",
"dashboard_title":"Senior monthly control dashboard","completion":"Preference completion","missing_preferences":"Not submitted","missing_email":"Missing email","all_complete":"Everyone submitted preferences.","email_ready":"Email channel configuration found","email_not_ready":"Email channel is not ready yet. The Senior control shows one clear fix and provides a channel test.","send_reminders":"SEND TODAY'S REMINDERS","reminders_result":"Reminder result","no_due_reminders":"No reminders are due today under the current settings.","email_log":"Email log",
"generation_title":"Schedule generation and publication","senior_only":"Only the senior scheduler can use this function.","generate_draft":"GENERATE / REGENERATE DRAFT","solver_wait":"The system is searching for the best solution...","draft_saved":"Draft created. The official schedule has not changed.","no_solution":"No feasible schedule could be found under the current hard rules.","publish":"PUBLISH AND LOCK","published":"Schedule published and baseline locked.","publication_mail":"Publication emails","no_draft":"There is no draft to publish.","draft_outdated":"Preferences, recurring rules, or bonus selection changed after the draft was generated. Regenerate the draft before publishing.","state":"Status","draft":"Draft","published_state":"Paskelbtas","not_created":"Not created","hard_errors":"Privalomų taisyklių klaidos","fairness_score":"Teisingumo įvertis","monthly_fairness":"Mėnesio teisingumas","cumulative_fairness":"Kaupiamasis teisingumas","fairness_hierarchy":"Grafiko vertinimo hierarchija","fairness_hierarchy_intro":"Prioritetų tvarka: absoliučios saugos ir darbo taisyklės → rezidentų „Negaliu dirbti“ (0 pažeidimų privaloma) → kuo lygesnis SPS RO, SPS UG ir savaitgalių paskirstymas → poilsis ir darbo krūvis → švenčių pasirinkimai → visų ne-Onko darbo vietų struktūrinis water-filling → kiti pageidavimai. Didesnis pateiktų pageidavimų skaičius nesuteikia didesnio prioriteto.","hard_validity":"Privalomų taisyklių atitiktis","hard_validity_pass":"0 privalomų taisyklių klaidų — tinkama","hard_validity_fail":"Yra privalomų taisyklių klaidų — skelbti negalima","fairness_monthly_explain":"Mėnesio teisingumas vertina tik pasirinktą mėnesį. Jis gali būti sąmoningai mažesnis, kai taisomas ankstesniais mėnesiais susikaupęs netolygumas.","fairness_cumulative_explain":"Kaupiamasis teisingumas apima visus anksčiau paskelbtus mėnesius ir pasirinktą mėnesį. Tai pagrindinis ilgalaikio grupės balanso rodiklis.","fairness_100_note":"100 % reiškia, kad sistemos paskirtas nesavanoriškas nepopuliarus krūvis ir kiti teisingumo komponentai yra optimaliai subalansuoti. Aiškiai savanoriškai pasirinkta penktadienio ar savaitgalio darbo data pati savaime teisingumo balo nemažina.","fairness_formula_month":"Mėnesio formulė: 100 − 18× savaitgalių skirtumas − 7× penktadienių skirtumas − 4× dvigubų pamainų skirtumas − 2× darbo dienų skirtumas.","fairness_formula_cumulative":"Kaupiamoji formulė tokia pati, tačiau kiekvienas skirtumas skaičiuojamas iš visų paskelbtų mėnesių sukauptų sumų.","fairness_breakdown":"Teisingumo išskaidymas","fairness_penalty":"Baudos taškai","fairness_scope":"Apimtis","fairness_metric":"Komponentas","fairness_spread":"Skirtumas (didž.−maž.)","fairness_history":"Teisingumo istorija","fairness_history_help":"Mėnesio teisingumas rodo balansą vieno mėnesio viduje; kaupiamasis teisingumas rodo, ar ilgainiui sistema artėja prie vienodo bendro krūvio.","fairness_ledger":"Sistemos teisingumo apskaita","actual_ledger":"Faktinio darbo apskaita","fairness_swap_neutral":"Abipusis savanoriškas apsikeitimas nekeičia sistemos teisingumo apskaitos: faktinis darbas pasikeičia, tačiau algoritmo paskirstymo balansas lieka tas pats.","fairness_forced_change":"Pateisinamas post-publication repair (liga, atostogos, force majeure, kritinis SPS pull-down) registruojamas ACTUAL audite, bet NEKEIČIA SYSTEM fairness / spread / postų istorijos. Savanoriški swapai taip pat fairness-neutral; keičiasi tik faktinis grafikas ir retrospektyvinis request satisfaction.","fairness_no_history":"Dar nėra pakankamai paskelbtų mėnesių teisingumo istorijai.","fairness_priority_table":"Kaip skaityti hierarchiją","fairness_level":"Lygis","fairness_goal":"Tikslas","fairness_interpretation":"Kaip interpretuoti","fairness_hard_goal":"ABSOLUTE HARD: 0 saugos / fizinio neįmanomumo pažeidimų","voluntary_unpopular_goal":"Vykdyti aiškiai savanoriškai pasirinktą nepopuliarų darbą tik išlaikant aukštesnes SYSTEM struktūrines taisykles, įskaitant Friday raw spread 0–1","voluntary_unpopular_explain":"Pageidautas penktadienis yra SOFT: generatorius jį vykdo tik tada, kai išlaiko penktadienių structural floor/ceil raw spread 0–1 ir aukštesnes HARD taisykles. Po publikavimo abipusis ACTUAL swapas gali šį balansą pakeisti neperrašydamas SYSTEM fairness.","other_preferences_goal":"SOFT: griežtai SOFT-1 → SOFT-2 → SOFT-3; kiekviename range pirmiausia horizontalus water-filling, po to likęs įmanomas išpildymas","other_preferences_explain":"SOFT pageidavimai optimizuojami tik po TRUE ABSOLUTE HARD, SYSTEM HARD postų lygybės, RESIDENT HARD ir likusio workload/fatigue fairness; aukštesnis SOFT rangas užrakinamas prieš pereinant į žemesnį.","fairness_cumulative_goal":"Antrinis ilgalaikis tikslas: po gero einamojo mėnesio balanso taisyti ankstesnės SYSTEM istorijos likutinę nelygybę","fairness_monthly_goal":"SPS RO, SPS UG, savaitgaliai IR penktadieniai SYSTEM grafike turi raw spread 0–1. Penktadieniai water-fill'inami pagal visų užpildytų penktadienio priskyrimų floor/ceil dalį. Visos ne-Onko darbo vietos taip pat water-fill'inamos iki raw spread 0–1 prieš SOFT; platesnis postų koridorius leidžiamas tik jei siauresnis įrodytas neįmanomas.","preference_avg":"Vidutinis pageidavimų išpildymas","weekend_spread":"Savaitgalių skirtumas",
"published_schedule":"Current published schedule","not_published":"No official schedule has been published for this month.","colors":"Permanent resident colors","download_xlsx":"DOWNLOAD FORMATTED SCHEDULE (.xlsx)","download_csv":"Download data list (.csv)",
"summary_title":"Resident summary","frozen_fairness":"Publication fairness","current_after_changes":"Current state after voluntary changes","fairness_frozen_note":"SYSTEM fairness, workplace spread, and future catch-up accounting are frozen from the publication baseline. Bilateral voluntary swaps and justified post-publication repairs (sickness, leave, force majeure, SPS pull-down) change the ACTUAL schedule but are EXCLUDED from fairness/spread/debt. Actual work and retrospective request satisfaction may be shown separately.","person":"Person","name":"Name","target":"Target","workload":"Workload","weekday_assignments":"Weekday assignments","weekday_days":"Distinct weekdays","weekend_assignments":"Weekend assignments","saturday_assignments":"Saturday assignments","sunday_assignments":"Sunday assignments","prior_weekends":"Prior weekends","cumulative_weekends":"Cumulative weekends","fridays":"Fridays","double_shifts":"12h workdays (AM+PM)","max_consecutive":"Max consecutive days","max_rolling7_hours":"Max hours / rolling 7d","max_calendar_week_hours":"Max calendar-week hours","free_days":"Free days","preference_score":"Preference fulfillment, %","planned_backups":"AUTO backup duties","effective_backups":"Current / effective backup duties",
"transparency_title":"Transparency","validity_heading":"1. Privalomų taisyklių patikra","validity_text":"0 klaidų reiškia, kad paskelbtas pradinis grafikas nepažeidė nė vienos privalomos taisyklės. Tai atitikties, o ne teisingumo procentas.","fairness_heading":"2. Grupės teisingumas","fairness_text":"Sistema skiria mėnesio ir kaupiamąjį teisingumą. Kaupiamasis teisingumas yra pagrindinis ilgalaikio balanso rodiklis, o mėnesio teisingumas apibūdina pasirinktą mėnesį.","fair_formula":"Abiem įverčiams naudojama ta pati formulė: 100 − 18× savaitgalių skirtumas − 7× penktadienių skirtumas − 4× dvigubų pamainų skirtumas − 2× darbo dienų skirtumas. Skiriasi tik apimtis: vienas mėnuo arba visų paskelbtų mėnesių suma.","metric_weekend":"Savaitgalių skirtumas","metric_friday":"Penktadienių skirtumas","metric_double":"Dvigubų pamainų skirtumas","metric_weekday":"Darbo dienų skirtumas",
"personal_vs_group":"Personal preference fulfillment versus group fairness","balance_ratio":"Balance ratio","ratio_help":"Balance ratio = smaller percentage / larger percentage. 1.00 means the two scores are at the same level; it is not an absolute quality measure.","baseline_personal":"Personal at publication","current_personal":"Personal now","not_applicable":"N/A","all_resident_scores":"All resident preference scores",
"backup_title":"Backup cover","backup_self_select":"Choose my monthly backup slots","backup_self_select_help":"Reservable mandatory groups are position-based: SPS RO on any day/block, SPS UG on any day/block, Centro UG 120 morning, and full 9h Onko RO. CENTRO RO is planned automatically as best-effort. Multiple slots may be selected; a reserved backup slot blocks overlapping normal work.","backup_claim_deadline":"Backup-choice deadline","backup_claim_saved":"Backup slots reserved.","backup_claim_released":"Backup choice released.","backup_claim_taken":"Another resident just took that slot. Please choose another.","backup_claim_locked":"The selection deadline has passed or the schedule is already published. Further backup-slot changes use the Swaps tab.","backup_claim_missing_penalty":"You have not selected any backup slot yet. The engine will still AUTO-assign backups using fairness-first water-filling; self-selection cannot create an unfair backup load.","backup_claim_yours":"Your reserved backups","backup_claim_board":"Backup reservations","backup_claim_free":"Free","backup_claim_auto_queue":"Automatic-assignment priority pool","backup_claim_auto_queue_help":"Residents who did not self-select any backup slot may enter the automatic-assignment pool first, but claims are only a tie-break. AUTO backups are water-filled by total backup load and never violate Cannot-work / RESIDENT HARD or ABSOLUTE HARD.","release_backup_claim":"RELEASE MY SELECTION","backup_claim_reminder_kind":"Backup-choice reminder","backup_swap_title":"Backup swaps","backup_swap_help":"After publication, you may propose swapping any planned mandatory backup slot. The swap is applied only if both remain eligible for the new slots.","my_backup_duty":"My backup slot","their_backup_duty":"Other resident's backup slot","request_backup_swap":"PROPOSE BACKUP SWAP","backup_swap_sent":"Backup swap proposal sent.","backup_swap_invalid":"This swap cannot be applied because at least one resident would be ineligible for the new backup slot.","backup_swap_accepted":"Backup swap accepted and applied.","backup_swap_rejected":"Backup swap rejected.","backup_definition":"Mandatory named backup is position-based: SPS RO on every day/block, SPS UG on every day/block, Centro UG 120 morning, and full 9h Onko RO. CENTRO RO is covered as widely as safe remaining capacity allows; missing CENTRO RO backup does not block publication. ABSOLUTE HARD and overlapping normal work are never allowed.","my_backup_schedule":"My backup schedule","no_backups":"This resident has no backup duties this month.","covered_assignment":"Covered resident and schedule","covered_person":"Covered resident","covered_schedule":"Covered shift","planned_backup":"Planned backup","actual_backup":"Actual backup","effective_backup":"Effective backup","backup_note":"Note","manage_backups":"Senior backup control","backup_coverage":"Backup coverage","working_person_days":"Required covered shifts","covered_person_days":"Shifts with named backup","backup_complete":"Every mandatory covered shift has a named backup.","backup_incomplete":"At least one mandatory covered shift lacks an eligible backup. The schedule cannot be published.","resync_backups":"REFRESH BACKUPS FROM CURRENT SCHEDULE","backup_synced":"Backups recalculated automatically from the current schedule.","backup_capacity_block":"The draft cannot be published if at least one mandatory covered shift has no ABSOLUTE-HARD-safe resident who is free during that block. Missing CENTRO RO best-effort coverage does not block publication. Any RESIDENT HARD / Unavailable violation blocks SYSTEM publication. V2.5.107 requires zero such violations in a generated SYSTEM draft.","cover_credit_type":"Automatically derived cover type","cover_6h":"MORNING 08:00–14:00 = 6h","cover_12h":"MORNING + AFTERNOON = two separate 6h events","cover_night12h":"NIGHT 20:00–08:00 = 12h","cover_credit_note":"The senior does not choose the type manually: it is derived from the concrete covered slot. Covering both MORNING and AFTERNOON creates two separate 6h cover events. NIGHT will be one 12h event once night slots exist in the scheduler.","actual_override":"Manual actual-backup override","mark_backup_completed":"MARK ACTUAL COVER COMPLETED","backup_completed":"Actual cover recorded. A rest credit was awarded to the covering resident; no debt is created for the covered resident.","undo_backup_completed":"UNDO ACTUAL COVER AND ITS CREDIT","backup_completion_undone":"Actual cover and the covering resident’s credit were reversed.","completed_backup":"Completed","credit_balances":"Rest credits","bonus_units":"Credits","bonus_shift_value":"Available shift reduction","rest_credit_bank":"Rest-credit bank","credit_type":"Credit type","credit_am":"MORNING — 6h","credit_pm":"AFTERNOON — 6h","credit_night":"NIGHT — 12h","use_credit_am":"Redeem MORNING rest credits","use_credit_pm":"Redeem AFTERNOON rest credits","credit_month_cap":"At most 2 daytime rest credits in total may be used in one month.","night_bank_only":"NIGHT credits are bank-only for now; they cannot reduce the current PGY1 daytime target.","netting_explain":"A credit is a one-way benefit for the resident who actually covers. The covered resident receives no debt and keeps any existing rest credits.","cover_effect_rest":"A new rest credit was awarded to the covering resident.","max_credit_error":"At most 2 daytime rest credits in total may be selected for one month.","backup_record":"Backup record","record_actual":"RECORD ACTUAL BACKUP","actual_saved":"Actual backup recorded. Backup statistics will use this resident.","clear_actual":"RESTORE PLANNED BACKUP","actual_cleared":"Actual override removed; the planned backup is effective again.","no_eligible_backup":"No eligible resident is available for this shift. The backup must be free during the same time block and ABSOLUTE-HARD-safe; RESIDENT HARD / Unavailable is mandatory for SYSTEM generation; a blocked resident is not eligible for that backup slot.",
"swap_title":"Voluntary swaps","swap_note":"Voluntary swaps change only the ACTUAL schedule. Before consent, each affected resident sees a consequence table. The swap is blocked by ABSOLUTE/operational and labour-time guardrails, exact monthly workload equality, and even Onko pairing (0/2/4...). Consecutive Onko may be an explicit ACK consequence, but parity may never be overridden. SYSTEM fairness remains frozen. Workplace/post fairness, modality mix, US exposure and diversity never block a mutually accepted ACTUAL swap.","repair_title":"Unplanned post-publication repairs","repair_help":"Sickness, leave, another justified absence, or force majeure changes only the ACTUAL schedule. The SYSTEM fairness baseline frozen at publication and fairness_history remain unchanged; the absent resident receives no fairness debt. V2.5.57: if the absent resident covered SPS RO / SPS UG, critical coverage is preserved first by pulling a same-block resident from a lower-priority NON-MANDATORY post; the optional donor post may remain empty. Only when no safe donor transfer exists is a resident free in that target block used as fallback. ABSOLUTE safety, overlap and mandatory coverage remain hard. Pull-downs and other justified repairs are EXCLUDED from SYSTEM fairness, workplace spread, and future catch-up accounting.","repair_assignment":"Assignment to replace","repair_replacement":"Covering resident","repair_reason":"Reason category","repair_reason_sickness":"Sickness","repair_reason_leave":"Leave","repair_reason_approved":"Other justified absence","repair_reason_force":"Force majeure / unexpected event","repair_note":"Internal note (optional)","apply_repair":"APPLY UNPLANNED REPAIR","repair_applied":"Repair applied to the ACTUAL schedule. SYSTEM fairness, workplace spread, and future catch-up remain tied to the publication baseline; the repair is excluded from them. Actual schedule and request satisfaction were recalculated.","repair_invalid":"This repair cannot be applied because of an operational safety / HARD rule","repair_no_candidate":"No safely eligible covering resident is available for this shift.","repair_history":"Unplanned repair history","repair_load":"Additional repair load this month","repair_load_help":"This is an operational audit counter only. It is NOT used in fairness, workplace spread, future catch-up, or future compensation. A critical pull-down is a station change during an already scheduled work block, not an extra fairness burden.","repair_fairness_neutral":"FAIRNESS NEUTRAL","repair_from":"Absent resident","repair_to":"Covered by","repair_date":"Date / shift","my_assignment":"My assignment","their_assignment":"Other person's assignment","request_swap":"PROPOSE SWAP","request_sent":"Swap request sent.","incoming":"Incoming requests","accept":"ACCEPT","reject":"REJECT","accepted":"Swap approved and applied.","accepted_pending":"Both residents accepted the swap. In the beta, the senior scheduler performs the final hard-rule validation and applies it.","finalize_swap":"APPLY APPROVED SWAP","swap_applied":"Swap applied, hard rules revalidated, and backups recalculated.","swap_finalize_failed":"The swap could not be applied because final validation would violate a hard rule.","rejected":"Swap rejected.","hard_reject":"Swap rejected because it would violate a hard rule.","history":"Swap history","pending":"Pending","approved":"Approved","rejected_status":"Rejected",
"calendar_title":"My calendar schedule","calendar_help":"Download a one-time .ics snapshot or subscribe once to a private calendar feed. The feed is refreshed after a new schedule is published and after important ACTUAL schedule changes. Backups are included when enabled in Settings.","download_ics":"DOWNLOAD MY SCHEDULE (.ics)","calendar_feed":"Private calendar subscription URL","calendar_feed_private":"Treat this URL like a password to your schedule; do not share it. It contains a long random token and is not shown to other residents.","calendar_google":"GOOGLE CALENDAR","calendar_apple":"APPLE CALENDAR","calendar_other":"OUTLOOK CALENDAR","calendar_google_help":"On a computer in Google Calendar: Add other calendars → From URL → paste the private URL shown below. This is a one-time setup.","calendar_apple_help":"Apple Calendar can subscribe directly. The button should open the Calendar subscription prompt.","calendar_other_help":"Outlook can subscribe to the same private iCalendar URL (Add calendar → Subscribe from web). For another app, use the .ics file or subscription URL if supported.",
"proof_title":"My schedule proof","proof_intro":"This visual check shows exactly what matched your needs and where mismatches remain. It displays final results rather than program code.","matches":"MATCHES","partial":"PARTIAL","mismatch":"DOES NOT MATCH","baseline":"At publication","current":"Now","hard_ok":"Hard unavailability respected","hard_bad":"Hard rule violated","soft_off_ok":"Requested days off","preferred_ok":"Preferred work dates","workload_ok":"Monthly workload target","style_component":"Work-style criterion","missed_dates":"Mismatched dates","criterion":"Criterion","result":"Result","score":"Fulfillment","explanation":"Explanation","proof_all_good":"Based on the submitted data, hard rules are respected and active soft preferences have no major mismatch.","proof_soft_issues":"Hard rules are respected, but some soft preferences were not fully fulfilled.","proof_hard_issue":"A hard-rule mismatch was detected and the schedule requires review.","no_active_preferences":"No active preference was submitted for this category.","swap_suggestion":"If no hard rule is violated but a soft preference remains unmet, you can look for a voluntary solution in the Swaps tab.",
"rules_title":"Schedule rules","read":"Read","edit":"Edit","save_rules":"SAVE RULE CHANGES","rules_saved":"Rules updated.","edit_senior_only":"Only the senior scheduler can edit the rules.","yes":"YES","no":"NO","reminder_kind":"Reminder","publication_kind":"Schedule publication","date":"Date","day":"Day","time":"Time","department":"Department","shift":"Shift","morning":"Morning","afternoon":"Afternoon","full_day":"Full day","status":"Status","details":"Details","sent":"Sent","failed":"Failed","skipped":"Skipped"
}}

TR["LT"].update({
"research":"Anketa","research_title":"Radiology Scheduler anketos langas","research_survey":"Anketa","research_dashboard":"Tyrimo skydas","research_phase":"Etapas","research_baseline":"Prieš naudojimą","research_followup":"Po naudojimo","research_likert_help":"1 = visiškai nesutinku · 5 = visiškai sutinku","research_submit":"IŠSAUGOTI ANKETĄ","research_saved":"Tyrimo anketa išsaugota.","research_privacy":"Atsakymai analizei saugomi atskirai nuo grafiko. Seniūnė SP mato tik grupės suvestines ir anoniminius komentarus; ŠR tyrimo lange gali matyti deidentifikuotus atsakymus. Individualūs vardai šiame lange nerodomi.","research_response_count":"Atsakymų skaičius","research_mean":"Vidurkis","research_change":"Pokytis","research_operational":"Operaciniai mėnesio rodikliai","research_survey_results":"Anketos rezultatai","research_deidentified":"Deidentifikuoti atsakymai","research_comments":"Anoniminiai komentarai","research_gm_note":"Seniūnės vaizde rodomi tik grupės lygio rezultatai, kad individualūs atsakymai nebūtų naudojami personalo vertinimui.","research_rs_note":"ŠR tyrėjo vaizde papildomai rodomi deidentifikuoti individualūs įrašai kokybės kontrolei ir vėlesnei analizei.","research_no_data":"Dar nėra tyrimo duomenų.","research_month_note":"Operaciniai rodikliai skaičiuojami pasirinktam mėnesiui; anketų suvestinė apima visus pateiktus įrašus.","research_stress":"Kiek stresą kelia grafiko sudarymo / keitimo procesas? (0–10)","research_changes":"Kiek kartų per mėnesį paprastai prašote ar atliekate grafiko pakeitimą?","research_contact":"Kaip dažnai reikia kreiptis į seniūnę dėl neaiškaus ar neteisingo grafiko?","research_problem":"Didžiausia dabartinio proceso problema","research_improve":"Ką labiausiai reikėtų pagerinti?","research_easy":"Sistema lengva naudotis","research_mobile":"Mobilioji versija lengva naudotis","research_actual":"Sistema rodo realų aktualų grafiką po pakeitimų","research_system_actual":"Naudingas skirtumas tarp sistemos pradinio ir faktinio grafiko","research_continue":"Norėčiau tęsti šios sistemos naudojimą vietoje ankstesnio metodo","research_access_denied":"Tyrimo skydas prieinamas tik ŠR ir SP","research_hard_errors":"Privalomų taisyklių klaidos","research_changed_assignments":"Pakeistos normalios pamainos","research_normal_swaps":"Normalių apsikeitimų","research_backup_swaps":"Dublių apsikeitimų","research_completed_covers":"Realiai įvykdyti pavadavimai"
})
TR["EN"].update({
"research":"Research","research_title":"Radiology Scheduler research window","research_survey":"Research survey","research_dashboard":"Research dashboard","research_phase":"Phase","research_baseline":"Before use","research_followup":"After use","research_likert_help":"1 = strongly disagree · 5 = strongly agree","research_submit":"SAVE RESEARCH SURVEY","research_saved":"Research survey saved.","research_privacy":"Research answers are stored separately from scheduling data. The current senior SP sees group summaries and anonymous comments only; the ŠR research view can inspect de-identified responses. Individual names are not shown in this window.","research_response_count":"Response count","research_mean":"Mean","research_change":"Change","research_operational":"Monthly operational metrics","research_survey_results":"Survey results","research_deidentified":"De-identified responses","research_comments":"Anonymous comments","research_gm_note":"The senior view shows group-level results only so individual responses are not used for personnel evaluation.","research_rs_note":"The ŠR researcher view additionally shows de-identified individual records for quality control and later analysis.","research_no_data":"No research data yet.","research_month_note":"Operational metrics use the selected month; survey summaries include all submitted responses.","research_stress":"How stressful is the scheduling / change process? (0–10)","research_changes":"How many schedule changes do you usually request or participate in per month?","research_contact":"How often do you need to contact the senior because the schedule is unclear or incorrect?","research_problem":"Biggest problem with the current process","research_improve":"What should be improved most?","research_easy":"The platform is easy to use","research_mobile":"The mobile version is easy to use","research_actual":"The platform reflects the real current schedule after changes","research_system_actual":"The SYSTEM versus ACTUAL distinction is useful","research_continue":"I would prefer to continue using this platform rather than return to the previous method","research_access_denied":"The research dashboard is available only to ŠR and SP","research_hard_errors":"HARD errors","research_changed_assignments":"Changed normal assignments","research_normal_swaps":"Normal swaps","research_backup_swaps":"Backup swaps","research_completed_covers":"Completed actual covers"
})

# V2.5.53 — wording aligned with critical exposure + weekly recovery constitution.
# The legacy 0–100 summary score remains visible for continuity, but it is NOT
# the optimizer hierarchy and never overrules the explicit lexicographic locks.
TR["LT"].update({
    "fairness_100_note":"100 % yra suvestinis diagnostinis rodiklis, o ne solverio prioritetų formulė. Pirmame mėnesyje be ankstesnės savaitgalių istorijos aiškiai pasirinktas savaitgalio „Pageidauju dirbti“ laikomas savanoriška nepopuliaria pamaina: įvykdytas savanoriškas vienetas gali būti išimtas iš DABARTINIO savaitgalio / SPS RO fairness skaičiaus, o likęs nesavanoriškas krūvis lieka 0–1 water-fill. Tikras RAW dirbtas savaitgalio krūvis vis tiek išsaugomas istorijoje ir veikia vėlesnių mėnesių balansą.",
    "fairness_formula_month":"Rodoma 0–100 formulė yra tęstinis suvestinis indikatorius, bet ji NĖRA solverio prioritetų hierarchija. SPS UG ir penktadienių exposure išlieka raw 0–1 struktūriniai guardrail. Pirmame mėnesyje su savaitgalio savanoriais SPS RO / savaitgalio 0–1 fairness taikomas likusiam NESAVANORIŠKAM krūviui; RAW matomas skirtumas rodomas atskirai ir nėra paslepiamas.",
    "fairness_formula_cumulative":"Kaupiamasis 0–100 indikatorius naudoja paskelbtų mėnesių sukauptą istoriją. Papildomai postų lygybė sekama atskirai per cumulative exposure ir FUTURE CATCH-UP, kad laikinas necritical nukrypimas būtų kompensuotas ateityje.",
    "voluntary_unpopular_goal":"Savaitgalio „Pageidauju dirbti“ pirmame mėnesyje be istorijos gauna atskirą savanoriškos nepopuliarios pamainos sluoksnį: po ABSOLUTE HARD ir RESIDENT HARD sistema pirmiausia bando išpildyti savanorius, o likusį nesavanorišką savaitgalių krūvį water-fillina 0–1.",
    "voluntary_unpopular_explain":"Visa diena „Pageidauju dirbti“ savaitgalį reiškia vieną savanorišką tos dienos vienetą, o ne automatiškai AM+PM. Jei grafikas nepriklausomai skiria ir antrą pusdienį, papildomas vienetas lieka fairness apskaitoje. RAW realiai dirbtos pamainos išsaugomos ir kitą mėnesį grįžta į cumulative balansą.",
    "other_preferences_explain":"Įprastas SOFT vis dar optimizuojamas tik aukštesniuose užraktuose ir negali pralaužti saugos, RESIDENT HARD, SPS UG, penktadienių ar kitų struktūrinių guardrail. Vienintelė aiški išimtis — pirmo mėnesio savanoriškas savaitgalio „Pageidauju dirbti“: savanoriškas vienetas gali būti išimtas iš dabartinio SPS RO / savaitgalio fairness skaičiaus, o likęs nesavanoriškas krūvis privalo likti 0–1.",
})
TR["EN"].update({
    "fairness_100_note":"100% is a summary diagnostic, not the solver priority formula. In the first month with no prior weekend history, an explicit weekend “prefer to work” is treated as volunteering for unpopular duty: an honored volunteer unit may sit outside the CURRENT weekend / SPS RO fairness count, while the remaining non-voluntary burden stays at 0–1 water-fill. The real RAW weekend workload is still saved to history and affects later-month balancing.",
    "fairness_formula_month":"The displayed 0–100 formula is a continuity summary indicator, NOT the solver hierarchy. SPS UG and Friday remain raw 0–1 structural guardrails. In a first-month volunteer case, SPS RO / weekend 0–1 fairness is applied to the remaining NON-voluntary burden; the RAW visible spread is reported separately rather than hidden.",
    "fairness_formula_cumulative":"The cumulative 0–100 indicator uses published-month history. Post equality is also tracked separately through cumulative exposure and FUTURE CATCH-UP so temporary noncritical imbalance is repaid later.",
    "voluntary_unpopular_goal":"A weekend “prefer to work” in the first no-history month gets a dedicated unpopular-duty volunteer layer: after ABSOLUTE HARD and RESIDENT HARD, the engine first tries to honor willing residents and water-fills the remaining non-voluntary weekend burden at 0–1.",
    "voluntary_unpopular_explain":"A whole-day weekend “prefer to work” means one voluntary day unit, not an automatic AM+PM double. If the schedule independently needs a second half-day, that extra unit remains in the fairness count. RAW worked exposure is retained in history and returns to cumulative balancing in later months.",
    "other_preferences_explain":"Ordinary SOFT is still optimized only inside higher safety, Resident-HARD, SPS UG, Friday and structural guardrails. The one explicit exception is the first-month weekend-work volunteer rule: an honored volunteer unit may sit outside current SPS RO / weekend fairness, while the remaining non-voluntary burden must stay at 0–1.",
})

# V2.5.55 — voluntary swap consequence/ACK layer with labour-time reality guardrails.
TR["LT"].update({
    "swap_48_warning":"Šis savanoriškas swapas turi pasekmių tavo krūviui / poilsiui. Peržiūrėk lentelę ir patvirtink tik jei sąmoningai sutinki.",
    "swap_48_ack":"Peržiūrėjau visas šio swapo pasekmes lentelėje, suprantu jas ir savanoriškai sutinku su šiuo konkrečiu apsikeitimu.",
    "swap_48_other":"Šis swapas turi pasekmių ir kitam rezidentui. Jis / ji jas matys savo lentelėje ir turės patvirtinti atskirai.",
    "swap_48_only_exception":"ACK galioja tik šiam konkrečiam savanoriškam swapui. Jis neapeina ABSOLUTE HARD, patvirtinto neatvykimo, overlap/coverage, ≤12 h/d., ≥11 h tarp darbo dienų, ≤6 darbo dienų/7 d. ar ≤60 h/7 d. Post-double recovery ir naujas savo RESIDENT HARD konfliktas swape gali būti sąmoningai priimti ir todėl rodomi kaip ACK, o ne automatinis blokas.",
    "swap_preview_invalid":"Šio apsikeitimo negalima siūlyti / priimti dėl kitos privalomos taisyklės: {reason}",
    "swap_48_reaccept":"Swapo pasekmės nuo pasiūlymo sukūrimo pasikeitė. Reikia naujo aiškaus rezidento patvirtinimo; pasiūlymą sukurkite iš naujo.",
    "swap_note":"Savanoriškas swapas keičia tik FAKTINĮ grafiką. Prieš sutikdamas kiekvienas paveiktas rezidentas mato pasekmių lentelę. Swapas blokuojamas dėl ABSOLUTE / operacinių ir darbo-laiko guardrailų: overlap, pateisinamas neatvykimas, >12 h/d., <11 h poilsio, >6 darbo dienų/7 d. ar aktyvaus swapo hard-cap viršijimo (iki 60 h/7 d.). 12 h double, >40/>48 h, 6 dienų seka, post-double recovery ir savo RESIDENT HARD override rodomi kaip ACK. SYSTEM fairness baseline nesikeičia.",
    "labour_hard_summary":"Generuojant: ≤12 val./d.; ≥11 val. tarp darbo dienų; bent 1 visiškai laisva diena per slenkančias 7 d.; ≤48 žinomų darbo valandų/7 d.; tikslas ~40 val./7 d.; po 2 iš eilės double kita diena PM arba poilsis. Po publikavimo voluntary swapui taikomas atskiras consequence + ACK režimas: >48 h nebėra automatinis blokas, tačiau >12 h/d., <11 h poilsio, >6 darbo dienų/7 d., aktyvus swap hard-cap ir ABSOLUTE/overlap/coverage lieka blokai.",
    "fairness_hierarchy_intro":"Generuojant pirmiausia saugomos absoliučios saugos ir darbo taisyklės. Tada SPS RO, SPS UG ir savaitgaliai paskirstomi kuo lygiau, saugomi rezidentų privalomi negalėjimai, poilsis ir darbo krūvis, o po to tenkinami kiti pageidavimai. Po publikavimo savanoriški apsikeitimai turi atskirą pasekmių lentelę ir aiškų patvirtinimą.",
})
TR["EN"].update({
    "swap_48_warning":"This voluntary swap changes your workload/rest pattern. Review the consequence table and confirm only if you knowingly accept it.",
    "swap_48_ack":"I reviewed all consequences in the table, understand them, and voluntarily agree to this specific swap.",
    "swap_48_other":"This swap also has consequences for the other resident. They will see their own table and must acknowledge it separately.",
    "swap_48_only_exception":"The acknowledgement applies only to this specific voluntary swap. It never bypasses ABSOLUTE HARD, approved absence, overlap/coverage, ≤12h/day, ≥11h daily rest, ≤6 workdays/7d or ≤60h/7d. Post-double recovery and a self-overridden Resident-HARD request may be knowingly accepted in a swap and are therefore shown as ACK warnings rather than automatic blocks.",
    "swap_preview_invalid":"This swap cannot be proposed / accepted because another mandatory rule would be violated: {reason}",
    "swap_48_reaccept":"The swap consequences changed after the request was created. A fresh explicit acknowledgement is required; recreate the swap request.",
    "swap_note":"Voluntary swaps change only the ACTUAL schedule. Each affected resident sees a consequence table before consent. The swap is blocked by ABSOLUTE/operational and labour-time guardrails: overlap, approved absence, >12h/day, <11h rest, >6 workdays/7d or the active swap hard cap (up to 60h/7d). New 12h doubles, >40/>48h load, six-day streaks, post-double recovery patterns and self-overridden Resident-HARD requests are ACK warnings. SYSTEM fairness remains frozen. Workplace/post fairness, modality mix, US exposure and diversity never block a mutually accepted ACTUAL swap.",
    "labour_hard_summary":"Generation: ≤12h/day; ≥11h between workdays; at least 1 fully free day per rolling 7d; ≤48 known hours/7d; target ~40h/7d; after 2 consecutive doubles the next day is PM-only or off. Post-publication voluntary swaps use a separate consequence + ACK mode: >48h is not an automatic blocker, while >12h/day, <11h rest, >6 workdays/7d, the active swap hard cap and ABSOLUTE/overlap/coverage remain blockers.",
    "fairness_hierarchy_intro":"During generation: ABSOLUTE safety/work rules → Cannot-work / RESIDENT HARD with zero violations → ADMIN RAW Saturday/Sunday/weekend water-fill at the tightest feasible spread (weekend preferences cannot bypass it) → Dream Team SP+ŠR+GE together on CENTRO RO at least once per represented workweek when mathematically feasible → SPS RO/SPS UG and every other workplace water-filled as tightly as feasible → recovery/workload fairness → remaining SOFT wishes. After publication, ACTUAL swaps may disturb SYSTEM fairness only after both residents consent and SP gives final approval.",
})

# V2.5.66 — multiple swaps are allowed, but one concrete shift can have only
# one active future offer at a time. Plain-language UX; DB trigger is the
# authoritative concurrency guard.
TR["LT"].update({
    "swap_shift_busy":"Ši pamaina jau įtraukta į kitą laukiantį apsikeitimą. Galite turėti kelis apsikeitimus vienu metu, tačiau kiekviena konkreti pamaina gali būti tik viename aktyviame pasiūlyme. Pirmiausia užbaikite arba atšaukite esamą pasiūlymą.",
    "backup_swap_shift_busy":"Šis dublio slotas jau įtrauktas į kitą laukiantį dublių apsikeitimą. Pasirinkite kitą slotą arba pirmiausia atšaukite / užbaikite esamą pasiūlymą.",
    "my_outgoing_swaps":"Mano laukiantys pasiūlymai",
    "cancel_my_swap":"ATŠAUKTI MANO PASIŪLYMĄ",
    "swap_cancelled":"Apsikeitimo pasiūlymas atšauktas.",
    "swap_cancel_failed":"Šio pasiūlymo atšaukti nepavyko — jis galėjo būti ką tik priimtas arba pakeistas.",
    "multiple_swap_help":"Galite turėti kelis aktyvius apsikeitimus, jei jie liečia skirtingas pamainas. Ta pati konkreti pamaina vienu metu gali būti tik viename aktyviame pasiūlyme.",
})
TR["EN"].update({
    "swap_shift_busy":"This shift is already part of another pending swap. You may have several swaps at the same time, but each concrete shift can be in only one active offer. Finish or cancel the existing offer first.",
    "backup_swap_shift_busy":"This backup duty is already part of another pending backup swap. Choose another duty or finish/cancel the existing offer first.",
    "my_outgoing_swaps":"My pending offers",
    "cancel_my_swap":"CANCEL MY OFFER",
    "swap_cancelled":"Swap offer cancelled.",
    "swap_cancel_failed":"This offer could not be cancelled — it may already have been accepted or changed.",
    "multiple_swap_help":"You may have several active swaps when they involve different shifts. The same concrete shift can be in only one active offer at a time.",
})


# V2.5.11 — role-specific 6-month research workflow
TR["LT"].update({
"research_role_resident":"Rezidentas · tyrimo dalyvis","research_role_researcher":"Rezidentas · tyrėjas","research_role_senior":"Rezidentė · seniūnė / grafikų sudarytoja",
"research_study_plan":"6 mėn. tyrimo planas","research_study_period":"Prospektyvus naudojimas: 2026 m. spalis – 2027 m. kovas","research_primary_outcomes":"Pagrindiniai rodikliai: privalomų taisyklių laikymasis, grafiko sudarymo laikas ir rankinių korekcijų skaičius.",
"research_checkpoint":"Tyrimo matavimo taškas","research_baseline_checkpoint":"Pradinis vertinimas · prieš paleidimą (2026 rugsėjis)","research_month3_checkpoint":"3 mėn. pakartotinis vertinimas (2026 gruodis)","research_month6_checkpoint":"6 mėn. pakartotinis vertinimas (2027 kovas)",
"research_checkpoint_done":"Užpildyta","research_checkpoint_locked":"Dar neaktyvuota","research_checkpoint_pending":"Neužpildyta","research_next_task":"Kitas tyrimo veiksmas","research_resident_note":"Rezidentui reikia tik trijų trumpų anketų per visą tyrimą: pradinio vertinimo, po ~3 mėn. ir po ~6 mėn.",
"research_completion":"Anketų užpildymas","research_expected":"Numatyta","research_export":"Eksportas analizei","research_download_surveys":"ATSISIŲSTI DEIDENTIFIKUOTAS ANKETAS (.csv)","research_download_monthly":"ATSISIŲSTI MĖNESIO RODIKLIUS (.csv)",
"research_monthly_table":"Spalis–kovas: automatiniai operaciniai rodikliai","research_scheduler_section":"SP / Seniūnės grafikų sudarymo darbo krūvis","research_scheduler_intro":"SP kaip dabartinė Seniūnė pildo du trumpus įrašus kiekvienam tyrimo mėnesiui: iškart paruošus grafiką ir mėnesiui pasibaigus. Operacinių apsikeitimų ir pavadavimų skaičius sistema renka automatiškai.",
"research_scheduler_month":"Tyrimo mėnuo","research_scheduler_checkpoint":"Seniūnės SP matavimo taškas","research_after_creation":"Iškart paruošus grafiką","research_after_month":"Mėnesiui pasibaigus",
"research_workflow_method":"Kaip šį mėnesį buvo sudarytas grafikas?","research_method_tool":"Tik sistema","research_method_excel":"Tik Excel","research_method_shadow":"Excel + sistema: lygiagretus palyginimas",
"research_total_minutes":"Bendras tavo aktyvus laikas grafikui paruošti (min.)","research_corrections":"Rankinių korekcijų / iteracijų skaičius","research_resident_contacts":"Kiek rezidentų kontaktų / derinimų reikėjo?","research_communication_minutes":"Kiek laiko užėmė komunikacija dėl grafiko? (min.)",
"research_scheduler_stress":"Grafiko sudarymo stresas (0–10)","research_fairness_confidence":"Pasitikėjimas, kad paskirstymas teisingas (1–5)","research_hard_confidence":"Pasitikėjimas, kad privalomos taisyklės išlaikytos (1–5)","research_scheduler_satisfaction":"Pasitenkinimas galutiniu grafiku (1–5)",
"research_excel_minutes":"Excel laikas (min.)","research_tool_minutes":"Sistemos laikas (min.)","research_excel_corrections":"Excel korekcijų sk.","research_tool_corrections":"Sistemos korekcijų / pakartotinių generavimų sk.",
"research_post_minutes":"Laikas po publikavimo pakeitimams / problemoms (min.)","research_post_interventions":"Kiek pakeitimų reikėjo tavo tiesioginio įsikišimo?","research_post_contacts":"Kiek žinučių / skambučių dėl grafiko gavai per mėnesį?",
"research_actual_confidence":"Pasitikėjimas, kad portalas rodė realų faktinį grafiką (1–5)","research_use_next":"Ar rinktumeisi sistemą kitam mėnesiui?","research_yes":"Taip","research_unsure":"Neaišku","research_no":"Ne",
"research_scheduler_notes":"Pastabos / kas užėmė daugiausia laiko","research_scheduler_saved":"Seniūnės SP tyrimo įrašas išsaugotas.","research_scheduler_status":"Seniūnės SP duomenų užpildymas",
"research_generation_telemetry":"Automatiniai generavimo rodikliai","research_generation_attempts":"Generavimo bandymų","research_solver_seconds":"Skaičiavimo laikas, s","research_generation_success":"Sėkmingų generavimų",
"research_observer_tab":"Tyrimo atsiliepimai","research_observer_intro":"Tik peržiūros teisės grafikui nesikeičia. Šiame lange galima tik pateikti tyrimo atsiliepimą apie stebėsenos patogumą.","research_observer_checkpoint":"Administratorės vertinimas",
"research_obs_actual":"Lengva nustatyti, kas realiai dirba kiekvieną pamainą.","research_obs_changes":"Lengva peržiūrėti pakeitimus po publikavimo.","research_obs_system_actual":"Sistemos pradinio ir faktinio grafiko skirtumas yra naudingas.","research_obs_privacy":"Sistema suteikia pakankamai matomumo neatskleisdama nereikalingų privačių rezidentų duomenų.","research_obs_log":"Pakeitimų žurnalas yra suprantamas.","research_obs_fairness":"Teisingumo informacija yra suprantama.","research_obs_trust":"Pasitikiu portale rodoma operacine informacija.","research_obs_missing":"Kokios informacijos trūksta, kai reikia suprasti realią skyriaus situaciją?","research_observer_saved":"Administratorės tyrimo feedback išsaugotas.",
"research_data_quality":"Duomenų pilnumas","research_missing_scheduler":"Trūksta Seniūnės SP įrašo","research_complete":"Pilna","research_researcher_only":"Ši išsami skiltis matoma tik ŠR tyrėjo paskyroje."
})
TR["EN"].update({
"research_role_resident":"Resident · study participant","research_role_researcher":"Resident · researcher","research_role_senior":"Resident · senior scheduler",
"research_study_plan":"6-month research plan","research_study_period":"Prospective use: October 2026 – March 2027","research_primary_outcomes":"Primary outcomes: HARD-rule compliance, scheduler time, and number of manual corrections.",
"research_checkpoint":"Research checkpoint","research_baseline_checkpoint":"Baseline · pre-launch (September 2026)","research_month3_checkpoint":"3-month follow-up (December 2026)","research_month6_checkpoint":"6-month follow-up (March 2027)",
"research_checkpoint_done":"Completed","research_checkpoint_locked":"Not active yet","research_checkpoint_pending":"Not completed","research_next_task":"Next research task","research_resident_note":"Residents complete only three short surveys during the whole study: baseline, ~3 months, and ~6 months.",
"research_completion":"Survey completion","research_expected":"Expected","research_export":"Analysis export","research_download_surveys":"DOWNLOAD DE-IDENTIFIED SURVEYS (.csv)","research_download_monthly":"DOWNLOAD MONTHLY METRICS (.csv)",
"research_monthly_table":"October–March: automatic operational metrics","research_scheduler_section":"SP / senior scheduler workload","research_scheduler_intro":"SP as the current senior completes two short records per study month: immediately after preparing the schedule and after the month ends. Operational swap / cover counts are collected automatically.",
"research_scheduler_month":"Study month","research_scheduler_checkpoint":"Senior SP checkpoint","research_after_creation":"Immediately after schedule preparation","research_after_month":"After the month is complete",
"research_workflow_method":"How was this month's schedule prepared?","research_method_tool":"Tool only","research_method_excel":"Excel only","research_method_shadow":"Excel + Tool parallel comparison",
"research_total_minutes":"Total active time you spent preparing the schedule (min)","research_corrections":"Manual corrections / iterations","research_resident_contacts":"Resident contacts / coordination episodes","research_communication_minutes":"Time spent on schedule communication (min)",
"research_scheduler_stress":"Schedule-preparation stress (0–10)","research_fairness_confidence":"Confidence allocation is fair (1–5)","research_hard_confidence":"Confidence HARD rules are respected (1–5)","research_scheduler_satisfaction":"Satisfaction with final schedule (1–5)",
"research_excel_minutes":"Excel time (min)","research_tool_minutes":"Tool time (min)","research_excel_corrections":"Excel corrections","research_tool_corrections":"Tool corrections / regenerations",
"research_post_minutes":"Time spent on post-publication changes/problems (min)","research_post_interventions":"Changes requiring your direct intervention","research_post_contacts":"Schedule-related messages/calls received during month",
"research_actual_confidence":"Confidence portal reflected the true ACTUAL schedule (1–5)","research_use_next":"Would you choose the Tool next month?","research_yes":"Yes","research_unsure":"Unsure","research_no":"No",
"research_scheduler_notes":"Notes / what consumed the most time","research_scheduler_saved":"Senior SP research record saved.","research_scheduler_status":"Senior SP data completion",
"research_generation_telemetry":"Automatic generation telemetry","research_generation_attempts":"Generation attempts","research_solver_seconds":"Solver time, s","research_generation_success":"Successful generations",
"research_observer_tab":"Research feedback","research_observer_intro":"Read-only scheduling rights remain unchanged. This tab only allows research feedback about the monitoring experience.","research_observer_checkpoint":"Administrator evaluation",
"research_obs_actual":"It is easy to determine who is actually working each shift.","research_obs_changes":"It is easy to review changes made after publication.","research_obs_system_actual":"The SYSTEM baseline versus ACTUAL distinction is useful.","research_obs_privacy":"The platform provides enough visibility without exposing unnecessary private resident information.","research_obs_log":"The change log is understandable.","research_obs_fairness":"The fairness information is understandable.","research_obs_trust":"I trust the operational information shown in the portal.","research_obs_missing":"What information is missing when you need to understand the real departmental staffing situation?","research_observer_saved":"Administrator research feedback saved.",
"research_data_quality":"Data completeness","research_missing_scheduler":"Missing senior SP record","research_complete":"Complete","research_researcher_only":"This detailed section is visible only to the ŠR researcher account."
})

# V2.5.96 — monthly baseline fairness + live ACTUAL ledger; no future catch-up.
TR["LT"].update({
    "metric_saturday":"Šeštadienių skirtumas",
    "metric_sunday":"Sekmadienių skirtumas",
    "fairness_hierarchy_intro":"Kiekvienas mėnuo prasideda nuo švaraus SYSTEM baseline: ABSOLUTE HARD → kritinis struktūrinis water-fill (SPS RO, SPS UG, ŠEŠTADIENIAI, SEKMADIENIAI ir penktadieniai) → RESIDENT HARD / poilsis / tikslus workload → aktyvūs pageidavimai ir darbo pobūdžio nustatymai → likusių postų optimizavimas. 12 val. darbo dienos pageidavimas gali perskirstyti jau reikalingas dvigubas pamainas ir dėl to išplėsti dublių skirtumą, jei HARD ir kritinis water-fill lieka validūs. Po publikavimo leidžiami ACTUAL pakeitimai gali water-fill pakeisti. Ankstesnių mėnesių fairness kitam mėnesiui catch-up nesukuria.",
    "fairness_monthly_explain":"SYSTEM mėnesio fairness rodo algoritmo baseline publikavimo momentu. ACTUAL mėnesio fairness perskaičiuojamas iš realaus dabartinio darbo po manual override'ų, swapų, repair ir realiai įvykdytų dublių. Platesnis ACTUAL spread leidžiamas ir rodomas, bet nėra perkeliamas kaip skola į kitą mėnesį.",
    "fairness_cumulative_explain":"Istorija yra tik auditas / stebėjimas. Ji NENAUDOJAMA kito mėnesio solverio kompensacijai ar catch-up paskyrimams.",
    "fairness_100_note":"Fairness procentas yra diagnostinis mėnesio balanso rodiklis. SYSTEM baseline turi laikytis generatoriaus water-fill; ACTUAL gali nukrypti po leidžiamų žmogaus sprendimų ir tada rodomas toks, koks yra realybėje.",
    "fairness_formula_cumulative":"Istoriniai mėnesių rodikliai rodomi palyginimui, tačiau nėra solverio įvestis ir nesukuria ateities future catch-up / catch-up.",
    "fairness_swap_neutral":"SYSTEM baseline lieka užšaldytas auditui, tačiau swapas ar manual override pakeičia ACTUAL fairness statistiką pagal realų darbą. Tai nėra ateities fairness skola.",
    "fairness_forced_change":"Post-publication repair ir realiai įvykdytas dublio cover keičia ACTUAL realaus darbo fairness statistiką. SYSTEM publikavimo baseline lieka nepakeistas auditui. Jokio future catch-up nėra.",
    "fairness_frozen_note":"SYSTEM = publikavimo momento algoritmo water-fill baseline. ACTUAL = realus dabartinis pasiskirstymas po override'ų, swapų, repair ir completed cover. Abu rodomi atskirai; ACTUAL istorija yra tik stebėjimui ir niekada nevaldo kito mėnesio generatoriaus.",
    "fairness_cumulative_goal":"Nėra fairness catch-up sluoksnio. Kiekvienas mėnuo vėl pradedamas nuo neutralaus water-fill baseline.",
    "voluntary_unpopular_goal":"Kritiniai nesavanoriško krūvio komponentai pirmiausia water-fill'inami. Tada aktyvūs individualūs pageidavimai ir darbo pobūdžio nustatymai tenkinami maksimaliai, kol nepažeidžiamos privalomos taisyklės ir kritiniai struktūriniai guardrailai.",
    "voluntary_unpopular_explain":"Aktyvus pageidavimas nėra lyginamas su neutraliu N/A kaip konkuruojančiu noru. Pvz., jei tik vienas rezidentas pasirenka dažniausiai 12 val. darbo dienas, sistema turi jam skirti kuo daugiau jau matematiškai reikalingų AM+PM dienų, jei tai nepažeidžia HARD ir kritinio water-fill. Po publikavimo abipusis swapas gali keisti ACTUAL balansą dar plačiau.",
    "other_preferences_explain":"Po privalomų taisyklių ir kritinio struktūrinio water-fill sistema aktyviai maksimalizuoja realiai pateiktų pageidavimų bei darbo pobūdžio nustatymų išpildymą. Neutralūs žmonės nekonkuruoja su aiškiai išreikštu pageidavimu. Darbo dienos trukmės pageidavimas gali perskirstyti dvigubas pamainas, tačiau negali kurti papildomo bendro dublių poreikio ar pažeisti ABSOLUTE/HARD bei kritinių šeštadienio/sekmadienio/SPS guardrailų.",
    "swap_note":"Savanoriškas swapas keičia ACTUAL grafiką ir ACTUAL fairness statistiką. SYSTEM baseline auditui neperrašomas. Water-fill nėra post-publication swapo blokatorius, todėl leidžiamas ACTUAL spread padidėjimas aiškiai atsispindi statistikoje. Jokio kito mėnesio catch-up dėl to nėra.",
    "repair_help":"Liga, atostogos ar force majeure keičia ACTUAL grafiką. Realaus darbo fairness perskaičiuojamas pagal tai, kas iš tikrųjų dirba; SYSTEM publikavimo baseline lieka auditui. Šis skirtumas niekada nekuria fairness skolos ar kito mėnesio catch-up.",
    "repair_applied":"Pakeitimas pritaikytas ACTUAL grafikui. ACTUAL fairness perskaičiuotas pagal realų darbą; SYSTEM baseline liko nepakeistas auditui. Future catch-up nėra.",
    "repair_load_help":"Operacinis audito skaitiklis. Realiai pakeistas darbas įeina į ACTUAL mėnesio statistiką, tačiau nėra paverčiamas ateities fairness skola.",
    "repair_fairness_neutral":"ACTUAL FAIRNESS PERSKAIČIUOTA",
})
TR["EN"].update({
    "metric_saturday":"Saturday spread",
    "metric_sunday":"Sunday spread",
    "fairness_hierarchy_intro":"Every month starts from a clean SYSTEM baseline: ABSOLUTE HARD → critical structural water-fill (SPS RO, SPS UG, SATURDAYS, SUNDAYS and Fridays) → RESIDENT HARD / rest / exact workload → active preferences and work-style settings → remaining workplace optimization. A 12-hour workday preference may redistribute the already-needed double-shift pool and widen double spread when HARD and critical water-fill remain valid. Allowed ACTUAL changes may diverge after publication. Prior-month fairness never creates future catch-up.",
    "fairness_monthly_explain":"SYSTEM monthly fairness is the algorithmic publication baseline. ACTUAL monthly fairness is recalculated from real current work after manual overrides, swaps, repairs and completed backup covers. A wider ACTUAL spread is allowed and shown, but never carried forward as a debt.",
    "fairness_cumulative_explain":"History is audit/monitoring only. It is NOT used by the next month's solver for compensation or catch-up assignments.",
    "fairness_100_note":"The fairness percentage is a diagnostic monthly balance indicator. SYSTEM must satisfy the generator water-fill baseline; ACTUAL may diverge after allowed human decisions and is shown exactly as reality stands.",
    "fairness_formula_cumulative":"Historical monthly metrics are displayed for comparison only; they are not solver input and create no future future catch-up/catch-up.",
    "fairness_swap_neutral":"SYSTEM remains frozen for audit, but a swap or manual override changes ACTUAL fairness statistics according to real work. This never becomes a future fairness debt.",
    "fairness_forced_change":"Post-publication repair and a completed backup cover change ACTUAL real-work fairness statistics. The SYSTEM publication baseline remains unchanged for audit. There is no future catch-up.",
    "fairness_frozen_note":"SYSTEM = the publication-time algorithmic water-fill baseline. ACTUAL = the real current distribution after overrides, swaps, repairs and completed covers. Both are shown separately; ACTUAL history is monitoring-only and never controls the next month's generator.",
    "fairness_cumulative_goal":"There is no fairness catch-up layer. Every month starts again from a neutral water-fill baseline.",
    "voluntary_unpopular_goal":"Critical involuntary burden is water-filled first. Active individual preferences and work-style settings are then fulfilled as much as possible while mandatory rules and critical structural guardrails remain valid.",
    "voluntary_unpopular_explain":"An active preference is not treated as competing with a neutral N/A. For example, if only one resident prefers mostly 12-hour days, the system should allocate as many of the already-required AM+PM days to that resident as feasible while preserving HARD and critical water-fill. Accepted post-publication swaps may alter ACTUAL balance further.",
    "other_preferences_explain":"After mandatory rules and critical structural water-fill are protected, the system actively maximizes fulfillment of expressed preferences and work-style settings. Neutral residents do not compete with an explicit preference. Workday-length preferences may redistribute double shifts, but cannot create extra total double demand or violate ABSOLUTE/HARD or critical Saturday/Sunday/SPS guardrails.",
    "swap_note":"A voluntary swap changes the ACTUAL schedule and ACTUAL fairness statistics. The SYSTEM baseline remains frozen for audit. Water-fill is not a post-publication swap blocker, so an allowed ACTUAL spread increase is visible in the statistics. No later-month catch-up is created.",
    "repair_help":"Sickness, leave or force majeure changes ACTUAL. Real-work fairness is recalculated from who actually works; SYSTEM remains the frozen publication baseline. The difference never creates a fairness debt or next-month catch-up.",
    "repair_applied":"Repair applied to ACTUAL. ACTUAL fairness was recalculated from real work; SYSTEM baseline remains frozen for audit. No future catch-up is created.",
    "repair_load_help":"Operational audit counter. Changed real work enters ACTUAL monthly statistics but is never converted into a future fairness debt.",
    "repair_fairness_neutral":"ACTUAL FAIRNESS UPDATED",
})

RESEARCH_ITEMS = {
"fair_undesirable":"The scheduling process distributes undesirable shifts fairly across residents.",
"fair_weekend":"Weekend duties are distributed fairly over time.",
"fair_friday":"Friday duties are distributed fairly over time.",
"fair_longitudinal":"The process accounts fairly for workload accumulated in previous months.",
"transparent_why":"I understand why I receive the shifts that I receive.",
"transparent_group":"It is easy to see whether the schedule is fair across the whole group.",
"control":"I have enough influence over my monthly schedule.",
"hard_respected":"My unavailable dates are reliably respected.",
"prefs_used":"My scheduling preferences are taken into account when feasible.",
"prefs_balance":"The process balances my preferences with fairness to other residents.",
"swap_easy":"It is easy to arrange a shift swap when I need one.",
"swap_clear":"After swaps, it is clear who is actually responsible for each shift.",
"backup_clear":"The backup / cover process is clear.",
"trust":"I trust the scheduling process.",
"satisfaction":"Overall, I am satisfied with the way our monthly schedule is created.",
"burden":"Schedule-related communication and corrections take too much of my time."
}
RESEARCH_ITEMS_LT = {
"fair_undesirable":"Nepatogios pamainos tarp rezidentų paskirstomos teisingai.",
"fair_weekend":"Savaitgalių budėjimai laikui bėgant paskirstomi teisingai.",
"fair_friday":"Penktadienių pamainos laikui bėgant paskirstomos teisingai.",
"fair_longitudinal":"Sistema teisingai atsižvelgia į ankstesnių mėnesių darbo krūvį.",
"transparent_why":"Suprantu, kodėl gaunu man paskirtas pamainas.",
"transparent_group":"Lengva matyti, ar grafikas teisingas visos grupės mastu.",
"control":"Turiu pakankamai įtakos savo mėnesio grafikui.",
"hard_respected":"Mano nurodytos datos, kai negaliu dirbti, patikimai išlaikomos.",
"prefs_used":"Kai įmanoma, į mano pageidavimus atsižvelgiama.",
"prefs_balance":"Mano pageidavimai subalansuojami su teisingumu kitiems rezidentams.",
"swap_easy":"Kai reikia, lengva susitarti dėl pamainos apsikeitimo.",
"swap_clear":"Po apsikeitimų aišku, kas realiai atsakingas už kiekvieną pamainą.",
"backup_clear":"Dublių / pavadavimo procesas yra aiškus.",
"trust":"Pasitikiu grafiko sudarymo procesu.",
"satisfaction":"Apskritai esu patenkintas(-a), kaip sudaromas mūsų mėnesio grafikas.",
"burden":"Su grafiku susijusi komunikacija ir taisymai užima per daug mano laiko."
}

STUDY_MONTHS = [(2026,10),(2026,11),(2026,12),(2027,1),(2027,2),(2027,3)]
RESIDENT_RESEARCH_CHECKPOINTS = [
    ("baseline", "baseline", 2026, 9),
    ("month3", "followup", 2026, 12),
    ("month6", "followup", 2027, 3),
]
OBSERVER_RESEARCH_CHECKPOINTS = [("month3", 2026, 12),("month6", 2027, 3)]
RESEARCH_EXPECTED_RESIDENTS = 16

def research_checkpoint_label(code):
    return {
        "baseline":tr("research_baseline_checkpoint"),
        "month3":tr("research_month3_checkpoint"),
        "month6":tr("research_month6_checkpoint"),
    }.get(code,code)

def research_checkpoint_for_storage(code):
    for c,phase,y,m in RESIDENT_RESEARCH_CHECKPOINTS:
        if c==code:
            return phase,y,m
    return "baseline",2026,9

def study_month_label(y,m):
    return f"{MONTHS[lang][m-1]} {y}"

MONTHS={"LT":["Sausis","Vasaris","Kovas","Balandis","Gegužė","Birželis","Liepa","Rugpjūtis","Rugsėjis","Spalis","Lapkritis","Gruodis"],"EN":list(calendar.month_name)[1:]}
WEEKDAYS={"LT":["Pr","An","Tr","Kt","Pn","Št","Sk"],"EN":["Mon","Tue","Wed","Thu","Fri","Sat","Sun"]}
WEEKDAY_FULL={"LT":["Pirmadienis","Antradienis","Trečiadienis","Ketvirtadienis","Penktadienis","Šeštadienis","Sekmadienis"],"EN":list(calendar.day_name)}

# Language must be resolved before translating the rest of the interface.
lang = st.sidebar.radio("Kalba / Language", ["LT","EN"], horizontal=True, key="lang_switch")
# V2.5.112 FINAL ADMIN POLICY OVERRIDES — these intentionally supersede older
# V2.5.104 volunteer-weekend wording retained in historical source comments.
TR["LT"].update({
    "fairness_hierarchy_intro":"NORMALUS SYSTEM grafikas: TRUE ABSOLUTE / sauga → 0 „Negaliu dirbti“ pažeidimų → mažiausias matematiškai įmanomas RAW šeštadienių, sekmadienių ir bendro savaitgalio water-fill → SPS RO / SPS UG ir kitų postų water-fill → Dream Team SP+ŠR+GE CENTRO RO kartą per savaitę → SOFT pageidavimai. TIK PO TO, jau nekeisdamas normalaus grafiko, kuriamas atskiras teorinis AUTO dublių/pavadavimo standby sluoksnis. Savaitgalio noras SYSTEM paskirstymo nekeičia; ACTUAL swapas gali pakeisti tik po abiejų rezidentų + SP patvirtinimo.",
    "fairness_100_note":"100% yra diagnostinis idealios pusiausvyros rodiklis. SYSTEM savaitgalių skirstymas remiasi RAW administraciniu water-fill, o ne rezidentų noru dirbti daugiau savaitgalių. Jei 0–1 matematiškai neįmanoma dėl HARD prieinamumo ir tikslaus krūvio, rodomas mažiausias įrodytas įmanomas spread.",
    "fairness_monthly_goal":"SYSTEM pirmiausia ieško mažiausio įmanomo RAW šeštadienių, sekmadienių ir bendro savaitgalio spread. SPS RO / SPS UG ir visi kiti postai taip pat water-fill'inami kuo lygiau, bet 0 „Negaliu dirbti“ pažeidimų lieka aukščiau fairness. Dream Team savaitinis CENTRO RO tikslas yra aukštas administracinis prioritetas postų paskirstymo fazėje.",
    "voluntary_unpopular_goal":"Savaitgalio darbo noras yra tik informacinis / audito SOFT signalas ir SYSTEM savaitgalių water-fill nekeičia.",
    "voluntary_unpopular_explain":"Rezidentas negali nupirkti papildomų šeštadienių ar sekmadienių pažymėdamas „Pageidauju dirbti“. SYSTEM paskirsto savaitgalius pagal mažiausią įmanomą RAW spread. Po publikavimo abipusis swapas gali balansą pakeisti tik po SP galutinio patvirtinimo.",
    "weekend_help":"ADMIN WATER-FILL: savaitgalių krypties pasirinkimas išjungtas SYSTEM generavimui. Naudok „Negaliu dirbti“ tik tikram neprieinamumui.",
    "preferred_help":"Darbo dienų pageidavimai yra SOFT. Savaitgalio pageidavimas registruojamas auditui, bet SYSTEM dėl jo negali skirti papildomų šeštadienių ar sekmadienių virš mažiausio įmanomo RAW water-fill.",
    "backup_swap_help":"Po publikavimo galima siūlyti dublio apsikeitimą. Jis pritaikomas tik po abiejų rezidentų sutikimo, HARD patikros ir SP galutinio APPROVE; SP gali ir DECLINE.",
})
TR["EN"].update({
    "fairness_hierarchy_intro":"NORMAL SYSTEM schedule: TRUE ABSOLUTE / safety → zero Cannot-work violations → tightest mathematically feasible RAW Saturday, Sunday and total-weekend water-fill → SPS RO / SPS UG and all-post water-fill → Dream Team SP+ŠR+GE at CENTRO RO once per week → SOFT wishes. ONLY AFTER the normal schedule is frozen, a separate theoretical AUTO backup/standby layer is built without changing it. Weekend preference cannot change SYSTEM allocation; ACTUAL swaps require both residents plus SP final approval.",
    "fairness_100_note":"100% is a diagnostic ideal-balance score. SYSTEM weekend allocation uses RAW administrative water-fill, not resident willingness to take extra weekends. If 0–1 is mathematically impossible because of HARD availability and exact workload, the tightest proven feasible spread is shown.",
    "fairness_monthly_goal":"SYSTEM first searches for the tightest feasible RAW Saturday, Sunday and total-weekend spread. SPS RO / SPS UG and every other workplace are also water-filled as evenly as possible, while zero Cannot-work violations remain above fairness. Weekly Dream Team CENTRO RO co-location is a high administrative priority in the workplace-placement phase.",
    "voluntary_unpopular_goal":"Weekend willingness is informational/audit-only and does not change SYSTEM weekend water-fill.",
    "voluntary_unpopular_explain":"A resident cannot buy extra Saturdays or Sundays by selecting Prefer to work. SYSTEM allocates weekends by the tightest feasible RAW spread. After publication a bilateral swap may change the balance only after SP final approval.",
    "weekend_help":"ADMIN WATER-FILL: resident weekend-direction selection is disabled for SYSTEM generation. Use Cannot-work only for genuine unavailability.",
    "preferred_help":"Weekday requests are SOFT. Weekend requests are recorded for audit, but SYSTEM cannot use them to give extra Saturdays or Sundays beyond the tightest feasible RAW water-fill.",
    "backup_swap_help":"After publication a backup swap may be proposed, but it is applied only after both residents consent, HARD validation passes, and SP gives final APPROVE; SP may DECLINE.",
})

def tr(k): return TR[lang][k]

def contrast_text(hex_color):
    h=hex_color.lstrip("#"); r,g,b=int(h[0:2],16),int(h[2:4],16),int(h[4:6],16)
    return "#111111" if (0.299*r+0.587*g+0.114*b)>160 else "#FFFFFF"

def badge(initials, include_name=True):
    p=next((x for x in DEFAULT_PEOPLE if x["initials"]==initials),None); c=PERSON_COLORS.get(initials,"#DDD")
    label=initials+(f" — {p['name']}" if include_name and p else "")
    return f'<span style="display:inline-block;background:{c};color:{contrast_text(c)};padding:5px 10px;border-radius:8px;font-weight:700;margin:2px 4px 2px 0;">{html.escape(label)}</span>'


def _person_name(initials):
    p=next((x for x in DEFAULT_PEOPLE if x["initials"]==initials),None)
    return p["name"] if p else initials


def _swap_shift_text(slot):
    if slot is None:
        return "—"
    return (
        f"{slot.day:02d} {WEEKDAYS[lang][slot.weekday]} · "
        f"{slot.department} · {block_label(slot.block)}"
    )


def _render_swap_people_line(person_a, person_b, arrow="↔"):
    st.markdown(
        f'<div style="font-size:1.08rem;display:flex;align-items:center;gap:10px;'
        f'flex-wrap:wrap;margin:4px 0 10px 0;">'
        f'{badge(person_a,include_name=True)}'
        f'<span style="font-size:1.45rem;font-weight:800;color:#666;">{html.escape(arrow)}</span>'
        f'{badge(person_b,include_name=True)}'
        f'</div>',
        unsafe_allow_html=True,
    )


def _render_shift_tile(title, person, slot, accent=None):
    c=accent or PERSON_COLORS.get(person,"#777777")
    st.markdown(
        f'<div style="border:2px solid {c};border-radius:14px;padding:12px 14px;'
        f'min-height:112px;background:linear-gradient(180deg,{c}18,rgba(255,255,255,0.03));">'
        f'<div style="font-size:.76rem;font-weight:800;letter-spacing:.04em;opacity:.72;'
        f'text-transform:uppercase;margin-bottom:7px;">{html.escape(title)}</div>'
        f'{badge(person,include_name=False)}'
        f'<div style="font-size:1.02rem;font-weight:700;margin-top:8px;">'
        f'{html.escape(_swap_shift_text(slot))}</div>'
        f'</div>',
        unsafe_allow_html=True,
    )


def _render_swap_request_card(req, idx, total, slot_map, incoming=False):
    a=str(req.get("person_a") or "")
    b=str(req.get("person_b") or "")
    sa=slot_map.get(int(req.get("slot_a") or -1))
    sb=slot_map.get(int(req.get("slot_b") or -1))
    c=PERSON_COLORS.get(a,"#777777")
    title=(
        f"GAUTA UŽKLAUSA {idx}/{total} · DB #{req.get('id')}"
        if incoming and lang=="LT" else
        f"INCOMING REQUEST {idx}/{total} · DB #{req.get('id')}"
        if incoming else
        f"MANO PASIŪLYMAS {idx}/{total} · DB #{req.get('id')}"
        if lang=="LT" else
        f"MY OFFER {idx}/{total} · DB #{req.get('id')}"
    )
    st.markdown(
        f'<div style="border-left:8px solid {c};border-radius:16px;padding:12px 16px;'
        f'background:rgba(127,127,127,.07);margin:8px 0 10px 0;">'
        f'<div style="font-size:.82rem;font-weight:900;letter-spacing:.055em;'
        f'text-transform:uppercase;">{html.escape(title)}</div>'
        f'</div>',
        unsafe_allow_html=True,
    )
    _render_swap_people_line(a,b,"→")
    c1,c2=st.columns(2)
    with c1:
        _render_shift_tile(
            ("SIŪLO / GIVES" if lang=="LT" else "OFFERS"),
            a,sa,PERSON_COLORS.get(a)
        )
    with c2:
        _render_shift_tile(
            ("PRAŠO / WANTS" if lang=="LT" else "REQUESTS"),
            b,sb,PERSON_COLORS.get(b)
        )



def _swap_hard_user_explanation(code, details=""):
    """Plain-language explanation for a true ACTUAL swap blocker."""
    lt={
        "ONKO_EVEN_PARITY":(
            "Onko porų taisyklė",
            "Po apsikeitimo bent vieno rezidento Onko skaičius taptų nelyginis. Onko ACTUAL grafike turi likti 0 / 2 / 4 / …"
        ),
        "ONKO_COVERAGE":(
            "Privalomas Onko padengimas",
            "Apsikeitimas sugadintų reikiamą bendrą Onko padengimą."
        ),
        "OVERLAPPING_ASSIGNMENTS":(
            "Persidengiančios pamainos",
            "Po apsikeitimo tam pačiam rezidentui tuo pačiu metu būtų paskirtos dvi persidengiančios pamainos."
        ),
        "MAX_HOURS_PER_DAY":(
            "Maksimali darbo trukmė per dieną",
            "Po apsikeitimo būtų viršyta ACTUAL swapo maksimali 12 val. darbo trukmė per vieną dieną."
        ),
        "MAX_WORKDAYS_7D":(
            "Maksimalus darbo dienų skaičius per 7 dienas",
            "Po apsikeitimo rezidentas dirbtų daugiau nei leidžiamos 6 darbo dienos per slenkantį 7 dienų langą."
        ),
        "MAX_HOURS_7D":(
            "Maksimali darbo trukmė per 7 dienas",
            "Po apsikeitimo būtų viršytas ACTUAL swapo absoliutus valandų limitas per slenkantį 7 dienų langą (iki 60 val.)."
        ),
        "MIN_DAILY_REST":(
            "Minimalus 11 val. paros poilsis",
            "Po apsikeitimo tarp dviejų darbo dienų liktų mažiau nei 11 val. nepertraukiamo poilsio."
        ),
        "POST_DUTY_REST":(
            "Privalomas poilsis po budėjimo",
            "Apsikeitimas paskirtų darbą dieną, kuri turi likti laisva po ilgo budėjimo."
        ),
        "ABSOLUTE_UNAVAILABILITY":(
            "Absoliutus nebuvimas",
            "Apsikeitimas paskirtų darbą per atostogas ar kitą absoliučiai pateisintą nebuvimą."
        ),
        "BLOCKED_SLOT":(
            "Uždaryta / neegzistuojanti pamaina",
            "Apsikeitimas bandytų užpildyti pamainą, kuri pagal klinikos modelį tuo metu yra uždaryta."
        ),
        "MANDATORY_COVERAGE":(
            "Privalomas klinikos padengimas",
            "Po apsikeitimo liktų neužpildyta privaloma klinikos pamaina."
        ),
        "MANDATORY_BACKUP_AVAILABILITY":(
            "Privalomas backup padengimas",
            "Po apsikeitimo kritinei pamainai neliktų nė vieno HARD-available, nepersidengiančio backup rezidento."
        ),
        "OTHER_OPERATIONAL_HARD":(
            "Kita operational HARD taisyklė",
            "Šis apsikeitimas pažeistų kitą neapeinamą saugos / darbo laiko / klinikos padengimo taisyklę."
        ),
    }
    en={
        "ONKO_EVEN_PARITY":("Even Onko parity","The swap would leave an odd Onko count. ACTUAL Onko must remain 0 / 2 / 4 / …"),
        "ONKO_COVERAGE":("Required Onko coverage","The swap would break required overall Onko coverage."),
        "OVERLAPPING_ASSIGNMENTS":("No overlapping shifts","The swap would give a resident overlapping assignments at the same time."),
        "MAX_HOURS_PER_DAY":("Maximum hours per day","The swap would exceed the ACTUAL 12-hour daily maximum."),
        "MAX_WORKDAYS_7D":("Maximum workdays in 7 days","The swap would exceed 6 working days in a rolling 7-day window."),
        "MAX_HOURS_7D":("Maximum hours in 7 days","The swap would exceed the ACTUAL absolute rolling-7-day limit (up to 60 hours)."),
        "MIN_DAILY_REST":("Minimum 11-hour daily rest","The swap would leave less than 11 uninterrupted hours of rest between workdays."),
        "POST_DUTY_REST":("Mandatory post-duty rest","The swap would assign work on a required post-duty rest day."),
        "ABSOLUTE_UNAVAILABILITY":("Absolute unavailability","The swap would assign work during vacation or another absolute justified absence."),
        "BLOCKED_SLOT":("Closed / blocked shift","The swap would fill a shift that is closed in the clinic model."),
        "MANDATORY_COVERAGE":("Mandatory clinical coverage","The swap would leave a mandatory clinical shift unfilled."),
        "MANDATORY_BACKUP_AVAILABILITY":("Required backup availability","The swap would leave a critical shift without any HARD-available non-overlapping backup resident."),
        "OTHER_OPERATIONAL_HARD":("Other operational HARD rule","The swap would violate another non-relaxable safety, work-time or clinical-coverage rule."),
    }
    table=lt if lang=="LT" else en
    return table.get(str(code),table["OTHER_OPERATIONAL_HARD"])


def _render_swap_hard_block(preview_stats, fallback_reason=""):
    rows=((preview_stats or {}).get("global",{}) or {}).get("swap_hard_block_rows") or []
    if not rows:
        st.error(
            ("Apsikeitimas negalimas: " if lang=="LT" else "Swap cannot be completed: ")
            + str(fallback_reason or "unknown reason")
        )
        return

    pretty=[]
    for r in rows:
        code=str(r.get("code") or "OTHER_OPERATIONAL_HARD")
        title,why=_swap_hard_user_explanation(code,r.get("details",""))
        pretty.append({
            ("HARD taisyklė" if lang=="LT" else "HARD rule"):title,
            ("Kodėl negalima" if lang=="LT" else "Why blocked"):why,
            ("Techninė detalė" if lang=="LT" else "Technical detail"):str(r.get("details") or ""),
        })

    first=pretty[0]
    st.error(
        (
            f"APSIKEITIMAS NEGALIMAS — pažeidžiama HARD taisyklė: {first['HARD taisyklė']}"
            if lang=="LT" else
            f"SWAP BLOCKED — HARD rule violated: {first['HARD rule']}"
        )
    )
    st.dataframe(pd.DataFrame(pretty),use_container_width=True,hide_index=True)


def _friday_waterfill_proof(stats):
    pdata=((stats or {}).get("people") or {})
    counts={i:int((d or {}).get("friday_assignments",0) or 0) for i,d in pdata.items()}
    vals=list(counts.values())
    total=int(sum(vals))
    n=len(vals)
    lo=int(total//n) if n else 0
    hi=int((total+n-1)//n) if n else 0
    spread=(max(vals)-min(vals)) if vals else 0
    passed=bool(vals and spread<=1 and all(lo<=v<=hi for v in vals))
    return {"counts":counts,"total":total,"n":n,"floor":lo,"ceil":hi,"spread":spread,"passed":passed}


def _delete_confirm(token, applied=False):
    state_key="_confirm_delete_schedule_action"
    label=("IŠTRINTI / UNDO" if lang=="LT" else "DELETE / UNDO") if applied else ("IŠTRINTI" if lang=="LT" else "DELETE")
    if st.button(label,key=f"delete_open_{token}",use_container_width=True):
        st.session_state[state_key]=token
    if st.session_state.get(state_key)!=token:
        return False
    st.warning(
        ("Patvirtinkite. Jei veiksmas jau pakeitė ACTUAL, sistema bandys saugiai grąžinti ankstesnę būseną. Jei tie slotai po to buvo pakeisti dar kartą, DELETE bus atmestas, kad nesugadintų naujesnio grafiko."
         if lang=="LT" else
         "Confirm. If this action already changed ACTUAL, the system will safely restore the previous state. If those slots changed again later, DELETE will be refused to protect newer changes.")
    )
    c1,c2=st.columns(2)
    with c1:
        yes=st.button("PATVIRTINTI DELETE" if lang=="LT" else "CONFIRM DELETE",type="primary",key=f"delete_yes_{token}",use_container_width=True)
    with c2:
        no=st.button("ATŠAUKTI" if lang=="LT" else "CANCEL",key=f"delete_no_{token}",use_container_width=True)
    if no:
        st.session_state.pop(state_key,None)
        st.rerun()
    if yes:
        st.session_state.pop(state_key,None)
        return True
    return False


def _prepare_swap_action_undo(row,y,m):
    meta=_swap_meta_decode(row.get("reason"))
    kind=str(meta.get("kind") or "")
    phase=str(meta.get("phase") or "")
    applied=bool(kind=="emergency_rescue" or phase=="applied")
    if not applied:
        return None,None

    fresh=refresh_result_payload(db.load_schedule(y,m,"current"),y,m)
    candidate=deepcopy(fresh)
    candidate.assignments=dict(fresh.assignments)
    sa=int(row["slot_a"]); sb=int(row["slot_b"])

    if kind=="emergency_rescue":
        if sa in candidate.assignments:
            raise RuntimeError("Emergency Rescue source slotas nebėra tuščias — po Rescue jis jau buvo pakeistas." if lang=="LT" else "Emergency Rescue source is no longer vacant; it changed after the Rescue.")
        if candidate.assignments.get(sb)!=row.get("person_a"):
            raise RuntimeError("Emergency Rescue target slotas po Rescue jau buvo pakeistas." if lang=="LT" else "Emergency Rescue target changed after the Rescue.")
        candidate.assignments[sa]=row.get("person_a")
        candidate.assignments[sb]=row.get("person_b")
        mode="voluntary_swap_actual"
    else:
        if candidate.assignments.get(sa)!=row.get("person_b") or candidate.assignments.get(sb)!=row.get("person_a"):
            raise RuntimeError("Šio swapo slotai po pritaikymo jau buvo pakeisti dar kartą — automatinis UNDO nebesaugus." if lang=="LT" else "These swap slots changed again after application; automatic UNDO is no longer safe.")
        candidate.assignments[sa]=row.get("person_a")
        candidate.assignments[sb]=row.get("person_b")
        mode="voluntary_swap_actual"

    candidate=revalidate_loaded_result(
        y,m,people_for_stored_result(candidate,y,m),candidate,
        backup_assignments=None,validation_mode=mode,
    )
    desired,backup_errors=plan_backups(y,m,candidate)
    if backup_errors:
        raise RuntimeError(("UNDO negalimas: nepavyksta atkurti privalomo backup plano: " if lang=="LT" else "UNDO blocked: required backup plan cannot be rebuilt: ")+str(backup_errors[0]))
    candidate=revalidate_loaded_result(
        y,m,people_for_stored_result(candidate,y,m),candidate,
        backup_assignments=desired,validation_mode=mode,
    )
    errs=list(((candidate.stats or {}).get("global") or {}).get("errors") or [])
    if errs:
        raise RuntimeError(("UNDO negalimas dėl operational HARD: " if lang=="LT" else "UNDO blocked by operational HARD: ")+str(errs[0]))
    return candidate,desired


def _delete_swap_row(row,y,m):
    candidate,desired=_prepare_swap_action_undo(row,y,m)
    saved=db.delete_swap_action_v2586(
        int(row["id"]),
        serialize_result(candidate) if candidate is not None else None,
        desired if desired is not None else None,
    )
    if saved.get("undone_actual"):
        try: persist_actual_satisfaction(y,m)
        except Exception: pass
        try: refresh_calendar_subscription_feeds([row.get("person_a"),row.get("person_b")])
        except Exception: pass
    return saved


def month_label(y,m): return f"{MONTHS[lang][m-1]} {y}"
def block_label(b): return {"AM":tr("morning"),"PM":tr("afternoon"),"FULL":tr("full_day")}[b]
def pretty_day(y,m,d): return f"{d:02d} {WEEKDAYS[lang][date(y,m,d).weekday()]}"
def safe_filename(s): return "".join(ch for ch in unicodedata.normalize("NFKD",s).encode("ascii","ignore").decode() if ch.isalnum() or ch in "_-")
def ics_escape(s): return str(s).replace("\\","\\\\").replace(";","\\;").replace(",","\\,").replace("\n","\\n")

def deadline_day():
    return int(rule_value("deadline_day"))


def deadline_for(y,m):
    dd=deadline_day()
    if m==1: return date(y-1,12,dd)
    return date(y,m-1,dd)

def preference_cutoff_for(y,m):
    """Exact cutoff: immediately after the configured last full submission day."""
    last_full_day=deadline_for(y,m)
    return datetime.combine(last_full_day+timedelta(days=1),time(0,0),tzinfo=ZoneInfo("Europe/Vilnius"))

def deadline_message(y,m):
    dl=deadline_for(y,m); today=date.today(); diff=(dl-today).days
    if diff>0: msg=tr("deadline_future").format(n=diff)
    elif diff==0: msg=tr("deadline_today")
    else: msg=tr("deadline_passed").format(n=abs(diff))
    return dl,msg,diff

def ensure_zero_preference_submissions_if_due(y,m):
    """Operator-side automatic completion of missing preference forms after cutoff.

    Missing residents become submitted with exactly zero monthly requests. This never
    creates HARD/SOFT wishes and is transparent via submission_source=deadline_zero.
    """
    cutoff=preference_cutoff_for(y,m)
    now_lt=datetime.now(ZoneInfo("Europe/Vilnius"))
    if now_lt < cutoff:
        return {"ok":True,"due":False,"count":0,"initials":[]}
    try:
        return db.auto_submit_zero_preferences_v2594(y,m,cutoff.isoformat())
    except Exception as e:
        # Do not break the whole app if a transient DB call fails; senior dashboard
        # will still show any genuinely missing rows and generation stays inspectable.
        return {"ok":False,"due":True,"count":0,"initials":[],"error":str(e)}

_SMTP_SECRET_KEYS={
    "SCHEDULER_SMTP_HOST":"host",
    "SCHEDULER_SMTP_PORT":"port",
    "SCHEDULER_SMTP_USER":"user",
    "SCHEDULER_SMTP_PASSWORD":"password",
    "SCHEDULER_SMTP_USE_TLS":"use_tls",
    "SCHEDULER_SMTP_USE_SSL":"use_ssl",
    "SCHEDULER_EMAIL_FROM":"from_email",
    "SCHEDULER_SMTP_PROVIDER":"provider",
}

def config_value(name, default=""):
    """Read deployment configuration from root Streamlit secrets, nested [smtp], or environment.

    V2.5.93 accepts both the historical SCHEDULER_* secret names and a simpler
    nested [smtp] block without ever exposing the password in the UI.
    """
    try:
        if name in st.secrets:
            return str(st.secrets[name])
        key=_SMTP_SECRET_KEYS.get(name)
        if key and "smtp" in st.secrets:
            block=st.secrets["smtp"]
            if key in block:
                return str(block[key])
    except Exception:
        pass
    return os.environ.get(name, default)

def get_supabase_client():
    if "supabase_client" not in st.session_state:
        url=config_value("SUPABASE_URL",DEFAULT_SUPABASE_URL)
        key=config_value("SUPABASE_PUBLISHABLE_KEY",DEFAULT_SUPABASE_PUBLISHABLE_KEY)
        st.session_state["supabase_client"]=create_client(url,key)
    return st.session_state["supabase_client"]

def authenticated_user(sb):
    try:
        r=sb.auth.get_user()
        return getattr(r,"user",None)
    except Exception:
        return None

def _auth_uid(user):
    return str(getattr(user,"id","") or "").strip()


def clear_cross_account_session_state(*, keep_client=True):
    """Remove user-specific Streamlit state on logout/account switch.

    The Supabase client may be kept so its authenticated session survives the
    rerun. Everything else is rebuilt for the exact auth.uid().
    """
    keep={"supabase_client"} if keep_client else set()
    # Language is UI-only and carries no resident data.
    keep.add("language")
    for key in list(st.session_state.keys()):
        if key not in keep:
            st.session_state.pop(key,None)


def enforce_auth_session_identity(user):
    """Fail closed if Streamlit state belongs to another authenticated account."""
    uid=_auth_uid(user)
    if not uid:
        return
    previous=str(st.session_state.get("_identity_auth_uid") or "")
    if previous and previous!=uid:
        clear_cross_account_session_state(keep_client=True)
    st.session_state["_identity_auth_uid"]=uid


def render_auth_gate():
    sb=get_supabase_client(); user=authenticated_user(sb)
    if user is not None:
        enforce_auth_session_identity(user)
        return sb,user
    st.title(tr("login_title"))
    login_tab,signup_tab=st.tabs([tr("login_title"),tr("signup_title")])
    with login_tab:
        with st.form("login_form"):
            email=st.text_input(tr("auth_email"),key="login_email")
            password=st.text_input(tr("password"),type="password",key="login_password")
            go=st.form_submit_button(tr("login"),type="primary",use_container_width=True)
            if go:
                try:
                    sb.auth.sign_in_with_password({"email":email.strip(),"password":password})
                    st.rerun()
                except Exception:
                    st.error(tr("auth_invalid"))
    with signup_tab:
        with st.form("signup_form"):
            email=st.text_input(tr("auth_email"),key="signup_email")
            p1=st.text_input(tr("password"),type="password",key="signup_password")
            p2=st.text_input(tr("password_repeat"),type="password",key="signup_password2")
            go=st.form_submit_button(tr("signup"),use_container_width=True)
            if go:
                if p1!=p2 or len(p1)<8:
                    st.error(tr("signup_password_mismatch"))
                else:
                    try:
                        response=sb.auth.sign_up({"email":email.strip(),"password":p1})
                        st.success(tr("signup_sent"))
                        st.caption(
                            "Registracijos bandymų skaičiaus mūsų sistemoje neribojame. Invite kodas sunaudojamas tik vėliau, kai sėkmingai susiejamas rezidento profilis."
                            if lang=="LT" else
                            "Our app does not limit registration attempts. The invite code is consumed only later, when the resident profile is successfully linked."
                        )
                        if getattr(response,"session",None) is not None:
                            st.rerun()
                    except Exception as e:
                        msg=str(e or "")
                        low=msg.lower()
                        if "email rate limit" in low or "rate limit exceeded" in low:
                            st.error(
                                "Supabase laikinai pasiekė Auth el. laiškų limitą. Tavo invite kodas NEPANAUDOTAS ir paskyra nuo šio bandymo neužsirakino. "
                                "Gali bandyti dar kartą vėliau tiek kartų, kiek reikia. Šio providerio valandinio email limito aplikacija pati nunulinti negali."
                                if lang=="LT" else
                                "Supabase temporarily reached the Auth email limit. Your invite code is NOT consumed and this attempt did not lock the account. "
                                "You can retry later as many times as needed. The app cannot reset the provider's hourly email quota itself."
                            )
                        else:
                            st.error(msg)
    st.stop()

def require_linked_profile(sb,user):
    db.set_client(sb)
    auth_uid=_auth_uid(user)
    if not auth_uid:
        st.error("AUTH IDENTITY ERROR: authenticated user has no UID.")
        st.stop()

    profile=db.current_profile(auth_uid)
    if profile:
        if str(profile.get("user_id") or "")!=auth_uid:
            st.error(
                "IDENTITY SAFETY LOCK: prisijungusi paskyra neatitinka rezidento profilio. Prieiga sustabdyta."
                if lang=="LT" else
                "IDENTITY SAFETY LOCK: authenticated account does not match the resident profile. Access stopped."
            )
            st.stop()
        if profile.get("approved") and (
            profile.get("initials") or profile.get("access_role")=="observer"
        ):
            return profile

    directory=db.directory()
    st.title(tr("claim_title"))
    resident_tab,observer_tab=st.tabs([tr("resident_claim_tab"),tr("observer_claim_tab")])

    with resident_tab:
        st.info(tr("claim_help"))
        with st.form("claim_profile_form"):
            initials=st.selectbox(tr("user"),list(directory),format_func=lambda i:f"{i} — {directory[i]['full_name']}")
            code=st.text_input(tr("invite_code"),type="password")
            go=st.form_submit_button(tr("claim"),type="primary")
            if go:
                try:
                    db.claim_profile(initials,code.strip())
                    st.rerun()
                except Exception:
                    st.error(tr("claim_failed"))

    with observer_tab:
        st.info(tr("observer_claim_help"))
        with st.form("claim_observer_form"):
            observer_code=st.text_input(tr("observer_invite_code"),type="password")
            observer_go=st.form_submit_button(tr("observer_claim"),type="primary")
            if observer_go:
                try:
                    db.claim_observer_profile(observer_code.strip())
                    st.success(tr("observer_access_ready"))
                    st.rerun()
                except Exception:
                    st.error(tr("claim_failed"))
    st.stop()

def recurring_dates_for_month(y,m,rows):
    _,ndays=calendar.monthrange(y,m)
    out={"unavailable":set(),"unavailable_am":set(),"unavailable_pm":set(),"soft_free":set(),"preferred":set()}
    for r in rows:
        if not r.get("active",True): continue
        wd=int(r.get("weekday",0)); typ=r.get("preference_type"); block=r.get("block","FULL")
        days={d for d in range(1,ndays+1) if date(y,m,d).weekday()==wd}
        if typ=="hard_unavailable":
            if block=="AM": out["unavailable_am"] |= days
            elif block=="PM": out["unavailable_pm"] |= days
            else: out["unavailable"] |= days
        elif typ=="soft_free": out["soft_free"] |= days
        elif typ=="preferred": out["preferred"] |= days
    return out

def historical_rotation_counts_before(y,m):
    """Count prior published SYSTEM workplace exposures for longitudinal catch-up."""
    out={p["initials"]:{cat:0 for cat in ROTATION_CATEGORIES} for p in DEFAULT_PEOPLE}
    try:
        rows=db.published_baselines_before(y,m)
    except Exception:
        return out
    for r in rows:
        try:
            yy=int(r["year"]); mm=int(r["month"])
            payload=r.get("baseline_json") or {}
            ass={int(k):v for k,v in (payload.get("assignments") or {}).items()}
            slot_map={s.idx:s for s in make_slots(yy,mm)}
            for sid,ini in ass.items():
                s=slot_map.get(int(sid))
                if s is None or ini not in out:
                    continue
                cat=rotation_category(s)
                if cat in out[ini]:
                    out[ini][cat]+=1
        except Exception:
            continue
    return out


def historical_holiday_counts_before(y,m):
    """Count prior published SYSTEM public-holiday assignments for longitudinal rotation."""
    out={p["initials"]:0 for p in DEFAULT_PEOPLE}
    try:
        rows=db.published_baselines_before(y,m)
    except Exception:
        return out
    for r in rows:
        try:
            yy=int(r["year"]); mm=int(r["month"])
            payload=r.get("baseline_json") or {}
            ass={int(k):v for k,v in (payload.get("assignments") or {}).items()}
            slot_map={sl.idx:sl for sl in make_slots(yy,mm)}
            for sid,ini in ass.items():
                sl=slot_map.get(int(sid))
                if sl is not None and ini in out and is_public_holiday(yy,mm,sl.day):
                    out[ini]+=1
        except Exception:
            continue
    return out


def _previous_month_effective_actual_assignments(y,m):
    """Immediately prior month's real/effective work for cross-boundary safety only."""
    py,pm=(y-1,12) if m==1 else (y,m-1)
    try:
        rows=db.list_published_schedules()
        row=next((r for r in reversed(rows) if int(r.get("year",0))==py and int(r.get("month",0))==pm),None)
        if not row or not row.get("current_json"):
            return py,pm,{}
        payload=row.get("current_json") or {}
        ass={int(k):v for k,v in (payload.get("assignments") or {}).items()}
        return py,pm,effective_actual_assignments(ass,db.list_backups(py,pm))
    except Exception:
        return py,pm,{}


def historical_weekend_tail_streak_before(y,m):
    """Prior-month ACTUAL weekend tail streak; spacing only, never catch-up."""
    out={p["initials"]:0 for p in DEFAULT_PEOPLE}
    py,pm,ass=_previous_month_effective_actual_assignments(y,m)
    if not ass:
        return out
    try:
        slot_map={s.idx:s for s in make_slots(py,pm)}
        anchors=sorted({(sl.day if sl.weekday==5 else sl.day-1) for sl in slot_map.values() if sl.weekday>=5})
        for ini in out:
            worked={(slot_map[sid].day if slot_map[sid].weekday==5 else slot_map[sid].day-1)
                    for sid,who in ass.items() if who==ini and sid in slot_map and slot_map[sid].weekday>=5}
            streak=0
            for a in reversed(anchors):
                if a in worked: streak+=1
                else: break
            out[ini]=streak
    except Exception:
        return {p["initials"]:0 for p in DEFAULT_PEOPLE}
    return out


def historical_previous_last_day_onko_before(y,m):
    """Prior-month ACTUAL last-day Onko state for cross-boundary safety only."""
    out={p["initials"]:False for p in DEFAULT_PEOPLE}
    py,pm,ass=_previous_month_effective_actual_assignments(y,m)
    if not ass:
        return out
    try:
        slot_map={s.idx:s for s in make_slots(py,pm)}
        last_day=calendar.monthrange(py,pm)[1]
        for sid,ini in ass.items():
            sl=slot_map.get(int(sid))
            if sl is not None and ini in out and sl.day==last_day and sl.department=="Onko RO centre":
                out[ini]=True
    except Exception:
        return {p["initials"]:False for p in DEFAULT_PEOPLE}
    return out


def historical_resident_hard_losses_before(y,m):
    """Legacy audit-only RESIDENT-HARD violation count from prior SYSTEM baselines. V2.5.107 never uses it as a generation input."""
    out={p["initials"]:0 for p in DEFAULT_PEOPLE}
    try:
        rows=db.published_baselines_before(y,m)
    except Exception:
        return out
    for r in rows:
        try:
            payload=r.get("baseline_json") or {}
            people_stats=((payload.get("stats") or {}).get("people") or {})
            for initials,d in people_stats.items():
                if initials in out:
                    out[initials]+=int((d or {}).get("resident_hard_losses",0) or 0)
        except Exception:
            continue
    return out


def _request_source(monthly_set, recurring_set, day):
    a=day in set(monthly_set or set()); b=day in set(recurring_set or set())
    if a and b: return "monthly+recurring"
    if a: return "monthly"
    if b: return "recurring"
    return "effective"


def _build_request_ledger(y,m,initials,p,s,rp,recurring_rows,claims,slot_lookup,
                          effective_unavail,effective_unavail_am,effective_unavail_pm,
                          effective_soft,effective_soft_am,effective_soft_pm,
                          effective_pref,effective_pref_am,effective_pref_pm):
    """One normalized resident-facing ledger across ALL structured input tables."""
    items=[]
    def add(kind,tier,day=None,block="FULL",source="monthly",value=None,score=True,**extra):
        rid=f"{source}:{kind}:{day if day is not None else '-'}:{block}:{len(items)}"
        row={"id":rid,"kind":kind,"tier":tier,"day":day,"block":block,"source":source,
             "value":value,"included_in_score":bool(score)}
        row.update(extra); items.append(row)

    # RESIDENT HARD: monthly and/or recurring `Negaliu dirbti`.
    monthly_u=set(p.get("unavailable",set())); monthly_am=set(p.get("unavailable_am",set())); monthly_pm=set(p.get("unavailable_pm",set()))
    for d in sorted(effective_unavail): add("resident_hard","RESIDENT_HARD",d,"FULL",_request_source(monthly_u,rp["unavailable"],d))
    for d in sorted(set(effective_unavail_am)-set(effective_unavail)): add("resident_hard","RESIDENT_HARD",d,"AM",_request_source(monthly_am,rp["unavailable_am"],d))
    for d in sorted(set(effective_unavail_pm)-set(effective_unavail)): add("resident_hard","RESIDENT_HARD",d,"PM",_request_source(monthly_pm,rp["unavailable_pm"],d))

    # Exact SOFT monthly/recurring wishes after the same override logic used by the engine.
    m_sf=set(p.get("soft_free",set())); m_sfa=set(p.get("soft_free_am",set())); m_sfp=set(p.get("soft_free_pm",set()))
    m_pr=set(p.get("preferred",set())); m_pra=set(p.get("preferred_am",set())); m_prp=set(p.get("preferred_pm",set()))
    for d in sorted(effective_soft): add("soft_free","SOFT1_TIME_PROTECTION",d,"FULL",_request_source(m_sf,rp["soft_free"],d))
    for d in sorted(effective_soft_am): add("soft_free","SOFT1_TIME_PROTECTION",d,"AM","monthly")
    for d in sorted(effective_soft_pm): add("soft_free","SOFT1_TIME_PROTECTION",d,"PM","monthly")
    for d in sorted(effective_pref): add("preferred","SOFT2_POSITIVE_PLACEMENT",d,"FULL",_request_source(m_pr,rp["preferred"],d))
    for d in sorted(effective_pref_am): add("preferred","SOFT2_POSITIVE_PLACEMENT",d,"AM","monthly")
    for d in sorted(effective_pref_pm): add("preferred","SOFT2_POSITIVE_PLACEMENT",d,"PM","monthly")

    # V2.5.102 persistent work-style settings are real SOFT inputs. Weekend
    # direction is optimized only after the raw Saturday/Sunday water-fill locks,
    # so it can choose the upper/lower fair layer but can never widen SYSTEM fairness.
    if int(s.get("weekday_preference",0) or 0): add("weekday_preference","SOFT3_SCHEDULE_SHAPE",source="account_settings",value=max(-2,min(2,int(s.get("weekday_preference",0) or 0))))
    # V2.5.112: weekend preference is administratively retired for SYSTEM.
    # Historical DB values are retained for audit compatibility but are not an active wish.
    if int(s.get("spread_preference",0) or 0): add("spread_preference","SOFT3_SCHEDULE_SHAPE",source="account_settings",value=int(s.get("spread_preference",0)))
    # Dedicated holiday inclination is one normalized SOFT unit only in months
    # that actually contain official public-holiday duty slots.
    if int(s.get("holiday_preference",0) or 0) and public_holiday_days_in_month(y,m):
        add("holiday_preference","SOFT_HOLIDAY",source="account_settings",value=max(-1,min(1,int(s.get("holiday_preference",0) or 0))))
    shift_len=max(0,min(3,int(s.get("shift_length_preference",0) or 0)))
    if shift_len:
        add("shift_length_preference","SOFT3_SCHEDULE_SHAPE",source="account_settings",value=shift_len)
    if bool(s.get("avoid_doubles",False)) and shift_len==0:
        add("avoid_doubles","SOFT1_TIME_PROTECTION",source="account_settings",value=True)

    # ABSOLUTE HARD / safety facts are audited but excluded from the preference % denominator.
    for d in sorted(set(p.get("vacation",set()))): add("vacation","ABSOLUTE_HARD",d,"FULL","monthly",score=False)
    for d in sorted(set(p.get("justified_absence",set()))): add("justified_absence","ABSOLUTE_HARD",d,"FULL","monthly",score=False)
    for d in sorted(set(p.get("long_duty",set()))): add("long_duty","ABSOLUTE_HARD",d,"FULL","monthly",score=False)

    # Self-selected backup commitments and valid rest-credit redemptions are structured resident choices.
    for claim in claims or []:
        cs=slot_lookup.get(int(claim.get("covered_slot")))
        if cs is not None:
            add("backup_claim","RESIDENT_CHOICE",cs.day,cs.block,"backup_claim",covered_slot=cs.idx,department=cs.department)
    for _ in range(max(0,int(p.get("backup_credits_am_to_use",0) or 0))): add("rest_credit","ENTITLEMENT",source="rest_credit",value="AM")
    for _ in range(max(0,int(p.get("backup_credits_pm_to_use",0) or 0))): add("rest_credit","ENTITLEMENT",source="rest_credit",value="PM")

    note=str(p.get("note") or "").strip()
    if note:
        add("note","INFO",source="free_text",value=note,score=False)
    return items

def load_people(y,m):
    prefs=db.all_preferences(y,m); settings=db.all_account_settings(); recurring=db.all_recurring_preferences(); people=[]
    # V2.5.96: fairness history is audit-only; never solver input for a new month.
    fairness_prior={}
    holiday_prior={}
    rotation_prior={}
    resident_hard_prior={}
    # Cross-month safety/spacing state remains because it is not fairness catch-up.
    weekend_tail=historical_weekend_tail_streak_before(y,m)
    previous_last_day_onko=historical_previous_last_day_onko_before(y,m)
    claim_rows=db.list_backup_claims(y,m)
    claims_by_initials={}
    for r in claim_rows:
        claims_by_initials.setdefault(r["initials"],[]).append(r)
    slot_lookup={s.idx:s for s in make_slots(y,m)}
    py,pm=(y-1,12) if m==1 else (y,m-1)
    prev_prefs=db.all_preferences(py,pm); prev_last=calendar.monthrange(py,pm)[1]
    for row in DEFAULT_PEOPLE:
        initials=row["initials"]; p=prefs.get(initials,{}); s=settings.get(initials,{})
        duty_days=set(p.get("long_duty",set()))
        if prev_last in set(prev_prefs.get(initials,{}).get("long_duty",set())):
            duty_days.add(0)
        rp=recurring_dates_for_month(y,m,recurring.get(initials,[]))
        short_soft=set(p.get("soft_free",set())); short_soft_am=set(p.get("soft_free_am",set())); short_soft_pm=set(p.get("soft_free_pm",set()))
        short_pref=set(p.get("preferred",set())); short_pref_am=set(p.get("preferred_am",set())); short_pref_pm=set(p.get("preferred_pm",set()))
        # Specific monthly requests override an opposite recurring whole-day pattern on that date.
        any_short_pref=short_pref|short_pref_am|short_pref_pm
        any_short_soft=short_soft|short_soft_am|short_soft_pm
        # V2.5.90 BASELINE WEEKEND VOLUNTEER OVERRIDE.
        # Recurring weekend "Noriu laisvos" remains blocked because it can dump
        # unavoidable weekend burden onto peers. The OPPOSITE signal is allowed:
        # recurring weekend "Pageidauju dirbti" is a voluntary unpopular-duty offer.
        recurring_soft_allowed={d for d in set(rp["soft_free"]) if date(y,m,d).weekday()<5}
        recurring_pref_allowed=set(rp["preferred"])
        effective_soft=(recurring_soft_allowed-any_short_pref)|short_soft
        effective_pref=(recurring_pref_allowed-any_short_soft)|short_pref
        # Long-term RESIDENT HARD cannot be overridden by an opposite monthly SOFT
        # request. Whole-day resident hard subsumes same-day AM/PM rows.
        effective_unavail=set(rp["unavailable"])|set(p.get("unavailable",set()))
        effective_unavail_am=(set(rp["unavailable_am"])|set(p.get("unavailable_am",set())))-effective_unavail
        effective_unavail_pm=(set(rp["unavailable_pm"])|set(p.get("unavailable_pm",set())))-effective_unavail
        effective_soft_am=short_soft_am
        effective_soft_pm=short_soft_pm
        effective_pref_am=short_pref_am
        effective_pref_pm=short_pref_pm
        credits_am=int(p.get("backup_credits_am_to_use",0))
        credits_pm=int(p.get("backup_credits_pm_to_use",0))
        credits_day=credits_am+credits_pm
        prior=fairness_prior.get(initials,{})
        reserved=set()
        for claim in claims_by_initials.get(initials,[]):
            cs=slot_lookup.get(int(claim["covered_slot"]))
            if cs is not None:
                reserved.add((cs.day,cs.block))
        request_items=_build_request_ledger(
            y,m,initials,p,s,rp,recurring.get(initials,[]),claims_by_initials.get(initials,[]),slot_lookup,
            effective_unavail,effective_unavail_am,effective_unavail_pm,
            effective_soft,effective_soft_am,effective_soft_pm,
            effective_pref,effective_pref_am,effective_pref_pm
        )
        people.append(Person(initials=initials,name=row["name"],target_adjustment=row.get("target_adjustment",0)-credits_day,
            unavailable=effective_unavail,
            unavailable_am=effective_unavail_am,
            unavailable_pm=effective_unavail_pm,
            vacation=set(p.get("vacation",set())),
            justified_absence=set(p.get("justified_absence",set())),
            long_duty=duty_days,reserved_backup=reserved,
            soft_free=effective_soft,soft_free_am=effective_soft_am,soft_free_pm=effective_soft_pm,
            preferred=effective_pref,preferred_am=effective_pref_am,preferred_pm=effective_pref_pm,
            weekday_preference=max(-2,min(2,int(s.get("weekday_preference",0) or 0))),weekend_preference=0,holiday_preference=max(-1,min(1,int(s.get("holiday_preference",0) or 0))),spread_preference=int(s.get("spread_preference",0)),
            shift_length_preference=max(0,min(3,int(s.get("shift_length_preference",0) or 0))),
            avoid_doubles=(max(0,min(3,int(s.get("shift_length_preference",0) or 0)))==1 or bool(s.get("avoid_doubles",False))),note=p.get("note",""),
            request_items=request_items,rest_credit_am_to_use=credits_am,rest_credit_pm_to_use=credits_pm,
            prior_weekend_count=0,
            prior_holiday_count=0,
            prior_friday_count=0,
            prior_double_count=0,
            prior_weekday_day_count=0,
            prior_rotation_counts={},
            prior_consecutive_weekend_streak=int(weekend_tail.get(initials,0)),
            prior_last_day_onko=bool(previous_last_day_onko.get(initials,False)),
            prior_resident_hard_loss_count=0))
    return people


def people_for_stored_result(result, y, m):
    """Use the immutable publication-time request snapshot whenever available.

    This is both semantically correct and operationally resilient: post-publication
    swap/revalidation must use the frozen request set, so it should not re-read
    preferences/account_settings/recurring_preferences on every Streamlit rerun.
    Legacy payloads without a snapshot fall back to live load_people().
    """
    frozen=people_from_request_snapshot(getattr(result,"request_snapshot",None))
    if frozen:
        return frozen
    return load_people(y,m)


def refresh_result_payload(payload, y, m, use_actual_backups=True):
    """Revalidate stored assignments against the CURRENT engine and ORIGINAL requests.

    SYSTEM/baseline views pass ``use_actual_backups=False`` so their publication-time
    backup snapshot stays frozen. ACTUAL/current views use the live backup table, so
    swaps and backup changes update realized request satisfaction without rewriting
    the original SYSTEM fairness ledger.
    """
    if not payload:
        return None
    stored=deserialize_result(payload)
    # V2.5.115 — theoretical backup duties are a separate standby layer.
    # SYSTEM/DRAFT request satisfaction is computed from NORMAL work only.
    # ACTUAL may pass live backup rows so a COMPLETED real-life cover can be
    # reflected as actual work; planned/activated-only standby still never counts.
    backup_override=[]
    if use_actual_backups:
        try:
            backup_override=db.list_backups(y,m)
        except Exception:
            backup_override=stored.backup_snapshot or []
    try:
        refreshed=revalidate_loaded_result(
            y,m,people_for_stored_result(stored,y,m),stored,
            backup_assignments=backup_override,
            validation_mode=("voluntary_swap_actual" if use_actual_backups else "generation"),
        )
    except TypeError as exc:
        # Deployment-safety guard: never let a mixed app/engine deployment crash
        # the whole Apsikeitimai page with a raw unexpected-keyword TypeError.
        if "validation_mode" in str(exc):
            raise RuntimeError(
                "APP_ENGINE_VERSION_MISMATCH: deployed app.py expects the V2.5.107 "
                "scheduler_engine.py API. Redeploy BOTH files from the same package."
            ) from exc
        raise
    if use_actual_backups:
        # V2.5.57: CURRENT stats are operational / satisfaction facts only.
        # All fairness, post spread, future catch-up and longitudinal catch-up remain
        # anchored to baseline_json/fairness_history and must not be inferred
        # from CURRENT after voluntary swaps or fairness-neutral repairs.
        gg=refreshed.stats.setdefault("global",{})
        gg["actual_operational_view"] = True
        gg["system_fairness_accounting_source"] = "PUBLISHED_BASELINE_JSON"
        gg["post_publication_changes_excluded_from_fairness"] = True
        gg["post_publication_changes_excluded_from_post_spread"] = True
        gg["post_publication_changes_excluded_from_post_debt"] = True
    return refreshed


def persist_actual_satisfaction(y,m,payload=None):
    """Persist the ACTUAL request/satisfaction snapshot after any post-publication change.

    The original request set remains frozen inside the schedule payload. This helper
    only recalculates the CURRENT/ACTUAL realization against that original set and
    the live backup table; only COMPLETED covers can change ACTUAL work, so later retrospective month review can read the final
    satisfaction percentages without touching the immutable SYSTEM baseline.
    """
    payload=payload or db.load_schedule(y,m,"current")
    if not payload:
        return None
    refreshed=refresh_result_payload(payload,y,m,use_actual_backups=True)
    db.save_current(y,m,serialize_result(refreshed))
    return refreshed


def credit_selection_errors(y,m):
    prefs=db.all_preferences(y,m); errors=[]
    for row in DEFAULT_PEOPLE:
        i=row["initials"]; p=prefs.get(i,{})
        use_am=int(p.get("backup_credits_am_to_use",0))
        use_pm=int(p.get("backup_credits_pm_to_use",0))
        if use_am+use_pm>2:
            errors.append({tr("person"):i,tr("details"):tr("max_credit_error")})
        avail_am=db.rest_credit_available_for_month(i,y,m,"AM")
        avail_pm=db.rest_credit_available_for_month(i,y,m,"PM")
        if use_am>avail_am:
            errors.append({tr("person"):i,tr("credit_am"):f"{use_am}/{avail_am}",tr("details"):tr("bonus_insufficient")})
        if use_pm>avail_pm:
            errors.append({tr("person"):i,tr("credit_pm"):f"{use_pm}/{avail_pm}",tr("details"):tr("bonus_insufficient")})
        if int(p.get("backup_credits_night_to_use",0))>0:
            errors.append({tr("person"):i,tr("credit_night"):p.get("backup_credits_night_to_use",0),tr("details"):tr("night_bank_only")})
    return errors


def _parse_iso_dt(value):
    if not value: return None
    try: return datetime.fromisoformat(str(value).replace("Z","+00:00"))
    except Exception: return None



def schedule_grid(y,m,result,status_rows=None):
    _,ndays=calendar.monthrange(y,m); rows={}
    for s in make_slots(y,m):
        key=f"{s.department} [{block_label(s.block)}]"; rows.setdefault(key,{d:"" for d in range(1,ndays+1)})
        rows[key][s.day]="BLOCK" if s.blocked else result.assignments.get(s.idx,"")
    for r in (status_rows or []):
        ini=str(r.get("initials") or "")
        try: day=int(r.get("day"))
        except Exception: continue
        if not ini or not (1<=day<=ndays): continue
        key=f"NEDIRBA · {ini}" if lang=="LT" else f"NOT WORKING · {ini}"
        rows.setdefault(key,{d:"" for d in range(1,ndays+1)})
        rows[key][day]=ini
    df=pd.DataFrame.from_dict(rows,orient="index"); df.columns=[f"{d:02d}\n{WEEKDAYS[lang][date(y,m,d).weekday()]}" for d in range(1,ndays+1)]
    return df

def style_schedule(df):
    def cs(v):
        if v=="BLOCK": return "background-color:#D9D9D9;color:#555;font-weight:700;"
        c=PERSON_COLORS.get(str(v)); return "" if not c else f"background-color:{c};color:{contrast_text(c)};font-weight:700;text-align:center;"
    return df.style.map(cs)

def style_rows(df):
    pc=tr("person")
    def rowstyle(row):
        c=PERSON_COLORS.get(str(row.get(pc,""))); css="" if not c else f"background-color:{c};color:{contrast_text(c)};"
        return [css]*len(row)
    return df.style.apply(rowstyle,axis=1)

def _stable_rank(*parts):
    raw="|".join(map(str,parts)).encode("utf-8")
    return int(hashlib.sha1(raw).hexdigest()[:12],16)


def _assignments_for_person_day(y,m,result,initials,day):
    rows=[]
    for s in make_slots(y,m):
        if s.day==day and result.assignments.get(s.idx)==initials:
            rows.append(s)
    return sorted(rows,key=lambda s:{"AM":0,"FULL":1,"PM":2}.get(s.block,9))


def _covered_shift_text(y,m,result,slot_id):
    slots={s.idx:s for s in make_slots(y,m)}
    s=slots.get(int(slot_id))
    if s is None:
        return "—"
    return f"{s.department} ({block_label(s.block)})"


def _eligible_backup_candidates(y,m,result,covered_slot,people=None,allow_resident_hard=False):
    """Residents who can cover one specific shift.

    Default is strict zero-RESIDENT-HARD-loss eligibility. The planner may invoke
    ``allow_resident_hard=True`` only as a last-resort fallback; ABSOLUTE HARD and
    overlapping normal work remain unbreakable in both modes.
    """
    if people is None:
        people=load_people(y,m)
    slots=make_slots(y,m)
    assigned=result.assignments.get(covered_slot.idx)
    candidates=[]
    for p in people:
        if p.initials==assigned:
            continue
        unavailable=(
            absolute_unavailable_for_block(p,covered_slot.day,covered_slot.block)
            if allow_resident_hard else
            hard_unavailable_for_block(p,covered_slot.day,covered_slot.block)
        )
        if unavailable:
            continue
        own=[
            sl for sl in slots
            if sl.day==covered_slot.day and result.assignments.get(sl.idx)==p.initials
        ]
        if any(blocks_overlap(sl.block,covered_slot.block) for sl in own):
            continue
        candidates.append(p.initials)
    return candidates




def _backup_override_note_decode(note):
    prefix="V2561_BACKUP_OVERRIDE:"
    raw=str(note or "")
    if raw.startswith(prefix):
        try:
            data=json.loads(raw[len(prefix):])
            return data if isinstance(data,dict) else {}
        except Exception:
            return {}
    return {"legacy_note":raw} if raw else {}


def _backup_override_note_encode(meta, legacy_note=""):
    prefix="V2561_BACKUP_OVERRIDE:"
    payload=dict(meta or {})
    if legacy_note and not payload.get("legacy_note"):
        payload["legacy_note"]=str(legacy_note)
    return prefix+json.dumps(payload,ensure_ascii=False,separators=(",",":"))


def _backup_shift_hours(slot):
    return 9.0 if slot and slot.block=="FULL" else 6.0


def _backup_block_hours(block):
    return ({"AM":(8.0,14.0),"FULL":(8.0,17.0),"PM":(14.0,20.0),"NIGHT":(20.0,32.0)}.get(str(block), (8.0,14.0)))


def _preview_manual_backup_takeover(y,m,result,covered_slot,initials,exclude_backup_id=None):
    """Preview the *actual work* created if ``initials`` is called to cover this backup.

    Generator fatigue targets are warnings here. True ABSOLUTE/operational and
    statutory-style safety guardrails stay blocking: justified absence / mandatory
    post-duty rest, overlap, >12h/day, <11h daily rest, >6 workdays/7d, <35h
    continuous weekly rest, and >60h/7d. 48h remains a visible warning because the
    exact legal interpretation depends on the active work-time regime/accounting period.
    """
    people=people_for_stored_result(result,y,m)
    byinit={p.initials:p for p in people}
    p=byinit.get(str(initials))
    rows=[]; blockers=[]; warnings=[]
    if covered_slot is None or p is None:
        msg="Nerastas dublio slotas arba rezidentas."
        return {"ok":False,"blockers":[msg],"warnings":[],"rows":[],"fingerprint":""}
    if result.assignments.get(covered_slot.idx)==p.initials:
        msg="Rezidentas negali dubliuoti savo paties pamainos."
        return {"ok":False,"blockers":[msg],"warnings":[],"rows":[],"fingerprint":""}

    # ABSOLUTE no-work state remains non-overrideable.
    if absolute_unavailable_for_block(p,covered_slot.day,covered_slot.block):
        blockers.append("Šią dieną rezidentui taikomas ABSOLUTE HARD nedarbo / privalomo poilsio apribojimas.")

    # Build actual-work intervals from SYSTEM/CURRENT normal assignments plus already
    # activated/completed backup cover, then add the proposed takeover once.
    slots=make_slots(y,m)
    normal=[s for s in slots if result.assignments.get(s.idx)==p.initials]
    intervals=[]
    def add_interval(day,block,label,hours=None):
        start,end=_backup_block_hours(block)
        # NIGHT may cross midnight; ordinary current PGY1 backup slots are AM/PM.
        abs_start=(int(day)-1)*24.0+start
        abs_end=(int(day)-1)*24.0+end
        intervals.append({"day":int(day),"block":str(block),"start":abs_start,"end":abs_end,
                          "hours":float(hours if hours is not None else (end-start)),"label":label})
    for s in normal:
        add_interval(s.day,s.block,f"NORMAL {s.department}",_backup_shift_hours(s))

    # Only activated/completed backup rows count as already-realized work.
    for br in db.list_backups(y,m):
        if exclude_backup_id is not None and int(br.get("id") or -1)==int(exclude_backup_id):
            continue
        eff=str(br.get("actual_backup") or br.get("planned_backup") or "")
        if eff!=p.initials or not (br.get("activated_at") or br.get("completed_at")):
            continue
        sid=int(br.get("covered_slot") or -1)
        bs=next((s for s in slots if s.idx==sid),None)
        if bs is not None:
            add_interval(bs.day,bs.block,f"AKTYVUOTAS DUBLIS {bs.department}",_backup_shift_hours(bs))

    # Snapshot the currently-realized work before adding the proposed backup.
    base_intervals=list(intervals)

    # Proposed backup may not overlap work already known to the system.
    ps,pe=_backup_block_hours(covered_slot.block)
    pstart=(covered_slot.day-1)*24.0+ps; pend=(covered_slot.day-1)*24.0+pe
    overlaps=[x for x in intervals if max(x["start"],pstart) < min(x["end"],pend)-1e-9]
    if overlaps:
        blockers.append("Dublis persidengtų su jau esančia darbo pamaina / aktyvuotu dubliu tuo pačiu metu.")
    else:
        add_interval(covered_slot.day,covered_slot.block,f"SIŪLOMAS DUBLIS {covered_slot.department}",_backup_shift_hours(covered_slot))

    # Voluntary self-override of RESIDENT HARD is a warning, not ABSOLUTE block.
    if resident_hard_unavailable_for_block(p,covered_slot.day,covered_slot.block):
        warnings.append("Rezidentas savo noru perimtų pamainą per anksčiau pateiktą „Negaliu dirbti“ laiką.")
        rows.append({"Rodiklis":"RESIDENT HARD","Dabar":"Prašė nedirbti","Po dublio":"Dirbtų šią pamainą",
                     "Taisyklė / paaiškinimas":"Galima tik aiškiai savanoriškai patvirtinus; ORIGINAL pageidavimas istorijoje lieka.","Būsena":"PERSPĖJIMAS"})

    base_byday={}
    for x in base_intervals:
        base_byday.setdefault(x["day"],[]).append(x)
    byday={}
    for x in intervals:
        byday.setdefault(x["day"],[]).append(x)
    before_normal_hours=sum(x["hours"] for x in base_byday.get(covered_slot.day,[]))
    after_day_hours=sum(x["hours"] for x in byday.get(covered_slot.day,[]))
    if after_day_hours>12.0+1e-9:
        blockers.append(f"Po dublio būtų {after_day_hours:g} val. darbo per dieną (maksimali darbo dienos / pamainos trukmė šiame saugos profilyje: 12 val.).")
        status="BLOKUOJA"
    elif after_day_hours>=12.0-1e-9 and after_day_hours>before_normal_hours+1e-9:
        warnings.append(f"Diena taptų 12 val. darbo diena ({after_day_hours:g} val.).")
        status="PERSPĖJIMAS"
    else:
        status="GERAI"
    rows.append({"Rodiklis":f"{covered_slot.day:02d} d. darbo valandos","Dabar":f"{before_normal_hours:g} val.","Po dublio":f"{after_day_hours:g} val.",
                 "Taisyklė / paaiškinimas":"Ne daugiau kaip 12 val. per darbo dieną / pamainą.","Būsena":status})

    # Rolling 7-day hours + workday counts.
    _,ndays=calendar.monthrange(y,m)
    base_hours_by_day={d:sum(x["hours"] for x in base_byday.get(d,[])) for d in range(1,ndays+1)}
    hours_by_day={d:sum(x["hours"] for x in byday.get(d,[])) for d in range(1,ndays+1)}
    worked={d for d,h in hours_by_day.items() if h>1e-9}
    max7=0.0; max7range=(1,min(7,ndays)); maxdays=0; maxdaysrange=max7range
    before_max7=0.0
    legal_rolling_cap=min(float(rule_value("swap_max_hours_rolling7")),60.0)
    bad60=[]; bad7days=[]
    for start in range(1,max(2,ndays-6+1)):
        end=min(ndays,start+6)
        h=sum(hours_by_day.get(d,0.0) for d in range(start,end+1))
        bh=sum(base_hours_by_day.get(d,0.0) for d in range(start,end+1))
        wd=sum(1 for d in range(start,end+1) if d in worked)
        before_max7=max(before_max7,bh)
        if h>max7: max7=h; max7range=(start,end)
        if wd>maxdays: maxdays=wd; maxdaysrange=(start,end)
        if h>legal_rolling_cap+1e-9: bad60.append((start,end,h))
        if wd>6: bad7days.append((start,end,wd))
    before_max=float(before_max7)
    if bad60:
        s,e,h=bad60[0]
        blockers.append(f"Po dublio būtų {h:g} val. per 7 dienas ({s}–{e} d.; aktyvi šio swapo/dublio saugos riba: {legal_rolling_cap:g} val.).")
        hstatus="BLOKUOJA"
    elif max7>48.0+1e-9:
        warnings.append(f"Po dublio didžiausias 7 dienų krūvis būtų {max7:g} val. ({max7range[0]}–{max7range[1]} d.).")
        hstatus="PERSPĖJIMAS"
    elif max7>40.0+1e-9:
        warnings.append(f"Po dublio didžiausias 7 dienų krūvis būtų {max7:g} val.; tai viršija ~40 val. planavimo tikslą.")
        hstatus="PERSPĖJIMAS"
    else:
        hstatus="GERAI"
    rows.append({"Rodiklis":"Didžiausias krūvis per 7 d.","Dabar":f"{before_max:g} val.","Po dublio":f"{max7:g} val. ({max7range[0]}–{max7range[1]} d.)",
                 "Taisyklė / paaiškinimas":f"~40 val. yra planavimo tikslas; >48 val. rodoma kaip aiškus įspėjimas; aktyvi manual swapo/dublio riba: {legal_rolling_cap:g} val./7 d.","Būsena":hstatus})
    if bad7days:
        s,e,wd=bad7days[0]
        blockers.append(f"Po dublio būtų dirbama {wd} dienas per {s}–{e} d. laikotarpį (leidžiama ne daugiau kaip 6 darbo dienas per 7 paeiliui einančias dienas).")
        dstatus="BLOKUOJA"
    elif maxdays==6:
        warnings.append(f"Susidarytų 6 darbo dienų seka / 6 darbo dienos per 7 d. ({maxdaysrange[0]}–{maxdaysrange[1]} d.).")
        dstatus="PERSPĖJIMAS"
    else:
        dstatus="GERAI"
    rows.append({"Rodiklis":"Darbo dienos per 7 d.","Dabar":str(((result.stats or {}).get("people",{}).get(p.initials,{}) or {}).get("max_consecutive_days",0) or 0),
                 "Po dublio":f"iki {maxdays} d. ({maxdaysrange[0]}–{maxdaysrange[1]} d.)",
                 "Taisyklė / paaiškinimas":"Ne daugiau kaip 6 darbo dienos per 7 paeiliui einančias dienas.","Būsena":dstatus})

    # Daily rest: only between separate calendar workdays; same-day AM+PM is one 12h day.
    intervals_sorted=sorted(intervals,key=lambda x:(x["start"],x["end"]))
    min_rest=999.0; min_pair=None
    for a,b in zip(intervals_sorted,intervals_sorted[1:]):
        if a["day"]==b["day"]:
            continue
        gap=b["start"]-a["end"]
        if gap<min_rest:
            min_rest=gap; min_pair=(a,b)
    if min_pair is not None and min_rest<11.0-1e-9:
        blockers.append(f"Tarp pamainų liktų tik {min_rest:g} val. poilsio (minimalus nepertraukiamas paros poilsis: 11 val.).")
        rstatus="BLOKUOJA"
    else:
        rstatus="GERAI"
    rows.append({"Rodiklis":"Trumpiausias poilsis tarp darbo dienų","Dabar":"—","Po dublio":("—" if min_pair is None else f"{min_rest:g} val."),
                 "Taisyklė / paaiškinimas":"Minimalus nepertraukiamas paros poilsis: 11 val.","Būsena":rstatus})

    # Weekly-rest legal note: with the current AM/PM model, the tool enforces the
    # operational proxy of no more than 6 workdays in any 7 plus >=11h daily rest.
    # The exact 35h statutory weekly-rest interpretation can depend on the active
    # work-time regime / accounting period and remains an employer-level legal check.

    # Fatigue patterns are warnings if the true safety blockers above remain satisfied.
    double_days=sorted(d for d,h in hours_by_day.items() if h>=12.0-1e-9)
    if covered_slot.day in double_days:
        if covered_slot.day-1 in double_days or covered_slot.day+1 in double_days:
            warnings.append("Dublis sudarytų dviejų 12 val. darbo dienų seką; generatorius to vengia, bet savanoriškai galima patvirtinti, jei absoliučios ribos išlaikytos.")
        if covered_slot.day-1 in double_days and covered_slot.day-2 in double_days:
            warnings.append("Tai būtų darbo diena po dviejų 12 val. dienų iš eilės; rodoma kaip nuovargio perspėjimas.")

    ok=not blockers
    fp_raw="|".join([p.initials,str(covered_slot.idx)]+sorted(blockers)+sorted(warnings)+[str(r) for r in rows])
    fp=hashlib.sha256(fp_raw.encode("utf-8")).hexdigest()[:16]
    return {"ok":ok,"blockers":blockers,"warnings":warnings,"rows":rows,"fingerprint":fp,
            "max_rolling7_hours":round(max7,1),"day_hours":round(after_day_hours,1),"max_workdays7":int(maxdays)}

def _operational_repair_validation(y,m,result,new_assignments):
    """Validate an emergency/admin repair without rewriting the published fairness baseline.

    Publication-time workload targets, Onko pair-evenness and weekend uniqueness are planning
    constraints. They must not prevent a justified post-publication repair. Safety, coverage,
    overlap, HARD unavailability, rest and known-hours rules remain enforced.
    """
    people=people_for_stored_result(result,y,m)
    slots=make_slots(y,m)
    # V2.5.85: use the published SYSTEM targets as the immutable workload-credit
    # ledger. Emergency Rescue changes placement only; no target units move.
    stats=validate_schedule(
        y,m,people,slots,new_assignments,result.targets,
        satisfaction_people=people,
        backup_assignments=(result.backup_snapshot or []),
        validation_mode="emergency_rescue",
    )
    errs=list(stats.get("global",{}).get("errors",[]))
    planning_only=[]; blocking=[]
    for e in errs:
        # Emergency post-publication repairs are allowed to create a new OPTIONAL
        # gap when a lower-priority resident is pulled into a mandatory SPS slot.
        # The published SYSTEM gap pattern / fairness baseline stays frozen; only
        # ABSOLUTE safety, mandatory coverage, overlap and operational feasibility
        # remain blocking here.
        gap_planning_only=(
            e.startswith("Gap dispersion violated")
            or e.startswith("Gap workplace dispersion violated")
            or e.startswith("Gap-day dispersion pattern outdated")
        )
        if (
            e.endswith("odd Onko count")
            or (e.startswith("Weekend ") and (
                "must have exactly 4 different people" in e
                or "resident weekend cap exceeded" in e
            ))
            or gap_planning_only
        ):
            planning_only.append(e)
        else:
            blocking.append(e)
    stats.setdefault("global",{})["planning_exceptions_after_publication"]=planning_only
    stats["global"]["errors"]=blocking
    stats["global"]["hard_errors"]=len(blocking)
    return stats


def _repair_candidate_check(y,m,result,slot_id,replacement,source_slot_id=None):
    """Validate one post-publication repair candidate.

    V2.5.56 supports a critical-cover *transfer*: for SPS RO / SPS UG, a resident
    already working an overlapping lower-priority optional post may be moved into
    the critical slot and the optional source slot is intentionally vacated.
    """
    slots={s.idx:s for s in make_slots(y,m)}
    target=slots.get(int(slot_id))
    if target is None or int(slot_id) not in result.assignments:
        return False,"Pamaina neužpildyta.",None
    if result.assignments[int(slot_id)]==replacement:
        return False,"Tas pats rezidentas jau paskirtas.",None
    source=None
    if source_slot_id is not None:
        source=slots.get(int(source_slot_id))
        if source is None:
            return False,"Nerasta donorinė neprivaloma pamaina.",None
    try:
        fresh_assign=(
            apply_emergency_critical_transfer(result.assignments,target,replacement,source)
            if is_emergency_critical_slot(target)
            else dict(result.assignments)
        )
        if not is_emergency_critical_slot(target):
            if source is not None:
                return False,"Donorinis perkėlimas taikomas tik kritiniam SPS postui.",None
            fresh_assign[int(slot_id)]=replacement
    except Exception as exc:
        return False,str(exc),None
    stats=_operational_repair_validation(y,m,result,fresh_assign)
    errs=stats.get("global",{}).get("errors",[])
    return (not errs, "" if not errs else errs[0], stats)


def _critical_repair_candidate_rows(y,m,result,target_slot,repair_load):
    """Rank critical-cover candidates using the V2.5.56 rescue hierarchy.

    1) same-block residents are pulled from lower-priority OPTIONAL posts;
    2) only if no such ABSOLUTE-safe transfer exists do we expose residents who
       are free in that block as fallback cover.
    Within a tier, reject any new Resident-HARD conflict during SYSTEM generation. V2.5.57 deliberately does NOT
    rank same-block pull-down donors by post spread or future catch-up: they were already
    scheduled to work that block, so changing station for emergency coverage is
    fairness-neutral and must not create a new structural fairness burden.
    """
    from_person=result.assignments.get(target_slot.idx)
    base_rh=int((result.stats or {}).get("global",{}).get("resident_hard_total_losses",0) or 0)
    slots=make_slots(y,m)
    pull_rows=[]; free_rows=[]
    for prow in DEFAULT_PEOPLE:
        cand=prow["initials"]
        if cand==from_person:
            continue
        sources=emergency_donor_source_slots(slots,result.assignments,target_slot,cand)
        source_options=sources or [None]
        for source in source_options:
            # If the resident is already working an overlapping non-donor post,
            # this is not a free fallback and must not create an overlap. Validation
            # below will reject it.
            ok,why,cstats=_repair_candidate_check(
                y,m,result,target_slot.idx,cand,(source.idx if source else None)
            )
            if not ok:
                continue
            cg=(cstats or {}).get("global",{})
            row={
                "initials":cand,
                "source_slot":(source.idx if source else None),
                "source_department":(source.department if source else ""),
                "source_block":(source.block if source else ""),
                "mode":"PULL_OPTIONAL" if source else "FREE_FALLBACK",
                "rh_total":int(cg.get("resident_hard_total_losses",0) or 0),
                "rh_delta":int(cg.get("resident_hard_total_losses",0) or 0)-base_rh,
                "rh_max":int(cg.get("resident_hard_max_loss_per_resident",0) or 0),
                "rh_cum_spread":int(cg.get("resident_hard_cumulative_spread",0) or 0),
                "stats":cstats,
            }
            (pull_rows if source else free_rows).append(row)
    rows=pull_rows if pull_rows else free_rows
    zero_rh=[r for r in rows if r["rh_delta"]<=0]
    rows=zero_rh or rows
    return sorted(rows,key=lambda r:(
        r["rh_total"],r["rh_max"],r["rh_cum_spread"],r["initials"]
    ))


def plan_backups(y,m,result):
    """Build SYSTEM backup duties with mandatory-first fairness water-filling.

    Required cover: SPS RO, SPS UG, Centro UG 120 AM and Onko RO. CENTRO RO is
    intentionally *best effort*: after all required cover exists, CENTRO RO duties
    are used only to lift low-backup residents toward the highest required-backup
    layer. We stop before optional cover would create a new unequal burden.

    This gives every eligible resident real backup exposure without manufacturing
    hundreds of optional duties merely to say CENTRO RO is "covered".
    """
    people=people_for_stored_result(result,y,m)
    initials=[p.initials for p in people]
    claims=db.list_backup_claims(y,m)
    claim_by_slot={int(r["covered_slot"]):r for r in claims}
    backup_load={i:0 for i in initials}; pair_load={}; same_time_load={}
    desired=[]; errors=[]; slots=make_slots(y,m)
    required=sorted(
        [sl for sl in slots if backup_required_slot(sl) and result.assignments.get(sl.idx)],
        key=lambda sl:(sl.day,{"AM":0,"FULL":1,"PM":2}.get(sl.block,9),sl.idx)
    )
    best_effort=sorted(
        [sl for sl in slots if backup_best_effort_slot(sl) and result.assignments.get(sl.idx) and sl not in required],
        key=lambda sl:(sl.day,{"AM":0,"FULL":1,"PM":2}.get(sl.block,9),sl.idx)
    )

    def assign_one(sl,is_required,optional_ceiling=None):
        covered=result.assignments.get(sl.idx)
        strict=list(_eligible_backup_candidates(y,m,result,sl,people,allow_resident_hard=False))
        if not strict:
            if is_required:
                errors.append({
                    "day":sl.day,"shift":block_label(sl.block),"department":sl.department,
                    "covered_person":covered,"covered_slot":sl.idx,
                    "reason":"no strict-eligible backup without mandatory-unavailability / overlap violation",
                })
            return False
        # Prefer a different person for simultaneous cover whenever possible.
        fresh=[b for b in strict if same_time_load.get((sl.day,sl.block,b),0)==0]
        eligible=fresh or strict
        claimed=(claim_by_slot.get(sl.idx) or {}).get("initials")
        def rank(b):
            return (
                backup_load.get(b,0),
                same_time_load.get((sl.day,sl.block,b),0),
                pair_load.get((b,covered),0),
                0 if claimed==b else 1,   # claim = tie-break, never fairness override
                _stable_rank(y,m,sl.day,sl.block,sl.idx,b,covered),
            )
        backup=min(eligible,key=rank)
        if optional_ceiling is not None and backup_load.get(backup,0)>=int(optional_ceiling):
            return False
        desired.append({
            "covered_slot":sl.idx,"covered_person":covered,"day":sl.day,"block":sl.block,
            "department":sl.department,"planned_backup":backup,"actual_backup":None,
            "resident_hard_relaxed":False,
            "claim_overridden_by_higher_priority":bool(claimed and claimed!=backup),
            "coverage_priority":"required" if is_required else "best_effort_fairness_filler",
            "note":"AUTO SYSTEM WATERFILL",
        })
        same_time_load[(sl.day,sl.block,backup)]=same_time_load.get((sl.day,sl.block,backup),0)+1
        backup_load[backup]=backup_load.get(backup,0)+1
        pair_load[(backup,covered)]=pair_load.get((backup,covered),0)+1
        return True

    # Phase 1: all important positions must have a named backup.
    for sl in required:
        assign_one(sl,True)

    # Phase 2: use CENTRO RO only as a low-load lift. Freeze the highest burden
    # created by required cover and do not let optional duties push anyone above it.
    required_ceiling=max(backup_load.values()) if backup_load else 0
    for sl in best_effort:
        assign_one(sl,False,optional_ceiling=required_ceiling)

    return desired,errors


def sync_backup_plan(y,m,result):
    desired,errors=plan_backups(y,m,result)
    db.sync_backups(y,m,desired)
    return desired,errors


def _backup_rows_for_result(y,m,result=None,override=None):
    """Return the operational backup rows, falling back to the generated draft snapshot.

    Before SYSTEM publication the database intentionally has no backup_assignments rows.
    V2.5.112 therefore treats result.backup_snapshot as the authoritative draft-time
    plan so Summary / Excel / personal schedule do not incorrectly show zero backups.
    """
    if override is not None:
        return [dict(x) for x in override]
    try:
        rows=db.list_backups(y,m)
    except Exception:
        rows=[]
    if rows:
        return [dict(x) for x in rows]
    snap=list(getattr(result,"backup_snapshot",None) or []) if result is not None else []
    return [dict(x) for x in snap]


def backup_counts(y,m,result=None):
    planned={p["initials"]:0 for p in DEFAULT_PEOPLE}
    effective={p["initials"]:0 for p in DEFAULT_PEOPLE}
    for r in _backup_rows_for_result(y,m,result):
        pb=r.get("planned_backup")
        if not pb:
            continue
        planned[pb]=planned.get(pb,0)+1
        eff=r.get("actual_backup") or pb
        effective[eff]=effective.get(eff,0)+1
    return planned,effective


def summary_df(result,y,m):
    planned,effective=backup_counts(y,m,result); rows=[]
    for i,d in result.stats.get("people",{}).items():
        row={
            tr("person"):i,tr("name"):d.get("name",""),tr("target"):d.get("target"),
            tr("workload"):d.get("workload_credit",d.get("workload")),
            tr("weekday_assignments"):d.get("weekday_assignments"),
            tr("weekday_days"):d.get("weekday_days"),
            tr("weekend_assignments"):d.get("weekend_assignments"),
            tr("saturday_assignments"):d.get("saturdays"),
            tr("sunday_assignments"):d.get("sundays"),
            tr("prior_weekends"):d.get("prior_weekend_count"),
            ("Šventės" if lang=="LT" else "Holidays"):d.get("holiday_assignments",0),
            ("Ankstesnės šventės" if lang=="LT" else "Prior holidays"):d.get("prior_holiday_count",0),
            tr("cumulative_weekends"):d.get("cumulative_weekend_count"),
            tr("fridays"):d.get("friday_assignments"),
            tr("double_shifts"):d.get("doubles"),
            tr("max_consecutive"):d.get("max_consecutive_days"),
            tr("max_rolling7_hours"):d.get("max_rolling7_hours"),
            tr("max_calendar_week_hours"):d.get("max_calendar_week_hours"),
            tr("free_days"):d.get("fully_free_days"),
            ("Skirtingos darbo vietos" if lang=="LT" else "Distinct workplaces"):d.get("distinct_rotations"),
            tr("preference_score"):d.get("preference_score"),
            tr("planned_backups"):planned.get(i,0),
            tr("effective_backups"):effective.get(i,0),
        }
        if d.get("workload_credit_policy")=="FROZEN_SYSTEM_LEDGER":
            row[("ACTUAL slotų svoris (ne target)" if lang=="LT" else "ACTUAL placement workload (not target)")]=d.get("actual_assignment_workload",d.get("workload"))
        # V2.5.15: exact monthly number of assignments in every workplace.
        rotation_counts=d.get("rotation_counts") or {}
        for cat in ROTATION_CATEGORIES:
            row[cat]=int(rotation_counts.get(cat,0) or 0)
        rows.append(row)
    return pd.DataFrame(rows)


def preference_scores_df(result):
    return pd.DataFrame([
        {
            tr("person"):i,
            tr("name"):d.get("name",""),
            ("RESIDENT HARD %" if lang=="LT" else "RESIDENT HARD %"):d.get("resident_hard_score"),
            ("SOFT %" if lang=="LT" else "SOFT %"):d.get("soft_preference_score"),
            tr("preference_score"):d.get("preference_score"),
            ("RESIDENT HARD pažeidimai" if lang=="LT" else "RESIDENT HARD violations"):d.get("resident_hard_losses",0),
        }
        for i,d in result.stats.get("people",{}).items()
    ])


def resident_wishes_audit_df(result):
    """Senior-facing pre-publication request audit for every resident.

    Uses the CURRENT candidate's own validated request ledger, so the senior can
    inspect exactly what the generated DRAFT would satisfy before publication.
    """
    rows=[]
    for initials,d in (result.stats.get("people",{}) or {}).items():
        details=list(d.get("request_detail_rows") or [])
        included=[r for r in details if r.get("included_in_score")]
        preferred=[r for r in included if r.get("kind")=="preferred"]
        soft_free=[r for r in included if r.get("kind")=="soft_free"]
        missed=[r for r in included if not r.get("fulfilled")]
        components=d.get("preference_components") or {}
        workstyle=components.get("shift_length_preference")
        if workstyle is None:
            workstyle=components.get("avoid_doubles")
        rotation_counts=d.get("rotation_counts") or {}
        theoretical_backup_count=sum(
            1 for br in (getattr(result,"backup_snapshot",None) or [])
            if str(br.get("actual_backup") or br.get("planned_backup") or "")==initials
        )

        def ratio(items):
            if not items:
                return "—"
            return f"{sum(1 for r in items if r.get('fulfilled'))}/{len(items)}"

        hard_req=int(d.get("resident_hard_requested",0) or 0)
        hard_ok=int(d.get("resident_hard_honored",0) or 0)
        hard_txt=(f"{hard_ok}/{hard_req}" if hard_req else "—")
        missed_preview=[]
        for r in missed[:3]:
            missed_preview.append(f"{r.get('date','—')} {r.get('block','—')} · {r.get('type','—')}")
        if len(missed)>3:
            missed_preview.append(f"+{len(missed)-3}")

        overall=d.get("overall_request_score")
        soft_score=d.get("soft_preference_score")
        rows.append({
            ("Žmogus" if lang=="LT" else "Person"):initials,
            ("Vardas" if lang=="LT" else "Name"):d.get("name",""),
            ("Target" if lang=="LT" else "Target"):d.get("target"),
            ("Krūvis" if lang=="LT" else "Workload"):d.get("workload_credit",d.get("workload")),
            ("RESIDENT HARD" if lang=="LT" else "RESIDENT HARD"):hard_txt,
            ("Noriu laisvos" if lang=="LT" else "Requested off"):ratio(soft_free),
            ("Pageidauju dirbti" if lang=="LT" else "Prefer to work"):ratio(preferred),
            ("SOFT %" if lang=="LT" else "SOFT %"):("—" if soft_score is None else soft_score),
            ("Bendras išpildymas %" if lang=="LT" else "Overall satisfaction %"):("—" if overall is None else overall),
            ("Workstyle %" if lang=="LT" else "Workstyle %"):("—" if workstyle is None else round(float(workstyle),1)),
            ("Šeštadieniai" if lang=="LT" else "Saturdays"):int(d.get("saturdays",0) or 0),
            ("Sekmadieniai" if lang=="LT" else "Sundays"):int(d.get("sundays",0) or 0),
            ("12h dienos (AM+PM)" if lang=="LT" else "12h workdays (AM+PM)"):int(d.get("doubles",0) or 0),
            ("Teoriniai AUTO dubliai" if lang=="LT" else "Theoretical AUTO backups"):int(theoretical_backup_count),
            "Onko RO":int(rotation_counts.get("Onko RO",0) or 0),
            "SPS RO":int(rotation_counts.get("SPS RO",0) or 0),
            "SPS UG":int(rotation_counts.get("SPS UG",0) or 0),
            ("Neįvykdyta" if lang=="LT" else "Missed"):len(missed),
            ("Neįvykdytų santrauka" if lang=="LT" else "Missed summary"):"; ".join(missed_preview) if missed_preview else "—",
            "__hard_losses":int(d.get("resident_hard_losses",0) or 0),
            "__score_sort":101.0 if overall is None else float(overall),
        })
    if not rows:
        return pd.DataFrame()
    rows.sort(key=lambda r:(-r["__hard_losses"], r["__score_sort"], -r[("Neįvykdyta" if lang=="LT" else "Missed")], r[("Žmogus" if lang=="LT" else "Person")]))
    for r in rows:
        r.pop("__hard_losses",None); r.pop("__score_sort",None)
    return pd.DataFrame(rows)



def generation_wish_summary(result):
    """Return explicit all-resident wish totals + a flat table of misses."""
    total=honored=missed=0
    hard_total=hard_missed=0
    rows=[]
    for initials,d in (result.stats.get("people",{}) or {}).items():
        details=[r for r in (d.get("request_detail_rows") or []) if r.get("included_in_score")]
        for r in details:
            total+=1
            ok=bool(r.get("fulfilled"))
            honored+=int(ok); missed+=int(not ok)
            if r.get("kind")=="resident_hard":
                hard_total+=1; hard_missed+=int(not ok)
            if ok:
                continue
            base=request_details_df([r],initials)
            if base.empty:
                continue
            rec=base.iloc[0].to_dict()
            rec={(("Žmogus" if lang=="LT" else "Person")):initials,
                 (("Vardas" if lang=="LT" else "Name")):d.get("name","")} | rec
            rows.append(rec)
    return {
        "total":int(total),"honored":int(honored),"missed":int(missed),
        "hard_total":int(hard_total),"hard_missed":int(hard_missed),
        "table":pd.DataFrame(rows),
    }



def resident_group_satisfaction_df(result):
    """Privacy-safe group view for ordinary residents.

    Only initials, name and overall fulfillment percentage are exposed.
    Peer HARD/SOFT counts, request dates/blocks, workstyle details, missed counts
    and request-level rows are deliberately not serialized to the resident UI.
    """
    rows=[]
    for initials,d in (result.stats.get("people",{}) or {}).items():
        overall=d.get("overall_request_score")
        rows.append({
            ("Žmogus" if lang=="LT" else "Person"):initials,
            ("Vardas" if lang=="LT" else "Name"):d.get("name",""),
            ("Bendras išpildymas %" if lang=="LT" else "Overall satisfaction %"):(
                "—" if overall is None else overall
            ),
        })
    return pd.DataFrame(rows)


def render_resident_wishes_audit(
    result, *, draft_mode=False, key_suffix="", senior_view=False
):
    """Role-aware request audit.

    Senior: full all-resident category + request-level audit.
    Resident: only group overall satisfaction percentages.
    """
    if senior_view:
        if draft_mode:
            st.markdown("### JUODRAŠČIO PAGEIDAVIMŲ AUDITAS" if lang=="LT" else "### DRAFT REQUEST AUDIT")
            st.caption(
                "Tai yra būtent dabar sugeneruoto JUODRAŠČIO rezultatas. Seniūnė gali įvertinti, ar pageidavimai maksimaliai išpildyti, prieš paspausdama PASKELBTI / PATVIRTINTI. Regeneravus lentelė persiskaičiuos iš naujo."
                if lang=="LT" else
                "This is the currently generated DRAFT. The senior can inspect whether requests are maximized before pressing PUBLISH / CONFIRM. Regeneration recalculates this table."
            )
        else:
            st.markdown("### SYSTEM pageidavimų auditas" if lang=="LT" else "### SYSTEM request audit")
            st.caption(
                "Ši lentelė rodo publikavimo momento SYSTEM rezultatą ir TIK tuo run metu užšaldytus pageidavimus. "
                "Vėliau pakeisti Nustatymai čia retroaktyviai nepridedami. Workstyle eilutėje rodomi konkretūs 6 h / 12 h / Onko skaičiai."
                if lang=="LT" else
                "This table shows the publication-time SYSTEM result and ONLY the requests frozen for that run. "
                "Settings changed later are not added retroactively. Workstyle rows show concrete 6 h / 12 h / Onko counts."
            )

        st.info(
            "V2.5.115: DUBLIAI = ATSKIRAS TEORINIS STANDBY SLUOKSNIS. Planuotas ar tik aktyvuotas dublis NĖRA darbo pamaina ir NIEKADA nemažina „Noriu laisvos“, „Negaliu dirbti“ ar kitų SYSTEM pageidavimų score. Tik COMPLETED realus pavadavimas gali atsirasti ACTUAL darbo audite."
            if lang=="LT" else
            "V2.5.115: BACKUPS are a separate theoretical standby layer. A planned or merely activated backup is NOT a work shift and NEVER reduces SYSTEM request satisfaction. Only a COMPLETED real-life cover may appear in ACTUAL work audit."
        )
        audit_df=resident_wishes_audit_df(result)
        if audit_df.empty:
            st.caption("Nėra rezidentų audito duomenų." if lang=="LT" else "No resident audit data.")
            return
        st.dataframe(audit_df,use_container_width=True,hide_index=True,height=610)

        people=list((result.stats.get("people",{}) or {}).keys())
        if not people:
            return
        selected=st.selectbox(
            "Detaliai patikrinti rezidentą" if lang=="LT" else "Inspect resident in detail",
            people,
            key=f"summary_request_person_{key_suffix}",
        )
        pdict=(result.stats.get("people",{}).get(selected,{}) or {})
        misses=list(pdict.get("unhonored_request_details") or [])
        if misses:
            st.markdown("#### Neįvykdyti prašymai" if lang=="LT" else "#### Missed requests")
            st.dataframe(request_details_df(misses,selected),use_container_width=True,hide_index=True)
        else:
            st.success(
                "Šiam rezidentui į score įtrauktų neįvykdytų prašymų nėra."
                if lang=="LT" else
                "This resident has no scored missed requests."
            )
        with st.expander(
            "Rodyti įvykdytus prašymus" if lang=="LT" else "Show honored requests",
            expanded=False
        ):
            honored=list(pdict.get("honored_request_details") or [])
            if honored:
                st.dataframe(
                    request_details_df(honored,selected),
                    use_container_width=True,hide_index=True
                )
            else:
                st.caption(
                    "Nėra score įtrauktų struktūruotų prašymų."
                    if lang=="LT" else
                    "No scored structured requests."
                )
        return

    # Resident profile: privacy-safe group view only.
    st.markdown(
        "### Grupės pageidavimų išpildymas"
        if lang=="LT" else
        "### Group request satisfaction"
    )
    st.caption(
        "Privatumo sumetimais čia rodoma tik kiekvieno rezidento bendra pageidavimų išpildymo procentinė reikšmė. Kitų rezidentų HARD/SOFT kiekiai, datos, workstyle ir konkretūs prašymai nėra rodomi. Savo detalų auditą matai savo asmeninėje patikroje."
        if lang=="LT" else
        "For confidentiality, this table shows only each resident's overall request-satisfaction percentage. Other residents' HARD/SOFT counts, dates, workstyle and individual requests are not shown. Your own detailed audit remains available in your personal proof view."
    )
    safe_df=resident_group_satisfaction_df(result)
    if safe_df.empty:
        st.caption("Nėra grupės statistikos." if lang=="LT" else "No group statistics.")
        return
    st.dataframe(safe_df,use_container_width=True,hide_index=True,height=610)



def _plain_request_sentence(r, initials=""):
    """Human-readable single-sentence explanation of one request result.

    V2.5.60 deliberately avoids the unexplained English word ``claim`` in the
    resident/senior UI.  The sentence states: what was requested, what the
    SYSTEM/ACTUAL schedule actually contains, and why the request is counted as
    honored or missed.
    """
    typ=str(r.get("type") or "Pageidavimas")
    date_txt=str(r.get("date") or "—")
    block=str(r.get("block") or "FULL")
    station=str(r.get("station") or "—")
    fulfilled=bool(r.get("fulfilled"))
    who=(f"{initials}: " if initials else "")
    if lang=="LT":
        if r.get("kind") in ("shift_length_preference","avoid_doubles") and r.get("workstyle_proof"):
            wp=r.get("workstyle_proof") or {}
            mode=int(wp.get("mode") or 0)
            if mode==3:
                threshold=int(wp.get("fulfilled_threshold_min") or 0)
                cohort=int(wp.get("prefer12_cohort_size") or 0)
                return (
                    f"{who}SYSTEM generavimo metu buvo užšaldytas pageidavimas „prefer 12 h“. "
                    f"Grafike skirta {int(wp.get('double_days') or 0)} 12 h dienų ir "
                    f"{int(wp.get('single_days') or 0)} 6 h dienų; Onko 9 h dienų — {int(wp.get('onko_9h_days') or 0)}. "
                    f"Visos grupės dublių intervalas šiame SYSTEM yra {int(wp.get('group_double_min') or 0)}–{int(wp.get('group_double_max') or 0)}. "
                    f"Kad 12 h workstyle būtų laikomas išpildytu, šiam run pakanka būti viršutiniame intervalo krašte: ≥{threshold} 12 h dienų. "
                    f"12 h pageidavimą turėjo {cohort} rezidentas(-ai). "
                    + ("Todėl pageidavimas ĮVYKDYTAS." if fulfilled else "Todėl pageidavimas NEĮVYKDYTAS ir reikia peržiūrėti solverio workstyle paskirstymą.")
                )
            if mode==1:
                return (
                    f"{who}SYSTEM generavimo metu buvo užšaldytas pageidavimas „prefer 6 h“. "
                    f"Skirta {int(wp.get('double_days') or 0)} 12 h dienų ir {int(wp.get('single_days') or 0)} 6 h dienų. "
                    f"Šio run 6 h workstyle riba yra ≤{int(wp.get('fulfilled_threshold_max') or 0)} 12 h dienų. "
                    + ("Pageidavimas ĮVYKDYTAS." if fulfilled else "Pageidavimas NEĮVYKDYTAS.")
                )
        if fulfilled:
            if r.get("kind")=="preferred":
                return f"{who}{date_txt} {block}: „{typ}“ — ĮVYKDYTA, nes grafike yra tinkama darbo pamaina ({station})."
            if r.get("kind") in ("resident_hard","soft_free"):
                return f"{who}{date_txt} {block}: „{typ}“ — ĮVYKDYTA, nes šiame bloke nėra persidengiančios normalios darbo pamainos."
            return f"{who}{date_txt} {block}: „{typ}“ — ĮVYKDYTA."
        if r.get("kind") in ("resident_hard","soft_free"):
            return f"{who}{date_txt} {block}: „{typ}“ — NEĮVYKDYTA, nes grafike šiame bloke yra paskyrimas: {station}."
        if r.get("kind")=="preferred":
            return f"{who}{date_txt} {block}: „{typ}“ — NEĮVYKDYTA, nes grafike nėra tinkamos darbo pamainos šiame bloke."
        return f"{who}{date_txt} {block}: „{typ}“ — NEĮVYKDYTA pagal parodytą rezultatą ({station})."
    if r.get("kind") in ("shift_length_preference","avoid_doubles") and r.get("workstyle_proof"):
        wp=r.get("workstyle_proof") or {}
        mode=int(wp.get("mode") or 0)
        if mode==3:
            return (
                f"{who}The frozen SYSTEM input was 'prefer 12 h'. "
                f"The schedule contains {int(wp.get('double_days') or 0)} 12 h days and "
                f"{int(wp.get('single_days') or 0)} 6 h days; Onko 9 h days: {int(wp.get('onko_9h_days') or 0)}. "
                f"Group double-day range is {int(wp.get('group_double_min') or 0)}–{int(wp.get('group_double_max') or 0)}; "
                f"this run counts the 12 h preference as honored at ≥{int(wp.get('fulfilled_threshold_min') or 0)} double-days. "
                + ("HONORED." if fulfilled else "NOT HONORED.")
            )
    if fulfilled:
        if r.get("kind")=="preferred":

            return f"{who}{date_txt} {block}: '{typ}' — HONORED because the schedule contains an eligible assignment ({station})."
        if r.get("kind") in ("resident_hard","soft_free"):
            return f"{who}{date_txt} {block}: '{typ}' — HONORED because no overlapping normal work shift exists in that block."
        return f"{who}{date_txt} {block}: '{typ}' — HONORED."
    if r.get("kind") in ("resident_hard","soft_free"):
        return f"{who}{date_txt} {block}: '{typ}' — NOT HONORED because the schedule contains: {station}."
    if r.get("kind")=="preferred":
        return f"{who}{date_txt} {block}: '{typ}' — NOT HONORED because no eligible assignment exists in that block."
    return f"{who}{date_txt} {block}: '{typ}' — NOT HONORED ({station})."


def _plain_verify_instruction(r, initials=""):
    date_txt=str(r.get("date") or "—")
    block=str(r.get("block") or "FULL")
    station=str(r.get("station") or "—")
    typ=str(r.get("type") or "request")
    if lang=="LT":
        if r.get("kind") in ("resident_hard","soft_free") and not r.get("fulfilled"):
            return f"Atverk {date_txt}, rask {initials or 'rezidentą'} ir {block} bloką. Ten turi matytis {station}. Jei tokio paskyrimo nėra, įrankio teiginys klaidingas."
        if r.get("kind")=="preferred" and not r.get("fulfilled"):
            return f"Atverk {date_txt}, rask {initials or 'rezidentą'} ir {block} bloką. Tame bloke neturi būti tinkamos darbo pamainos. Jei ji yra, įrankio teiginys klaidingas."
        if r.get("fulfilled"):
            return f"Atverk {date_txt}, rask {initials or 'rezidentą'} ir {block} bloką ir patikrink, ar grafikas atitinka sakinį kairėje."
        return "Patikrink konkretų nurodytą įrašą prieš SYSTEM grafiką / Post Matrix."
    if r.get("kind") in ("resident_hard","soft_free") and not r.get("fulfilled"):
        return f"Open {date_txt}, find {initials or 'the resident'} and the {block} block. {station} must be present; otherwise the tool statement is wrong."
    if r.get("kind")=="preferred" and not r.get("fulfilled"):
        return f"Open {date_txt}, find {initials or 'the resident'} and the {block} block. There should be no eligible assignment in that block."
    return "Verify the specific statement against the SYSTEM grid / Post Matrix."



def _workstyle_request_text(r):
    wp=r.get("workstyle_proof") or {}
    mode=int(wp.get("mode") or r.get("requested_value") or 0)
    if lang=="LT":
        return {
            1:"Prefer 6 h darbo dienas — FROZEN SYSTEM input",
            2:"Mišrus 6 h / 12 h darbo stilius — FROZEN SYSTEM input",
            3:"Prefer 12 h darbo dienas — FROZEN SYSTEM input",
        }.get(mode,"Darbo dienos trukmės pageidavimas — FROZEN SYSTEM input")
    return {
        1:"Prefer 6 h workdays — FROZEN SYSTEM input",
        2:"Mixed 6 h / 12 h workstyle — FROZEN SYSTEM input",
        3:"Prefer 12 h workdays — FROZEN SYSTEM input",
    }.get(mode,"Workday-length preference — FROZEN SYSTEM input")


def _workstyle_schedule_text(r):
    wp=r.get("workstyle_proof") or {}
    if not wp:
        return r.get("station","—")
    mode=int(wp.get("mode") or 0)
    base=(
        f"12 h: {int(wp.get('double_days') or 0)} d.; "
        f"6 h: {int(wp.get('single_days') or 0)} d.; "
        f"Onko 9 h: {int(wp.get('onko_9h_days') or 0)} d."
    )
    if mode==3:
        base+=(
            f"; grupės 12 h dienų intervalas {int(wp.get('group_double_min') or 0)}–"
            f"{int(wp.get('group_double_max') or 0)}; išpildymo riba ≥"
            f"{int(wp.get('fulfilled_threshold_min') or 0)}"
        )
    elif mode==1:
        base+=f"; išpildymo riba ≤{int(wp.get('fulfilled_threshold_max') or 0)} 12 h dienų"
    return base


def _workstyle_verify_text(r, initials=""):
    wp=r.get("workstyle_proof") or {}
    mode=int(wp.get("mode") or 0)
    if lang=="LT":
        if mode==3:
            return (
                f"Grafike suskaičiuok {initials or 'rezidento'} dienas, kur yra ir AM, ir PM normalios pamainos: "
                f"jų turi būti {int(wp.get('double_days') or 0)}. Šio SYSTEM grupės max yra "
                f"{int(wp.get('group_double_max') or 0)}, o 12 h pageidavimo išpildymo riba ≥"
                f"{int(wp.get('fulfilled_threshold_min') or 0)}. Onko 9 h į 12 h dublių skaičių neįtraukiamas."
            )
        return "Patikrink 6 h / 12 h dienų skaičių prieš frozen SYSTEM grafiką; Onko 9 h skaičiuojamas atskirai."
    if mode==3:
        return (
            f"Count {initials or 'the resident'} days with both AM and PM normal assignments: "
            f"there should be {int(wp.get('double_days') or 0)}. Group maximum is "
            f"{int(wp.get('group_double_max') or 0)} and the 12 h fulfillment threshold is ≥"
            f"{int(wp.get('fulfilled_threshold_min') or 0)}. Onko 9 h does not count as a 12 h double-day."
        )
    return "Verify the 6 h / 12 h day counts against the frozen SYSTEM schedule."


def request_details_df(rows, initials=""):
    """Plain-language resident/senior request audit table.

    V2.5.87: workstyle rows are rendered from the FROZEN SYSTEM request snapshot
    and contain concrete 6 h / 12 h / Onko counts rather than a generic FULL row.
    """
    out=[]
    for r in rows or []:
        fulfilled=bool(r.get("fulfilled"))
        is_workstyle=bool(
            r.get("kind") in ("shift_length_preference","avoid_doubles")
            and r.get("workstyle_proof")
        )
        requested=(
            _workstyle_request_text(r)
            if is_workstyle else
            f"{r.get('type','—')} · {r.get('date','—')} · {r.get('block','—')}"
        )
        shown=(
            _workstyle_schedule_text(r)
            if is_workstyle else
            (r.get("station","—") if r.get("station","—")!="—" else ("Nėra persidengiančio paskyrimo" if lang=="LT" else "No overlapping assignment"))
        )
        verify=(
            _workstyle_verify_text(r,initials)
            if is_workstyle else
            _plain_verify_instruction(r,initials)
        )
        fix_hint=(r.get("swap_hint","—") if not fulfilled else "—")
        if not fulfilled and r.get("kind")=="preferred":
            try:
                _rd=str(r.get("date") or "")
                _dt=date.fromisoformat(_rd) if _rd and _rd!="—" else None
            except Exception:
                _dt=None
            if _dt is not None and _dt.weekday()>=5:
                fix_hint=(
                    "Sistema pirmiausia bandė įvykdyti tavo savanorišką savaitgalio pasirinkimą, bet aukštesnės grafiko taisyklės to neleido. "
                    "Sugeneravus grafiką gali pasiūlyti asmeninį apsikeitimą Apsikeitimų lange."
                    if lang=="LT" else
                    "The engine first tried to honor your voluntary weekend choice, but higher scheduling rules prevented it. "
                    "After generation you can propose a personal swap in the Swaps window."
                )
        out.append({
            ("Lygis" if lang=="LT" else "Level"):r.get("priority","—"),
            ("Ko prašei" if lang=="LT" else "What was requested"):requested,
            ("Ką rodo grafikas" if lang=="LT" else "What the schedule shows"):shown,
            ("Rezultatas" if lang=="LT" else "Result"):("ĮVYKDYTA" if fulfilled else "NEĮVYKDYTA") if lang=="LT" else ("HONORED" if fulfilled else "NOT HONORED"),
            ("Aiškus paaiškinimas" if lang=="LT" else "Plain-language explanation"):_plain_request_sentence(r,initials),
            ("Kaip patikrinti" if lang=="LT" else "How to verify"):verify,
            ("Jei nori taisyti" if lang=="LT" else "If you want to fix it"):fix_hint,
        })
    return pd.DataFrame(out)


def workplace_exposure_df(y,m,result):
    """Exact monthly assignment counts per workplace, computed directly from assignments.

    This intentionally does not trust stored stats so historical schedules created
    before the rotation-fairness feature still show correct workplace counts.
    """
    slots={s.idx:s for s in make_slots(y,m)}
    names={p["initials"]:p["name"] for p in DEFAULT_PEOPLE}
    rows=[]
    for i in [p["initials"] for p in DEFAULT_PEOPLE]:
        counts={cat:0 for cat in ROTATION_CATEGORIES}
        for sid,who in result.assignments.items():
            if who!=i:
                continue
            s=slots.get(int(sid))
            if s is None:
                continue
            cat=rotation_category(s)
            if cat in counts:
                counts[cat]+=1
        row={
            tr("person"):i,
            tr("name"):names.get(i,""),
        }
        for cat in ROTATION_CATEGORIES:
            row[cat]=int(counts[cat])
        row[("Skirtingi postai" if lang=="LT" else "Distinct workplaces")]=sum(1 for v in counts.values() if v>0)
        rows.append(row)
    return pd.DataFrame(rows)


def schedule_list_df(y,m,result):
    out=[]
    for s in make_slots(y,m):
        who=result.assignments.get(s.idx,"")
        if not who and not s.blocked:
            continue
        out.append({
            tr("date"):f"{y}-{m:02d}-{s.day:02d}",
            tr("day"):WEEKDAY_FULL[lang][s.weekday],
            tr("department"):s.department,
            tr("shift"):block_label(s.block),
            tr("workload"):s.workload2/2,
            tr("person"):who if who else "—",
        })
    return pd.DataFrame(out)


def backup_table(y,m,result,backup_rows_override=None):
    slots={s.idx:s for s in make_slots(y,m)}
    rows=[]
    backup_rows=_backup_rows_for_result(y,m,result,backup_rows_override)
    for r in backup_rows:
        sid=int(r["covered_slot"])
        s=slots.get(sid)
        if s is None:
            continue
        covered=result.assignments.get(sid,"")
        rows.append({
            "ID":r.get("id","DRAFT"),
            tr("date"):f"{y}-{m:02d}-{s.day:02d}",
            tr("covered_person"):covered,
            tr("department"):s.department,
            tr("shift"):block_label(s.block),
            tr("covered_schedule"):_covered_shift_text(y,m,result,sid),
            tr("planned_backup"):r["planned_backup"],
            tr("actual_backup"):r.get("actual_backup") or "",
            tr("effective_backup"):r.get("actual_backup") or r.get("planned_backup",""),
            tr("backup_note"):r.get("note",("AUTO SYSTEM" if lang=="EN" else "AUTO SYSTEM")),
        })
    return pd.DataFrame(rows)


def backup_grid(y,m,result,initials):
    _,ndays=calendar.monthrange(y,m)
    slots={s.idx:s for s in make_slots(y,m)}
    by_day={d:[] for d in range(1,ndays+1)}
    for r in _backup_rows_for_result(y,m,result):
        eff=r.get("actual_backup") or r.get("planned_backup")
        if eff!=initials:
            continue
        sid=int(r["covered_slot"])
        s=slots.get(sid)
        if s is None:
            continue
        covered=result.assignments.get(sid,"")
        by_day[s.day].append((s.block,covered,s.department))

    for d in by_day:
        by_day[d]=sorted(
            by_day[d],
            key=lambda x:({"AM":0,"FULL":1,"PM":2}.get(x[0],9),x[1],x[2])
        )

    max_rows=max([len(v) for v in by_day.values()] or [0])
    if max_rows==0:
        vals={d:"" for d in range(1,ndays+1)}
        df=pd.DataFrame([vals],index=[tr("backups")])
    else:
        rows=[]; idx=[]
        for k in range(max_rows):
            vals={}
            for d in range(1,ndays+1):
                if k<len(by_day[d]):
                    block,covered,dept=by_day[d][k]
                    vals[d]=f"{covered}\n{block_label(block)} · {dept}"
                else:
                    vals[d]=""
            rows.append(vals)
            idx.append(f"{tr('backups')} {k+1}" if max_rows>1 else tr("backups"))
        df=pd.DataFrame(rows,index=idx)

    df.columns=[
        f"{d:02d}\n{WEEKDAYS[lang][date(y,m,d).weekday()]}"
        for d in range(1,ndays+1)
    ]
    def cs(v):
        if not v:
            return ""
        ini=str(v).split("\n")[0]
        c=PERSON_COLORS.get(ini)
        return "" if not c else (
            f"background-color:{c};color:{contrast_text(c)};"
            "font-weight:700;text-align:center;white-space:pre-wrap;"
        )
    return df.style.map(cs)


def personal_schedule_df(y,m,result,initials):
    """Personal NORMAL-WORK ledger only.

    V2.5.115 deliberately does not mix theoretical backup/standby duties into the
    resident's actual work table. Backups are rendered immediately below in their
    own dedicated backup layer/grid and only become real work after COMPLETED cover.
    """
    hours={"AM":"08:00–14:00","PM":"14:00–20:00","FULL":"08:00–17:00"}; rows=[]
    slots=make_slots(y,m)
    for sl in slots:
        if result.assignments.get(sl.idx)!=initials:
            continue
        rows.append({
            tr("date"):f"{y}-{m:02d}-{sl.day:02d}",
            tr("day"):WEEKDAY_FULL[lang][sl.weekday],
            tr("time"):hours.get(sl.block,block_label(sl.block)),
            tr("department"):sl.department,
            tr("shift"):block_label(sl.block),
        })
    if not rows:
        return pd.DataFrame(rows)
    df=pd.DataFrame(rows)
    df["__sort_date"]=pd.to_datetime(df[tr("date")],errors="coerce")
    df=df.sort_values(["__sort_date",tr("time")],kind="stable").drop(columns=["__sort_date"])
    return df.reset_index(drop=True)


def build_ics(y,m,result,initials):
    slots={s.idx:s for s in make_slots(y,m)}
    p=next((x for x in DEFAULT_PEOPLE if x["initials"]==initials),None)
    name=p["name"] if p else initials
    color=PERSON_COLORS.get(initials,"#777777")
    tz=ZoneInfo("Europe/Vilnius")
    generated=datetime.now(timezone.utc).strftime("%Y%m%dT%H%M%SZ")
    bt={"AM":(time(8),time(14)),"PM":(time(14),time(20)),"FULL":(time(8),time(17))}
    calname=f"{name} — Radiologija" if lang=="LT" else f"{name} — Radiology"
    lines=[
        "BEGIN:VCALENDAR","VERSION:2.0",
        "PRODID:-//Radiology Scheduler//V2.5.68//EN",
        "CALSCALE:GREGORIAN","METHOD:PUBLISH",
        f"X-WR-CALNAME:{ics_escape(calname)}",
        "X-WR-TIMEZONE:Europe/Vilnius",
        f"COLOR:{color}",f"X-APPLE-CALENDAR-COLOR:{color}"
    ]

    # Normal assignments.
    for sid,who in sorted(
        result.assignments.items(),
        key=lambda kv:(slots[kv[0]].day,{"AM":0,"FULL":1,"PM":2}.get(slots[kv[0]].block,9),kv[0])
    ):
        if who!=initials:
            continue
        s=slots[sid]
        stt,endt=bt[s.block]
        start=datetime.combine(date(y,m,s.day),stt,tzinfo=tz).astimezone(timezone.utc).strftime("%Y%m%dT%H%M%SZ")
        end=datetime.combine(date(y,m,s.day),endt,tzinfo=tz).astimezone(timezone.utc).strftime("%Y%m%dT%H%M%SZ")
        title=("Radiologija — " if lang=="LT" else "Radiology — ")+s.department
        lines += [
            "BEGIN:VEVENT",
            f"UID:{y}{m:02d}{s.day:02d}-{sid}-{safe_filename(initials)}@radiology-scheduler",
            f"DTSTAMP:{generated}",f"DTSTART:{start}",f"DTEND:{end}",
            f"SUMMARY:{ics_escape(title)}",
            f"DESCRIPTION:{ics_escape(name+' | '+block_label(s.block))}",
            "TRANSP:OPAQUE",f"COLOR:{color}","STATUS:CONFIRMED","END:VEVENT"
        ]

    # Shift-level backup duties are optional in the personal calendar.
    include_backups = bool(db.get_account_settings(initials).get("include_backups_in_calendar", False))
    if include_backups:
        for r in db.list_backups(y,m):
            eff=r.get("actual_backup") or r.get("planned_backup")
            if eff!=initials:
                continue
            sid=int(r["covered_slot"])
            s=slots.get(sid)
            covered=result.assignments.get(sid,"")
            if s is None or not covered:
                continue
            stt,endt=bt[s.block]
            start=datetime.combine(date(y,m,s.day),stt,tzinfo=tz).astimezone(timezone.utc).strftime("%Y%m%dT%H%M%SZ")
            end=datetime.combine(date(y,m,s.day),endt,tzinfo=tz).astimezone(timezone.utc).strftime("%Y%m%dT%H%M%SZ")
            if lang=="LT":
                title=f"Dublis — {covered} — {s.department}"
                desc=f"Jei prireiktų, pavaduojate {covered}: {s.department}, {block_label(s.block)}."
            else:
                title=f"Backup — {covered} — {s.department}"
                desc=f"If cover is needed, you back up {covered}: {s.department}, {block_label(s.block)}."
            lines += [
                "BEGIN:VEVENT",
                f"UID:backup-{y}{m:02d}{s.day:02d}-{sid}-{safe_filename(initials)}@radiology-scheduler",
                f"DTSTAMP:{generated}",f"DTSTART:{start}",f"DTEND:{end}",
                f"SUMMARY:{ics_escape(title)}",f"DESCRIPTION:{ics_escape(desc)}",
                "TRANSP:TRANSPARENT","STATUS:TENTATIVE","END:VEVENT"
            ]
    lines.append("END:VCALENDAR")
    return ("\r\n".join(lines)+"\r\n").encode("utf-8")



def build_calendar_subscription_ics(initials, published_rows=None):
    """Build one stable subscription calendar from every published ACTUAL month.

    A subscription must not lose the previous month when a new month is published.
    Therefore the feed is assembled from all published current_json payloads, not
    merely the month currently open in the UI.
    """
    rows = published_rows if published_rows is not None else db.list_published_schedules()
    month_blobs=[]
    seen=set()
    for row in rows:
        try:
            yy=int(row.get("year")); mm=int(row.get("month"))
            if (yy,mm) in seen or not row.get("current_json"):
                continue
            rr=deserialize_result(row["current_json"])
            month_blobs.append((yy,mm,rr))
            seen.add((yy,mm))
        except Exception:
            continue
    month_blobs.sort(key=lambda x:(x[0],x[1]))

    p=next((x for x in DEFAULT_PEOPLE if x["initials"]==initials),None)
    name=p["name"] if p else initials
    color=PERSON_COLORS.get(initials,"#777777")
    settings=db.get_account_settings(initials)
    cal_lang=str(settings.get("preferred_language") or "LT").upper()
    if cal_lang not in ("LT","EN"):
        cal_lang="LT"
    generated=datetime.now(timezone.utc).strftime("%Y%m%dT%H%M%SZ")
    calname=f"{name} — Radiologija" if cal_lang=="LT" else f"{name} — Radiology"
    lines=[
        "BEGIN:VCALENDAR","VERSION:2.0",
        "PRODID:-//Radiology Scheduler//V2.5.68//EN",
        "CALSCALE:GREGORIAN","METHOD:PUBLISH",
        f"X-WR-CALNAME:{ics_escape(calname)}",
        "X-WR-TIMEZONE:Europe/Vilnius",
        f"COLOR:{color}",f"X-APPLE-CALENDAR-COLOR:{color}"
    ]
    bt={"AM":(time(8),time(14)),"PM":(time(14),time(20)),"FULL":(time(8),time(17))}
    tz=ZoneInfo("Europe/Vilnius")
    include_backups=bool(settings.get("include_backups_in_calendar",False))
    block_txt_lt={"AM":"Rytas","PM":"Popietė","FULL":"Visa diena"}
    block_txt_en={"AM":"Morning","PM":"Afternoon","FULL":"Full day"}

    for yy,mm,result in month_blobs:
        slots={sl.idx:sl for sl in make_slots(yy,mm)}
        for sid,who in sorted(result.assignments.items(),key=lambda kv:(slots[kv[0]].day,{"AM":0,"FULL":1,"PM":2}.get(slots[kv[0]].block,9),kv[0])):
            if who!=initials:
                continue
            sl=slots[sid]
            stt,endt=bt[sl.block]
            start=datetime.combine(date(yy,mm,sl.day),stt,tzinfo=tz).astimezone(timezone.utc).strftime("%Y%m%dT%H%M%SZ")
            end=datetime.combine(date(yy,mm,sl.day),endt,tzinfo=tz).astimezone(timezone.utc).strftime("%Y%m%dT%H%M%SZ")
            block_txt=(block_txt_lt if cal_lang=="LT" else block_txt_en).get(sl.block,sl.block)
            title=("Radiologija — " if cal_lang=="LT" else "Radiology — ")+sl.department
            lines += [
                "BEGIN:VEVENT",
                f"UID:{yy}{mm:02d}{sl.day:02d}-{sid}-{safe_filename(initials)}@radiology-scheduler",
                f"DTSTAMP:{generated}",f"DTSTART:{start}",f"DTEND:{end}",
                f"SUMMARY:{ics_escape(title)}",f"DESCRIPTION:{ics_escape(name+' | '+block_txt)}",
                "TRANSP:OPAQUE",f"COLOR:{color}","STATUS:CONFIRMED","END:VEVENT"
            ]
        if include_backups:
            for br in db.list_backups(yy,mm):
                eff=br.get("actual_backup") or br.get("planned_backup")
                if eff!=initials:
                    continue
                sid=int(br["covered_slot"]); sl=slots.get(sid); covered=result.assignments.get(sid,"")
                if sl is None or not covered:
                    continue
                stt,endt=bt[sl.block]
                start=datetime.combine(date(yy,mm,sl.day),stt,tzinfo=tz).astimezone(timezone.utc).strftime("%Y%m%dT%H%M%SZ")
                end=datetime.combine(date(yy,mm,sl.day),endt,tzinfo=tz).astimezone(timezone.utc).strftime("%Y%m%dT%H%M%SZ")
                block_txt=(block_txt_lt if cal_lang=="LT" else block_txt_en).get(sl.block,sl.block)
                if cal_lang=="LT":
                    title=f"Dublis — {covered} — {sl.department}"; desc=f"Jei prireiktų, pavaduojate {covered}: {sl.department}, {block_txt}."
                else:
                    title=f"Backup — {covered} — {sl.department}"; desc=f"If cover is needed, you back up {covered}: {sl.department}, {block_txt}."
                lines += [
                    "BEGIN:VEVENT",f"UID:backup-{yy}{mm:02d}{sl.day:02d}-{sid}-{safe_filename(initials)}@radiology-scheduler",
                    f"DTSTAMP:{generated}",f"DTSTART:{start}",f"DTEND:{end}",
                    f"SUMMARY:{ics_escape(title)}",f"DESCRIPTION:{ics_escape(desc)}",
                    "TRANSP:TRANSPARENT","STATUS:TENTATIVE","END:VEVENT"
                ]
    lines.append("END:VCALENDAR")
    return ("\r\n".join(lines)+"\r\n").encode("utf-8")


def refresh_calendar_subscription_feeds(initials_list=None):
    """Best-effort refresh of private subscription feeds after schedule changes.

    Publication must never fail only because a third-party calendar service is
    unavailable, so errors are returned for UI/audit instead of being raised.
    """
    people=list(initials_list) if initials_list is not None else [p["initials"] for p in DEFAULT_PEOPLE]
    rows=db.list_published_schedules()
    out=[]
    for ini in dict.fromkeys(people):
        try:
            feed=build_calendar_subscription_ics(ini,published_rows=rows)
            url=db.publish_calendar_feed(ini,feed)
            out.append({"initials":ini,"ok":True,"url":url,"error":""})
        except Exception as exc:
            out.append({"initials":ini,"ok":False,"url":"","error":str(exc)})
    return out



def build_xlsx(y,m,result,document_status=None,backup_rows_override=None):
    out=BytesIO(); wb=xlsxwriter.Workbook(out,{"in_memory":True}); ws=wb.add_worksheet("Grafikas" if lang=="LT" else "Schedule"); sm=wb.add_worksheet("Suvestinė" if lang=="LT" else "Summary"); bk=wb.add_worksheet("Dubliai" if lang=="LT" else "Backups")
    dark="#1F2937"; light="#F3F4F6"; weekend="#E5E7EB"; border="#D1D5DB"; title=wb.add_format({"bold":True,"font_size":16,"font_color":"#FFFFFF","bg_color":dark}); header=wb.add_format({"bold":True,"bg_color":light,"border":1,"border_color":border,"align":"center","valign":"vcenter","text_wrap":True}); wh=wb.add_format({"bold":True,"bg_color":weekend,"border":1,"border_color":border,"align":"center","valign":"vcenter","text_wrap":True}); cell=wb.add_format({"border":1,"border_color":border,"text_wrap":True}); blocked=wb.add_format({"border":1,"border_color":border,"bg_color":"#BFC4CA","align":"center"})
    pf={i:wb.add_format({"bold":True,"bg_color":c,"font_color":contrast_text(c),"border":1,"border_color":border,"align":"center","valign":"vcenter","text_wrap":True}) for i,c in PERSON_COLORS.items()}
    _,nd=calendar.monthrange(y,m); last=1+nd; status_prefix=(str(document_status).strip()+" — ") if document_status else ""
    ws.merge_range(0,0,0,last,status_prefix+("Rezidentų grafikas — " if lang=="LT" else "Resident schedule — ")+month_label(y,m),title); ws.write(1,0,tr("department"),header); ws.write(1,1,tr("shift"),header)
    for d in range(1,nd+1): ws.write(1,1+d,f"{d:02d}\n{WEEKDAYS[lang][date(y,m,d).weekday()]}",wh if date(y,m,d).weekday()>=5 else header)
    rowkeys=[]; slots=make_slots(y,m)
    for s in slots:
        k=(s.department,s.block)
        if k not in rowkeys: rowkeys.append(k)
    by={(s.department,s.block,s.day):s for s in slots}
    for r,(dept,block) in enumerate(rowkeys,start=2):
        ws.write(r,0,dept,cell); ws.write(r,1,block_label(block),cell)
        for d in range(1,nd+1):
            s=by.get((dept,block,d))
            if not s: ws.write_blank(r,1+d,None,cell)
            elif s.blocked: ws.write(r,1+d,"—",blocked)
            else:
                who=result.assignments.get(s.idx,""); ws.write(r,1+d,who,pf[who] if who else cell)
    ws.set_column(0,0,23); ws.set_column(1,1,13); ws.set_column(2,last,7); ws.freeze_panes(2,2)
    sdf=summary_df(result,y,m)
    for c,col in enumerate(sdf.columns): sm.write(0,c,col,header)
    for rr,(_,row) in enumerate(sdf.iterrows(),start=1):
        fmt=pf.get(str(row[tr("person")]),cell)
        for c,v in enumerate(row): sm.write(rr,c,"" if pd.isna(v) else v,fmt)
    sm.set_column(0,1,20); sm.set_column(2,len(sdf.columns)-1,17); sm.freeze_panes(1,0)

    # Shift-level backup sheet: rows are backup residents, days are columns.
    # Each cell states the exact covered colleague, block and department.
    bk.merge_range(0,0,0,nd,("Dublių grafikas — " if lang=="LT" else "Backup schedule — ")+month_label(y,m),title)
    bk.write(1,0,tr("person"),header)
    for d in range(1,nd+1):
        bk.write(1,d,f"{d:02d}\n{WEEKDAYS[lang][date(y,m,d).weekday()]}",wh if date(y,m,d).weekday()>=5 else header)
    effective_map={}
    slot_map={s.idx:s for s in slots}
    backup_rows=_backup_rows_for_result(y,m,result,backup_rows_override)
    for r in backup_rows:
        sid=int(r["covered_slot"]); s=slot_map.get(sid)
        if s is None: continue
        eff=r["actual_backup"] or r["planned_backup"]
        covered=result.assignments.get(sid,"")
        effective_map.setdefault((eff,s.day),[]).append((s.block,covered,s.department))
    rr=2
    for p in DEFAULT_PEOPLE:
        i=p["initials"]
        max_rows=max([len(effective_map.get((i,d),[])) for d in range(1,nd+1)] or [0])
        max_rows=max(1,max_rows)
        for k in range(max_rows):
            label=i if max_rows==1 else f"{i} {k+1}"
            bk.write(rr,0,label,pf[i])
            for d in range(1,nd+1):
                vals=sorted(effective_map.get((i,d),[]),key=lambda x:({"AM":0,"FULL":1,"PM":2}.get(x[0],9),x[1],x[2]))
                if k>=len(vals):
                    bk.write_blank(rr,d,None,cell)
                else:
                    block,covered,dept=vals[k]
                    bk.write(rr,d,f"{covered}\n{block_label(block)} · {dept}",pf.get(covered,cell))
            rr+=1
    bk.set_column(0,0,10); bk.set_column(1,nd,20); bk.freeze_panes(2,1)

    detail_start=rr+2; bdf=backup_table(y,m,result,backup_rows_override=backup_rows)
    if not bdf.empty:
        bk.write(detail_start,0,tr("details"),title)
        for c,col in enumerate(bdf.columns): bk.write(detail_start+1,c,col,header)
        for rr,(_,row) in enumerate(bdf.iterrows(),start=detail_start+2):
            for c,v in enumerate(row): bk.write(rr,c,"" if pd.isna(v) else v,cell)
    wb.close(); out.seek(0); return out.getvalue()

def render_schedule_download_buttons(y,m,result,*,status_label,file_prefix,key_prefix,backup_rows_override=None):
    """Explicit schedule exports next to every operational schedule view.

    Streamlit's dataframe toolbar exposes CSV only. This helper makes the same
    displayed schedule downloadable as a formatted Excel workbook as well, while
    retaining an explicit CSV option for users who prefer flat data.
    """
    if result is None:
        return
    excel_label=("ATSISIŲSTI EXCEL (.xlsx)" if lang=="LT" else "DOWNLOAD EXCEL (.xlsx)")
    csv_label=("ATSISIŲSTI CSV (.csv)" if lang=="LT" else "DOWNLOAD CSV (.csv)")
    c_excel,c_csv=st.columns(2)
    with c_excel:
        st.download_button(
            excel_label,
            build_xlsx(y,m,result,document_status=status_label,backup_rows_override=backup_rows_override),
            file_name=f"{file_prefix}_{y}_{m:02d}.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            use_container_width=True,
            key=f"{key_prefix}_xlsx_{y}_{m}",
        )
    with c_csv:
        st.download_button(
            csv_label,
            schedule_list_df(y,m,result).to_csv(index=False).encode("utf-8-sig"),
            file_name=f"{file_prefix}_{y}_{m:02d}.csv",
            mime="text/csv",
            use_container_width=True,
            key=f"{key_prefix}_csv_{y}_{m}",
        )


def _truthy_cfg(value,default=False):
    if value is None or str(value).strip()=="": return bool(default)
    return str(value).strip().lower() in ("1","true","yes","on","y")

def smtp_config():
    return _smtp_config_core(config_value)

def smtp_diagnostics():
    cfg=smtp_config()
    return cfg,_smtp_missing_core(cfg)

def smtp_ready():
    _cfg,missing=smtp_diagnostics()
    return not missing

def smtp_probe():
    cfg,missing=smtp_diagnostics()
    if missing:
        return False,("Trūksta konfigūracijos: " if lang=="LT" else "Missing configuration: ")+", ".join(missing)
    return _smtp_probe_core(cfg)

def send_email(to_addr,subject,body,ics_bytes=None,ics_name=None):
    cfg,missing=smtp_diagnostics()
    if missing:
        detail=("El. pašto kanalas nesukonfigūruotas: " if lang=="LT" else "Email channel is not configured: ")+", ".join(missing)
        return False,detail
    return _send_email_core(cfg,to_addr,subject,body,ics_text=ics_bytes,ics_name=ics_name)

def deliver_lifecycle_notification(event_key,event_type,initials,y,m,to_email,subject,body,ics_bytes=None,ics_name=None):
    """Idempotent lifecycle delivery: durable DB outbox first, immediate SMTP second.

    If SMTP fails, the row remains FAILED and the background worker can retry it.
    If the event was already sent, reruns do not duplicate the email.
    """
    ics_text=None
    if ics_bytes is not None:
        ics_text=ics_bytes.decode("utf-8") if isinstance(ics_bytes,(bytes,bytearray)) else str(ics_bytes)
    row=db.enqueue_notification_v25100(
        event_key,event_type,initials,y,m,to_email,subject,body,
        scheduled_for=datetime.now(timezone.utc).isoformat(),ics_text=ics_text,ics_name=ics_name
    )
    if str(row.get("status") or "")=="sent":
        return True,"already sent"
    if not str(to_email or "").strip():
        return False,tr("missing_email")
    ok,detail=send_email(to_email,subject,body,ics_bytes,ics_name)
    try:
        db.mark_notification_delivery_v25100(int(row.get("id")),ok,detail)
    except Exception:
        pass
    return ok,detail

def retry_failed_lifecycle_notifications(y,m,event_types=None):
    if isinstance(event_types,str): event_types={event_types}
    elif event_types is not None: event_types=set(event_types)
    rows=db.failed_notification_outbox_v25100(y,m,100)
    if event_types is not None:
        rows=[r for r in rows if str(r.get("event_type") or "") in event_types]
    results=[]
    for row in rows:
        email=str(row.get("to_email") or "").strip()
        if not email:
            # Refresh from current resident settings before declaring it blocked.
            email=str((db.get_account_settings(str(row.get("initials") or "")) or {}).get("email") or "").strip()
            if email:
                row=db.enqueue_notification_v25100(
                    row.get("event_key"),row.get("event_type"),row.get("initials"),y,m,email,
                    row.get("subject") or "",row.get("body") or "",
                    scheduled_for=datetime.now(timezone.utc).isoformat(),ics_text=row.get("ics_text"),ics_name=row.get("ics_name")
                )
        if not email:
            results.append((row.get("initials"),"blocked",tr("missing_email"))); continue
        ics=(row.get("ics_text") or None)
        ok,detail=send_email(email,row.get("subject") or "",row.get("body") or "",ics,row.get("ics_name"))
        try: db.mark_notification_delivery_v25100(int(row.get("id")),ok,detail)
        except Exception: pass
        results.append((row.get("initials"),"sent" if ok else "failed",detail))
    return results

def send_backup_activation_email(y,m,result,backup_row):
    eff=backup_row.get("actual_backup") or backup_row.get("planned_backup")
    settings=db.get_account_settings(eff)
    if not settings.get("backup_email_alerts", True):
        return True, "alerts disabled"
    email=(settings.get("email") or "").strip()
    if not email:
        return False, tr("missing_email")
    slots={s.idx:s for s in make_slots(y,m)}
    sid=int(backup_row["covered_slot"]); s=slots.get(sid)
    covered=result.assignments.get(sid,"") if s else ""
    if not s:
        return False, "slot not found"
    subject=(f"Dublis aktyvuotas — {covered} — {s.department}" if lang=="LT" else f"Backup activated — {covered} — {s.department}")
    when=f"{y}-{m:02d}-{s.day:02d} · {block_label(s.block)}"
    body=(f"Sveiki,\n\nJūsų dublio pareiga aktyvuota.\n{when}\nDubliuojate: {covered}\nPadalinys: {s.department}\n\nPrašome susisiekti su seniūne / skyriumi ir patvirtinti veiksmus.\n"
          if lang=="LT" else
          f"Hello,\n\nYour backup duty has been activated.\n{when}\nCovered resident: {covered}\nDepartment: {s.department}\n\nPlease contact the senior scheduler / department and confirm next steps.\n")
    return send_email(email,subject,body)



def send_swap_request_email(y,m,request_row):
    """Best-effort operational email. DB request already exists before this runs."""
    target=str(request_row.get("person_b") or "")
    proposer=str(request_row.get("person_a") or "")
    request_id=int(request_row.get("id") or 0)
    settings=db.get_account_settings(target)
    email=(settings.get("email") or "").strip()
    kind=f"swap_request_{request_id}"
    send_date=date.today().isoformat()

    if not email:
        detail=tr("missing_email")
        try: db.record_email(target,kind,y,m,send_date,"failed",detail)
        except Exception: pass
        return False,detail

    slot_map={s.idx:s for s in make_slots(y,m)}
    sa=slot_map.get(int(request_row.get("slot_a") or -1))
    sb=slot_map.get(int(request_row.get("slot_b") or -1))
    subject=(
        f"Naujas apsikeitimo prašymas nuo {proposer}"
        if lang=="LT" else
        f"New swap request from {proposer}"
    )
    body=(
        f"Sveiki,\n\nGavote naują apsikeitimo prašymą nuo {proposer} ({_person_name(proposer)}).\n\n"
        f"Jis/ji siūlo: {_swap_shift_text(sa)}\n"
        f"Mainais prašo jūsų pamainos: {_swap_shift_text(sb)}\n\n"
        f"Prašymo DB numeris: #{request_id}\n"
        f"Prisijunkite prie Shift Happens → Apsikeitimai ir PRIIMKITE arba ATMESKITE prašymą.\n"
        if lang=="LT" else
        f"Hello,\n\nYou received a new swap request from {proposer} ({_person_name(proposer)}).\n\n"
        f"They offer: {_swap_shift_text(sa)}\n"
        f"They request your shift: {_swap_shift_text(sb)}\n\n"
        f"Database request number: #{request_id}\n"
        f"Open Shift Happens → Swaps and ACCEPT or REJECT the request.\n"
    )
    public=config_value("SCHEDULER_PUBLIC_URL","").strip()
    if public:
        body += (f"\nPortalas: {public}\n" if lang=="LT" else f"\nPortal: {public}\n")
    ok,detail=send_email(email,subject,body)
    try:
        db.record_email(target,kind,y,m,send_date,"sent" if ok else "failed",detail)
    except Exception:
        pass
    return ok,detail



def send_backup_swap_request_email(y,m,request_row):
    target=str(request_row.get("target") or "")
    proposer=str(request_row.get("requester") or "")
    request_id=int(request_row.get("id") or 0)
    settings=db.get_account_settings(target)
    email=(settings.get("email") or "").strip()
    kind=f"backup_swap_request_{request_id}"
    send_date=date.today().isoformat()

    if not email:
        detail=tr("missing_email")
        try: db.record_email(target,kind,y,m,send_date,"failed",detail)
        except Exception: pass
        return False,detail

    subject=(
        f"Naujas dublio apsikeitimo prašymas nuo {proposer}"
        if lang=="LT" else
        f"New backup swap request from {proposer}"
    )
    body=(
        f"Sveiki,\n\nGavote naują DUBLIO apsikeitimo prašymą nuo {proposer} ({_person_name(proposer)}).\n"
        f"Prašymo DB numeris: #{request_id}.\n\n"
        f"Prisijunkite prie Shift Happens → Apsikeitimai ir priimkite arba atmeskite prašymą.\n"
        if lang=="LT" else
        f"Hello,\n\nYou received a new BACKUP swap request from {proposer} ({_person_name(proposer)}).\n"
        f"Database request number: #{request_id}.\n\n"
        f"Open Shift Happens → Swaps and accept or reject the request.\n"
    )
    public=config_value("SCHEDULER_PUBLIC_URL","").strip()
    if public:
        body += (f"\nPortalas: {public}\n" if lang=="LT" else f"\nPortal: {public}\n")
    ok,detail=send_email(email,subject,body)
    try:
        db.record_email(target,kind,y,m,send_date,"sent" if ok else "failed",detail)
    except Exception:
        pass
    return ok,detail


def publication_emails(y,m,result):
    settings=db.all_account_settings(); results=[]; public=config_value("SCHEDULER_PUBLIC_URL","").strip()
    for p in DEFAULT_PEOPLE:
        i=p["initials"]; email=settings.get(i,{}).get("email","").strip(); send_date=date.today().isoformat()
        if not email: results.append((i,"skipped",tr("missing_email"))); continue
        subject=f"{month_label(y,m)} grafikas patvirtintas" if lang=="LT" else f"{month_label(y,m)} schedule approved"
        body=(f"Sveiki,\n\nKito mėnesio tvarkaraštis patvirtintas. Prisegtame .ics faile yra jūsų normalios pamainos; dubliai įtraukiami tik jei tai įjungėte Nustatymuose. Atidarykite failą ir įsidėkite grafiką į savo kalendorių.\n" if lang=="LT" else f"Hello,\n\nThe next-month schedule has been approved. The attached .ics file contains your normal shifts; backup duties are included only if you enabled them in Settings. Open it to add the schedule to your calendar.\n")
        if public: body += (f"\nPortalas: {public}\n" if lang=="LT" else f"\nPortal: {public}\n")
        ok,detail=send_email(email,subject,body,build_ics(y,m,result,i),f"{safe_filename(i)}_{y}_{m:02d}.ics"); status="sent" if ok else "failed"; db.record_email(i,"publication",y,m,send_date,status,detail); results.append((i,status,detail))
    return results

def _parse_iso_dt(value):
    if not value:
        return None
    try:
        return datetime.fromisoformat(str(value).replace("Z","+00:00"))
    except Exception:
        return None


def _vilnius_now():
    return datetime.now(ZoneInfo("Europe/Vilnius"))


def _workflow_card(title,body,state="draft"):
    palette={
        "draft":("#E8F0FE","#174EA6","#8AB4F8"),
        "swap_open":("#FFF4E5","#8A4B08","#F6B26B"),
        "expired":("#FDECEC","#9C1C1C","#E57373"),
        "swap_closed":("#FFF8D9","#6B5600","#E4C441"),
        "final":("#E6F4EA","#176B36","#81C995"),
    }
    bg,fg,border=palette.get(state,palette["draft"])
    st.markdown(
        f"""<div style=\"border:2px solid {border};background:{bg};color:{fg};border-radius:16px;padding:16px 18px;margin:8px 0 14px 0;\">
        <div style=\"font-weight:800;font-size:1.05rem;letter-spacing:.02em;\">{html.escape(str(title))}</div>
        <div style=\"margin-top:6px;line-height:1.45;\">{html.escape(str(body))}</div></div>""",
        unsafe_allow_html=True,
    )


def _resident_email_preflight():
    settings=db.all_account_settings()
    missing=[]
    for pp in DEFAULT_PEOPLE:
        ini=pp["initials"]
        if not str((settings.get(ini,{}) or {}).get("email") or "").strip():
            missing.append(ini)
    return missing


def render_operator_email_smtp_admin(current_operator):
    """Compact operator email readiness UI; technical detail stays in Advanced mode."""
    settings=db.all_account_settings()
    missing=_resident_email_preflight()
    cfg,smtp_missing=smtp_diagnostics()
    ready=bool(not missing and not smtp_missing)

    with st.expander(
        "El. pašto kanalas" if lang=="LT" else "Email channel",
        expanded=not ready,
    ):
        if ready:
            st.success("El. pašto konfigūracija paruošta, o visi 16 rezidentų turi gavėjo adresą." if lang=="LT" else "Email configuration is present and all 16 residents have recipient addresses.")
        else:
            problems=[]
            if smtp_missing: problems.append("siuntėjo konfigūracija" if lang=="LT" else "sender configuration")
            if missing: problems.append(("gavėjo adresai: "+", ".join(missing)) if lang=="LT" else ("recipient addresses: "+", ".join(missing)))
            st.warning(("Dar neparuošta: " if lang=="LT" else "Not ready yet: ")+"; ".join(problems))

        if missing:
            if st.button(
                "UŽPILDYTI TRŪKSTAMUS IŠ PRISIJUNGIMO PASKYRŲ" if lang=="LT" else "FILL MISSING FROM LOGIN ACCOUNTS",
                use_container_width=True,key="autofill_notification_emails_v25100"
            ):
                try:
                    res=db.autofill_notification_emails_v2593()
                    st.success((f"Užpildyta: {int(res.get('filled',0))}." if lang=="LT" else f"Filled: {int(res.get('filled',0))}."))
                    st.rerun()
                except Exception as exc:
                    st.error(str(exc))

            people_options=[pp["initials"] for pp in DEFAULT_PEOPLE]
            default_ini=missing[0] if missing else people_options[0]
            c1,c2=st.columns([1,2])
            with c1:
                email_ini=st.selectbox("Rezidentas" if lang=="LT" else "Resident",people_options,index=people_options.index(default_ini),key="operator_email_ini_v25100")
            with c2:
                existing_email=str((settings.get(email_ini,{}) or {}).get("email") or "").strip()
                email_value=st.text_input("Pranešimų el. paštas" if lang=="LT" else "Notification email",value=existing_email,key=f"operator_email_value_v25100_{email_ini}")
            if st.button("IŠSAUGOTI ADRESĄ" if lang=="LT" else "SAVE ADDRESS",use_container_width=True,key="operator_save_email_v25100"):
                if not re.fullmatch(r"[^@\s]+@[^@\s]+\.[^@\s]+",email_value.strip()):
                    st.error("Neteisingas el. pašto formatas." if lang=="LT" else "Invalid email format.")
                else:
                    try:
                        db.set_resident_notification_email_v2593(email_ini,email_value.strip())
                        st.rerun()
                    except Exception as exc:
                        st.error(str(exc))

        c1,c2=st.columns(2)
        with c1:
            if st.button("PATIKRINTI KANALĄ" if lang=="LT" else "CHECK CHANNEL",use_container_width=True,disabled=bool(smtp_missing),key="smtp_probe_v25100"):
                ok,detail=smtp_probe()
                st.session_state["_smtp_probe_v25100"]=(ok,detail,datetime.now().isoformat(timespec="seconds"))
        with c2:
            operator_email=str((settings.get(current_operator,{}) or {}).get("email") or "").strip()
            if st.button("SIŲSTI TESTĄ MAN" if lang=="LT" else "SEND TEST TO ME",use_container_width=True,disabled=bool(smtp_missing or not operator_email),key="smtp_test_v25100"):
                ok,detail=send_email(
                    operator_email,
                    "Shift Happens — el. pašto testas" if lang=="LT" else "Shift Happens — email test",
                    "Testas sėkmingas. Shift Happens gali siųsti realius grafiko pranešimus." if lang=="LT" else "Test successful. Shift Happens can send real schedule notifications."
                )
                st.session_state["_smtp_probe_v25100"]=(ok,detail or "test email sent",datetime.now().isoformat(timespec="seconds"))

        probe=st.session_state.get("_smtp_probe_v25100")
        if probe:
            ok,detail,when=probe
            if ok: st.success(("Kanalas veikia ✓ · " if lang=="LT" else "Channel works ✓ · ")+when)
            else: st.error(("Kanalo patikra nepavyko: " if lang=="LT" else "Channel check failed: ")+str(detail))

        if advanced_mode:
            st.caption("Techninė konfigūracija" if lang=="LT" else "Technical configuration")
            smtp_rows=[
                {"Parametras" if lang=="LT" else "Setting":"Host","Reikšmė" if lang=="LT" else "Value":cfg.get("host") or "—"},
                {"Parametras" if lang=="LT" else "Setting":"Port","Reikšmė" if lang=="LT" else "Value":cfg.get("port")},
                {"Parametras" if lang=="LT" else "Setting":"From","Reikšmė" if lang=="LT" else "Value":cfg.get("from_email") or "—"},
                {"Parametras" if lang=="LT" else "Setting":"Login","Reikšmė" if lang=="LT" else "Value":cfg.get("user") or "—"},
                {"Parametras" if lang=="LT" else "Setting":"Security","Reikšmė" if lang=="LT" else "Value":"SSL" if cfg.get("use_ssl") else "STARTTLS" if cfg.get("use_tls") else "plain"},
            ]
            st.dataframe(pd.DataFrame(smtp_rows),use_container_width=True,hide_index=True)
            st.caption("Slaptažodis niekada nerodomas. Naudokite Streamlit Secrets [smtp] bloką; Gmail atveju — App Password, ne įprastą paskyros slaptažodį." if lang=="LT" else "The password is never displayed. Use the Streamlit Secrets [smtp] block; for Gmail use an App Password, not the normal account password.")

            try:
                audit=db.list_resident_email_admin_audit_v2593(30)
                if audit:
                    with st.expander("Adresų pakeitimų auditas" if lang=="LT" else "Address change audit",expanded=False):
                        rows=[]
                        for r in audit:
                            rows.append({
                                "Laikas" if lang=="LT" else "Time":r.get("created_at"),
                                "Rezidentas" if lang=="LT" else "Resident":r.get("initials"),
                                "Senas" if lang=="LT" else "Old":r.get("old_email"),
                                "Naujas" if lang=="LT" else "New":r.get("new_email"),
                                "Šaltinis" if lang=="LT" else "Source":r.get("source"),
                                "Operatorius" if lang=="LT" else "Operator":r.get("actor_initials"),
                            })
                        st.dataframe(pd.DataFrame(rows),use_container_width=True,hide_index=True)
            except Exception:
                pass


def preferences_open_emails(y,m):
    settings=db.all_account_settings(); results=[]; public=config_value("SCHEDULER_PUBLIC_URL","").strip()
    cutoff_text=preference_cutoff_for(y,m).strftime("%Y-%m-%d %H:%M")
    for pp in DEFAULT_PEOPLE:
        i=pp["initials"]; email=str((settings.get(i,{}) or {}).get("email") or "").strip()
        subject=(f"{month_label(y,m)} pageidavimai atidaryti" if lang=="LT" else f"{month_label(y,m)} preferences are open")
        body=(
            f"Sveiki,\n\nAtidarytas {month_label(y,m)} grafiko pageidavimų etapas. "
            f"Pageidavimus pateikite iki {cutoff_text} Lietuvos laiku.\n"
            if lang=="LT" else
            f"Hello,\n\nThe {month_label(y,m)} preference stage is open. "
            f"Please submit your preferences by {cutoff_text} Lithuania time.\n"
        )
        if public: body += f"\nShift Happens: {public}\n"
        ok,detail=deliver_lifecycle_notification(
            f"preferences_open:{y}-{m:02d}","preferences_open",i,y,m,email,subject,body
        )
        results.append((i,"sent" if ok else "failed",detail))
    return results


def render_notification_delivery_status(y,m):
    """Lifecycle delivery dashboard: one compact table + retry only failures."""
    try:
        rows=db.notification_summary_v25100(y,m)
    except Exception as exc:
        if advanced_mode:
            st.caption(("Pranešimų outbox dar nepasiekiamas: " if lang=="LT" else "Notification outbox unavailable: ")+str(exc))
        return
    labels={
        "preferences_open":"Pageidavimai atidaryti" if lang=="LT" else "Preferences open",
        "preferences_reminder":"Pageidavimų priminimai" if lang=="LT" else "Preference reminders",
        "swap_open":"Apsikeitimų etapas" if lang=="LT" else "Swap stage",
        "final":"Baigta / FINAL" if lang=="LT" else "Done / FINAL",
    }
    st.markdown("#### Pranešimų etapai" if lang=="LT" else "#### Notification stages")
    if not rows:
        st.caption("Šiam mėnesiui dar nėra išsiųstų etapų pranešimų." if lang=="LT" else "No lifecycle notifications have been sent for this month yet.")
    else:
        view=[]
        for r in rows:
            view.append({
                "Etapas" if lang=="LT" else "Stage":labels.get(r.get("event_type"),r.get("event_type")),
                "Išsiųsta" if lang=="LT" else "Sent":int(r.get("sent",0)),
                "Laukia" if lang=="LT" else "Pending":int(r.get("pending",0)),
                "Nepavyko" if lang=="LT" else "Failed":int(r.get("failed",0))+int(r.get("blocked",0)),
                "Iš viso" if lang=="LT" else "Total":int(r.get("total",0)),
            })
        st.dataframe(pd.DataFrame(view),use_container_width=True,hide_index=True)

    failed=db.failed_notification_outbox_v25100(y,m,100)
    if failed:
        st.warning((f"Nepavykusių / užblokuotų pranešimų: {len(failed)}." if lang=="LT" else f"Failed / blocked notifications: {len(failed)}."))
        failed_types=sorted({str(r.get("event_type") or "other") for r in failed})
        chosen=st.selectbox(
            "Kurį etapą kartoti" if lang=="LT" else "Stage to retry",failed_types,
            format_func=lambda x:labels.get(x,x),key=f"retry_notification_type_{y}_{m}"
        )
        if st.button("PAKARTOTI TIK ŠIO ETAPO NEPAVYKUSIEMS" if lang=="LT" else "RETRY FAILED RECIPIENTS FOR THIS STAGE",use_container_width=True,key=f"retry_notifications_{y}_{m}"):
            res=retry_failed_lifecycle_notifications(y,m,chosen)
            if res: st.dataframe(localized_delivery_rows(res),use_container_width=True,hide_index=True)
            else: st.success("Nebėra ką kartoti." if lang=="LT" else "Nothing left to retry.")

def _manual_override_diff_rows(record,y,m):
    slots={s.idx:s for s in make_slots(y,m)}
    before=deserialize_result(record.get("before_json")) if record.get("before_json") else None
    after=deserialize_result(record.get("after_json")) if record.get("after_json") else None
    rows=[]
    for sid in (int(record.get("slot_a",0)),int(record.get("slot_b",0))):
        sl=slots.get(sid)
        if not sl: continue
        rows.append({
            "Data / postas" if lang=="LT" else "Date / post":f"{y}-{m:02d}-{sl.day:02d} · {sl.department} · {block_label(sl.block)}",
            "Prieš" if lang=="LT" else "Before":(before.assignments.get(sid) if before else "—"),
            "Po" if lang=="LT" else "After":(after.assignments.get(sid) if after else "—"),
        })
    return rows


def render_manual_override_review_checkpoint(y,m):
    """Persistent checkpoint: manual changes must be reviewed before PRELIMINARY or FINAL."""
    pending=db.list_unreviewed_manual_overrides_v2593(y,m)
    if not pending:
        return 0
    _workflow_card(
        "REIKIA PERŽIŪRĖTI RANKINIUS PAKEITIMUS" if lang=="LT" else "MANUAL CHANGES REQUIRE REVIEW",
        (f"Neperžiūrėtų pakeitimų: {len(pending)}. Kitas etapas ir FINAL užblokuoti, kol patvirtinsite pokyčių peržiūrą." if lang=="LT" else f"Unreviewed changes: {len(pending)}. The next phase and FINAL are blocked until the changes are reviewed."),
        "expired"
    )
    current_payload=db.load_schedule(y,m,"current")
    current_result=refresh_result_payload(current_payload,y,m,use_actual_backups=True) if current_payload else None
    hard=int(((current_result.stats or {}).get("global",{}) if current_result else {}).get("hard_errors",999))
    if hard==0:
        st.success("Dabartinis ACTUAL po korekcijų: HARD klaidų 0." if lang=="LT" else "Current ACTUAL after corrections: 0 HARD errors.")
    else:
        st.error((f"Dabartinis ACTUAL po korekcijų turi HARD klaidų: {hard}." if lang=="LT" else f"Current ACTUAL after corrections has HARD errors: {hard}."))
    for r in pending:
        with st.container(border=True):
            st.markdown(f"**#{r.get('id')} · {r.get('person_a')} ↔ {r.get('person_b')}**")
            st.caption(f"{r.get('created_at')} · {r.get('actor_initials')} · {r.get('reason')}")
            diff=_manual_override_diff_rows(r,y,m)
            if diff: st.dataframe(pd.DataFrame(diff),use_container_width=True,hide_index=True)
            ack=st.checkbox(
                "Peržiūrėjau pakeitimą ir dabartinį ACTUAL rezultatą." if lang=="LT" else "I reviewed this change and the current ACTUAL result.",
                key=f"review_override_ack_{r.get('id')}"
            )
            if st.button(
                "PATVIRTINTI POKYČIO PERŽIŪRĄ" if lang=="LT" else "CONFIRM CHANGE REVIEW",
                use_container_width=True,disabled=not ack,key=f"review_override_btn_{r.get('id')}"
            ):
                try:
                    db.review_manual_override_v2593(int(r["id"]))
                    st.rerun()
                except Exception as exc:
                    st.error(str(exc))
    return len(pending)


def swap_window_open_emails(y,m,result,deadline,window_days):
    settings=db.all_account_settings(); results=[]; public=config_value("SCHEDULER_PUBLIC_URL","").strip()
    local_deadline=deadline.astimezone(ZoneInfo("Europe/Vilnius")) if deadline.tzinfo else deadline.replace(tzinfo=ZoneInfo("Europe/Vilnius"))
    deadline_text=local_deadline.strftime("%Y-%m-%d %H:%M")
    for pp in DEFAULT_PEOPLE:
        i=pp["initials"]; email=str((settings.get(i,{}) or {}).get("email") or "").strip(); send_date=date.today().isoformat(); kind=f"preliminary_swap_window_{y}_{m:02d}"
        if not email:
            detail=tr("missing_email")
            try: db.record_email(i,kind,y,m,send_date,"failed",detail)
            except Exception: pass
            results.append((i,"failed",detail)); continue
        subject=(f"Preliminarus {month_label(y,m)} grafikas paskelbtas" if lang=="LT" else f"Preliminary {month_label(y,m)} schedule published")
        body=(
            f"Sveiki,\n\nPreliminarus {month_label(y,m)} grafikas paskelbtas. "
            f"Individualius apsikeitimo prašymus galite pateikti sistemoje iki {deadline_text} Lietuvos laiku.\n\n"
            f"Po termino naujų apsikeitimo prašymų teikti nebus galima, išskyrus individualų administratoriaus suteiktą leidimą. "
            f"Galutinė versija bus paskelbta atskiru pranešimu.\n"
            if lang=="LT" else
            f"Hello,\n\nThe preliminary {month_label(y,m)} schedule has been published. "
            f"Individual swap requests may be submitted in the system until {deadline_text} Lithuania time.\n\n"
            f"After the deadline, new swap requests are closed unless individual late access is granted by an operator. "
            f"The final version will be announced separately.\n"
        )
        if public: body += (f"\nPortalas: {public}\n" if lang=="LT" else f"\nPortal: {public}\n")
        ok,detail=deliver_lifecycle_notification(
            f"swap_open:{y}-{m:02d}","swap_open",i,y,m,email,subject,body,
            build_ics(y,m,result,i),f"PRELIMINARUS_{safe_filename(i)}_{y}_{m:02d}.ics"
        )
        try: db.record_email(i,kind,y,m,send_date,"sent" if ok else "failed",detail)
        except Exception: pass
        results.append((i,"sent" if ok else "failed",detail))
    return results


def late_swap_access_email(y,m,grant):
    initials=str(grant.get("initials") or ""); settings=db.get_account_settings(initials); email=str(settings.get("email") or "").strip()
    if not email: return False,tr("missing_email")
    exp=_parse_iso_dt(grant.get("expires_at")); exp_local=exp.astimezone(ZoneInfo("Europe/Vilnius")) if exp else None
    exp_text=exp_local.strftime("%Y-%m-%d %H:%M") if exp_local else str(grant.get("expires_at") or "")
    remaining=max(0,int(grant.get("max_requests",1) or 1)-int(grant.get("requests_used",0) or 0))
    subject=(f"Suteikta papildoma apsikeitimo prieiga — iki {exp_text}" if lang=="LT" else f"Additional swap access granted — until {exp_text}")
    body=(f"Sveiki,\n\nJums suteikta individuali papildoma apsikeitimo prieiga {month_label(y,m)} grafikui.\nGalioja iki: {exp_text}.\nGalite sukurti iki {remaining} naujo(-ų) apsikeitimo prašymo(-ų).\nPrisijunkite į Shift Happens → Apsikeitimai.\n" if lang=="LT" else f"Hello,\n\nYou have been granted individual additional swap access for the {month_label(y,m)} schedule.\nValid until: {exp_text}.\nYou may create up to {remaining} new swap request(s).\nOpen Shift Happens → Swaps.\n")
    public=config_value("SCHEDULER_PUBLIC_URL","").strip()
    if public: body += (f"\nPortalas: {public}\n" if lang=="LT" else f"\nPortal: {public}\n")
    ok,detail=send_email(email,subject,body)
    try: db.record_email(initials,f"late_swap_access_{grant.get('id','')}",y,m,date.today().isoformat(),"sent" if ok else "failed",detail)
    except Exception: pass
    return ok,detail


def final_schedule_emails(y,m,result):
    settings=db.all_account_settings(); results=[]; public=config_value("SCHEDULER_PUBLIC_URL","").strip()
    for pp in DEFAULT_PEOPLE:
        i=pp["initials"]; email=str((settings.get(i,{}) or {}).get("email") or "").strip(); kind=f"final_schedule_{y}_{m:02d}"; send_date=date.today().isoformat()
        if not email:
            results.append((i,"failed",tr("missing_email"))); continue
        subject=(f"Galutinis {month_label(y,m)} grafikas paskelbtas" if lang=="LT" else f"Final {month_label(y,m)} schedule published")
        body=(
            f"Sveiki,\n\nGalutinis {month_label(y,m)} tvarkaraštis paskelbtas ir pateiktas administracijai. "
            f"Įprasti ir pavėluoti apsikeitimai uždaryti.\n\nPrisegtas jūsų galutinis .ics grafikas.\n"
            if lang=="LT" else
            f"Hello,\n\nThe final {month_label(y,m)} schedule has been published and submitted to administration. "
            f"Ordinary and late swaps are now closed.\n\nYour final .ics schedule is attached.\n"
        )
        if public: body += (f"\nPortalas: {public}\n" if lang=="LT" else f"\nPortal: {public}\n")
        ok,detail=deliver_lifecycle_notification(
            f"final:{y}-{m:02d}","final",i,y,m,email,subject,body,
            build_ics(y,m,result,i),f"FINAL_{safe_filename(i)}_{y}_{m:02d}.ics"
        )
        try: db.record_email(i,kind,y,m,send_date,"sent" if ok else "failed",detail)
        except Exception: pass
        results.append((i,"sent" if ok else "failed",detail))
    return results


def preliminary_swap_window_for(y,m):
    """Default operational window for the target month: previous month 14th 08:00 → 17th 00:00."""
    if int(m)==1:
        py,pm=int(y)-1,12
    else:
        py,pm=int(y),int(m)-1
    tz=ZoneInfo("Europe/Vilnius")
    return (
        datetime(py,pm,14,8,0,tzinfo=tz),
        datetime(py,pm,17,0,0,tzinfo=tz),
    )


def manual_override_emails(y,m,result,initials_list):
    settings=db.all_account_settings(); results=[]; public=config_value("SCHEDULER_PUBLIC_URL","").strip()
    for i in dict.fromkeys(str(x) for x in initials_list if x):
        email=str((settings.get(i,{}) or {}).get("email") or "").strip(); kind=f"manual_schedule_override_{y}_{m:02d}"; send_date=date.today().isoformat()
        if not email:
            results.append((i,"failed",tr("missing_email"))); continue
        subject=(f"{month_label(y,m)} grafiko korekcija" if lang=="LT" else f"{month_label(y,m)} schedule correction")
        body=(
            f"Sveiki,\n\nTvarkaraščio administratorius atliko rankinę {month_label(y,m)} grafiko korekciją, kuri palietė jūsų pamainas. "
            f"Patikrinkite atnaujintą grafiką sistemoje. Prisegtas atnaujintas .ics failas.\n"
            if lang=="LT" else
            f"Hello,\n\nA schedule operator made a manual correction to the {month_label(y,m)} schedule affecting your shifts. "
            f"Please review the updated schedule in the system. An updated .ics file is attached.\n"
        )
        if public: body += (f"\nPortalas: {public}\n" if lang=="LT" else f"\nPortal: {public}\n")
        ok,detail=send_email(email,subject,body,build_ics(y,m,result,i),f"ATNAUJINTAS_{safe_filename(i)}_{y}_{m:02d}.ics")
        try: db.record_email(i,kind,y,m,send_date,"sent" if ok else "failed",detail)
        except Exception: pass
        results.append((i,"sent" if ok else "failed",detail))
    return results


def publish_system_baseline_for_swap_window(y,m):
    """Validate the draft, freeze SYSTEM/ACTUAL, backups and fairness history. No email is sent here."""
    draft_payload=db.load_schedule(y,m,"draft")
    if not draft_payload:
        return {"ok":False,"error":tr("no_draft")}
    credit_err=credit_selection_errors(y,m)
    if credit_err:
        return {"ok":False,"error":tr("bonus_insufficient"),"rows":credit_err}
    draft_result=refresh_result_payload(draft_payload,y,m,use_actual_backups=False)
    current_people=load_people(y,m); expected_targets=calculate_targets(y,m,current_people)
    current_snapshot=serialize_people_request_snapshot(current_people)
    if expected_targets!=draft_result.targets or (draft_result.request_snapshot and current_snapshot!=draft_result.request_snapshot):
        return {"ok":False,"error":tr("draft_outdated")}
    frozen_people=people_from_request_snapshot(draft_result.request_snapshot) or current_people
    # V2.5.115: validate the NORMAL schedule independently of theoretical backups.
    revalidated=validate_schedule(y,m,current_people,make_slots(y,m),draft_result.assignments,expected_targets,satisfaction_people=frozen_people,backup_assignments=[])
    if revalidated["global"].get("hard_errors",0):
        return {"ok":False,"error":tr("draft_outdated"),"rows":revalidated["global"].get("errors",[])}
    draft_result.targets=expected_targets
    desired,backup_errors=plan_backups(y,m,draft_result)
    if backup_errors:
        return {"ok":False,"error":tr("backup_capacity_block"),"rows":backup_errors}
    draft_result.backup_snapshot=[dict(x) for x in desired]
    draft_result=revalidate_loaded_result(y,m,current_people,draft_result,backup_assignments=[])
    _pgg=draft_result.stats.setdefault("global",{})
    _pgg["theoretical_backup_layer"]=True
    _pgg["theoretical_backup_layer_errors"]=[]
    _pgg["theoretical_backup_layer_complete"]=True
    if draft_result.stats.get("global",{}).get("hard_errors",0):
        return {"ok":False,"error":tr("draft_outdated"),"rows":draft_result.stats["global"].get("errors",[])}
    db.save_draft(y,m,serialize_result(draft_result))
    if not db.publish_draft(y,m):
        return {"ok":False,"error":tr("no_draft")}
    month_prefs=db.all_preferences(y,m)
    for person in DEFAULT_PEOPLE:
        i=person["initials"]; pp=month_prefs.get(i,{})
        db.set_rest_credit_redemptions(i,y,m,int(pp.get("backup_credits_am_to_use",0)),int(pp.get("backup_credits_pm_to_use",0)))
    db.sync_backups(y,m,desired)
    payload=db.load_schedule(y,m,"current"); result=refresh_result_payload(payload,y,m,use_actual_backups=True)
    db.save_current(y,m,serialize_result(result))
    baseline=deserialize_result(db.load_schedule(y,m,"baseline"))
    db.sync_fairness_history(y,m,baseline.stats["people"])
    feeds=refresh_calendar_subscription_feeds()
    return {"ok":True,"result":result,"feeds":feeds}


def render_operator_manual_override(y,m,current_payload,lifecycle_state):
    """Direct SP/ŠR pre-FINAL ACTUAL correction tool.

    Resident consent is not required, but ACTUAL safety/coverage HARD checks remain
    mandatory. SYSTEM is never rewritten by this tool.
    """
    st.markdown("### Seniūnės / administratoriaus rankinis koregavimas" if lang=="LT" else "### Operator manual correction")
    st.caption(
        "Galima naudoti bet kuriuo metu iki FINAL. Keičiama tik ACTUAL versija; užšaldytas SYSTEM lieka nepakeistas tyrimui. "
        "Rezidentų sutikimas šiam administraciniam veiksmui nereikalingas, tačiau saugos ir operacinės HARD taisyklės neapeinamos."
        if lang=="LT" else
        "Available at any time before FINAL. Only ACTUAL changes; the frozen SYSTEM baseline remains unchanged for research. "
        "Resident consent is not required for this administrative action, but safety and operational HARD rules cannot be bypassed."
    )
    if lifecycle_state=="final":
        st.info("FINAL versija užrakinta — rankinis koregavimas nebegalimas." if lang=="LT" else "The FINAL version is locked — manual correction is no longer available.")
        return
    if not current_payload:
        st.info(
            "Norint pradėti rankinį koregavimą, pirmiausia reikia užšaldyti sugeneruotą SYSTEM kaip pradinę ACTUAL versiją. El. laiškai šiame žingsnyje nesiunčiami."
            if lang=="LT" else
            "To begin manual correction, first freeze the generated SYSTEM as the initial ACTUAL version. No email is sent at this step."
        )
        return

    fresh=refresh_result_payload(current_payload,y,m)
    slots={s.idx:s for s in make_slots(y,m)}
    assigned=[sid for sid in fresh.assignments if sid in slots]
    assigned.sort(key=lambda sid:(slots[sid].day,{"AM":0,"FULL":1,"PM":2}.get(slots[sid].block,9),slots[sid].department,sid))
    if len(assigned)<2:
        st.warning("Nepakanka dviejų užpildytų pamainų apsikeitimui." if lang=="LT" else "At least two filled shifts are required.")
        return

    def slot_label(sid):
        sl=slots[int(sid)]; who=fresh.assignments.get(int(sid),"—")
        return f"{sl.day:02d} · {sl.department} · {block_label(sl.block)} · {who}"

    c1,c2=st.columns(2)
    with c1:
        sid_a=int(st.selectbox("Pirma pamaina" if lang=="LT" else "First shift",assigned,format_func=slot_label,key=f"op_manual_a_{y}_{m}"))
    with c2:
        choices_b=[x for x in assigned if int(x)!=sid_a]
        sid_b=int(st.selectbox("Antra pamaina" if lang=="LT" else "Second shift",choices_b,format_func=slot_label,key=f"op_manual_b_{y}_{m}"))
    person_a=str(fresh.assignments.get(sid_a) or "")
    person_b=str(fresh.assignments.get(sid_b) or "")
    reason=st.text_input("Koregavimo priežastis (privaloma auditui)" if lang=="LT" else "Correction reason (required for audit)",key=f"op_manual_reason_{y}_{m}")

    ok,msg,pstats,needed=preview_swap(
        y,m,people_for_stored_result(fresh,y,m),fresh,sid_a,sid_b,
        backup_assignments=db.list_backups(y,m)
    )
    if not ok:
        st.error(("Koregavimas negalimas: " if lang=="LT" else "Correction blocked: ")+str(msg))
        block_rows=((pstats or {}).get("global",{}) or {}).get("swap_hard_block_rows") or []
        if block_rows:
            st.dataframe(pd.DataFrame(block_rows),use_container_width=True,hide_index=True)
        return

    p1,p2=st.columns(2)
    with p1:
        st.markdown(f"**{person_a} → {slot_label(sid_b).rsplit(' · ',1)[0]}**")
    with p2:
        st.markdown(f"**{person_b} → {slot_label(sid_a).rsplit(' · ',1)[0]}**")
    warnings=((pstats or {}).get("global",{}) or {}).get("swap_warning_rows") or {}
    warning_rows=[]
    for who,rows in warnings.items():
        for row in (rows or []):
            rr=dict(row); rr.setdefault("resident",who); warning_rows.append(rr)
    if warning_rows:
        st.warning("Yra pasekmių įspėjimų. Administratorius gali tęsti tik aiškiai juos patvirtinęs." if lang=="LT" else "There are consequence warnings. The operator may continue only after explicit acknowledgement.")
        st.dataframe(pd.DataFrame(warning_rows),use_container_width=True,hide_index=True)

    ack=st.checkbox(
        "Patvirtinu šią rankinę korekciją ir jos parodytas pasekmes." if lang=="LT" else "I confirm this manual correction and the displayed consequences.",
        key=f"op_manual_ack_{y}_{m}_{sid_a}_{sid_b}"
    )
    affected_settings=db.all_account_settings()
    affected_missing=[who for who in (person_a,person_b) if not str((affected_settings.get(who,{}) or {}).get("email") or "").strip()]
    can_apply=bool(reason.strip() and ack and smtp_ready() and not affected_missing)
    if not smtp_ready():
        st.error("SMTP nesukonfigūruotas — paveikti rezidentai negalėtų gauti korekcijos pranešimo." if lang=="LT" else "SMTP is not configured — affected residents could not receive the correction notification.")
    if affected_missing:
        st.error(("Trūksta paveiktų rezidentų email: " if lang=="LT" else "Missing email for affected residents: ")+", ".join(affected_missing))
    if st.button("PRITAIKYTI RANKINĘ KOREKCIJĄ" if lang=="LT" else "APPLY MANUAL CORRECTION",type="primary",use_container_width=True,disabled=not can_apply,key=f"op_manual_apply_{y}_{m}_{sid_a}_{sid_b}"):
        try:
            apply_result=refresh_result_payload(db.load_schedule(y,m,"current"),y,m)
            if apply_result.assignments.get(sid_a)!=person_a or apply_result.assignments.get(sid_b)!=person_b:
                st.error("Grafikas pasikeitė po peržiūros. Pasirinkite pamainas iš naujo." if lang=="LT" else "The schedule changed after preview. Select the shifts again.")
                st.stop()
            ok2,msg2,_=attempt_swap(
                y,m,people_for_stored_result(apply_result,y,m),apply_result,sid_a,sid_b,
                backup_assignments=db.list_backups(y,m),acknowledged_fingerprints=needed
            )
            if not ok2:
                st.error(("Koregavimas nebetaikomas: " if lang=="LT" else "Correction no longer applies: ")+str(msg2)); st.stop()
            db.apply_manual_schedule_override_v2592(
                y,m,serialize_result(apply_result),sid_a,sid_b,person_a,person_b,reason.strip()
            )
            sync_backup_plan(y,m,apply_result)
            persist_actual_satisfaction(y,m)
            refresh_calendar_subscription_feeds([person_a,person_b])
            mails=manual_override_emails(y,m,apply_result,[person_a,person_b])
            failed=[x for x in mails if x[1]!="sent"]
            st.session_state["_finalization_flash"]=(
                "warning" if failed else "success",
                (f"Rankinė korekcija pritaikyta. Pranešimai: {len(mails)-len(failed)}/{len(mails)} išsiųsta."
                 if lang=="LT" else
                 f"Manual correction applied. Notifications: {len(mails)-len(failed)}/{len(mails)} sent.")
            )
            st.rerun()
        except Exception as exc:
            st.error(str(exc))

    history=db.list_manual_schedule_overrides_v2592(y,m)
    if history:
        with st.expander("Rankinių korekcijų auditas" if lang=="LT" else "Manual correction audit",expanded=False):
            rows=[]
            for r in history[:20]:
                rows.append({
                    ("Laikas" if lang=="LT" else "Time"):r.get("created_at"),
                    ("Operatorius" if lang=="LT" else "Operator"):r.get("actor_initials"),
                    ("Pakeitimas" if lang=="LT" else "Change"):f"{r.get('person_a')} ↔ {r.get('person_b')}",
                    ("Slotai" if lang=="LT" else "Slots"):f"#{r.get('slot_a')} ↔ #{r.get('slot_b')}",
                    ("Priežastis" if lang=="LT" else "Reason"):r.get("reason"),
                })
            st.dataframe(pd.DataFrame(rows),use_container_width=True,hide_index=True)


def _lt_days_left_phrase(n: int) -> str:
    n=int(n)
    if n==1:
        return "1 diena"
    if 2<=n<=9:
        return f"{n} dienos"
    return f"{n} dienų"


def send_due_reminders(y,m):
    """Manual fallback for today's preference reminder queue.

    The background worker performs the same check automatically. One consolidated
    reminder replaces the old duplicate preference + backup-claim reminder pair.
    """
    dl=deadline_for(y,m); today=date.today(); prefs=db.all_preferences(y,m); settings=db.all_account_settings(); results=[]
    for p in DEFAULT_PEOPLE:
        i=p["initials"]; s=settings.get(i,{})
        if i in prefs or not s.get("notifications_on",True):
            continue
        start_day=max(1,min(deadline_day(),int(s.get("reminder_start_day",8) or 8)))
        start=date(dl.year,dl.month,start_day)
        if not (start<=today<=dl):
            continue
        email=str(s.get("email") or "").strip()
        left=max(0,(dl-today).days)
        cutoff_text=preference_cutoff_for(y,m).strftime("%Y-%m-%d %H:%M")
        subject=(
            f"Šiandien paskutinė diena pateikti {month_label(y,m)} pageidavimus"
            if lang=="LT" and left==0 else
            f"Liko {_lt_days_left_phrase(left)} iki {month_label(y,m)} pageidavimų pateikimo pabaigos"
            if lang=="LT" else
            f"Today is the last day to submit {month_label(y,m)} preferences"
            if left==0 else
            f"{left} day(s) left to submit {month_label(y,m)} preferences"
        )
        body=(
            f"Sveiki,\n\nJūsų {month_label(y,m)} pageidavimai dar nepateikti. Tikslus terminas: {cutoff_text} Lietuvos laiku.\n"
            if lang=="LT" else
            f"Hello,\n\nYour {month_label(y,m)} preferences have not been submitted. Exact deadline: {cutoff_text} Lithuania time.\n"
        )
        public=config_value("SCHEDULER_PUBLIC_URL","").strip()
        if public:
            body += (f"\nShift Happens: {public}\n")
        ok,detail=deliver_lifecycle_notification(
            f"preferences_reminder:{y}-{m:02d}:{today.isoformat()}",
            "preferences_reminder",i,y,m,email,subject,body
        )
        try:
            db.record_email(i,"reminder",y,m,today.isoformat(),"sent" if ok else "failed",detail)
        except Exception:
            pass
        results.append((i,"sent" if ok else "failed",detail))
    return results


def balance_ratio(a,b):
    if a is None or b is None: return None
    a=float(a); b=float(b)
    if max(a,b)==0: return 1.0
    return round(min(a,b)/max(a,b),2)

def _fairness_from_spreads(w,f,d,wd):
    return max(0.0,100.0-18.0*w-7.0*f-4.0*d-2.0*wd)


def fairness_trend_df(rows):
    """Build monthly and cumulative fairness history from finalized monthly ledger rows."""
    if not rows:
        return pd.DataFrame()
    periods=sorted({(int(r["year"]),int(r["month"])) for r in rows})
    initials=sorted({r["initials"] for r in rows})
    cumulative={i:{"weekend":0,"friday":0,"double":0,"weekday_day":0} for i in initials}
    out=[]
    for y,m in periods:
        month_rows=[r for r in rows if int(r["year"])==y and int(r["month"])==m]
        by={r["initials"]:r for r in month_rows}
        monthly_vals={k:[] for k in ("weekend","friday","double","weekday_day")}
        for i in initials:
            r=by.get(i,{})
            vals={
                "weekend":int(r.get("weekend_assignments",0)),
                "friday":int(r.get("friday_assignments",0)),
                "double":int(r.get("doubles",0)),
                "weekday_day":int(r.get("weekday_days",0)),
            }
            for k,v in vals.items():
                monthly_vals[k].append(v)
                cumulative[i][k]+=v
        def spread(vals): return max(vals)-min(vals) if vals else 0
        mw,mf,md,mwd=[spread(monthly_vals[k]) for k in ("weekend","friday","double","weekday_day")]
        cw=spread([cumulative[i]["weekend"] for i in initials])
        cf=spread([cumulative[i]["friday"] for i in initials])
        cd=spread([cumulative[i]["double"] for i in initials])
        cwd=spread([cumulative[i]["weekday_day"] for i in initials])
        out.append({
            "Period":f"{y}-{m:02d}",
            "Monthly fairness":round(_fairness_from_spreads(mw,mf,md,mwd),1),
            "Cumulative fairness":round(_fairness_from_spreads(cw,cf,cd,cwd),1),
        })
    return pd.DataFrame(out)


def live_fairness_snapshot(y,m,result,include_completed_covers=True):
    """Live monthly real-work fairness; never used as future solver input."""
    if result is None:
        return {"people":{},"effective_assignments":{},"global":{}}
    backups=[]
    if include_completed_covers:
        try:
            backups=db.list_backups(y,m)
        except Exception:
            backups=[]
    return calculate_live_fairness_snapshot(
        y,m,result.assignments,
        people_initials=[p["initials"] for p in DEFAULT_PEOPLE],
        backup_assignments=backups,
    )


def system_actual_fairness_trend_df(up_to_year=None,up_to_month=None):
    """Monthly SYSTEM vs ACTUAL fairness history; descriptive only."""
    try:
        rows=db.list_published_schedules()
    except Exception:
        return pd.DataFrame()
    limit=(int(up_to_year)*12+int(up_to_month)) if up_to_year is not None and up_to_month is not None else None
    out=[]
    ids=[p["initials"] for p in DEFAULT_PEOPLE]
    for r in rows:
        y=int(r.get("year",0)); m=int(r.get("month",0))
        if limit is not None and y*12+m>limit:
            continue
        try:
            cp=r.get("current_json") or {}
            bp=r.get("baseline_json") or cp
            cr=deserialize_result(cp); br=deserialize_result(bp)
            backups=db.list_backups(y,m)
            sg=calculate_live_fairness_snapshot(y,m,br.assignments,people_initials=ids,backup_assignments=[])["global"]
            ag=calculate_live_fairness_snapshot(y,m,cr.assignments,people_initials=ids,backup_assignments=backups)["global"]
            out.append({
                "Period":f"{y}-{m:02d}",
                "SYSTEM monthly fairness":float(sg.get("monthly_fairness_score",0.0)),
                "ACTUAL monthly fairness":float(ag.get("monthly_fairness_score",0.0)),
                "Delta (ACTUAL-SYSTEM)":round(float(ag.get("monthly_fairness_score",0.0))-float(sg.get("monthly_fairness_score",0.0)),1),
                "SYSTEM post imbalance":int(sg.get("rotation_monthly_imbalance",0)),
                "ACTUAL post imbalance":int(ag.get("rotation_monthly_imbalance",0)),
                "Completed covers":int(ag.get("completed_cover_transfers",0)),
            })
        except Exception:
            continue
    return pd.DataFrame(out)


def actual_workplace_exposure_df(y,m,result):
    snap=live_fairness_snapshot(y,m,result,include_completed_covers=True)
    view=deepcopy(result)
    view.assignments=dict(snap.get("effective_assignments") or result.assignments)
    return workplace_exposure_df(y,m,view)


def component_status(score):
    if score is None: return tr("not_applicable")
    if score>=80: return tr("matches")
    if score>=50: return tr("partial")
    return tr("mismatch")

def localized_delivery_rows(rows):
    smap={"sent":tr("sent"),"failed":tr("failed"),"skipped":tr("skipped")}
    return pd.DataFrame([(i,smap.get(status,status),detail) for i,status,detail in rows], columns=[tr("person"),tr("status"),tr("details")])

def localized_email_log(rows):
    smap={"sent":tr("sent"),"failed":tr("failed"),"skipped":tr("skipped")}
    kmap={"reminder":tr("reminder_kind"),"publication":tr("publication_kind"),"backup_claim_reminder":tr("backup_claim_reminder_kind")}
    return pd.DataFrame([{
        "ID":r["id"], tr("person"):r["initials"], tr("details"):kmap.get(r["kind"],r["kind"]),
        tr("date"):r["send_date"], tr("status"):smap.get(r["status"],r["status"]),
        tr("comment"):r["detail"], tr("updated"):r["sent_at"]
    } for r in rows])

# --- Identity, role and month ---
st.sidebar.title("Shift Happens"); st.sidebar.caption("PGY-1 Radiology")

def observer_assignment_changes_df(y,m,baseline,current):
    if baseline is None or current is None:
        return pd.DataFrame()
    slots={s.idx:s for s in make_slots(y,m)}
    rows=[]
    all_ids=sorted(set(baseline.assignments)|set(current.assignments))
    for sid in all_ids:
        before=baseline.assignments.get(sid,"")
        after=current.assignments.get(sid,"")
        if before==after:
            continue
        s=slots.get(sid)
        if s is None:
            continue
        rows.append({
            tr("date"):f"{y}-{m:02d}-{s.day:02d}",
            tr("department"):s.department,
            tr("shift"):block_label(s.block),
            tr("observer_from"):before or "—",
            tr("observer_to"):after or "—",
        })
    return pd.DataFrame(rows)


def observer_swap_df(rows, backup=False):
    out=[]
    for r in rows:
        if backup:
            pair=f"{r.get('requester','')} ↔ {r.get('target','')}"
            slots=f"#{r.get('requester_slot')} ↔ #{r.get('target_slot')}"
        else:
            pair=f"{r.get('person_a','')} ↔ {r.get('person_b','')}"
            slots=f"#{r.get('slot_a')} ↔ #{r.get('slot_b')}"
        out.append({
            "ID":r.get("id"),
            tr("person"):pair,
            tr("shift"):slots,
            tr("status"):r.get("status",""),
            tr("updated"):r.get("responded_at") or r.get("created_at") or "",
        })
    return pd.DataFrame(out)


def observer_backup_df(y,m,current):
    slots={s.idx:s for s in make_slots(y,m)}
    rows=[]
    for r in db.list_backups(y,m):
        s=slots.get(int(r.get("covered_slot",0)))
        if s is None:
            continue
        rows.append({
            tr("date"):f"{y}-{m:02d}-{s.day:02d}",
            tr("shift"):block_label(s.block),
            tr("covered_person"):(current.assignments.get(s.idx,"") if current else r.get("covered_person","")),
            tr("observer_planned_backup"):r.get("planned_backup") or "—",
            tr("observer_actual_backup"):r.get("actual_backup") or "—",
            tr("observer_activated"):tr("yes") if r.get("activated_at") else tr("no"),
            tr("observer_completed"):tr("yes") if r.get("completed_at") else tr("no"),
        })
    return pd.DataFrame(rows)


def render_observer_portal(profile,auth_user):
    st.sidebar.markdown(
        f'<div style="padding:10px 12px;border:1px solid #d1d5db;border-radius:10px;'
        f'font-weight:700;">{html.escape(tr("observer_read_only"))} · '
        f'{html.escape(tr("observer_role"))}</div>',
        unsafe_allow_html=True
    )
    st.sidebar.caption(getattr(auth_user,"email",profile.get("email","")))
    if st.sidebar.button(tr("logout"),use_container_width=True,key="observer_logout"):
        try: st.session_state["supabase_client"].auth.sign_out()
        except Exception: pass
        st.session_state.pop("supabase_client",None)
        st.rerun()

    default_y,default_m=next_month(date.today())
    y=int(st.sidebar.number_input(tr("year"),2026,2100,default_y,1,key="observer_year"))
    m=int(st.sidebar.selectbox(
        tr("month"),list(range(1,13)),index=default_m-1,
        format_func=lambda x:MONTHS[lang][x-1],key="observer_month"
    ))

    st.title(tr("observer_portal"))
    st.warning(f"{tr('observer_read_only')} — {tr('observer_scope_note')}")
    st.caption(tr("observer_privacy_note"))

    tab_overview,tab_schedule,tab_changes,tab_fairness,tab_backups,tab_research,tab_rules=st.tabs([
        tr("observer_overview"),tr("observer_schedule"),tr("observer_changes"),
        tr("observer_fairness"),tr("observer_backups"),tr("research_observer_tab"),tr("observer_rules")
    ])

    current_payload=db.load_schedule(y,m,"current")
    baseline_payload=db.load_schedule(y,m,"baseline")
    current=refresh_result_payload(current_payload,y,m) if current_payload else None
    baseline=refresh_result_payload(baseline_payload or current_payload,y,m,use_actual_backups=False) if current_payload else None
    normal_swaps=db.list_swap_requests(y,m,None)
    backup_swaps=db.list_backup_swap_requests(y,m,None)
    changes=observer_assignment_changes_df(y,m,baseline,current) if current else pd.DataFrame()

    with tab_overview:
        if not current:
            st.info(tr("observer_no_schedule"))
        else:
            g=baseline.stats.get("global",{})
            sl=live_fairness_snapshot(y,m,baseline,include_completed_covers=False)["global"]
            al=live_fairness_snapshot(y,m,current,include_completed_covers=True)["global"]
            approved=sum(1 for r in normal_swaps+backup_swaps if r.get("status")=="approved")
            pending=sum(1 for r in normal_swaps+backup_swaps if r.get("status")=="pending")
            c1,c2,c3,c4,c5=st.columns(5)
            c1.metric(tr("hard_errors"),g.get("hard_errors",0))
            c2.metric("SYSTEM fairness",f"{sl.get('monthly_fairness_score',0)}%")
            c3.metric("ACTUAL fairness",f"{al.get('monthly_fairness_score',0)}%",delta=f"{al.get('monthly_fairness_score',0)-sl.get('monthly_fairness_score',0):+.1f}")
            c4.metric(tr("observer_change_count"),len(changes))
            c5.metric(tr("observer_pending_swaps"),pending)
            st.caption(tr("observer_change_log_help"))
            if changes.empty:
                st.success(tr("observer_no_changes"))
            else:
                st.dataframe(changes,use_container_width=True,hide_index=True)
            st.markdown(f"### {tr('fairness_hierarchy')}")
            st.dataframe(pd.DataFrame([
                {tr("fairness_level"):"1. ABSOLUTE HARD",tr("fairness_goal"):tr("fairness_hard_goal")},
                {tr("fairness_level"):"2. RESIDENT HARD",tr("fairness_goal"):("0 pažeidimų privaloma; fairness optimizuojamas tik likusioje validžioje erdvėje" if lang=="LT" else "Zero violations required; fairness is optimized only inside the valid zero-loss space")},
                {tr("fairness_level"):"3. POSTAI / WORKLOAD",tr("fairness_goal"):tr("fairness_monthly_goal")},
                {tr("fairness_level"):"4. MAX-MIN SOFT",tr("fairness_goal"):tr("other_preferences_goal")},
                {tr("fairness_level"):"5. ACTUAL AUDIT",tr("fairness_goal"):tr("fairness_cumulative_goal")},
            ]),use_container_width=True,hide_index=True)

    with tab_schedule:
        if not current:
            st.info(tr("observer_no_schedule"))
        else:
            st.markdown(f"### {tr('observer_baseline_schedule')}")
            st.caption(tr("fairness_swap_neutral"))
            st.dataframe(style_schedule(schedule_grid(y,m,baseline)),use_container_width=True,height=560)
            st.divider()
            st.markdown(f"### {tr('observer_actual_schedule')}")
            st.caption(tr("observer_change_log_help"))
            st.dataframe(style_schedule(schedule_grid(y,m,current)),use_container_width=True,height=560)

    with tab_changes:
        if not current:
            st.info(tr("observer_no_schedule"))
        else:
            st.markdown(f"### {tr('observer_normal_swaps')}")
            if normal_swaps:
                st.dataframe(observer_swap_df(normal_swaps),use_container_width=True,hide_index=True)
            else:
                st.caption("—")
            st.markdown(f"### {tr('observer_backup_swaps')}")
            if backup_swaps:
                st.dataframe(observer_swap_df(backup_swaps,backup=True),use_container_width=True,hide_index=True)
            else:
                st.caption("—")
            st.markdown(f"### {tr('observer_change_count')}")
            if changes.empty:
                st.success(tr("observer_no_changes"))
            else:
                st.dataframe(changes,use_container_width=True,hide_index=True)

    with tab_fairness:
        if not current:
            st.info(tr("observer_no_schedule"))
        else:
            g=baseline.stats.get("global",{})
            sl=live_fairness_snapshot(y,m,baseline,include_completed_covers=False)["global"]
            al=live_fairness_snapshot(y,m,current,include_completed_covers=True)["global"]
            c1,c2,c3=st.columns(3)
            c1.metric(tr("hard_validity"),tr("hard_validity_pass") if g.get("hard_errors",0)==0 else tr("hard_validity_fail"))
            c2.metric("SYSTEM fairness",f"{sl.get('monthly_fairness_score',0)}%")
            c3.metric("ACTUAL fairness",f"{al.get('monthly_fairness_score',0)}%",delta=f"{al.get('monthly_fairness_score',0)-sl.get('monthly_fairness_score',0):+.1f}")
            st.caption(tr("fairness_swap_neutral"))
            breakdown=[
                ("SYSTEM",tr("metric_saturday"),sl.get("saturday_monthly_spread",0)),
                ("ACTUAL",tr("metric_saturday"),al.get("saturday_monthly_spread",0)),
                ("SYSTEM",tr("metric_sunday"),sl.get("sunday_monthly_spread",0)),
                ("ACTUAL",tr("metric_sunday"),al.get("sunday_monthly_spread",0)),
                ("SYSTEM",tr("metric_friday"),sl.get("friday_monthly_spread",0)),
                ("ACTUAL",tr("metric_friday"),al.get("friday_monthly_spread",0)),
                ("SYSTEM",tr("metric_double"),sl.get("double_monthly_spread",0)),
                ("ACTUAL",tr("metric_double"),al.get("double_monthly_spread",0)),
                ("SYSTEM",tr("metric_weekday"),sl.get("weekday_day_monthly_spread",0)),
                ("ACTUAL",tr("metric_weekday"),al.get("weekday_day_monthly_spread",0)),
            ]
            st.dataframe(pd.DataFrame([{tr("fairness_scope"):scope,tr("fairness_metric"):metric,tr("fairness_spread"):spread} for scope,metric,spread in breakdown]),use_container_width=True,hide_index=True)
            trend=system_actual_fairness_trend_df(y,m)
            if not trend.empty:
                chart=trend.set_index("Period")
                st.line_chart(chart[["SYSTEM monthly fairness","ACTUAL monthly fairness"]],height=300)
                st.dataframe(trend,use_container_width=True,hide_index=True)
                st.caption("Tik auditas — jokio future catch-up." if lang=="LT" else "Audit only — no future catch-up.")
            else:
                st.caption(tr("fairness_no_history"))

    with tab_backups:
        if not current:
            st.info(tr("observer_no_schedule"))
        else:
            bdf=observer_backup_df(y,m,current)
            if bdf.empty:
                st.caption("—")
            else:
                st.dataframe(bdf,use_container_width=True,hide_index=True)

    with tab_research:
        st.info(tr("research_observer_intro"))
        cp_options=[c for c,_,_ in OBSERVER_RESEARCH_CHECKPOINTS]
        cp=st.selectbox(tr("research_observer_checkpoint"),cp_options,format_func=research_checkpoint_label,key="observer_research_cp")
        yy,mm=next((yy,mm) for c,yy,mm in OBSERVER_RESEARCH_CHECKPOINTS if c==cp)
        existing=db.get_my_observer_research_checkpoint(yy,mm,cp) or {}
        olda=existing.get("answers") or {}; oldf=existing.get("free_text") or {}
        obs_items={
            "actual":tr("research_obs_actual"),"changes":tr("research_obs_changes"),
            "system_actual":tr("research_obs_system_actual"),"privacy":tr("research_obs_privacy"),
            "change_log":tr("research_obs_log"),"fairness":tr("research_obs_fairness"),"trust":tr("research_obs_trust")
        }
        with st.form(f"observer_research_{cp}"):
            ans={}
            for key,label in obs_items.items():
                ans[key]=st.slider(label,1,5,int(olda.get(key,3) or 3),1,key=f"obs_{cp}_{key}")
            missing=st.text_area(tr("research_obs_missing"),value=oldf.get("missing","") or "")
            if st.form_submit_button(tr("research_submit"),type="primary"):
                db.submit_observer_research_checkpoint(yy,mm,cp,ans,{"missing":missing})
                st.success(tr("research_observer_saved"))

    with tab_rules:
        st.markdown(db.get_manual(lang))


sb,auth_user=render_auth_gate()
profile=require_linked_profile(sb,auth_user)
if profile.get("access_role")=="observer":
    render_observer_portal(profile,auth_user)
    st.stop()

active_user=profile["initials"]
resident_ok=True

# V2.5.34: install the ACTIVE versioned Rule Profile before any month calculations,
# scheduling, backup planning or rule-dependent UI is rendered.
try:
    _active_rule_row=db.get_active_rule_profile()
    if _active_rule_row:
        st.session_state["_last_good_rule_profile"]=dict(_active_rule_row)
except Exception as exc:
    _active_rule_row=st.session_state.get("_last_good_rule_profile")
    if not _active_rule_row:
        st.error(
            "Laikinas ryšio su duomenų baze sutrikimas. Aktyvių taisyklių nepavyko saugiai perskaityti. "
            "Jokio swapo / grafiko pakeitimo neįvykdžiau. Atnaujinkite puslapį po kelių sekundžių."
            if lang=="LT" else
            "Temporary database connection problem. The active rule profile could not be read safely. "
            "No swap/schedule change was performed. Refresh the page in a few seconds."
        )
        st.caption(f"{exc.__class__.__name__}: {exc}")
        st.stop()

if (_active_rule_row or {}).get("_read_fallback")=="memory_cache":
    st.warning(
        "Trumpam nutrūko DB ryšys — naudojama paskutinė šiame procese sėkmingai perskaityta aktyvi taisyklių versija."
        if lang=="LT" else
        "Database connection briefly dropped — using the last successfully read active rule profile from this process."
    )

_active_rule_config=(_active_rule_row or {}).get("config") or DEFAULT_RULE_PROFILE
try:
    ACTIVE_RULES=set_runtime_rules(_active_rule_config)
    ACTIVE_RULE_PROFILE_VERSION=int((_active_rule_row or {}).get("version_no") or 1)
except Exception:
    # Safety fallback: invalid DB config can never silently corrupt scheduling.
    ACTIVE_RULES=set_runtime_rules(DEFAULT_RULE_PROFILE)
    ACTIVE_RULE_PROFILE_VERSION=0

try:
    directory_map=db.directory()
    if directory_map:
        st.session_state["_last_good_directory"]=directory_map
except Exception as exc:
    directory_map=st.session_state.get("_last_good_directory")
    if not directory_map:
        st.error(
            "Laikinas DB ryšio sutrikimas. Rezidentų sąrašo nepavyko perskaityti; atnaujinkite puslapį."
            if lang=="LT" else
            "Temporary database connection problem. The resident directory could not be read; refresh the page."
        )
        st.caption(f"{exc.__class__.__name__}: {exc}")
        st.stop()
people_map={p["initials"]:p for p in DEFAULT_PEOPLE}
st.sidebar.markdown(badge(active_user),unsafe_allow_html=True)
research_role = "researcher" if active_user==RESEARCHER_INITIALS else "senior" if active_user==SENIOR_INITIALS else "resident"
role_key={"researcher":"research_role_researcher","senior":"research_role_senior","resident":"research_role_resident"}[research_role]
st.sidebar.caption(tr(role_key))
st.sidebar.caption(getattr(auth_user,"email",profile.get("email","")))
if st.sidebar.button(tr("logout"),use_container_width=True):
    try: sb.auth.sign_out()
    except Exception: pass
    clear_cross_account_session_state(keep_client=False)
    st.rerun()

# V2.5.90: one visible interface per account. There is no profile switch.
# SP = operational Seniūnė.
# ŠR = resident-facing account with embedded researcher + senior/admin capabilities.
# MG and all others = resident-only.
is_seniune_account=(active_user==SENIOR_INITIALS)
is_researcher_account=(active_user==RESEARCHER_INITIALS)
has_senior_functions=(is_seniune_account or is_researcher_account)
admin_ok=has_senior_functions
# Legacy name retained as a capability flag for existing protected senior-only blocks.
senior_mode=has_senior_functions

ui_simple=("Paprastas" if lang=="LT" else "Simple")
ui_advanced=("Išplėstinis" if lang=="LT" else "Advanced")
ui_mode=st.sidebar.radio(
    ("Sąsajos režimas" if lang=="LT" else "Interface mode"),
    [ui_simple,ui_advanced],
    index=0,
    key="ui_mode_v2530",
    help=("Paprastas: kasdieniai veiksmai ir tik svarbiausi rezultatai. Išplėstinis: pilna fairness, guardrail ir solverio diagnostika."
          if lang=="LT" else
          "Simple: daily actions and only the most important results. Advanced: full fairness, guardrail and solver diagnostics.")
)
advanced_mode=(ui_mode==ui_advanced)
# V2.5.92 lifecycle controls: SP primary; ŠR contingency only in Išplėstinis.
lifecycle_operator_ui=(is_seniune_account or (is_researcher_account and advanced_mode))
st.sidebar.caption(
    ("Paprastas režimas yra numatytasis." if not advanced_mode and lang=="LT" else
     "Simple mode is the default." if not advanced_mode else
     "Rodoma pilna techninė informacija." if lang=="LT" else
     "Full technical information is visible.")
)

default_y,default_m=next_month(date.today()); year=int(st.sidebar.number_input(tr("year"),2026,2100,default_y,1)); month=int(st.sidebar.selectbox(tr("month"),list(range(1,13)),index=default_m-1,format_func=lambda x:MONTHS[lang][x-1]))
wd=weekday_count(year,month); bt=standard_target(year,month)
if advanced_mode:
    st.sidebar.metric(tr("weekdays"),wd)
    st.sidebar.metric(tr("base_target"),bt)
    st.sidebar.caption(
        f"{wd} × {float(rule_value('target_daily_hours')):g} / "
        f"{float(rule_value('target_shift_hours')):g} → {bt}"
    )
    st.sidebar.caption(
        ("Aktyvus taisyklių profilis" if lang=="LT" else "Active Rule Profile")
        + f": v{ACTIVE_RULE_PROFILE_VERSION}"
    )

st.title(tr("app_title"))
if advanced_mode:
    st.caption(tr("app_caption"))
    st.info(
        ("IŠPLĖSTINIS REŽIMAS — rodoma pilna fairness, guardrail, solverio ir tyrimo diagnostika."
         if lang=="LT" else
         "ADVANCED MODE — full fairness, guardrail, solver and research diagnostics are visible.")
    )
else:
    st.caption("Paprastas režimas" if lang=="LT" else "Simple mode")

# V2.5.94: after the exact cutoff, SP/ŠR views automatically materialize
# every still-missing active resident as a submitted zero-request form.
_zero_pref_autosubmit={"ok":True,"due":False,"count":0,"initials":[]}
if lifecycle_operator_ui:
    _zero_pref_autosubmit=ensure_zero_preference_submissions_if_due(year,month)
    if int(_zero_pref_autosubmit.get("count",0) or 0)>0:
        names=", ".join(_zero_pref_autosubmit.get("initials") or [])
        st.info(
            (f"Po pageidavimų termino automatiškai užfiksuotos 0 pageidavimų anketos: {names}."
             if lang=="LT" else
             f"After the preference deadline, zero-request submissions were recorded automatically for: {names}.")
        )

if advanced_mode:
    with st.expander(("IŠPLĖSTINIS LANGAS" if lang=="LT" else "ADVANCED WINDOW"), expanded=True):
        adv_state=db.get_schedule_state(year,month)
        adv_payload=db.load_schedule(year,month,"baseline") or db.load_schedule(year,month,"draft")
        a1,a2,a3,a4=st.columns(4)
        a1.metric(("Būsena" if lang=="LT" else "State"),
                  ("Paskelbtas" if adv_state.get("has_published") else "Juodraštis" if adv_state.get("has_draft") else "Nesukurtas"))
        if adv_payload:
            adv_res=refresh_result_payload(adv_payload,year,month,use_actual_backups=False)
            adv_g=(adv_res.stats or {}).get("global",{})
            a2.metric("HARD",adv_g.get("hard_errors","—"))
            a3.metric(("Mėnesio fairness" if lang=="LT" else "Monthly fairness"),f"{adv_g.get('monthly_fairness_score',adv_g.get('fairness_score','—'))}%")
            a4.metric(("Solverio etapas" if lang=="LT" else "Solver stage"),adv_g.get("solve_stage","—"))
            st.caption(
                (f"Guardrails: {len(adv_g.get('fairness_guardrails') or {})} · Preference pre-check: {adv_g.get('preference_normalization_count',0)}"
                 if lang=="LT" else
                 f"Guardrails: {len(adv_g.get('fairness_guardrails') or {})} · Preference pre-check: {adv_g.get('preference_normalization_count',0)}")
            )
        else:
            a2.metric("HARD","—"); a3.metric(("Mėnesio fairness" if lang=="LT" else "Monthly fairness"),"—"); a4.metric(("Solverio etapas" if lang=="LT" else "Solver stage"),"—")

names=[]
if active_user==RESEARCHER_INITIALS and advanced_mode:
    names.append("AVAILABLE GPT + HUMAN vs MY ENGINE")
if senior_mode:
    names.append(tr("senior_dashboard"))
names += [tr("preferences"),tr("settings")]
if senior_mode:
    names.append(tr("generation"))
names.append(tr("schedule"))
if advanced_mode:
    names += [tr("summary"),tr("transparency"),tr("credits_debts")]
names += [tr("backups"),tr("swaps"),tr("calendar"),tr("research")]
if advanced_mode:
    names.append(tr("proof"))
if senior_mode:
    names.append(tr("senior_guide"))
names.append(tr("rules"))

# ŠR keeps the isolated research-shadow generator, but it is a tab inside the
# same single account interface rather than a second profile/window.
research_shadow_label=("TYRĖJO SUDARYMAS" if lang=="LT" else "RESEARCHER WORKBENCH")
if active_user==RESEARCHER_INITIALS and advanced_mode:
    names.append(research_shadow_label)

tabs=st.tabs(names)
research_shadow_tab_index=(
    names.index(research_shadow_label)
    if active_user==RESEARCHER_INITIALS and advanced_mode and research_shadow_label in names
    else None
)

# ŠR research comparison is an advanced analytical tool, so it does not occupy
# mobile/simple navigation.
pos=1 if (active_user==RESEARCHER_INITIALS and advanced_mode) else 0
if active_user==RESEARCHER_INITIALS and advanced_mode:
    st.sidebar.success("Research: FIRST TAB — AVAILABLE GPT + HUMAN vs MY ENGINE")
elif active_user==RESEARCHER_INITIALS and not advanced_mode:
    st.sidebar.caption("Tyrėjo įrankiai → Išplėstinis režimas" if lang=="LT" else "Researcher tools → Advanced mode")
if st.session_state.get("_save_flash"):
    st.success("✓ " + str(st.session_state.pop("_save_flash")))


def flash_saved(message):
    """Store a one-shot success message and rerun safely."""
    st.session_state["_save_flash"] = str(message)
    st.rerun()


def render_recurring_preferences_editor(initials: str):
    """Persistent recurring preferences live in the Preferences tab, not Settings."""
    st.divider()
    st.markdown(f"### {tr('long_term')}")
    st.caption(tr("long_term_help"))
    st.caption(
        "Ši dalis nėra pririšta prie vieno mėnesio: taisyklės automatiškai persikelia į visus būsimus dar neužšaldytus grafikus, kol jas pakeisite arba išjungsite. Jau paskelbto SYSTEM grafiko jos retroaktyviai nekeičia."
        if lang=="LT" else
        "This section is not tied to one month: the rules automatically carry into every future schedule that is not yet frozen until you change or disable them. They never rewrite an already published SYSTEM schedule."
    )
    st.caption(
        "Savaitgalio „Pageidauju dirbti“ leidžiamas kaip savanoriškas nepopuliarios pamainos pasirinkimas. Generuojant SYSTEM, šeštadienio ir sekmadienio water-fill išlieka struktūriškai lygus; po publikavimo abipusiai savanoriški swapai gali pakeisti ACTUAL pasiskirstymą ir jo spread."
        if lang=="LT" else
        "Weekend 'prefer to work' is allowed as volunteering for unpopular duty. During SYSTEM generation, Saturday/Sunday water-fill remains structurally equal; after publication, bilateral voluntary swaps may change the ACTUAL distribution and its spread."
    )
    existing_rec={int(r["weekday"]):r for r in db.get_recurring_preferences(initials)}
    rule_to_label={"hard_unavailable":tr("rec_hard"),"soft_free":tr("rec_soft"),"preferred":tr("rec_preferred"),"none":tr("rec_none")}
    label_to_rule={v:k for k,v in rule_to_label.items()}
    block_to_label={"FULL":tr("full_day"),"AM":tr("morning"),"PM":tr("afternoon")}
    label_to_block={v:k for k,v in block_to_label.items()}
    rec_rows=[]
    for wd_i in range(7):
        rr=existing_rec.get(wd_i,{}); typ=rr.get("preference_type","none"); block=rr.get("block","FULL")
        rec_rows.append({tr("weekday_name"):WEEKDAY_FULL[lang][wd_i],tr("recurring_rule"):rule_to_label.get(typ,tr("rec_none")),tr("recurring_time"):block_to_label.get(block,tr("full_day")),"_weekday":wd_i})
    rec_df=pd.DataFrame(rec_rows)
    edited=st.data_editor(rec_df,column_config={tr("recurring_rule"):st.column_config.SelectboxColumn(options=list(rule_to_label.values())),tr("recurring_time"):st.column_config.SelectboxColumn(options=list(block_to_label.values())),"_weekday":None},disabled=[tr("weekday_name")],hide_index=True,use_container_width=True,key=f"recurring_{initials}_{lang}")
    if st.button(tr("save_long_term"),type="primary",key=f"save_recurring_{initials}_{lang}"):
        payload=[]; invalid_weekend_soft=[]
        for _,r in edited.iterrows():
            wd=int(r["_weekday"]); typ=label_to_rule.get(r[tr("recurring_rule")],"none"); block=label_to_block.get(r[tr("recurring_time")],"FULL")
            if wd>=5 and typ=="soft_free":
                invalid_weekend_soft.append(WEEKDAY_FULL[lang][wd]); continue
            payload.append({"weekday":wd,"preference_type":typ,"block":block})
        if invalid_weekend_soft:
            st.error(
                "Pasikartojantis savaitgalio „Noriu laisvos“ nepriimamas, nes perkeltų privalomą savaitgalių krūvį kitiems. Jei iš tikrųjų negalite dirbti kiekvieną tokį savaitgalį, naudokite RESIDENT HARD. Tačiau savaitgalio „Pageidauju dirbti“ leidžiamas."
                if lang=="LT" else
                "Recurring weekend 'want off' is not accepted because it would shift unavoidable weekend burden to peers. Use RESIDENT HARD if you truly cannot work those weekends. Weekend 'prefer to work' is allowed."
            )
        else:
            db.save_recurring_preferences(initials,payload); flash_saved(tr("long_term_saved"))

# --- Senior dashboard ---
if senior_mode:
    with tabs[pos]:
        st.subheader(tr("dashboard_title")); prefs=db.all_preferences(year,month); sets=db.all_account_settings(); missing=[p["initials"] for p in DEFAULT_PEOPLE if p["initials"] not in prefs]; missing_mail=[p["initials"] for p in DEFAULT_PEOPLE if not sets.get(p["initials"],{}).get("email","").strip()]
        c1,c2,c3=st.columns(3); c1.metric(tr("completion"),f"{len(prefs)}/{len(DEFAULT_PEOPLE)}"); c2.metric(tr("missing_preferences"),len(missing)); c3.metric(tr("missing_email"),len(missing_mail))
        if missing: st.warning(f"{tr('missing_preferences')}: {', '.join(missing)}")
        else: st.success(tr("all_complete"))
        auto_zero=[i for i,x in prefs.items() if x.get("submission_source")=="deadline_zero"]
        if auto_zero:
            st.caption(
                ("TAIP = anketa užfiksuota kaip pateikta. Automatiniai 0 pageidavimų įrašai po termino: "
                 if lang=="LT" else
                 "YES = form recorded as submitted. Automatic zero-request submissions after the deadline: ")
                + ", ".join(sorted(auto_zero))
            )
        if missing_mail: st.warning(f"{tr('missing_email')}: {', '.join(missing_mail)}")
        st.info(tr("email_ready") if smtp_ready() else tr("email_not_ready"))
        if lifecycle_operator_ui:
            c_mail1,c_mail2=st.columns(2)
            with c_mail1:
                if st.button("SIŲSTI PAGEIDAVIMŲ ETAPO PRANEŠIMĄ" if lang=="LT" else "SEND PREFERENCES-OPEN NOTICE",use_container_width=True,disabled=(not smtp_ready() or bool(missing_mail)),key=f"pref_open_notice_{year}_{month}"):
                    res=preferences_open_emails(year,month)
                    st.dataframe(localized_delivery_rows(res),use_container_width=True,hide_index=True)
            with c_mail2:
                if st.button("PALEISTI ŠIANDIENOS PRIMINIMUS" if lang=="LT" else "RUN TODAY'S REMINDERS",use_container_width=True,disabled=not smtp_ready(),key=f"today_reminders_{year}_{month}"):
                    res=send_due_reminders(year,month)
                    if not res: st.info(tr("no_due_reminders"))
                    else: st.dataframe(localized_delivery_rows(res),use_container_width=True,hide_index=True)
            render_notification_delivery_status(year,month)
        if advanced_mode:
            logs=db.get_email_log(year,month)
            if logs:
                st.markdown(f"### {tr('email_log')}"); st.dataframe(localized_email_log(logs),use_container_width=True,hide_index=True)
    pos+=1

# --- Preferences ---
with tabs[pos]:
    st.subheader(f"{tr('my_preferences')} — {month_label(year,month)}")
    dl,dlmsg,dldiff=deadline_message(year,month); cutoff_exact=preference_cutoff_for(year,month)
    st.markdown(f'<div class="deadline-card"><b>{tr("deadline")}: {cutoff_exact.strftime("%Y-%m-%d %H:%M")}</b><br>{html.escape(dlmsg)}<br><span style="color:#6b7280">{html.escape(tr("deadline_note"))}</span></div>',unsafe_allow_html=True)
    st.caption(
        "Šio mėnesio konkretūs pageidavimai galioja tik pasirinktam grafikui. Kito mėnesio forma pildoma iš naujo; ilgalaikiai pasikartojantys pageidavimai žemiau išlieka, kol juos pakeisite."
        if lang=="LT" else
        "Month-specific requests apply only to the selected schedule. The next month starts with a new monthly form; long-term recurring preferences below persist until you change them."
    )
    if not resident_ok:
        st.error(tr("bad_pin"))
    else:
        cutoff_pref=preference_cutoff_for(year,month)
        now_lt=datetime.now(ZoneInfo("Europe/Vilnius"))
        own_deadline_open=now_lt < cutoff_pref
        pref_state=db.get_schedule_state(year,month)
        pref_system_frozen=bool(pref_state.get("has_published"))

        preference_target=active_user
        operator_manual_mode=False
        operator_reason_kind=""
        operator_reason_detail=""

        if lifecycle_operator_ui:
            st.markdown(
                """<div style="border:2px solid #7C9BFF;background:rgba(79,112,255,.08);
                border-radius:16px;padding:14px 16px;margin:4px 0 12px 0;">
                <b>Operatoriaus pageidavimų įvedimas</b><br>
                <span style="opacity:.82">Default — jūsų pačių anketa. Jei rezidentas negalėjo pateikti pats,
                pasirinkite jo inicialus ir suveskite informaciją jo vardu. Paskyros identitetas nekeičiamas,
                veiksmas audituojamas.</span></div>""",
                unsafe_allow_html=True,
            )
            target_order=[active_user]+[p["initials"] for p in DEFAULT_PEOPLE if p["initials"]!=active_user]
            name_map={p["initials"]:p["name"] for p in DEFAULT_PEOPLE}
            preference_target=st.selectbox(
                "Pildyti už:",
                target_order,
                index=0,
                format_func=lambda i:f"{i} — {name_map.get(i,i)}",
                key=f"pref_target_{year}_{month}_{active_user}",
            )
            st.markdown(
                f'<div style="border:1px solid rgba(124,155,255,.55);border-radius:12px;padding:10px 12px;">'
                f'<b>Pasirinktas rezidentas:</b> {badge(preference_target,include_name=True)}</div>',
                unsafe_allow_html=True,
            )
            operator_manual_mode=(preference_target!=active_user) or (not own_deadline_open)
            if pref_system_frozen:
                st.error(
                    "SYSTEM jau užšaldytas. Pageidavimų keisti nebegalima; operacinius pakeitimus darykite per Grafikas → rankinis koregavimas."
                    if lang=="LT" else
                    "SYSTEM is already frozen. Preferences can no longer be changed; use Schedule → manual correction for operational changes."
                )
        else:
            st.markdown(badge(active_user),unsafe_allow_html=True)
            if not own_deadline_open:
                st.warning(
                    f"Pageidavimų terminas baigėsi {cutoff_pref.strftime('%Y-%m-%d %H:%M')} Lietuvos laiku. "
                    "Anketos po termino rezidentas pats keisti nebegali. Jei būtina pataisa, kreipkitės į Seniūnę."
                    if lang=="LT" else
                    f"The preference deadline closed at {cutoff_pref.strftime('%Y-%m-%d %H:%M')} Lithuania time. "
                    "Residents can no longer edit the form themselves; contact the senior scheduler if a correction is required."
                )

        cur=db.get_preference(year,month,preference_target) or {}
        source=cur.get("submission_source","")
        submitter=cur.get("submitted_by_initials","")
        if cur:
            if source=="deadline_zero":
                st.info("Pateikta: TAIP — automatiškai užfiksuota 0 pageidavimų anketa po termino." if lang=="LT" else "Submitted: YES — automatic zero-request form after the deadline.")
            elif source=="operator_manual":
                st.info((f"Pateikta: TAIP — manualiai įvedė {submitter or 'operatorius'}." if lang=="LT" else f"Submitted: YES — manually entered by {submitter or 'operator'}."))
            else:
                st.success("Pateikta: TAIP — rezidento anketa." if lang=="LT" else "Submitted: YES — resident submission.")

        days=list(range(1,calendar.monthrange(year,month)[1]+1))
        avail_am=db.rest_credit_available_for_month(preference_target,year,month,"AM")
        avail_pm=db.rest_credit_available_for_month(preference_target,year,month,"PM")
        st.markdown(f"### {tr('short_term')}")
        with st.form(f"prefs_{year}_{month}_{active_user}_{preference_target}"):
            if lifecycle_operator_ui and operator_manual_mode:
                st.markdown("#### Manualaus įvedimo auditas" if lang=="LT" else "#### Manual-entry audit")
                operator_reason_kind=st.selectbox(
                    "Priežastis" if lang=="LT" else "Reason",
                    [
                        "Pateikta telefonu" if lang=="LT" else "Submitted by phone",
                        "Techninė / ryšio problema" if lang=="LT" else "Technical / connectivity issue",
                        "Sveikatos / neatvykimo situacija" if lang=="LT" else "Health / absence situation",
                        "Pavėluotas operatoriaus įvedimas" if lang=="LT" else "Late operator entry",
                        "Kita" if lang=="LT" else "Other",
                    ],
                    key=f"pref_operator_reason_{year}_{month}_{active_user}_{preference_target}",
                )
                operator_reason_detail=st.text_input(
                    "Trumpa pastaba (nebūtina)" if lang=="LT" else "Short note (optional)",
                    key=f"pref_operator_reason_detail_{year}_{month}_{active_user}_{preference_target}",
                )
            st.markdown(f"### {tr('hard_unavailable')}")
            st.caption(tr("hard_help"))
            h1,h2,h3=st.columns(3)
            with h1:
                unavailable=st.multiselect(
                    tr("hard_all_day"),days,default=sorted(cur.get("unavailable",set())),
                    format_func=lambda d:pretty_day(year,month,d)
                )
            with h2:
                unavailable_am=st.multiselect(
                    tr("hard_morning"),days,default=sorted(cur.get("unavailable_am",set())),
                    format_func=lambda d:pretty_day(year,month,d)
                )
            with h3:
                unavailable_pm=st.multiselect(
                    tr("hard_afternoon"),days,default=sorted(cur.get("unavailable_pm",set())),
                    format_func=lambda d:pretty_day(year,month,d)
                )
            st.caption(tr("hard_partial_note"))
            st.markdown(f"### {tr('soft_free')}")
            st.caption(tr("soft_help"))
            sf1,sf2,sf3=st.columns(3)
            with sf1:
                soft=st.multiselect(tr("hard_all_day"),days,default=sorted(cur.get("soft_free",set())),format_func=lambda d:pretty_day(year,month,d),key=f"soft_full_{year}_{month}_{active_user}_{preference_target}")
            with sf2:
                soft_am=st.multiselect(tr("hard_morning"),days,default=sorted(cur.get("soft_free_am",set())),format_func=lambda d:pretty_day(year,month,d),key=f"soft_am_{year}_{month}_{active_user}_{preference_target}")
            with sf3:
                soft_pm=st.multiselect(tr("hard_afternoon"),days,default=sorted(cur.get("soft_free_pm",set())),format_func=lambda d:pretty_day(year,month,d),key=f"soft_pm_{year}_{month}_{active_user}_{preference_target}")

            st.markdown(f"### {tr('preferred')}")
            st.caption(tr("preferred_help"))
            pf1,pf2,pf3=st.columns(3)
            with pf1:
                pref=st.multiselect(tr("hard_all_day"),days,default=sorted(cur.get("preferred",set())),format_func=lambda d:pretty_day(year,month,d),key=f"pref_full_{year}_{month}_{active_user}_{preference_target}")
            with pf2:
                pref_am=st.multiselect(tr("hard_morning"),days,default=sorted(cur.get("preferred_am",set())),format_func=lambda d:pretty_day(year,month,d),key=f"pref_am_{year}_{month}_{active_user}_{preference_target}")
            with pf3:
                pref_pm=st.multiselect(tr("hard_afternoon"),days,default=sorted(cur.get("preferred_pm",set())),format_func=lambda d:pretty_day(year,month,d),key=f"pref_pm_{year}_{month}_{active_user}_{preference_target}")

            st.markdown(f"### {tr('vacation')}")
            st.caption(tr("vacation_help"))
            vacation=st.multiselect(
                tr("vacation"),days,default=sorted(cur.get("vacation",set())),
                format_func=lambda d:pretty_day(year,month,d),key=f"vacation_{year}_{month}_{active_user}_{preference_target}"
            )

            note=st.text_area(tr("note"),value=cur.get("note",""),placeholder=tr("note_ph"))
            cr1,cr2=st.columns(2)
            with cr1:
                bonus_am_use=st.number_input(tr("use_credit_am"),min_value=0,max_value=max(0,avail_am),value=min(int(cur.get("backup_credits_am_to_use",0)),max(0,avail_am)),step=1,help=tr("bonus_target_effect"))
            with cr2:
                bonus_pm_use=st.number_input(tr("use_credit_pm"),min_value=0,max_value=max(0,avail_pm),value=min(int(cur.get("backup_credits_pm_to_use",0)),max(0,avail_pm)),step=1,help=tr("bonus_target_effect"))
            st.divider()
            st.markdown(f"### {tr('legal_safety_inputs')}")
            st.caption(tr("labour_hard_summary"))
            l1,l2=st.columns(2)
            with l1:
                justified_absence=st.multiselect(
                    tr("justified_absence"),days,default=sorted(cur.get("justified_absence",set())),
                    format_func=lambda d:pretty_day(year,month,d),help=tr("justified_absence_help")
                )
            with l2:
                long_duty=st.multiselect(
                    tr("long_duty"),days,default=sorted(cur.get("long_duty",set())),
                    format_func=lambda d:pretty_day(year,month,d),help=tr("long_duty_help")
                )
            st.caption(tr("labour_scope_note"))
            resident_edit_blocked=(not lifecycle_operator_ui and not own_deadline_open)
            operator_edit_blocked=(lifecycle_operator_ui and pref_system_frozen)
            submitted=st.form_submit_button(
                ("IŠSAUGOTI UŽ " + preference_target if lifecycle_operator_ui and operator_manual_mode and lang=="LT"
                 else "SAVE FOR " + preference_target if lifecycle_operator_ui and operator_manual_mode
                 else tr("save")),
                type="primary",
                disabled=resident_edit_blocked or operator_edit_blocked,
            )
            if submitted:
                whole=set(unavailable); am=set(unavailable_am); pm=set(unavailable_pm)
                sf=set(soft); sf_am=set(soft_am); sf_pm=set(soft_pm)
                pr=set(pref); pr_am=set(pref_am); pr_pm=set(pref_pm)
                vacation_set=set(vacation)
                absence=set(justified_absence)|vacation_set
                hard_pref_conflict = bool(
                    (whole|absence) & (pr|pr_am|pr_pm) or
                    am & (pr|pr_am) or
                    pm & (pr|pr_pm)
                )
                soft_pref_conflict = bool(
                    sf & (pr|pr_am|pr_pm) or
                    pr & (sf|sf_am|sf_pm) or
                    sf_am & pr_am or sf_pm & pr_pm
                )
                if whole & (am|pm):
                    st.error(tr("hard_overlap"))
                elif sf & (sf_am|sf_pm):
                    st.error(tr("soft_overlap"))
                elif pr & (pr_am|pr_pm):
                    st.error(tr("preferred_overlap"))
                elif vacation_set & set(justified_absence):
                    st.error(tr("vacation_overlap"))
                elif hard_pref_conflict:
                    st.error(tr("hard_conflict"))
                elif soft_pref_conflict:
                    st.error(tr("soft_conflict"))
                elif int(bonus_am_use)+int(bonus_pm_use)>2:
                    st.error(tr("max_credit_error"))
                else:
                    pref_payload={
                        "unavailable":unavailable,
                        "unavailable_am":unavailable_am,
                        "unavailable_pm":unavailable_pm,
                        "justified_absence":justified_absence,
                        "vacation":vacation,
                        "long_duty":long_duty,
                        "soft_free":soft,
                        "soft_free_am":soft_am,
                        "soft_free_pm":soft_pm,
                        "preferred":pref,
                        "preferred_am":pref_am,
                        "preferred_pm":pref_pm,
                        "note":note,
                        "backup_credits_am_to_use":int(bonus_am_use),
                        "backup_credits_pm_to_use":int(bonus_pm_use),
                        "backup_credits_night_to_use":0,
                    }
                    try:
                        if lifecycle_operator_ui and operator_manual_mode:
                            audit_reason=operator_reason_kind.strip()
                            if operator_reason_detail.strip():
                                audit_reason += " — " + operator_reason_detail.strip()
                            db.save_preference_for_resident_v2595(
                                year,month,preference_target,pref_payload,audit_reason
                            )
                            draft_note=(
                                " Jei šiam mėnesiui jau buvo sugeneruotas DRAFT, jį reikia regeneruoti."
                                if db.get_schedule_state(year,month).get("has_draft") else ""
                            )
                            flash_saved(
                                (f"{preference_target} pageidavimai įvesti operatoriaus vardu ir audituoti.{draft_note}"
                                 if lang=="LT" else
                                 f"{preference_target} preferences were entered by the operator and audited."
                                 + (" Regenerate the existing DRAFT." if draft_note else ""))
                            )
                        else:
                            db.save_preference(year,month,active_user,pref_payload)
                            flash_saved(tr("saved"))
                    except Exception as e:
                        msg=str(e)
                        if "PREFERENCE_DEADLINE_CLOSED" in msg:
                            st.error("Pageidavimų terminas jau uždarytas. Susisiekite su Seniūne." if lang=="LT" else "The preference deadline is closed. Contact the senior scheduler.")
                        elif "PREFERENCE_INPUT_FROZEN_AFTER_SYSTEM" in msg:
                            st.error("SYSTEM jau užšaldytas — pageidavimų keisti nebegalima." if lang=="LT" else "SYSTEM is frozen — preferences can no longer be changed.")
                        else:
                            st.error(msg)
        if preference_target==active_user:
            render_recurring_preferences_editor(active_user)
        elif lifecycle_operator_ui:
            st.divider()
            st.caption(
                "Ilgalaikius pasikartojančius pageidavimus kiekvienas rezidentas valdo savo Pageidavimų lange. Operatorinis įvedimas aukščiau keičia tik pasirinktą konkretų mėnesį."
                if lang=="LT" else
                "Each resident manages long-term recurring preferences in their own Preferences tab. The operator entry above changes only the selected month."
            )
    if senior_mode:
        st.divider(); st.markdown(f"### {tr('all_preferences')}"); prefs=db.all_preferences(year,month); sets=db.all_account_settings(); recurring_all=db.all_recurring_preferences(); nd=calendar.monthrange(year,month)[1]; rows=[]
        for p in DEFAULT_PEOPLE:
            x=prefs.get(p["initials"],{})
            vol=(len(x.get("unavailable",set()))+len(x.get("unavailable_am",set()))+
                 len(x.get("unavailable_pm",set()))+len(x.get("justified_absence",set()))+len(x.get("vacation",set()))+
                 len(x.get("long_duty",set()))+len(x.get("soft_free",set()))+
                 len(x.get("soft_free_am",set()))+len(x.get("soft_free_pm",set()))+
                 len(x.get("preferred",set()))+len(x.get("preferred_am",set()))+len(x.get("preferred_pm",set())))
            flag=tr("review") if vol>=max(10,round(nd/3)) else tr("normal")
            rows.append({
                tr("person"):p["initials"],tr("name"):p["name"],
                tr("submitted"):tr("yes") if x else tr("no"),
                ("Pateikimo būdas" if lang=="LT" else "Submission source"):(
                    ("Automatiškai — 0 pageidavimų" if lang=="LT" else "Automatic — 0 requests")
                    if x and x.get("submission_source")=="deadline_zero"
                    else ((f"Manualiai — {x.get('submitted_by_initials') or 'operatorius'}" if lang=="LT"
                           else f"Manual — {x.get('submitted_by_initials') or 'operator'}")
                          if x and x.get("submission_source")=="operator_manual"
                          else ("Rezidentas" if lang=="LT" else "Resident")
                          if x else "—")
                ),
                tr("preference_load"):f"{vol} — {flag}",
                tr("hard_dates"):", ".join(map(str,sorted(x.get("unavailable",set())))),
                tr("hard_am_dates"):", ".join(map(str,sorted(x.get("unavailable_am",set())))),
                tr("hard_pm_dates"):", ".join(map(str,sorted(x.get("unavailable_pm",set())))),
                tr("vacation"):", ".join(map(str,sorted(x.get("vacation",set())))),
                tr("justified_absence"):", ".join(map(str,sorted(x.get("justified_absence",set())))),
                tr("long_duty"):", ".join(map(str,sorted(x.get("long_duty",set())))),
                tr("soft_dates"):", ".join(map(str,sorted(x.get("soft_free",set())))),
                tr("soft_am_dates"):", ".join(map(str,sorted(x.get("soft_free_am",set())))),
                tr("soft_pm_dates"):", ".join(map(str,sorted(x.get("soft_free_pm",set())))),
                tr("preferred_dates"):", ".join(map(str,sorted(x.get("preferred",set())))),
                tr("preferred_am_dates"):", ".join(map(str,sorted(x.get("preferred_am",set())))),
                tr("preferred_pm_dates"):", ".join(map(str,sorted(x.get("preferred_pm",set())))),
                tr("long_term"):"; ".join(f"{WEEKDAY_FULL[lang][int(r['weekday'])]}: {r['preference_type']} {r.get('block','FULL')}" for r in recurring_all.get(p["initials"],[])),
                tr("use_credit_am"):x.get("backup_credits_am_to_use",0),
                tr("use_credit_pm"):x.get("backup_credits_pm_to_use",0),
                tr("comment"):x.get("note",""),
                tr("holiday_pref"):({-1:tr("holiday_rest"),0:tr("holiday_neutral"),1:tr("holiday_work")}.get(int(sets.get(p["initials"],{}).get("holiday_preference",0) or 0),tr("holiday_neutral"))),
                tr("email"):sets.get(p["initials"],{}).get("email","")
            })
        st.dataframe(style_rows(pd.DataFrame(rows)),use_container_width=True,hide_index=True); st.caption(tr("visibility_flag"))
pos+=1

# --- Settings ---
with tabs[pos]:
    st.subheader(tr("settings_title")); st.markdown(badge(active_user),unsafe_allow_html=True)
    if not resident_ok: st.error(tr("bad_pin"))
    else:
        s=db.get_account_settings(active_user)
        st.caption(
            "Darbo pobūdžio nustatymai yra ilgalaikiai: jie automatiškai taikomi kiekvienam būsimam dar neužšaldytam mėnesiui, kol pats juos pakeisite. Jie nėra iš naujo nustatomi kiekvieną mėnesį."
            if lang=="LT" else
            "Work-style settings are persistent: they automatically apply to every future schedule that is not yet frozen until you change them. They do not reset each month."
        )
        with st.form(f"settings_{active_user}"):
            shift_len_options=[tr("shift_length_6"),tr("shift_length_12"),tr("shift_length_mixed"),tr("shift_length_any")]
            shift_len_value_to_label={0:tr("shift_length_any"),1:tr("shift_length_6"),2:tr("shift_length_mixed"),3:tr("shift_length_12")}
            shift_len_label_to_value={v:k for k,v in shift_len_value_to_label.items()}
            current_shift_len=max(0,min(3,int(s.get("shift_length_preference",0) or 0)))
            shift_len_label=st.selectbox(
                tr("shift_length_pref"),shift_len_options,
                index=shift_len_options.index(shift_len_value_to_label.get(current_shift_len,tr("shift_length_any"))),
                help=tr("shift_length_help")
            )
            shift_len_pref=shift_len_label_to_value.get(shift_len_label,0)
            st.caption(tr("shift_length_help"))
            st.divider()
            email=st.text_input(tr("email"),value=s.get("email","")); st.caption(tr("email_required"))
            st.caption("Pageidavimuose paliekami konkretūs, aiškūs poreikiai. Bendras savaitgalių vengimas nekeičia privalomo lygaus savaitgalių paskirstymo visai grupei. Jei reikia konkrečios laisvos dienos, pažymėkite ją Pageidavimų lange.")
            holiday_options=[tr("holiday_rest"),tr("holiday_neutral"),tr("holiday_work")]
            holiday_value_to_label={-1:tr("holiday_rest"),0:tr("holiday_neutral"),1:tr("holiday_work")}
            holiday_label_to_value={v:k for k,v in holiday_value_to_label.items()}
            holiday_pref_label=st.selectbox(
                tr("holiday_pref"),holiday_options,
                index=holiday_options.index(holiday_value_to_label.get(max(-1,min(1,int(s.get("holiday_preference",0) or 0))),tr("holiday_neutral"))),
                help=tr("holiday_pref_help")
            )
            hp=holiday_label_to_value.get(holiday_pref_label,0)
            wp=st.slider(tr("weekday_pref"),-2,2,int(s.get("weekday_preference",0) or 0),help=tr("weekday_help"))
            wep=0
            st.info(
                "SAVAITGALIAI — ADMIN WATER-FILL. Rezidentas nebegali pasirinkti „noriu daugiau savaitgalių / sekmadienių“. SYSTEM automatiškai ieško mažiausio matematiškai įmanomo šeštadienių, sekmadienių ir bendro savaitgalio krūvio skirtumo. Tik tikras „Negaliu dirbti“ gali apriboti konkrečią dieną. Po publikavimo savanoriškas swapas gali pakeisti ACTUAL balansą, bet tik po abiejų rezidentų sutikimo ir SP galutinio patvirtinimo."
                if lang=="LT" else
                "WEEKENDS — ADMIN WATER-FILL. Residents can no longer choose to receive more weekends/Sundays. SYSTEM searches for the tightest mathematically feasible Saturday, Sunday and total-weekend spread. Only genuine Cannot-work availability can block a specific date. A voluntary post-publication swap may alter ACTUAL balance, but only after both residents consent and SP gives final approval."
            )
            sp=st.slider(tr("spread_pref"),-2,2,int(s.get("spread_preference",0)),help=tr("spread_help"))
            # `avoid_doubles` is retained in the database for backward compatibility;
            # the new four-way workday-length selector is the resident-facing source of truth.
            avoid=(shift_len_pref==1)
            st.markdown(f"### {tr('calendar')}")
            include_bk=st.checkbox(tr("include_backups_calendar"),value=bool(s.get("include_backups_in_calendar",False)))
            st.markdown(f"### {tr('notifications')}"); st.caption(tr("notification_default"))
            notif=st.toggle(tr("notifications_on"),value=bool(s.get("notifications_on",True)))
            dl=deadline_for(year,month)
            st.caption((f"Asmeniniai priminimai skaičiuojami iki šio grafiko pageidavimų termino: {dl:%Y-%m-%d}." if lang=="LT" else f"Personal reminders count down to this schedule’s preference deadline: {dl:%Y-%m-%d}."))
            backup_email=st.toggle(tr("backup_email_alerts"),value=bool(s.get("backup_email_alerts",True)))
            phone=st.text_input(tr("phone_optional"),value=s.get("phone_e164",""),placeholder="+3706XXXXXXX")
            backup_sms=st.toggle(tr("backup_sms_alerts"),value=bool(s.get("backup_sms_alerts",False)),disabled=True)
            st.caption(tr("sms_future")); st.caption(tr("smtp_admin_note"))
            start=st.number_input(tr("reminder_start"),1,13,int(s.get("reminder_start_day",8)),1,help=tr("reminder_help"))
            if st.form_submit_button(tr("save"),type="primary"):
                if not email.strip() or "@" not in email or "." not in email.split("@")[-1]: st.error(tr("email_required"))
                elif phone.strip() and not re.fullmatch(r"\+[1-9][0-9]{7,14}",phone.strip()): st.error("+3706XXXXXXX")
                else:
                    try:
                        db.save_account_settings(active_user,{"email":email,"weekday_preference":wp,"weekend_preference":wep,"holiday_preference":hp,"spread_preference":sp,"shift_length_preference":shift_len_pref,"avoid_doubles":avoid,"notifications_on":notif,"reminder_start_day":int(start),"preferred_language":lang,"include_backups_in_calendar":include_bk,"backup_email_alerts":backup_email,"phone_e164":(phone.strip() or None),"backup_sms_alerts":False})
                        refresh_calendar_subscription_feeds([active_user])
                    except Exception as exc:
                        st.error(
                            "Nustatymų išsaugoti nepavyko. Duomenys nebuvo pakeisti. "
                            "Jei klaida kartojasi, administratorius turi patikrinti account_settings schemą."
                        )
                        if advanced_mode:
                            st.caption(f"{type(exc).__name__}: {exc}")
                    else:
                        flash_saved(tr("settings_saved"))
pos+=1

def _hard_error_explanation(raw, lang="LT"):
    s=str(raw)
    lt=[
        ("Mandatory slot unfilled", "Neužpildytas administraciškai privalomas SPS RO / SPS UG / savaitgalio slotas."),
        ("Gap dispersion violated", "Tą pačią dieną liko daugiau nei 1 reali neužpildyta optional vieta."),
        ("Gap workplace dispersion violated", "Neužpildytos optional vietos per daug susikoncentravo vienoje postų grupėje."),
        ("Gap-day dispersion pattern outdated", "Skylės nėra išmėtytos per mėnesį pagal dabartinę tolygaus paskirstymo taisyklę."),
        ("workload", "Rezidento mėnesio darbo krūvis neatitinka jo tikslaus targeto."),
        ("odd Onko", "Onko skaičius turi būti lyginis kiekvienam rezidentui: 0, 2, 4... Nelyginis 1/3/5 yra ABSOLUTE HARD klaida."),
        ("overlapping assignments", "Tam pačiam žmogui paskirtos laike persidengiančios pamainos."),
        ("HARD-unavailable", "Pamaina paskirta tuo metu, kai žmogus pažymėtas kaip HARD negalintis dirbti."),
        ("backup resident", "Privalomai dubliuojamai pamainai nėra laisvo ir HARD-prieinamo dublio."),
        ("weekend cap exceeded", "Pažeistas savaitgalio maksimalaus krūvio / unikalumo limitas."),
        ("rest between days", "Pažeistas minimalus poilsio laikas tarp darbo dienų."),
        ("workdays/7d cap", "Viršytas maksimalus darbo dienų skaičius per slenkantį 7 dienų langą."),
        ("hours/7d cap", "Viršytas maksimalus darbo valandų skaičius per slenkantį 7 dienų langą."),
        ("hours/day", "Viršytas maksimalus darbo valandų skaičius per dieną."),
    ]
    en=[
        ("Mandatory slot unfilled", "A mandatory SPS RO / SPS UG / weekend slot is not filled."),
        ("Gap dispersion violated", "The number of optional unfilled rows on a day does not match the workload-adjusted monthly plan."),
        ("Gap workplace dispersion violated", "Optional gaps are too concentrated in one workplace group."),
        ("Gap-day dispersion pattern outdated", "Gap dates do not match the current evenly distributed monthly pattern."),
        ("workload", "A resident's monthly workload does not match the exact target."),
        ("odd Onko", "Each resident must have an even Onko count: 0, 2, 4... Odd 1/3/5 is an ABSOLUTE HARD error."),
        ("overlapping assignments", "A resident has overlapping assignments."),
        ("HARD-unavailable", "A shift is assigned during HARD unavailability."),
        ("backup resident", "A required covered shift has no free HARD-eligible backup."),
        ("weekend cap exceeded", "Weekend uniqueness / resident cap is violated."),
        ("rest between days", "Minimum rest between workdays is violated."),
        ("workdays/7d cap", "Maximum workdays in a rolling 7-day window is exceeded."),
        ("hours/7d cap", "Maximum hours in a rolling 7-day window is exceeded."),
        ("hours/day", "Maximum hours per day is exceeded."),
    ]
    for key,msg in (lt if lang=="LT" else en):
        if key.lower() in s.lower():
            return msg
    return s


def render_hard_error_explainer(g, lang="LT", key_suffix=""):
    errors=list(g.get("errors") or [])
    title=("* Privalomų taisyklių klaidų paaiškinimas" if lang=="LT"
           else "* HARD-rule error explanation")
    with st.expander(title, expanded=bool(errors)):
        if not errors:
            st.success(
                "0 = dabartinis grafikas praeina visas šiuo metu aktyvias HARD patikras."
                if lang=="LT" else
                "0 = the current schedule passes all active HARD checks."
            )
        else:
            st.error(
                f"Rasta {len(errors)} HARD klaidų." if lang=="LT"
                else f"{len(errors)} HARD errors found."
            )
            for n,e in enumerate(errors,1):
                st.markdown(f"**{n}.** {_hard_error_explanation(e,lang)}")
                if advanced_mode:
                    st.caption(str(e))


def solve_schedule_isolated(year, month, people, time_limit=90.0):
    """Run the production generator in a disposable Python process.

    V2.5.106 safety/reliability boundary: SciPy/HiGHS normally respects its MILP
    time limit, but a native solver call cannot be force-cancelled safely from the
    Streamlit thread if the underlying library stalls. The operational Generate /
    Rebuild path therefore runs the engine in a fresh child process. A hard OS-level
    watchdog can terminate a stuck worker, and one clean-process retry is allowed.

    The worker receives only a frozen request snapshot + active rule profile; it
    does not connect to Supabase and cannot mutate the existing draft. The caller
    saves a new draft only after a verified SolveResult(ok=True) returns.
    """
    worker=BASE / "solver_runner.py"
    if not worker.exists():
        return SolveResult(False,"ISOLATED SOLVER WORKER MISSING — deploy solver_runner.py with this release.")

    frozen=serialize_people_request_snapshot(people)
    payload={
        "year":int(year),
        "month":int(month),
        "time_limit":float(time_limit),
        "rules":get_runtime_rules(),
        "people_snapshot":frozen,
        "expected_engine_api":EXPECTED_ENGINE_API_VERSION,
    }
    last_result=None
    attempt_log=[]
    # First run is deliberately bounded well above the real September regression
    # (~25-35 s on the direct regression; slower clean workers can take longer).
    # The second run is a clean process with additional room for cloud variance.
    watchdogs=(100.0,130.0)
    for attempt_no,watchdog in enumerate(watchdogs,1):
        with tempfile.TemporaryDirectory(prefix="shift_happens_solver_") as td:
            ip=Path(td)/"input.json"
            op=Path(td)/"output.json"
            ip.write_text(json.dumps(payload,ensure_ascii=False),encoding="utf-8")
            t0=perf_counter()
            proc=None
            try:
                proc=subprocess.Popen(
                    [sys.executable,str(worker),str(ip),str(op)],
                    cwd=str(BASE),
                    stdout=subprocess.PIPE,
                    stderr=subprocess.PIPE,
                    text=True,
                    start_new_session=True,
                )
                stdout_text,stderr_text=proc.communicate(timeout=watchdog)
                elapsed=perf_counter()-t0
            except subprocess.TimeoutExpired:
                elapsed=perf_counter()-t0
                try:
                    if proc is not None and proc.poll() is None:
                        if os.name=="posix":
                            os.killpg(proc.pid,signal.SIGKILL)
                        else:
                            proc.kill()
                    if proc is not None:
                        proc.communicate(timeout=5.0)
                except Exception:
                    try:
                        if proc is not None: proc.kill()
                    except Exception:
                        pass
                attempt_log.append({"attempt":attempt_no,"watchdog_seconds":watchdog,"elapsed_seconds":round(elapsed,3),"outcome":"WATCHDOG_TIMEOUT"})
                continue
            except Exception as exc:
                elapsed=perf_counter()-t0
                try:
                    if proc is not None and proc.poll() is None:
                        if os.name=="posix": os.killpg(proc.pid,signal.SIGKILL)
                        else: proc.kill()
                except Exception:
                    pass
                attempt_log.append({"attempt":attempt_no,"watchdog_seconds":watchdog,"elapsed_seconds":round(elapsed,3),"outcome":"PROCESS_ERROR","detail":str(exc)[:300]})
                continue

            if proc.returncode!=0 or not op.exists():
                attempt_log.append({
                    "attempt":attempt_no,"watchdog_seconds":watchdog,"elapsed_seconds":round(elapsed,3),
                    "outcome":"WORKER_EXIT","returncode":int(proc.returncode),
                    "stderr":str(stderr_text or "")[-500:],
                })
                continue
            try:
                raw=json.loads(op.read_text(encoding="utf-8"))
                if str(raw.get("engine_api"))!=EXPECTED_ENGINE_API_VERSION:
                    raise ValueError(f"worker engine {raw.get('engine_api')} != expected {EXPECTED_ENGINE_API_VERSION}")
                result=deserialize_result(raw.get("result") or {})
            except Exception as exc:
                attempt_log.append({"attempt":attempt_no,"watchdog_seconds":watchdog,"elapsed_seconds":round(elapsed,3),"outcome":"BAD_WORKER_OUTPUT","detail":str(exc)[:300]})
                continue

            attempt_log.append({"attempt":attempt_no,"watchdog_seconds":watchdog,"elapsed_seconds":round(elapsed,3),"outcome":"OK" if result.ok else "NO_VERIFIED_CANDIDATE"})
            last_result=result
            if result.ok:
                g=(result.stats or {}).setdefault("global",{})
                g["solver_process_mode"]="ISOLATED_SUBPROCESS"
                g["isolated_worker_attempts"]=attempt_no
                g["isolated_worker_attempt_log"]=attempt_log
                g["isolated_worker_watchdog_passed"]=True
                return result
            # Definitive model/validator errors should be shown immediately. Only
            # retry the explicit no-candidate/time-limit class in a clean process.
            retryable=("PREFERENCE-AWARE GENERATION DID NOT FINISH" in str(result.message or ""))
            if not retryable:
                return result

    if last_result is not None:
        last_result.message=(
            "ISOLATED GENERATION DID NOT RETURN A VERIFIED CANDIDATE AFTER CLEAN-PROCESS RETRY. "
            "The request set was not proven infeasible; the existing draft remains unchanged."
        )
        return last_result
    return SolveResult(
        False,
        "ISOLATED GENERATION WATCHDOG STOPPED THE SOLVER ON BOTH CLEAN-PROCESS ATTEMPTS. "
        "No infeasibility was inferred and the existing draft remains unchanged.",
        request_snapshot=frozen,
    )


def _draft_quality_tuple(result):
    """Lexicographic V2.5.49 draft comparison — lower is better.

    V2.5.107 requires every candidate to have 0 RESIDENT-HARD misses. Within that
    mandatory zero-loss space, ordinary fairness/SOFT refinements decide the winner.
    """
    g=(result.stats or {}).get("global",{})
    hard=int(g.get("hard_errors",9999))
    rh_total=int(g.get("resident_hard_total_losses",9999))
    rh_max=int(g.get("resident_hard_max_loss_per_resident",9999))
    rh_cum=int(g.get("resident_hard_cumulative_spread",9999))
    worst_post=int(g.get("worst_monthly_post_spread",9999))
    soft_min=g.get("min_soft_preference_score")
    soft_mean=g.get("mean_soft_preference_score")
    overall_mean=g.get("mean_preference_score")
    monthly_fair=float(g.get("monthly_fairness_score",g.get("fairness_score",0)) or 0)
    cumulative_fair=float(g.get("cumulative_fairness_score",g.get("fairness_score",0)) or 0)
    return (
        hard,
        rh_total,
        rh_max,
        rh_cum,
        worst_post,
        -(float(soft_min) if soft_min is not None else 100.0),
        -(float(soft_mean) if soft_mean is not None else 100.0),
        -monthly_fair,
        -cumulative_fair,
        -(float(overall_mean) if overall_mean is not None else 100.0),
    )

# --- Generation ---
if senior_mode:
    with tabs[pos]:
        st.subheader(tr("generation_title")); state=db.get_schedule_state(year,month); status=tr("published_state") if state["has_published"] else tr("draft") if state["has_draft"] else tr("not_created"); st.metric(tr("state"),status)
        lifecycle_generation=db.get_schedule_lifecycle(year,month)
        generation_locked=bool(state.get("has_published")) or str(lifecycle_generation.get("state") or "") in ("working","swap_open","swap_closed","final")
        if generation_locked:
            st.warning("SYSTEM jau užšaldytas šiame lifecycle etape. Operacinio juodraščio regeneruoti / gerinti nebegalima; swapai keičia tik ACTUAL, o FINAL nekeičia SYSTEM." if lang=="LT" else "SYSTEM is frozen at this lifecycle stage. The operational draft can no longer be regenerated/improved; swaps change ACTUAL only and FINAL does not rewrite SYSTEM.")
        prefs=db.all_preferences(year,month); missing=[p["initials"] for p in DEFAULT_PEOPLE if p["initials"] not in prefs]
        if missing: st.warning(f"{tr('missing_preferences')}: {', '.join(missing)}")
        c1,c2=st.columns(2)
        with c1:
            if active_user==SENIOR_INITIALS:
                try:
                    _weston_now=db.weston_beer_stats_v25110(year,month)
                    st.caption(
                        f"1 click = 1 WESTON beer ŠR · tavo skola ŠR: {_weston_now.get('total_beers',0)}"
                        if lang=="LT" else
                        f"1 click = 1 WESTON beer owed to ŠR · your debt to ŠR: {_weston_now.get('total_beers',0)}"
                    )
                except Exception:
                    _weston_now={"total_beers":0,"month_beers":0}
            if st.button(tr("generate_draft"),type="primary",use_container_width=True,disabled=generation_locked):
                if active_user==SENIOR_INITIALS:
                    try:
                        _weston_after=db.record_weston_beer_click_v25110(year,month)
                        st.caption(
                            f"WESTON +1. Dabar ŠR esi skolinga: {_weston_after.get('total_beers',0)}."
                            if lang=="LT" else
                            f"WESTON +1. You now owe ŠR: {_weston_after.get('total_beers',0)}."
                        )
                    except Exception:
                        st.warning(
                            "WESTON skaitiklio nepavyko įrašyti; grafiko generavimas tęsiamas."
                            if lang=="LT" else
                            "The WESTON counter could not be recorded; schedule generation will continue."
                        )
                credit_err=credit_selection_errors(year,month)
                if credit_err:
                    st.error(tr("bonus_insufficient")); st.dataframe(pd.DataFrame(credit_err),use_container_width=True,hide_index=True)
                else:
                    t0=perf_counter()
                    with st.spinner(tr("solver_wait")): result=solve_schedule_isolated(year,month,load_people(year,month),time_limit=90)
                    elapsed=perf_counter()-t0
                    try:
                        gg=(result.stats or {}).get("global",{}) if result.ok else {}
                        db.record_research_generation_event(
                            year,month,elapsed,result.ok,gg.get("hard_errors"),
                            gg.get("monthly_fairness_score",gg.get("fairness_score")),
                            gg.get("cumulative_fairness_score",gg.get("fairness_score")),
                            gg.get("mean_preference_score")
                        )
                    except Exception:
                        pass
                    if result.ok:
                        # V2.5.115: first freeze/revalidate the NORMAL schedule by itself.
                        # The theoretical backup plan is built only afterwards as a separate
                        # standby layer and can never change wish scores or normal-schedule HARD.
                        result=revalidate_loaded_result(
                            year,month,people_for_stored_result(result,year,month),result,
                            backup_assignments=[]
                        )
                        if result.stats.get("global",{}).get("hard_errors",0):
                            st.error(tr("draft_outdated"))
                            _berr=result.stats.get("global",{}).get("errors",[])
                            if _berr: st.dataframe(pd.DataFrame(_berr),use_container_width=True,hide_index=True)
                        else:
                            desired,backup_errors=plan_backups(year,month,result)
                            result.backup_snapshot=[dict(x) for x in desired]
                            _gg=result.stats.setdefault("global",{})
                            _gg["theoretical_backup_layer"]=True
                            _gg["theoretical_backup_layer_errors"]=list(backup_errors)
                            _gg["theoretical_backup_layer_complete"]=(len(backup_errors)==0)
                            db.save_draft(year,month,serialize_result(result))
                            st.success(tr("draft_saved"))
                            if backup_errors:
                                st.warning(
                                    "NORMALUS grafikas išsaugotas ir jo pageidavimų auditas lieka validus. Atskirame teorinių dublių sluoksnyje dar trūksta kelių standby paskyrimų; tai nėra darbo grafiko ar pageidavimų pažeidimas."
                                    if lang=="LT" else
                                    "The NORMAL schedule was saved and its request audit remains valid. The separate theoretical backup layer still has standby gaps; these are not work-schedule or preference violations."
                                )
                                st.dataframe(pd.DataFrame(backup_errors),use_container_width=True,hide_index=True)
                            _bc=backup_counts(year,month,result)[0]
                            _vals=list(_bc.values())
                            st.caption(
                                (f"TEORINIS AUTO dublių sluoksnis: {sum(_vals)} standby pareigų · rezidentų spread {max(_vals)-min(_vals) if _vals else 0}. Jos NĖRA darbo pamainos."
                                 if lang=="LT" else
                                 f"THEORETICAL AUTO backup layer: {sum(_vals)} standby duties · resident spread {max(_vals)-min(_vals) if _vals else 0}. They are NOT work shifts.")
                            )
                            norm=(result.stats or {}).get("global",{}).get("preference_normalization",[])
                            if norm:
                                st.caption(
                                    f"Preference pre-check: {len(norm)} redundant / impossible / engine-covered SOFT signalai "
                                    "nebuvo antrą kartą įtraukti į optimizerį."
                                )
                            if "fallback" in (result.message or "").lower():
                                st.warning(
                                    "Globalus fairness MILP nespėjo pilnai užsibaigti, bet sistema prieš išsaugodama pritaikė local fairness repair loop. "
                                    "Grafikas yra HARD-valid; „PERTIKRINTI / GERINTI“ gali bandyti jį dar pagerinti."
                                )
                            st.rerun()
                    else:
                        _msg=result.message if getattr(result,"message",None) else tr("no_solution")
                        if ("PREFERENCE-AWARE GENERATION DID NOT FINISH" in str(_msg) or "ISOLATED GENERATION" in str(_msg)):
                            st.warning(
                                "Solveris neįrodė, kad grafikas neįmanomas — jis tiesiog negavo patvirtinto kandidato net po automatinio retry. "
                                "Esamas juodraštis, jei yra, nepakeistas. Galima spausti GENERUOTI / PERKURTI dar kartą."
                                if lang=="LT" else
                                "The solver did not prove the schedule infeasible; it simply did not obtain a verified candidate even after automatic retry. "
                                "Any existing draft is preserved. You can run GENERATE / REBUILD again."
                            )
                        else:
                            st.error(_msg)
        draft_for_improve=db.load_schedule(year,month,"draft")
        if draft_for_improve:
            current_draft=refresh_result_payload(draft_for_improve,year,month,use_actual_backups=False)
            st.caption(
                "Juodraštis jau egzistuoja. Gali jį pertikrinti ir ieškoti geresnio varianto. "
                "Esamas grafikas nebus prarastas, jei naujas kandidatas blogesnis."
            )
            if st.button("PERTIKRINTI / GERINTI GRAFIKĄ",use_container_width=True,key=f"improve_{year}_{month}",disabled=generation_locked):
                t0=perf_counter()
                with st.spinner("Ieškau geresnio varianto pagal nustatytą prioritetų tvarką: privalomos taisyklės → kuo lygesnis darbo vietų paskirstymas → pageidavimai..."):
                    candidate=solve_schedule_isolated(year,month,load_people(year,month),time_limit=90)
                elapsed=perf_counter()-t0
                if not candidate.ok:
                    st.warning("Esamas validus juodraštis paliktas nepakeistas. Naujo geresnio kandidato rasti nepavyko: "+str(candidate.message))
                else:
                    candidate=revalidate_loaded_result(year,month,people_for_stored_result(candidate,year,month),candidate,backup_assignments=[])
                    _cand_backups,_cand_backup_errors=plan_backups(year,month,candidate)
                    candidate.backup_snapshot=[dict(x) for x in _cand_backups]
                    _cgg=candidate.stats.setdefault("global",{})
                    _cgg["theoretical_backup_layer"]=True
                    _cgg["theoretical_backup_layer_errors"]=list(_cand_backup_errors)
                    _cgg["theoretical_backup_layer_complete"]=(len(_cand_backup_errors)==0)
                    old_q=_draft_quality_tuple(current_draft)
                    new_q=_draft_quality_tuple(candidate)
                    if new_q < old_q:
                        db.save_draft(year,month,serialize_result(candidate))
                        st.success(
                            "Rastas geresnis NORMALUS grafikas ir juodraštis pakeistas. "
                            "Teorinis dublių sluoksnis vertinamas atskirai ir niekada nekeičia pageidavimų score."
                            if lang=="LT" else
                            "A better NORMAL schedule was found and saved. The theoretical backup layer is evaluated separately and never changes request scores."
                        )
                        if _cand_backup_errors:
                            st.warning("Atskirame standby dublių sluoksnyje liko neuždengtų vietų." if lang=="LT" else "The separate standby backup layer still has uncovered duties.")
                        st.rerun()
                    else:
                        st.success(
                            "Pertikrinta. Naujas normalus grafikas nebuvo geresnis pagal užfiksuotą hierarchiją, todėl esamas juodraštis paliktas."
                            if lang=="LT" else
                            "Rechecked. The new normal schedule was not better under the locked hierarchy, so the existing draft was kept."
                        )

        with c2:
            st.info(
                "SYSTEM patvirtinimas ir apsikeitimų lango atidarymas perkeltas į Grafikas → Grafiko tvirtinimas. "
                "Taip visas mėnesio lifecycle valdomas vienoje Grafiko tvirtinimo vietoje."
                if lang=="LT" else
                "SYSTEM confirmation and opening the swap window moved to Schedule → Finalization. "
                "This keeps the whole monthly lifecycle in one Schedule control center."
            )
            st.caption(("Sugeneruok / pagerink juodraštį čia, tada eik į Grafikas." if lang=="LT" else "Generate/improve the draft here, then open Schedule."))
        draftp=db.load_schedule(year,month,"draft")
        if draftp:
            dr=refresh_result_payload(draftp,year,month,use_actual_backups=False)
            g=dr.stats["global"]
            c1,c2,c3,c4=st.columns(4)
            c1.metric(tr("hard_errors")+" *",g["hard_errors"])
            fair_valid=(int(g.get("hard_errors",0))==0)
            c2.metric(tr("cumulative_fairness"),f"{g.get('cumulative_fairness_score',g['fairness_score'])}%" if fair_valid else "—")
            c3.metric(tr("monthly_fairness"),f"{g.get('monthly_fairness_score',g['fairness_score'])}%" if fair_valid else "—")
            c4.metric(tr("preference_avg"),tr("not_applicable") if g["mean_preference_score"] is None else f"{g['mean_preference_score']}%")

            # V2.5.107: every generation result immediately explains what wishes
            # were and were not achieved. Senior users should never need to infer
            # misses from a percentage alone.
            _wish=generation_wish_summary(dr)
            wa,wb,wc,wd=st.columns(4)
            wa.metric(("Aktyvūs pageidavimai" if lang=="LT" else "Active wishes"),_wish["total"])
            wb.metric(("Įvykdyta" if lang=="LT" else "Honored"),_wish["honored"])
            wc.metric(("Neįvykdyta" if lang=="LT" else "Missed"),_wish["missed"])
            wd.metric(("Negaliu dirbti pažeidimai" if lang=="LT" else "Cannot-work violations"),_wish["hard_missed"])
            if _wish["hard_missed"]:
                st.error(
                    "KRITINĖ KLAIDA: sugeneruotas juodraštis turi „Negaliu dirbti“ pažeidimą. V2.5.107 tokio juodraščio skelbti negalima."
                    if lang=="LT" else
                    "CRITICAL ERROR: the generated draft contains a Cannot-work violation. V2.5.107 must not publish such a draft."
                )
            elif _wish["missed"]==0:
                st.success(
                    "VISI AKTYVŪS PAGEIDAVIMAI ĮVYKDYTI — „Negaliu dirbti“ pažeidimų: 0."
                    if lang=="LT" else
                    "ALL ACTIVE WISHES MET — Cannot-work violations: 0."
                )
            else:
                st.warning(
                    f"Neįvykdyta {_wish['missed']} iš {_wish['total']} aktyvių pageidavimų. „Negaliu dirbti“ pažeidimų: 0. Žemiau tiksliai parodyta, kas neįvykdyta."
                    if lang=="LT" else
                    f"{_wish['missed']} of {_wish['total']} active wishes were not met. Cannot-work violations: 0. The exact misses are shown below."
                )
                st.markdown("#### NEĮVYKDYTI PAGEIDAVIMAI" if lang=="LT" else "#### UNMET WISHES")
                st.dataframe(_wish["table"],use_container_width=True,hide_index=True)

            # V2.5.112 admin guarantees: expose them instead of making SP inspect the grid.
            _dt_done=int(g.get("dream_team_centro_weeks",0) or 0)
            _dt_target=int(g.get("dream_team_centro_target_weeks",0) or 0)
            _wcap=g.get("admin_weekend_spread_cap_used")
            _bp,_be=backup_counts(year,month,dr)
            _bvals=list(_bp.values())
            _bspread=(max(_bvals)-min(_bvals)) if _bvals else 0
            _ga,_gb,_gc=st.columns(3)
            _ga.metric("Dream Team @ CENTRO RO",f"{_dt_done}/{_dt_target}")
            _gb.metric(("Mažiausias įmanomas savaitgalio spread" if lang=="LT" else "Tightest feasible weekend spread"),_wcap if _wcap is not None else "—")
            _gc.metric(("AUTO dublių spread" if lang=="LT" else "AUTO backup spread"),_bspread)
            if _dt_target and _dt_done<_dt_target:
                st.warning("Dream Team (SP + ŠR + GE) nepavyko sutalpinti į CENTRO RO kiekvieną savaitę nepažeidžiant aukštesnių taisyklių; žemiau rodomas maksimalus pasiektas kiekis." if lang=="LT" else "Dream Team (SP + ŠR + GE) could not be placed together at CENTRO RO every week without violating higher-ranked rules; the maximum achieved count is shown above.")
            else:
                st.success("Dream Team SP + ŠR + GE: CENTRO RO bent kartą kiekvieną savaitę — ĮVYKDYTA." if lang=="LT" else "Dream Team SP + ŠR + GE: together at CENTRO RO at least once every week — MET.")
            st.caption(
                (f"AUTO dubliai sukurti visoms svarbioms pozicijoms kartu su SYSTEM juodraščiu. Iš viso {sum(_bvals)} pareigų; rezidentų skaičiai: " + ", ".join(f"{i}={_bp.get(i,0)}" for i in sorted(_bp)))
                if lang=="LT" else
                (f"AUTO backups were created for the important positions together with the SYSTEM draft. Total {sum(_bvals)} duties; resident counts: " + ", ".join(f"{i}={_bp.get(i,0)}" for i in sorted(_bp)))
            )

            render_hard_error_explainer(g,lang,key_suffix=f"gen_{year}_{month}")
            st.caption(
                ("Teisingumas: 100% = idealus / beveik idealus balansas pagal postus, savaitgalius, penktadienius, doubles ir darbo dienų spread. "
                 "Rodomas score perskaičiuojamas gyvai pagal dabartinį engine."
                 if lang=="LT" else
                 "Fairness: 100% = ideal / near-ideal balance across workplaces, weekends, Fridays, doubles and weekday spread. "
                 "The score is recalculated live by the current engine.")
            )
            st.dataframe(style_schedule(schedule_grid(year,month,dr)),use_container_width=True,height=520)
            st.caption(
                "Lentelės viršuje Streamlit siūlo CSV. Žemiau visada pateikiamas ir pilnas spalvotas Excel failas."
                if lang=="LT" else
                "Streamlit offers CSV in the table toolbar. A full formatted Excel workbook is always available below as well."
            )
            render_schedule_download_buttons(
                year,month,dr,
                status_label=("SYSTEM JUODRAŠTIS" if lang=="LT" else "SYSTEM DRAFT"),
                file_prefix="SYSTEM_juodrastis" if lang=="LT" else "SYSTEM_draft",
                key_prefix="generation_draft_export",
            )

        # V2.5.22 — senior-only safe month reset.
        state_now=db.get_schedule_state(year,month)
        if state_now.get("has_draft") or state_now.get("has_published"):
            st.divider()
            with st.expander("PAVOJINGA ZONA — ištrinti ir perdaryti mėnesio grafiką", expanded=False):
                st.warning(
                    "Naudok tik jei šio mėnesio grafikas buvo sugeneruotas / paskelbtas per klaidą ir turi būti sudarytas iš naujo. "
                    "Bus pašalintas draft, paskelbtas SYSTEM/ACTUAL grafikas, šio mėnesio fairness_history, suplanuoti dubliai, "
                    "apsikeitimų užklausos ir administraciniai repair įrašai. "
                    "Rezidentų pageidavimai, HARD apribojimai, recurring pageidavimai ir pasirinkti credit redemptions LIEKA — "
                    "todėl galėsi iškart generuoti iš naujo iš tų pačių inputų."
                )
                st.caption(
                    "Saugiklis: jei šiame grafike jau buvo realiai užbaigtas dublio/pavadavimo įvykis, reset bus blokuojamas, "
                    "kad netyčia neištrintume realios darbo/credit istorijos."
                )
                confirm_text=f"RESET {year}-{month:02d}"
                typed=st.text_input(
                    f"Patvirtinimui įrašyk tiksliai: {confirm_text}",
                    key=f"reset_confirm_{year}_{month}"
                )
                reset_disabled=(typed.strip()!=confirm_text)
                if st.button(
                    "IŠTRINTI ŠIO MĖNESIO GRAFIKĄ IR PRADĖTI IŠ NAUJO",
                    type="primary",
                    disabled=reset_disabled,
                    use_container_width=True,
                    key=f"reset_month_{year}_{month}"
                ):
                    try:
                        reset_result=db.reset_month_schedule(year,month)
                        st.session_state.pop("shadow_result",None)
                        st.success(
                            f"{year}-{month:02d} grafikas ištrintas. Pageidavimai ir HARD inputai išsaugoti. "
                            "Dabar gali generuoti grafiką iš naujo."
                        )
                        st.rerun()
                    except Exception as e:
                        msg=str(e)
                        if "RESET_BLOCKED_FINAL_SCHEDULE" in msg:
                            st.error("RESET BLOKUOTAS: šio mėnesio grafikas jau patvirtintas kaip FINAL administracijai. FINAL snapshotas yra nekeičiamas.")
                        elif "RESET_BLOCKED_COMPLETED_BACKUP_ACTIVITY" in msg:
                            st.error(
                                "RESET BLOKUOTAS: šiame mėnesyje jau yra realiai užbaigtas dublio/pavadavimo įvykis. "
                                "Tokio mėnesio automatiškai trinti nesaugu."
                            )
                        elif "SENIOR_ONLY" in msg:
                            st.error("Šią funkciją gali naudoti tik seniūnės/senior paskyra.")
                        else:
                            st.error(f"Reset nepavyko: {e}")
    pos+=1
elif active_user=="ŠR":
    # ŠR has the same Sudarymas tab position, but it is rendered later as a
    # completely isolated RESEARCH SHADOW generator after the research import
    # helpers are defined.
    pos+=1

# --- Schedule ---
with tabs[pos]:
    st.subheader(f"{tr('published_schedule')} — {month_label(year,month)}")
    payload=db.load_schedule(year,month,"current")
    draft_payload=db.load_schedule(year,month,"draft")
    lifecycle=db.get_schedule_lifecycle(year,month)

    if lifecycle_operator_ui:
        st.markdown("## GRAFIKO TVIRTINIMAS" if lang=="LT" else "## SCHEDULE CONTROL")
        if is_researcher_account:
            st.info(
                "Kontingencinis valdymas aktyvus Išplėstiniame režime. Veiksmai atliekami ir audituojami kaip ŠR; SP paskyra niekada neperimama."
                if lang=="LT" else
                "Contingency control is active in Advanced mode. Actions are performed and audited as ŠR; the SP account is never impersonated."
            )

        state=str(lifecycle.get("state") or ("working" if payload else "draft"))
        now_lt=_vilnius_now()
        deadline=_parse_iso_dt(lifecycle.get("swap_deadline"))
        deadline_lt=deadline.astimezone(ZoneInfo("Europe/Vilnius")) if deadline else None
        expired=bool(state=="swap_open" and deadline_lt and now_lt>=deadline_lt)
        blockers=db.finalization_blockers_v2591(year,month)
        smtp_ok=smtp_ready(); missing_mail=_resident_email_preflight()
        prelim_start,prelim_end=preliminary_swap_window_for(year,month)
        render_operator_email_smtp_admin(active_user)
        # Refresh readiness after the operator email/SMTP panel (a successful edit reruns).
        smtp_ok=smtp_ready(); missing_mail=_resident_email_preflight()

        if state=="final":
            _workflow_card(
                "FINAL — PATEIKTA ADMINISTRACIJAI" if lang=="LT" else "FINAL — SUBMITTED TO ADMINISTRATION",
                "Administracijai skirta versija užrakinta. Įprasti, pavėluoti ir rankiniai prieš-FINAL pakeitimai uždaryti."
                if lang=="LT" else
                "The administration version is locked. Ordinary, late and pre-FINAL manual changes are closed.",
                "final"
            )
            fp=lifecycle.get("final_json")
            if fp:
                fr=refresh_result_payload(fp,year,month,use_actual_backups=True)
                _fc1,_fc2=st.columns(2)
                with _fc1:
                    st.download_button(
                        "ATSISIŲSTI FINAL EXCEL ADMINISTRACIJAI" if lang=="LT" else "DOWNLOAD FINAL EXCEL FOR ADMINISTRATION",
                        build_xlsx(year,month,fr,document_status="FINAL — ADMINISTRACIJAI" if lang=="LT" else "FINAL — FOR ADMINISTRATION",backup_rows_override=lifecycle.get("final_backups") or []),
                        file_name=f"FINAL_grafikas_{year}_{month:02d}.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                        type="primary",use_container_width=True,key=f"final_xlsx_{year}_{month}"
                    )
                with _fc2:
                    st.download_button(
                        "ATSISIŲSTI FINAL CSV" if lang=="LT" else "DOWNLOAD FINAL CSV",
                        schedule_list_df(year,month,fr).to_csv(index=False).encode("utf-8-sig"),
                        file_name=f"FINAL_grafikas_{year}_{month:02d}.csv",
                        mime="text/csv",use_container_width=True,key=f"final_csv_{year}_{month}"
                    )
                if smtp_ready() and st.button("PAKARTOTI TIK NEPAVYKUSIUS FINAL PRANEŠIMUS" if lang=="LT" else "RETRY FAILED FINAL NOTIFICATIONS",use_container_width=True,key=f"resend_final_mail_{year}_{month}"):
                    mailres=retry_failed_lifecycle_notifications(year,month,"final")
                    if mailres: st.dataframe(localized_delivery_rows(mailres),use_container_width=True,hide_index=True)
                    else: st.success("Visi FINAL pranešimai jau pristatyti." if lang=="LT" else "All FINAL notifications are already delivered.")
        else:
            if state=="draft":
                _workflow_card(
                    "SYSTEM JUODRAŠTIS — PRIVATUS" if lang=="LT" else "SYSTEM DRAFT — PRIVATE",
                    "Sugeneruotas grafikas dar nepaskelbtas rezidentams. Generavimas pats el. laiškų nesiunčia."
                    if lang=="LT" else
                    "The generated schedule has not been published to residents. Generation itself sends no email.",
                    "draft"
                )
            elif state=="working":
                _workflow_card(
                    "SYSTEM UŽŠALDYTAS — ACTUAL REDAGUOJAMAS" if lang=="LT" else "SYSTEM FROZEN — ACTUAL EDITABLE",
                    "SYSTEM išsaugotas tyrimui. SP arba ŠR gali koreguoti ACTUAL iki FINAL; rezidentams preliminarus grafikas dar nebūtinai paskelbtas."
                    if lang=="LT" else
                    "SYSTEM is preserved for research. SP or ŠR may correct ACTUAL until FINAL; the preliminary schedule does not have to be published to residents.",
                    "draft"
                )
            elif state=="swap_open":
                _workflow_card(
                    ("PRELIMINARUS PASKELBTAS — APSIKEITIMAI ATIDARYTI" if not expired else "APSIKEITIMŲ TERMINAS BAIGĖSI") if lang=="LT" else ("PRELIMINARY PUBLISHED — SWAPS OPEN" if not expired else "RESIDENT SWAP DEADLINE PASSED"),
                    ((f"Rezidentai gali kurti naujus apsikeitimo prašymus iki {deadline_lt:%Y-%m-%d %H:%M} Lietuvos laiku." if not expired else f"Nuo {deadline_lt:%Y-%m-%d %H:%M} nauji rezidentų prašymai automatiškai blokuojami. Operatorius vis dar gali koreguoti ACTUAL arba suteikti individualią pavėluotą prieigą.") if lang=="LT" else (f"Residents may create new swap requests until {deadline_lt:%Y-%m-%d %H:%M} Lithuania time." if not expired else f"New resident requests are automatically blocked after {deadline_lt:%Y-%m-%d %H:%M}. The operator may still correct ACTUAL or grant individual late access.")),
                    "swap_open" if not expired else "expired"
                )
            elif state=="swap_closed":
                _workflow_card(
                    "REZIDENTŲ APSIKEITIMAI UŽDARYTI" if lang=="LT" else "RESIDENT SWAPS CLOSED",
                    "Naujų rezidentų prašymų kurti negalima. SP/ŠR rankinis ACTUAL koregavimas lieka aktyvus iki FINAL."
                    if lang=="LT" else
                    "Residents cannot create new requests. SP/ŠR manual ACTUAL correction remains available until FINAL.",
                    "swap_closed"
                )

            # Freeze SYSTEM for private operator correction without notifying residents.
            if not payload and draft_payload:
                if st.button(
                    "UŽŠALDYTI SYSTEM IR ATIDARYTI ACTUAL KOREGAVIMĄ (BE EMAIL)" if lang=="LT" else "FREEZE SYSTEM AND OPEN ACTUAL CORRECTION (NO EMAIL)",
                    use_container_width=True,key=f"prepare_working_{year}_{month}"
                ):
                    try:
                        published=publish_system_baseline_for_swap_window(year,month)
                        if not published.get("ok"):
                            st.error(str(published.get("error"))); rows=published.get("rows")
                            if rows: st.dataframe(pd.DataFrame(rows if isinstance(rows,list) else [rows]),use_container_width=True,hide_index=True)
                            st.stop()
                        db.ensure_working_schedule_v2592(year,month)
                        st.session_state["_finalization_flash"]=("success","SYSTEM užšaldytas. ACTUAL paruoštas rankiniam koregavimui; email nesiųsti." if lang=="LT" else "SYSTEM frozen. ACTUAL is ready for manual correction; no email was sent.")
                        st.rerun()
                    except Exception as exc: st.error(str(exc))

            # Operator manual correction is permanently available until FINAL.
            if payload:
                render_operator_manual_override(year,month,payload,state)

            # V2.5.93 persistent post-override review checkpoint.
            unreviewed_manual=render_manual_override_review_checkpoint(year,month) if payload else 0

            st.divider()
            st.markdown("### Preliminarus paskelbimas ir rezidentų apsikeitimai" if lang=="LT" else "### Preliminary publication and resident swaps")
            st.caption(
                (f"Standartinis langas: {prelim_start:%Y-%m-%d %H:%M} → {prelim_end:%Y-%m-%d %H:%M} Lietuvos laiku. Pageidavimai renkami iki {prelim_start:%Y-%m-%d} 00:00 (ankstesnė diena imtinai)."
                 if lang=="LT" else
                 f"Standard window: {prelim_start:%Y-%m-%d %H:%M} → {prelim_end:%Y-%m-%d %H:%M} Lithuania time. Preferences are collected until 00:00 on {prelim_start:%Y-%m-%d} (the preceding day is the last full day).")
            )
            if not smtp_ok:
                st.error("SMTP nesukonfigūruotas — preliminaraus ar FINAL etapo aktyvuoti negalima, nes nebūtų realių pranešimų." if lang=="LT" else "SMTP is not configured — preliminary or FINAL activation is blocked because real notifications could not be sent.")
            if missing_mail:
                st.error(("Trūksta rezidentų el. pašto adresų: " if lang=="LT" else "Missing resident email addresses: ")+", ".join(missing_mail))

            if state in ("draft","working"):
                within_prelim=bool(prelim_start<=now_lt<prelim_end)
                if now_lt<prelim_start:
                    st.info((f"Preliminarų etapą bus galima aktyvuoti nuo {prelim_start:%Y-%m-%d %H:%M}." if lang=="LT" else f"The preliminary phase can be activated from {prelim_start:%Y-%m-%d %H:%M}."))
                elif now_lt>=prelim_end:
                    st.warning((f"Standartinis apsikeitimų langas šiam mėnesiui jau pasibaigė ({prelim_end:%Y-%m-%d %H:%M}). Galite pereiti tiesiai į FINAL." if lang=="LT" else f"The standard swap window for this month has already ended ({prelim_end:%Y-%m-%d %H:%M}). You may proceed directly to FINAL."))
                can_prelim=bool((payload or draft_payload) and smtp_ok and not missing_mail and within_prelim and int(unreviewed_manual)==0)
                if unreviewed_manual:
                    st.warning("Preliminarus etapas užblokuotas, kol peržiūrėsite rankinius pakeitimus." if lang=="LT" else "Preliminary publication is blocked until manual changes are reviewed.")
                if st.button(
                    "PASKELBTI PRELIMINARŲ GRAFIKĄ IR LEISTI APSIKEITIMUS" if lang=="LT" else "PUBLISH PRELIMINARY SCHEDULE AND OPEN SWAPS",
                    type="primary",use_container_width=True,disabled=not can_prelim,key=f"open_prelim_{year}_{month}"
                ):
                    try:
                        if not payload:
                            published=publish_system_baseline_for_swap_window(year,month)
                            if not published.get("ok"):
                                st.error(str(published.get("error"))); rows=published.get("rows")
                                if rows: st.dataframe(pd.DataFrame(rows if isinstance(rows,list) else [rows]),use_container_width=True,hide_index=True)
                                st.stop()
                            result_open=published["result"]
                        else:
                            result_open=refresh_result_payload(payload,year,month)
                        db.ensure_working_schedule_v2592(year,month)
                        db.open_swap_window_v2591(year,month,prelim_end.astimezone(timezone.utc).isoformat())
                        mailres=swap_window_open_emails(year,month,result_open,prelim_end,max(1,(prelim_end.date()-now_lt.date()).days))
                        failed=[x for x in mailres if x[1]!="sent"]
                        st.session_state["_finalization_flash"]=("warning" if failed else "success",(f"Preliminarus grafikas paskelbtas. Pranešimai: {len(mailres)-len(failed)}/{len(mailres)} išsiųsta." if lang=="LT" else f"Preliminary schedule published. Notifications: {len(mailres)-len(failed)}/{len(mailres)} sent."))
                        st.rerun()
                    except Exception as exc: st.error(str(exc))
            elif state=="swap_open":
                m1,m2,m3,m4=st.columns(4)
                m1.metric("Pending",blockers.get("pending_normal",0)); m2.metric("Waiting operator",blockers.get("waiting_senior_apply",0)); m3.metric("Pending backup",blockers.get("pending_backup",0)); m4.metric("Late access",blockers.get("active_late_grants",0))
                if smtp_ready() and deadline_lt and st.button("PAKARTOTI TIK NEPAVYKUSIUS PRELIMINARIUS PRANEŠIMUS" if lang=="LT" else "RETRY FAILED PRELIMINARY NOTIFICATIONS",use_container_width=True,key=f"resend_prelim_{year}_{month}"):
                    mailres=retry_failed_lifecycle_notifications(year,month,"swap_open")
                    if mailres: st.dataframe(localized_delivery_rows(mailres),use_container_width=True,hide_index=True)
                    else: st.success("Visi preliminarūs pranešimai jau pristatyti." if lang=="LT" else "All preliminary notifications are already delivered.")
                if not expired:
                    with st.expander("Uždaryti rezidentų apsikeitimus anksčiau" if lang=="LT" else "Close resident swaps early",expanded=False):
                        if st.button("UŽDARYTI DABAR" if lang=="LT" else "CLOSE NOW",use_container_width=True,key=f"close_swaps_{year}_{month}"):
                            try: db.close_swap_window_v2591(year,month); st.rerun()
                            except Exception as exc: st.error(str(exc))

            # Late access is an exception after the ordinary deadline / explicit close.
            if (state=="swap_open" and expired) or state=="swap_closed":
                st.markdown("#### Individuali pavėluota prieiga" if lang=="LT" else "#### Individual late access")
                settings_all=db.all_account_settings(); options=[pp["initials"] for pp in DEFAULT_PEOPLE]
                l1,l2,l3=st.columns(3)
                with l1: late_ini=st.selectbox("Rezidentas" if lang=="LT" else "Resident",options,key=f"late_ini_{year}_{month}")
                with l2: late_hours=int(st.number_input("Galioja valandų" if lang=="LT" else "Valid hours",min_value=1,max_value=168,value=24,step=1,key=f"late_hours_{year}_{month}"))
                with l3: late_limit=int(st.number_input("Naujų prašymų limitas" if lang=="LT" else "New-request limit",min_value=1,max_value=5,value=1,step=1,key=f"late_limit_{year}_{month}"))
                late_reason=st.text_input("Priežastis (auditui)" if lang=="LT" else "Reason (audit)",key=f"late_reason_{year}_{month}")
                late_email=str((settings_all.get(late_ini,{}) or {}).get("email") or "").strip()
                if not late_email: st.warning((f"{late_ini} neturi email nustatymuose." if lang=="LT" else f"{late_ini} has no email in settings."))
                if st.button("SUTEIKTI INDIVIDUALIĄ PRIEIGĄ" if lang=="LT" else "GRANT INDIVIDUAL ACCESS",use_container_width=True,disabled=(not smtp_ready() or not late_email or not late_reason.strip()),key=f"grant_late_{year}_{month}"):
                    try:
                        exp=_vilnius_now()+timedelta(hours=late_hours); grant=db.grant_late_swap_access_v2591(year,month,late_ini,exp.astimezone(timezone.utc).isoformat(),late_limit,late_reason); ok,detail=late_swap_access_email(year,month,grant); st.session_state["_finalization_flash"]=("success" if ok else "warning",("Prieiga suteikta ir pranešimas išsiųstas." if ok and lang=="LT" else "Access granted and notification sent." if ok else f"Prieiga suteikta, bet email nepavyko: {detail}")); st.rerun()
                    except Exception as exc: st.error(str(exc))
                grants=db.list_late_swap_access_v2591(year,month)
                for gr in grants:
                    exp=_parse_iso_dt(gr.get("expires_at")); active=bool(not gr.get("revoked_at") and exp and exp>datetime.now(timezone.utc) and int(gr.get("requests_used",0))<int(gr.get("max_requests",0)))
                    if not active: continue
                    exp_txt=exp.astimezone(ZoneInfo("Europe/Vilnius")).strftime("%Y-%m-%d %H:%M")
                    with st.container(border=True):
                        st.markdown(f"**{gr.get('initials')}** · iki {exp_txt} · {gr.get('requests_used',0)}/{gr.get('max_requests',1)}")
                        if gr.get("reason"): st.caption(str(gr.get("reason")))
                        if st.button("ATŠAUKTI PRIEIGĄ" if lang=="LT" else "REVOKE ACCESS",key=f"revoke_late_{gr.get('id')}"):
                            try: db.revoke_late_swap_access_v2591(int(gr["id"])); st.rerun()
                            except Exception as exc: st.error(str(exc))

            # FINAL is always an explicit option; the preliminary phase may be skipped.
            st.divider()
            st.markdown("### Galutinis patvirtinimas" if lang=="LT" else "### Final confirmation")
            candidate_payload=payload or draft_payload
            candidate_result=refresh_result_payload(candidate_payload,year,month,use_actual_backups=bool(payload)) if candidate_payload else None
            if candidate_result is not None:
                candidate_is_actual=bool(payload)
                render_schedule_download_buttons(
                    year,month,candidate_result,
                    status_label=(("ACTUAL — PRIEŠ FINAL" if lang=="LT" else "ACTUAL — BEFORE FINAL") if candidate_is_actual else ("SYSTEM JUODRAŠTIS — PRIEŠ TVIRTINIMĄ" if lang=="LT" else "SYSTEM DRAFT — BEFORE CONFIRMATION")),
                    file_prefix=("ACTUAL_pries_FINAL" if candidate_is_actual else "SYSTEM_juodrastis_pries_tvirtinima"),
                    key_prefix="finalization_candidate_export",
                )
            hard=int(((candidate_result.stats or {}).get("global",{}) if candidate_result else {}).get("hard_errors",999))
            blockers=db.finalization_blockers_v2591(year,month)
            unresolved=(int(blockers.get("pending_normal",0))+int(blockers.get("waiting_senior_apply",0))+int(blockers.get("pending_backup",0))+int(blockers.get("active_late_grants",0))+int(blockers.get("unreviewed_manual_overrides",0)))
            if hard!=0 or unresolved:
                st.error((f"FINAL blokuotas: HARD={hard}; pending={blockers.get('pending_normal',0)}; waiting operator={blockers.get('waiting_senior_apply',0)}; pending backup={blockers.get('pending_backup',0)}; active late={blockers.get('active_late_grants',0)}; neperžiūrėti rankiniai={blockers.get('unreviewed_manual_overrides',0)}." if lang=="LT" else f"FINAL blocked: HARD={hard}; pending={blockers.get('pending_normal',0)}; waiting operator={blockers.get('waiting_senior_apply',0)}; pending backup={blockers.get('pending_backup',0)}; active late={blockers.get('active_late_grants',0)}; unreviewed manual={blockers.get('unreviewed_manual_overrides',0)}."))
            if state=="swap_open" and not expired:
                st.warning("FINAL patvirtinimas dabar iš karto uždarys dar aktyvų rezidentų apsikeitimų langą." if lang=="LT" else "Confirming FINAL now will immediately close the still-active resident swap window.")
            confirm=st.checkbox("Patvirtinu, kad dabartinis ACTUAL grafikas yra galutinė administracijai teikiama versija." if lang=="LT" else "I confirm that the current ACTUAL schedule is the final version for administration.",key=f"final_confirm_{year}_{month}")
            final_ready=bool(candidate_payload and hard==0 and unresolved==0 and smtp_ok and not missing_mail and confirm)
            if st.button("PATVIRTINTI FINAL, PATEIKTI IR PARUOŠTI EXCEL" if lang=="LT" else "CONFIRM FINAL, SUBMIT AND PREPARE EXCEL",type="primary",use_container_width=True,disabled=not final_ready,key=f"make_final_{year}_{month}"):
                try:
                    if not payload:
                        published=publish_system_baseline_for_swap_window(year,month)
                        if not published.get("ok"):
                            st.error(str(published.get("error"))); rows=published.get("rows")
                            if rows: st.dataframe(pd.DataFrame(rows if isinstance(rows,list) else [rows]),use_container_width=True,hide_index=True)
                            st.stop()
                    db.ensure_working_schedule_v2592(year,month)
                    payload_now=db.load_schedule(year,month,"current")
                    rr=refresh_result_payload(payload_now,year,month)
                    hard_now=int((rr.stats or {}).get("global",{}).get("hard_errors",999))
                    if hard_now!=0:
                        st.error(f"FINAL validacija nepraėjo: HARD={hard_now}"); st.stop()
                    db.finalize_schedule_v2592(year,month,serialize_result(rr))
                    mails=final_schedule_emails(year,month,rr); failed=[x for x in mails if x[1]!="sent"]
                    st.session_state["_finalization_flash"]=("warning" if failed else "success",(f"FINAL užrakintas ir paruoštas administracijai. Pranešimai: {len(mails)-len(failed)}/{len(mails)} išsiųsta." if lang=="LT" else f"FINAL locked and prepared for administration. Notifications: {len(mails)-len(failed)}/{len(mails)} sent."))
                    st.rerun()
                except Exception as exc: st.error(str(exc))

        flash=st.session_state.pop("_finalization_flash",None)
        if flash:
            level,msg=flash
            if level=="success": st.success(msg)
            else: st.warning(msg)
        st.divider()

    if not payload:
        if draft_payload and not lifecycle_operator_ui: st.info(tr("not_published"))
        elif not draft_payload: st.info(tr("not_published"))
    else:
        state=str(lifecycle.get("state") or "")
        display_payload=(lifecycle.get("final_json") if state=="final" and lifecycle.get("final_json") else payload)
        result=refresh_result_payload(display_payload,year,month,use_actual_backups=(state!="final"))
        _status_rows=[]
        if lifecycle_operator_ui:
            try: _status_rows=db.list_schedule_day_statuses_v2597(year,month)
            except Exception: _status_rows=[]
        st.markdown(f"### {tr('colors')}")
        st.markdown("".join(badge(p["initials"],False) for p in DEFAULT_PEOPLE),unsafe_allow_html=True)
        st.dataframe(style_schedule(schedule_grid(year,month,result,status_rows=_status_rows if lifecycle_operator_ui else None)),use_container_width=True,height=720)
        if state=="final": st.success("Rodomas administracijai užrakintas FINAL snapshot." if lang=="LT" else "Showing the administration-locked FINAL snapshot.")

        if lifecycle_operator_ui and state!="final":
            st.markdown("### Operacinė nedarbo dienos žyma" if lang=="LT" else "### Operational non-working day marker")
            st.caption(
                "Seniūnė gali pažymėti nedarbingumą, kvalifikacijos kėlimą ar sveikatinimosi dieną. Žyma išima to rezidento ACTUAL pamainas tą dieną, bet nekeičia užšaldyto SYSTEM baseline ar mėnesio target kredito. Jei lieka kritinė skylė, ją reikia uždengti per dublį / rankinį koregavimą. Tiksli priežastis rodoma tik operatoriui."
                if lang=="LT" else
                "The scheduler can mark sick leave, qualification/training, or a wellness day. The marker removes that resident's ACTUAL assignments for the day but does not rewrite the frozen SYSTEM baseline or monthly workload credit. Any critical gap must then be covered through backup/manual correction. The exact reason is operator-only."
            )
            _kind_labels={
                "sick_leave":"Nedarbingumas" if lang=="LT" else "Sick leave",
                "qualification":"Kvalifikacijos kėlimas" if lang=="LT" else "Qualification / training",
                "wellness":"Sveikatinimosi diena" if lang=="LT" else "Wellness day",
            }
            with st.form(f"day_status_{year}_{month}"):
                _c1,_c2,_c3=st.columns(3)
                _ini=_c1.selectbox(tr("person"),[p["initials"] for p in DEFAULT_PEOPLE],format_func=lambda i:f"{i} — {people_map[i]['name']}")
                _day=_c2.selectbox(tr("date"),list(range(1,calendar.monthrange(year,month)[1]+1)),format_func=lambda d:f"{year}-{month:02d}-{d:02d}")
                _kind=_c3.selectbox("Tipas" if lang=="LT" else "Type",list(_kind_labels),format_func=lambda k:_kind_labels[k])
                _note=st.text_input("Pastaba (tik operatoriui)" if lang=="LT" else "Note (operator only)")
                _save_status=st.form_submit_button("PAŽYMĖTI NEDARBO DIENĄ" if lang=="LT" else "MARK NON-WORKING DAY",type="primary")
            if _save_status:
                try:
                    _cur_payload=db.load_schedule(year,month,"current")
                    if not _cur_payload: raise RuntimeError("ACTUAL grafikas dar nesukurtas")
                    _cur=deserialize_result(_cur_payload)
                    _slots={sl.idx:sl for sl in make_slots(year,month)}
                    _removed=[]
                    for _sid,_who in list((_cur.assignments or {}).items()):
                        _sl=_slots.get(int(_sid))
                        if _who==_ini and _sl is not None and int(_sl.day)==int(_day):
                            _removed.append(int(_sid)); _cur.assignments.pop(int(_sid),None)
                    db.set_schedule_day_status_v2597(year,month,int(_day),_ini,_kind,_note)
                    _fresh=revalidate_loaded_result(year,month,people_for_stored_result(_cur,year,month),_cur,backup_assignments=db.list_backups(year,month),validation_mode="voluntary_swap_actual")
                    db.save_current(year,month,serialize_result(_fresh))
                    sync_backup_plan(year,month,_fresh)
                    persist_actual_satisfaction(year,month)
                    try: refresh_calendar_subscription_feeds([_ini])
                    except Exception: pass
                    st.success((f"Pažymėta. Iš ACTUAL pašalinta pamainų: {len(_removed)}." if lang=="LT" else f"Marked. ACTUAL assignments removed: {len(_removed)}."))
                    st.rerun()
                except Exception as exc:
                    st.error(str(exc))
            _status_rows=db.list_schedule_day_statuses_v2597(year,month)
            if _status_rows:
                _display=[]
                for _r in _status_rows:
                    _display.append({
                        "ID":_r.get("id"),tr("person"):_r.get("initials"),tr("date"):f"{year}-{month:02d}-{int(_r.get('day')):02d}",
                        ("Tipas" if lang=="LT" else "Type"):_kind_labels.get(_r.get("status_kind"),_r.get("status_kind")),
                        tr("comment"):_r.get("note") or "",
                    })
                st.dataframe(style_rows(pd.DataFrame(_display)),use_container_width=True,hide_index=True)
                _ids=[int(r.get("id")) for r in _status_rows]
                _del=st.selectbox("Pašalinti žymą" if lang=="LT" else "Remove marker",[None]+_ids,format_func=lambda x:"—" if x is None else f"#{x}")
                if st.button("PAŠALINTI ŽYMĄ" if lang=="LT" else "REMOVE MARKER",disabled=_del is None):
                    db.clear_schedule_day_status_v2597(int(_del))
                    st.warning("Žyma pašalinta, tačiau anksčiau iš ACTUAL pašalintos pamainos automatiškai negrąžinamos." if lang=="LT" else "Marker removed; previously removed ACTUAL assignments are not restored automatically.")
                    st.rerun()
        if state!="final":
            st.caption("Tai dabartinis ACTUAL grafikas. Administracijai skirtas FINAL Excel atsiras tik po galutinio operatoriaus patvirtinimo." if lang=="LT" else "This is the current ACTUAL schedule. The administration FINAL Excel appears only after final operator confirmation.")
            render_schedule_download_buttons(
                year,month,result,
                status_label="ACTUAL",
                file_prefix="ACTUAL_grafikas",
                key_prefix="actual_schedule_export",
            )
pos+=1

# --- Summary ---
if advanced_mode:
    with tabs[pos]:
        st.subheader(tr("summary_title"))
        currentp=db.load_schedule(year,month,"current")
        basep=db.load_schedule(year,month,"baseline")
        draftp=db.load_schedule(year,month,"draft")
        # A published row keeps draft_json. Treat it as a pending candidate only
        # when it differs from the frozen publication baseline. This lets a senior
        # generate a new candidate while an older schedule is still published and
        # inspect the NEW draft before replacing the operational baseline.
        pending_draft=bool(draftp and (not basep or draftp!=basep))
        draft_mode=bool(draftp and (not currentp or pending_draft))

        if draft_mode:
            base=refresh_result_payload(draftp,year,month,use_actual_backups=False)
            current=base
            g=base.stats["global"]
            system_live=live_fairness_snapshot(year,month,base,include_completed_covers=False)
            actual_live=None
            st.warning(
                "JUODRAŠČIO SUVESTINĖ — DAR NEPASKELBTA. Visi žemiau esantys rodikliai priklauso naujausiam sugeneruotam kandidatui; PASKELBTAS GRAFIKAS dar nepakeistas."
                if lang=="LT" else
                "DRAFT SUMMARY — NOT PUBLISHED. All metrics below belong to the newest generated candidate; the PUBLISHED schedule has not been replaced yet."
            )
        elif currentp:
            current=refresh_result_payload(currentp,year,month)
            base=refresh_result_payload(basep or currentp,year,month,use_actual_backups=False)
            g=base.stats["global"]
            system_live=live_fairness_snapshot(year,month,base,include_completed_covers=False)
            actual_live=live_fairness_snapshot(year,month,current,include_completed_covers=True)
            st.success("SYSTEM + ACTUAL SUVESTINĖ — PASKELBTA" if lang=="LT" else "SYSTEM + ACTUAL SUMMARY — PUBLISHED")
        else:
            st.info(
                "Dar nėra nei sugeneruoto juodraščio, nei paskelbto grafiko. Pirmiausia Sudarymas lange paspausk GENERUOTI."
                if lang=="LT" else
                "There is no generated draft or published schedule yet. First press GENERATE in the Generation tab."
            )
            base=None; current=None; g=None; system_live=None; actual_live=None

        if base is not None:
            if active_user in (SENIOR_INITIALS,WESTON_CREDITOR_INITIALS):
                try:
                    _weston_personal=db.weston_beer_stats_v25110(year,month)
                    st.markdown("### WESTON beer ledger")
                    _wb1,_wb2=st.columns(2)
                    if active_user==SENIOR_INITIALS:
                        _wb1.metric(("Tavo WESTON skola ŠR" if lang=="LT" else "Your WESTON debt to ŠR"),int(_weston_personal.get("total_beers",0)))
                        _wb2.metric(("Šio mėnesio skola" if lang=="LT" else "Debt this month"),int(_weston_personal.get("month_beers",0)))
                        st.caption("1 GENERUOTI / PERKURTI paspaudimas = +1 WESTON, kurį esi skolinga ŠR. Rosita, matematikos neapgausi." if lang=="LT" else "1 GENERATE / REGENERATE click = +1 WESTON you owe ŠR. Rosita, mathematics keeps receipts.")
                    else:
                        _wb1.metric(("SP tau skolinga WESTON" if lang=="LT" else "WESTONs SP owes you"),int(_weston_personal.get("total_beers",0)))
                        _wb2.metric(("Šį mėnesį uždirbai" if lang=="LT" else "Earned this month"),int(_weston_personal.get("month_beers",0)))
                        st.caption("Kiekvienas SP GENERUOTI / PERKURTI paspaudimas = +1 WESTON tau. Skola ir tavo laimėjimas skaičiuojami iš to paties nekintamo ledgerio." if lang=="LT" else "Every SP GENERATE / REGENERATE click = +1 WESTON owed to you. Her debt and your gain are the same persistent ledger.")
                except Exception:
                    pass
            if advanced_mode:
                if draft_mode:
                    c1,c2,c3,c4=st.columns(4)
                    c1.metric(tr("hard_errors")+" *",g["hard_errors"])
                    c2.metric(("JUODRAŠČIO fairness" if lang=="LT" else "DRAFT fairness"),f"{system_live['global'].get('monthly_fairness_score',0)}%")
                    c3.metric(("Postų disbalansas" if lang=="LT" else "Post imbalance"),system_live["global"].get("rotation_monthly_imbalance",0))
                    c4.metric(tr("preference_avg"),tr("not_applicable") if g["mean_preference_score"] is None else f"{g['mean_preference_score']}%")
                else:
                    sg=system_live["global"]; ag=actual_live["global"]
                    c1,c2,c3,c4,c5,c6=st.columns(6)
                    c1.metric(tr("hard_errors")+" *",g["hard_errors"])
                    c2.metric("SYSTEM fairness",f"{sg.get('monthly_fairness_score',0)}%")
                    c3.metric("ACTUAL fairness",f"{ag.get('monthly_fairness_score',0)}%",delta=f"{ag.get('monthly_fairness_score',0)-sg.get('monthly_fairness_score',0):+.1f}")
                    c4.metric(("SYSTEM postų disbalansas" if lang=="LT" else "SYSTEM post imbalance"),sg.get("rotation_monthly_imbalance",0))
                    c5.metric(("ACTUAL postų disbalansas" if lang=="LT" else "ACTUAL post imbalance"),ag.get("rotation_monthly_imbalance",0))
                    c6.metric(("Realiai pavadavo" if lang=="LT" else "Completed covers"),ag.get("completed_cover_transfers",0))
                if draft_mode:
                    st.caption(
                        "JUODRAŠTIS: fairness, postų spread ir pageidavimų score yra pre-publication auditui. Jie taps SYSTEM baseline tik paspaudus PASKELBTI / PATVIRTINTI."
                        if lang=="LT" else
                        "DRAFT: fairness, workplace spread and request scores are pre-publication audit metrics. They become the SYSTEM baseline only after PUBLISH / CONFIRM."
                    )
                else:
                    st.caption(tr("fairness_frozen_note"))
            else:
                s1,s2,s3=st.columns(3)
                s1.metric(("Grafikas" if lang=="LT" else "Schedule"),("VALIDUS" if g["hard_errors"]==0 else "REIKIA PATIKROS"))
                s2.metric(("Pageidavimai" if lang=="LT" else "Preferences"),tr("not_applicable") if g["mean_preference_score"] is None else f"{g['mean_preference_score']}%")
                s3.metric(("Mėnesio fairness" if lang=="LT" else "Monthly fairness"),f"{g.get('monthly_fairness_score',g['fairness_score'])}%")
                st.caption(
                    ("Techninius spread, cumulative ir guardrail rodiklius rasi Išplėstiniame režime."
                     if lang=="LT" else
                     "Technical spread, cumulative and guardrail metrics are available in Advanced mode.")
                )

            render_resident_wishes_audit(
                base,
                draft_mode=draft_mode,
                key_suffix=f"{year}_{month}_{'draft' if draft_mode else 'system'}",
                senior_view=bool(senior_mode),
            )
            st.divider()

            st.markdown(
                ("### JUODRAŠČIO postų pasiskirstymas per mėnesį" if draft_mode else "### SYSTEM postų pasiskirstymas per mėnesį")
                if lang=="LT" else
                ("### DRAFT monthly workplace distribution" if draft_mode else "### SYSTEM monthly workplace distribution")
            )
            st.caption(
                (
                    "Tai naujausio JUODRAŠČIO postų matrica. Ji leidžia prieš publikavimą patikrinti structural water-fill, SPS/Onko ir diversity. Regeneravus ji gali pasikeisti; fairness_history dar neįrašoma."
                    if draft_mode else
                    "SYSTEM matrica yra publikavimo momento algoritmo water-fill baseline. Ji lieka užšaldyta auditui; realus pasiskirstymas rodomas ACTUAL matricoje žemiau."
                )
                if lang=="LT" else
                (
                    "This is the newest DRAFT workplace matrix. It can be audited for structural water-fill, SPS/Onko and diversity before publication. Regeneration may change it; fairness_history is not written yet."
                    if draft_mode else
                    "The SYSTEM matrix is the publication-time algorithmic water-fill baseline. It stays frozen for audit; the real distribution is shown in the ACTUAL matrix below."
                )
            )
            st.dataframe(
                style_rows(workplace_exposure_df(year,month,base)),
                use_container_width=True,
                hide_index=True,
                height=650,
            )

            if not draft_mode:
                st.markdown("### ACTUAL realus postų pasiskirstymas" if lang=="LT" else "### ACTUAL real-work workplace distribution")
                st.caption(
                    "Tai gyva fairness statistika: manual override'ai, swapai ir repair keičia ACTUAL iškart; backup cover perskiriamas dubliui tik pažymėjus COMPLETED. Šie skirtumai niekada nekuria kito mėnesio catch-up."
                    if lang=="LT" else
                    "This is the live fairness ledger: manual overrides, swaps and repairs change ACTUAL immediately; backup exposure transfers only when cover is marked COMPLETED. These differences never create next-month catch-up."
                )
                st.dataframe(style_rows(actual_workplace_exposure_df(year,month,current)),use_container_width=True,hide_index=True,height=650)

            if advanced_mode:
                spread_rows=[]
                sys_spreads=(system_live.get("global",{}).get("rotation_monthly_spreads") or {})
                act_spreads=((actual_live or system_live).get("global",{}).get("rotation_monthly_spreads") or {})
                for cat in ROTATION_CATEGORIES:
                    spread_rows.append({
                        ("Postas" if lang=="LT" else "Workplace"):cat,
                        "SYSTEM spread":sys_spreads.get(cat,0),
                        "ACTUAL spread":act_spreads.get(cat,0),
                        ("Pokytis" if lang=="LT" else "Delta"):int(act_spreads.get(cat,0))-int(sys_spreads.get(cat,0)),
                    })
                st.markdown("#### SYSTEM → ACTUAL postų spread" if lang=="LT" else "#### SYSTEM → ACTUAL workplace spread")
                st.caption(
                    "SYSTEM yra mėnesio water-fill baseline. ACTUAL gali nukrypti po leidžiamų pakeitimų; nukrypimas rodomas, bet kitą mėnesį nekompensuojamas."
                    if lang=="LT" else
                    "SYSTEM is the monthly water-fill baseline. ACTUAL may diverge after allowed changes; the divergence is shown but never compensated next month."
                )
                st.dataframe(pd.DataFrame(spread_rows),use_container_width=True,hide_index=True)

                guardrails=(g.get("fairness_guardrails") or {})
                if guardrails:
                    gr_rows=[]
                    for key,val in guardrails.items():
                        gr_rows.append({
                            ("Guardrail" if lang=="LT" else "Guardrail"):key,
                            ("Fairness baseline spread" if lang=="LT" else "Fairness baseline spread"):val.get("baseline_spread"),
                            ("Leistinas pablogėjimas" if lang=="LT" else "Allowed degradation"):val.get("tolerance"),
                            ("Maksimalus spread po SOFT" if lang=="LT" else "Max spread after SOFT"):val.get("ceiling"),
                        })
                    st.markdown("#### Fairness guardrails")
                    st.dataframe(pd.DataFrame(gr_rows),use_container_width=True,hide_index=True)
                    st.caption(
                        ("Šios ribos realiai įdėtos į solverio modelį prieš TRUE SOFT optimizavimą."
                         if lang=="LT" else
                         "These limits are real solver constraints added before TRUE SOFT optimization.")
                    )

                with st.expander(
                    ("JUODRAŠČIO krūvio rodikliai" if draft_mode else "SYSTEM fairness krūvio rodikliai")
                    if lang=="LT" else
                    ("DRAFT workload metrics" if draft_mode else "SYSTEM fairness workload metrics"),
                    expanded=False,
                ):
                    st.caption(
                        ("Šie skaičiai priklauso dabartiniam juodraščiui ir skirti auditui prieš publikavimą." if draft_mode else "Šie skaičiai priklauso publikavimo SYSTEM bazei; post-publication ACTUAL pakeitimai jų nekeičia, nes tai yra būtent SYSTEM publikavimo baseline.")
                        if lang=="LT" else
                        ("These figures belong to the current draft and are for pre-publication audit." if draft_mode else "These figures belong to the publication SYSTEM baseline; post-publication ACTUAL changes do not change them because these figures are the publication-time SYSTEM baseline.")
                    )
                    _summary_stats=summary_df(base,year,month)
                    if draft_mode:
                        # Backup obligations are finalized at publish time; hide DB-backed
                        # backup columns here so an older published month's backup rows
                        # cannot be mistaken for draft data.
                        _drop=[tr("planned_backups"),tr("effective_backups")]
                        _summary_stats=_summary_stats.drop(columns=[c for c in _drop if c in _summary_stats.columns],errors="ignore")
                    st.dataframe(style_rows(_summary_stats),use_container_width=True,hide_index=True)
                if (not draft_mode) and current.assignments != base.assignments:
                    with st.expander("FAKTINIAI operaciniai krūvio rodikliai (NE fairness)" if lang=="LT" else "ACTUAL operational workload metrics (NOT fairness)", expanded=False):
                        st.dataframe(style_rows(summary_df(current,year,month)),use_container_width=True,hide_index=True)
            else:
                with st.expander(("Rodyti paprastą mėnesio krūvio santrauką" if lang=="LT" else "Show simple monthly workload summary"), expanded=False):
                    simple_df=summary_df(base,year,month)
                    keep=[c for c in simple_df.columns if c in ["Žmogus","Vardas","Pamainos","Tikslas","Person","Name","Assignments","Target"]]
                    st.dataframe(simple_df[keep] if keep else simple_df,use_container_width=True,hide_index=True)
            if draft_mode:
                st.info(
                    "Jei šita Suvestinė netenkina: grįžk į Sudarymas → PERTIKRINTI / GERINTI arba GENERUOTI iš naujo. PASKELBTAS GRAFIKAS nepasikeis, kol aiškiai nepaspausi PASKELBTI / PATVIRTINTI."
                    if lang=="LT" else
                    "If this Summary is not satisfactory: return to Generation → IMPROVE / RECHECK or GENERATE again. The PUBLISHED schedule does not change until you explicitly press PUBLISH / CONFIRM."
                )
    pos+=1

# --- Transparency ---
if advanced_mode:
    with tabs[pos]:
        st.subheader(tr("transparency_title")); currentp=db.load_schedule(year,month,"current"); basep=db.load_schedule(year,month,"baseline")
        if not currentp: st.info(tr("not_published"))
        else:
            current=refresh_result_payload(currentp,year,month); base=refresh_result_payload(basep or currentp,year,month,use_actual_backups=False)
            g=base.stats["global"]; gb=base.stats["global"]
            system_live=live_fairness_snapshot(year,month,base,include_completed_covers=False)
            actual_live=live_fairness_snapshot(year,month,current,include_completed_covers=True)
            sg=system_live["global"]; ag=actual_live["global"]

            # V2.5.114 — one persistent WESTON ledger, mirrored as debt for SP
            # and receivable/gain for ŠR so both sides see the same running total.
            if active_user in (SENIOR_INITIALS,WESTON_CREDITOR_INITIALS):
                try:
                    _weston_mirror=db.weston_beer_stats_v25110(year,month)
                    st.markdown("### WESTON")
                    _wm1,_wm2=st.columns(2)
                    if active_user==SENIOR_INITIALS:
                        _wm1.metric(("Skola ŠR — iš viso" if lang=="LT" else "Debt to ŠR — lifetime"),int(_weston_mirror.get("total_beers",0)))
                        _wm2.metric(("Skola ŠR — šį mėnesį" if lang=="LT" else "Debt to ŠR — this month"),int(_weston_mirror.get("month_beers",0)))
                        st.caption("SP: kiekvienas tavo GENERUOTI / PERKURTI paspaudimas prideda +1 WESTON skolą ŠR." if lang=="LT" else "SP: every GENERATE / REGENERATE click adds +1 WESTON owed to ŠR.")
                    else:
                        _wm1.metric(("SP tau skolinga — iš viso" if lang=="LT" else "SP owes you — lifetime"),int(_weston_mirror.get("total_beers",0)))
                        _wm2.metric(("Tavo WESTON prieaugis — šį mėnesį" if lang=="LT" else "Your WESTON gain — this month"),int(_weston_mirror.get("month_beers",0)))
                        st.caption("ŠR: SP paspaudžia GENERUOTI / PERKURTI → tau +1 WESTON. Tas pats skaičius rodomas SP kaip skola." if lang=="LT" else "ŠR: SP presses GENERATE / REGENERATE → +1 WESTON owed to you. The same number is shown to SP as debt.")
                except Exception:
                    pass

            st.markdown(f"### {tr('fairness_hierarchy')}")
            st.caption(tr("fairness_hierarchy_intro"))
            h1,h2,h3,h4=st.columns(4)
            h1.metric(tr("hard_validity"),tr("hard_validity_pass") if g["hard_errors"]==0 else tr("hard_validity_fail"))
            h2.metric("SYSTEM fairness",f"{sg.get('monthly_fairness_score',0)}%")
            h3.metric("ACTUAL fairness",f"{ag.get('monthly_fairness_score',0)}%",delta=f"{ag.get('monthly_fairness_score',0)-sg.get('monthly_fairness_score',0):+.1f}")
            h4.metric(("ACTUAL postų disbalansas" if lang=="LT" else "ACTUAL post imbalance"),ag.get("rotation_monthly_imbalance",0))

            if advanced_mode:
                hierarchy_df=pd.DataFrame([
                    {tr("fairness_level"):"1. ABSOLUTE HARD",tr("fairness_goal"):tr("fairness_hard_goal"),tr("fairness_interpretation"):tr("hard_validity_pass") if g["hard_errors"]==0 else tr("hard_validity_fail")},
                    {tr("fairness_level"):"2. RESIDENT HARD",tr("fairness_goal"):("0 „Negaliu dirbti“ pažeidimų privaloma; jokio trade-off su fairness ar SOFT" if lang=="LT" else "Zero Unavailable violations are mandatory; never trade them for fairness or SOFT"),tr("fairness_interpretation"):(f"Šį mėnesį praradimų: {g.get('resident_hard_total_losses',0)}; paveikta rezidentų: {g.get('resident_hard_residents_affected',0)}; max vienam: {g.get('resident_hard_max_loss_per_resident',0)}; cumulative spread: {g.get('resident_hard_cumulative_spread',0)}" if lang=="LT" else f"Losses this month: {g.get('resident_hard_total_losses',0)}; residents affected: {g.get('resident_hard_residents_affected',0)}; max per resident: {g.get('resident_hard_max_loss_per_resident',0)}; cumulative spread: {g.get('resident_hard_cumulative_spread',0)}")},
                    {tr("fairness_level"):"3. POSTAI / WORKLOAD / FATIGUE",tr("fairness_goal"):tr("fairness_monthly_goal"),tr("fairness_interpretation"):tr("fairness_monthly_explain")},
                    {tr("fairness_level"):"4. MAX-MIN SOFT",tr("fairness_goal"):tr("other_preferences_goal"),tr("fairness_interpretation"):tr("other_preferences_explain")},
                    {tr("fairness_level"):"5. ACTUAL PO PUBLIKAVIMO",tr("fairness_goal"):("Swapai / override'ai gali pralaužti water-fill; realus spread tiksliai rodomas" if lang=="LT" else "Swaps/overrides may break water-fill; the real spread is reported exactly"),tr("fairness_interpretation"):tr("fairness_cumulative_explain")},
                ])
                st.dataframe(hierarchy_df,use_container_width=True,hide_index=True)
                st.caption(tr("fairness_100_note"))
                ledger_df=pd.DataFrame([
                    {tr("fairness_scope"):tr("fairness_ledger"),tr("fairness_interpretation"):tr("fairness_swap_neutral")},
                    {tr("fairness_scope"):tr("actual_ledger"),tr("fairness_interpretation"):tr("swap_note")},
                ])
                st.dataframe(ledger_df,use_container_width=True,hide_index=True)

                st.divider(); st.markdown(f"### {tr('fairness_breakdown')}")
                breakdown=[
                    ("SYSTEM",tr("metric_saturday"),sg.get("saturday_monthly_spread",0)),
                    ("ACTUAL",tr("metric_saturday"),ag.get("saturday_monthly_spread",0)),
                    ("SYSTEM",tr("metric_sunday"),sg.get("sunday_monthly_spread",0)),
                    ("ACTUAL",tr("metric_sunday"),ag.get("sunday_monthly_spread",0)),
                    ("SYSTEM",tr("metric_friday"),sg.get("friday_monthly_spread",0)),
                    ("ACTUAL",tr("metric_friday"),ag.get("friday_monthly_spread",0)),
                    ("SYSTEM",tr("metric_double"),sg.get("double_monthly_spread",0)),
                    ("ACTUAL",tr("metric_double"),ag.get("double_monthly_spread",0)),
                    ("SYSTEM",tr("metric_weekday"),sg.get("weekday_day_monthly_spread",0)),
                    ("ACTUAL",tr("metric_weekday"),ag.get("weekday_day_monthly_spread",0)),
                ]
                st.dataframe(pd.DataFrame([{tr("fairness_scope"):scope,tr("fairness_metric"):metric,tr("fairness_spread"):spread} for scope,metric,spread in breakdown]),use_container_width=True,hide_index=True)
                st.caption(tr("fairness_monthly_explain"))
                st.caption(tr("fairness_cumulative_explain"))

                st.divider(); st.markdown(f"### {tr('fairness_history')}")
                st.caption(tr("fairness_history_help"))
                trend=system_actual_fairness_trend_df(year,month)
                if trend.empty:
                    st.caption(tr("fairness_no_history"))
                else:
                    chart=trend.set_index("Period")
                    st.line_chart(chart[["SYSTEM monthly fairness","ACTUAL monthly fairness"]],height=280)
                    st.dataframe(trend,use_container_width=True,hide_index=True)
                    st.caption("Istorija tik stebėjimui — solveris jos nenaudoja kitam mėnesiui." if lang=="LT" else "History is monitoring-only — the solver never uses it for the next month.")

                st.divider(); st.markdown(f"### {tr('personal_vs_group')}")
                bp=base.stats["people"].get(active_user,{}).get("preference_score"); cp=current.stats["people"].get(active_user,{}).get("preference_score")
                ratio=balance_ratio(cp,ag.get("monthly_fairness_score"))
                a,b,c=st.columns(3); a.metric(tr("baseline_personal"),tr("not_applicable") if bp is None else f"{bp}%"); b.metric(tr("current_personal"),tr("not_applicable") if cp is None else f"{cp}%"); c.metric(tr("balance_ratio"),tr("not_applicable") if ratio is None else f"{ratio:.2f}")
                st.caption(tr("ratio_help")); st.markdown(f"### {tr('all_resident_scores')}"); st.dataframe(style_rows(preference_scores_df(current)),use_container_width=True,hide_index=True)

                # V2.5.49 resident request ledger: exact missed type/date/block/station,
                # so the resident immediately knows what kind of swap could repair it.
                pd_base=(base.stats.get("people",{}).get(active_user,{}) or {})
                pd_now=(current.stats.get("people",{}).get(active_user,{}) or {})
                st.divider()
                st.markdown("### Mano pageidavimų išpildymas — detalės" if lang=="LT" else "### My request satisfaction — details")
                st.caption(
                    ("SYSTEM = paskelbimo momentas ir užšaldytas ORIGINAL pageidavimų rinkinys. ACTUAL = dabartinis realus grafikas po apsikeitimų; pageidavimų išpildymas perskaičiuojamas prieš tą patį ORIGINAL rinkinį, todėl vėliau galima patikimai palyginti paskelbtą ir galutinį mėnesio rezultatą."
                     if lang=="LT" else
                     "SYSTEM uses the frozen ORIGINAL request set at publication. ACTUAL is the current real schedule after swaps; satisfaction is recalculated against that same ORIGINAL set, allowing a valid published-versus-final retrospective comparison.")
                )
                r1,r2,r3,r4=st.columns(4)
                r1.metric("RESIDENT HARD — SYSTEM",tr("not_applicable") if pd_base.get("resident_hard_score") is None else f"{pd_base.get('resident_hard_score')}%")
                r2.metric("RESIDENT HARD — ACTUAL",tr("not_applicable") if pd_now.get("resident_hard_score") is None else f"{pd_now.get('resident_hard_score')}%")
                r3.metric("SOFT — ACTUAL",tr("not_applicable") if pd_now.get("soft_preference_score") is None else f"{pd_now.get('soft_preference_score')}%")
                r4.metric(("VISI PRAŠYMAI — ACTUAL" if lang=="LT" else "ALL REQUESTS — ACTUAL"),tr("not_applicable") if pd_now.get("overall_request_score") is None else f"{pd_now.get('overall_request_score')}%")

                hard_misses=pd_now.get("resident_hard_conflicts") or []
                soft_misses=pd_now.get("soft_request_misses") or []
                if hard_misses:
                    st.error((f"RESIDENT HARD: neįvykdyta {len(hard_misses)} tavo „Negaliu dirbti“ prašymų. Kiekviena eilutė dabar skaitoma kaip paprastas sakinys: **ko prašei → ką sistema paskyrė → kodėl laikoma neįvykdyta → kaip tai patikrinti grafike → kokio swapo ieškoti**."
                              if lang=="LT" else
                              f"RESIDENT HARD: {len(hard_misses)} Unavailable request(s) were not honored. Each row now states: **request → schedule result → why it is a miss → how to verify it → what swap to look for**."))
                    if lang=="LT":
                        st.caption("Pavyzdys: „Noriu laisvos 18 d. PM“ + grafike yra „SPS UG PM“ = pageidavimas NEĮVYKDYTAS. Jei atsidaręs 18 d. grafike SPS UG PM nematai, lentelės teiginys yra klaida.")
                    st.dataframe(request_details_df(hard_misses,active_user),use_container_width=True,hide_index=True)
                else:
                    st.success("Visi tavo RESIDENT HARD prašymai išpildyti." if lang=="LT" else "All of your RESIDENT HARD requests are honored.")

                if soft_misses:
                    st.markdown("#### Neįvykdyti SOFT pageidavimai" if lang=="LT" else "#### Unhonored SOFT requests")
                    st.dataframe(request_details_df(soft_misses,active_user),use_container_width=True,hide_index=True)
                else:
                    st.caption("Neįvykdytų struktūruotų SOFT pageidavimų nėra." if lang=="LT" else "There are no unhonored structured SOFT requests.")

                with st.expander("Išpildyti pageidavimai" if lang=="LT" else "Honored requests"):
                    honored=pd_now.get("honored_request_details") or []
                    if honored:
                        st.dataframe(request_details_df(honored,active_user),use_container_width=True,hide_index=True)
                    else:
                        st.caption(tr("not_applicable"))

                st.markdown("#### Grupės RESIDENT HARD našta" if lang=="LT" else "#### Group RESIDENT HARD burden")
                rr1,rr2,rr3,rr4=st.columns(4)
                rr1.metric(("Pažeidimų iš viso" if lang=="LT" else "Total violations"),g.get("resident_hard_total_losses",0))
                rr2.metric(("Paveikta rezidentų" if lang=="LT" else "Residents affected"),g.get("resident_hard_residents_affected",0))
                rr3.metric(("Max pažeidimų vienam" if lang=="LT" else "Max violations per resident"),g.get("resident_hard_max_loss_per_resident",0))
                rr4.metric(("Cumulative spread" if lang=="LT" else "Cumulative spread"),g.get("resident_hard_cumulative_spread",0))
            else:
                bp=base.stats["people"].get(active_user,{}).get("preference_score")
                cp=current.stats["people"].get(active_user,{}).get("preference_score")
                st.caption(
                    ("Čia rodoma tik trumpa santrauka. Pilnas fairness breakdown, istorija ir visų rezidentų palyginimas yra Išplėstiniame režime."
                     if lang=="LT" else
                     "Only a short summary is shown here. Full fairness breakdown, history and all-resident comparison are in Advanced mode.")
                )
                p1,p2=st.columns(2)
                p1.metric(("Tavo pageidavimų išpildymas" if lang=="LT" else "Your preference fulfillment"),tr("not_applicable") if cp is None else f"{cp}%")
                p2.metric(("Mėnesio fairness" if lang=="LT" else "Monthly fairness"),f"{g.get('monthly_fairness_score',g['fairness_score'])}%")
    pos+=1

# --- Credits ---
if advanced_mode:
    with tabs[pos]:
        st.subheader(tr("credit_balances"))
        st.caption(tr("netting_explain"))
        rest_bank=db.rest_credit_balances(active_user)
        rb1,rb2,rb3=st.columns(3)
        rb1.metric(tr("credit_am"),rest_bank.get("AM",0))
        rb2.metric(tr("credit_pm"),rest_bank.get("PM",0))
        rb3.metric(tr("credit_night"),rest_bank.get("NIGHT",0))
        st.caption(tr("credit_month_cap")); st.caption(tr("night_bank_only"))
        if senior_mode:
            st.divider(); st.markdown(f"### {tr('credit_balances')} — {tr('summary')}")
            rest_all=db.all_rest_credit_balances()
            rows=[]
            for person in DEFAULT_PEOPLE:
                i=person["initials"]
                rows.append({
                    tr("person"):i,
                    f"{tr('rest_credit_bank')} AM":rest_all.get(i,{}).get("AM",0),
                    f"{tr('rest_credit_bank')} PM":rest_all.get(i,{}).get("PM",0),
                    f"{tr('rest_credit_bank')} NIGHT":rest_all.get(i,{}).get("NIGHT",0),
                })
            st.dataframe(pd.DataFrame(rows),use_container_width=True,hide_index=True)
        pos+=1

# --- Backups ---
with tabs[pos]:
    st.subheader(tr("backup_title")); st.write(tr("backup_definition")); currentp=db.load_schedule(year,month,"current")
    all_backup_slots=[s for s in make_slots(year,month) if backup_required_slot(s)]
    weekend_slots=[s for s in all_backup_slots if s.weekday >= 5]
    centro120_slots=[s for s in all_backup_slots if s.weekday < 5 and s.department.startswith("Centro UG 120") and s.block=="AM"]
    onko_slots=[s for s in all_backup_slots if s.weekday < 5 and s.department=="Onko RO centre"]
    sps_ro_slots=[s for s in all_backup_slots if s.weekday < 5 and s.department.startswith("SPS RO")]
    sps_ug_slots=[s for s in all_backup_slots if s.weekday < 5 and s.department.startswith("SPS UG")]

    claims=db.list_backup_claims(year,month)
    claim_by_slot={int(r["covered_slot"]):r for r in claims}
    my_claims=[r for r in claims if r["initials"]==active_user]
    my_claim_ids={int(r["covered_slot"]) for r in my_claims}
    claim_deadline=deadline_for(year,month)
    can_self_claim=(date.today() <= claim_deadline and not currentp and resident_ok)

    st.markdown("### Mano dublių rezervacijos" if lang=="LT" else "### My backup reservations")
    st.caption(
        ("Rezervuojami privalomo dengimo dubliai pagal poziciją: SPS RO bet kurią dieną / bloką, SPS UG bet kurią dieną / bloką, Centro UG 120 rytas ir Onko RO pilna diena. "
         "CENTRO RO dengiama automatiškai best-effort ir jos rezervuoti nereikia. Gali pasirinkti kelis slotus; rezervacija blokuoja persidengiančią normalią pamainą."
         if lang=="LT" else
         "Reservable mandatory backup groups are position-based: SPS RO on any day/block, SPS UG on any day/block, Centro UG 120 morning, and full-day Onko RO. "
         "CENTRO RO is planned automatically as best-effort and does not need self-reservation. You may choose multiple slots; a reservation blocks overlapping normal work.")
    )
    st.metric(tr("backup_claim_deadline"),claim_deadline.isoformat())

    slot_lookup={s.idx:s for s in all_backup_slots}
    def claim_label(sid):
        s=slot_lookup[int(sid)]
        owner=claim_by_slot.get(int(sid),{}).get("initials")
        status=owner if owner else tr("backup_claim_free")
        return f"{s.day:02d} {WEEKDAYS[lang][s.weekday]} · {s.department} · {block_label(s.block)} · {status}"

    if my_claim_ids:
        st.success(
            ("Rezervuota: " if lang=="LT" else "Reserved: ")
            + str(len(my_claim_ids))
            + (" dublių slotų" if lang=="LT" else " backup slots")
        )
    else:
        st.warning(
            "Dar nepasirinkai nė vieno dublio sloto." if lang=="LT"
            else "You have not selected any backup slot yet."
        )

    if can_self_claim:
        own_person=next(p for p in load_people(year,month) if p.initials==active_user)

        def selectable_ids(group):
            return [
                s.idx for s in group
                if (
                    (not claim_by_slot.get(s.idx) or claim_by_slot[s.idx]["initials"]==active_user)
                    and not hard_unavailable_for_block(own_person,s.day,s.block)
                )
            ]

        selected=[]
        group_specs=[
            ("Savaitgaliai · SPS RO budėjimai" if lang=="LT" else "Weekends · SPS RO duty",weekend_slots,"weekend"),
            ("Centro UG 120 · rytas" if lang=="LT" else "Centro UG 120 · morning",centro120_slots,"centro120"),
            ("Onko RO · pilna 9 val. pamaina" if lang=="LT" else "Onko RO · full 9h shift",onko_slots,"onko"),
            ("Darbo dienos · SPS RO" if lang=="LT" else "Weekdays · SPS RO",sps_ro_slots,"spsro"),
            ("Darbo dienos · SPS UG" if lang=="LT" else "Weekdays · SPS UG",sps_ug_slots,"spsug"),
        ]
        for title,group,key in group_specs:
            ids=selectable_ids(group)
            defaults=[sid for sid in ids if sid in my_claim_ids]
            with st.expander(title,expanded=(key=="weekend")):
                picked=st.multiselect(
                    ("Pasirink dublių slotus" if lang=="LT" else "Choose backup slots"),
                    ids,
                    default=defaults,
                    format_func=claim_label,
                    key=f"backup_claims_{key}_{year}_{month}_{active_user}",
                )
                selected.extend(picked)

        if st.button(
            "IŠSAUGOTI DUBLIŲ PASIRINKIMUS" if lang=="LT" else "SAVE BACKUP SELECTIONS",
            type="primary",use_container_width=True,key="save_backup_claims_v2532"
        ):
            try:
                db.replace_backup_claims(year,month,active_user,selected)
                st.success("Dublių pasirinkimai išsaugoti." if lang=="LT" else "Backup selections saved.")
                st.rerun()
            except Exception:
                st.error(
                    "Bent vieną pasirinktą slotą ką tik rezervavo kitas rezidentas. Atnaujinau sąrašą — pasirink iš naujo."
                    if lang=="LT" else
                    "At least one selected slot was just claimed by another resident. The list has been refreshed; choose again."
                )
                st.rerun()
    else:
        st.caption(tr("backup_claim_locked"))

    if advanced_mode:
        board=[]
        for s in all_backup_slots:
            r=claim_by_slot.get(s.idx)
            group=("Centro 120" if s.department.startswith("Centro UG 120") else "Onko RO" if s.department=="Onko RO centre" else "SPS RO" if s.department.startswith("SPS RO") else "SPS UG")
            board.append({
                ("Grupė" if lang=="LT" else "Group"):group,
                tr("date"):f"{year}-{month:02d}-{s.day:02d}",
                tr("department"):s.department,
                tr("shift"):block_label(s.block),
                tr("person"):(r["initials"] if r else tr("backup_claim_free")),
                tr("updated"):(r["claimed_at"] if r else "")
            })
        st.markdown("### Visų dublių rezervacijų lenta" if lang=="LT" else "### All backup reservations")
        st.dataframe(pd.DataFrame(board),use_container_width=True,hide_index=True)

    if senior_mode:
        nonclaimers=[p["initials"] for p in DEFAULT_PEOPLE if p["initials"] not in {r["initials"] for r in claims}]
        st.markdown(f"### {tr('backup_claim_auto_queue')}"); st.caption(tr("backup_claim_auto_queue_help")); st.write(", ".join(nonclaimers) if nonclaimers else "—")
    if not currentp:
        st.info(tr("not_published"))
    else:
        result=refresh_result_payload(currentp,year,month)
        if senior_mode:
            desired,backup_errors=sync_backup_plan(year,month,result)
        else:
            desired=[]; backup_errors=[]
        current_rows=db.list_backups(year,month)
        _slot_map_all={s.idx:s for s in make_slots(year,month)}
        expected=sum(1 for s in _slot_map_all.values() if backup_required_slot(s) and result.assignments.get(s.idx))
        required_actual=sum(1 for r in current_rows if (int(r.get("covered_slot")) in _slot_map_all and backup_required_slot(_slot_map_all[int(r.get("covered_slot"))])))
        centro_total=sum(1 for s in _slot_map_all.values() if backup_best_effort_slot(s) and result.assignments.get(s.idx))
        centro_covered=sum(1 for r in current_rows if (int(r.get("covered_slot")) in _slot_map_all and backup_best_effort_slot(_slot_map_all[int(r.get("covered_slot"))])))
        st.markdown(f"### {tr('backup_coverage')}")
        c1,c2,c3=st.columns(3)
        c1.metric(tr("working_person_days"),expected)
        c2.metric(tr("covered_person_days"),required_actual)
        c3.metric("CENTRO RO best-effort",f"{centro_covered}/{centro_total}")
        if backup_errors:
            st.error(tr("backup_incomplete")); st.dataframe(pd.DataFrame(backup_errors),use_container_width=True,hide_index=True)
        elif required_actual==expected:
            st.success(tr("backup_complete"))
        else:
            st.warning(tr("backup_incomplete"))

        st.markdown(f"### {tr('my_backup_schedule')}"); st.dataframe(backup_grid(year,month,result,active_user),use_container_width=True)
        bdf=backup_table(year,month,result); my=bdf[bdf[tr("effective_backup")]==active_user] if not bdf.empty else bdf
        if my.empty: st.caption(tr("no_backups"))
        else: st.dataframe(my,use_container_width=True,hide_index=True)

        if senior_mode:
            st.divider(); st.markdown(f"### {tr('manage_backups')}")
            if st.button(tr("resync_backups"),type="primary"):
                desired,errs=sync_backup_plan(year,month,result)
                if errs: st.error(tr("backup_incomplete"))
                else: st.success(tr("backup_synced")); st.rerun()
            allb=db.list_backups(year,month)
            if allb:
                st.markdown(f"### {tr('actual_override')}"); lookup={r["id"]:r for r in allb}; ids=list(lookup)
                slot_map={s.idx:s for s in make_slots(year,month)}
                def ridlabel(rid):
                    r=lookup[rid]; sid=int(r["covered_slot"]); s=slot_map.get(sid); covered=result.assignments.get(sid,"")
                    if s is None: return f"#{rid}"
                    return f"#{rid} · {s.day:02d} {WEEKDAYS[lang][s.weekday]} · {covered} · {s.department} · {block_label(s.block)} · {r['planned_backup']}"
                rid=st.selectbox(tr("backup_record"),ids,format_func=ridlabel); rr=lookup[rid]; sid=int(rr["covered_slot"]); covered_slot=slot_map.get(sid)
                people=people_for_stored_result(result,year,month); byinit={p.initials:p for p in people}
                # V2.5.61: manual/actual backup selection is broader than the automatic plan.
                # A resident may voluntarily take the backup even when this creates a fatigue / >48 h
                # warning or self-overrides RESIDENT HARD. ABSOLUTE unavailability and overlap remain
                # excluded from the candidate list.
                eligible=_eligible_backup_candidates(year,month,result,covered_slot,people,allow_resident_hard=True) if covered_slot is not None else []
                actual_opts=[""]+eligible
                current_actual=rr["actual_backup"] or ""
                if current_actual and current_actual not in actual_opts: actual_opts.append(current_actual)
                actual=st.selectbox(tr("actual_backup"),actual_opts,index=actual_opts.index(current_actual) if current_actual in actual_opts else 0,format_func=lambda i:"—" if i=="" else f"{i} — {people_map[i]['name']}",disabled=bool(rr.get("completed_at")))

                preview=None
                if actual and covered_slot is not None:
                    preview=_preview_manual_backup_takeover(year,month,result,covered_slot,actual,exclude_backup_id=rid)
                    st.markdown("#### Dublio perėmimo pasekmės" if lang=="LT" else "#### Backup takeover consequences")
                    if preview.get("rows"):
                        st.dataframe(pd.DataFrame(preview["rows"]),use_container_width=True,hide_index=True)
                    for msg in preview.get("blockers",[]):
                        st.error(("Negalima patvirtinti: " if lang=="LT" else "Cannot confirm: ")+msg)
                    if preview.get("warnings"):
                        st.warning(("Galima tik po aiškaus savanoriško patvirtinimo:\n- " if lang=="LT" else "Allowed only after explicit voluntary acknowledgement:\n- ")+"\n- ".join(preview["warnings"]))
                        st.caption(
                            "40/48 val. ir recovery perspėjimai nėra automatinis draudimas savanoriškam dubliui. Tačiau >12 val./d., aktyvios 7 d. ribos viršijimas, <11 val. paros poilsio, >6 darbo dienos/7 d., pateisinamas neatvykimas ar persidengimas lieka blokuojami."
                            if lang=="LT" else
                            "40/48h and recovery warnings do not automatically block a voluntary backup takeover. >12h/day, breach of the active 7-day cap, <11h daily rest, >6 workdays/7d, justified absence or overlap remain blocking."
                        )

                c1,c2=st.columns(2)
                with c1:
                    if actual and preview and preview.get("warnings") and preview.get("ok"):
                        confirm_manual=st.checkbox(
                            "Rezidentas supranta parodytas pasekmes ir SAVANORIŠKAI sutinka perimti šį dublį." if lang=="LT" else "The resident understands the shown consequences and VOLUNTARILY agrees to take this backup.",
                            key=f"backup_override_ack_{rid}_{actual}_{preview.get('fingerprint')}"
                        )
                        if st.button("PATVIRTINTI VIS TIEK" if lang=="LT" else "CONFIRM ANYWAY",type="primary",use_container_width=True,disabled=(not confirm_manual or bool(rr.get("completed_at")))):
                            meta={"version":"2.5.61","kind":"MANUAL_BACKUP_OVERRIDE","resident":actual,"covered_slot":int(sid),"fingerprint":preview.get("fingerprint"),"warnings":preview.get("warnings",[]),"rows":preview.get("rows",[]),"acknowledged_at":datetime.now(timezone.utc).isoformat()}
                            oldmeta=_backup_override_note_decode(rr.get("note"))
                            db.set_actual_backup(rid,actual or None,note=_backup_override_note_encode(meta,oldmeta.get("legacy_note","")))
                            persist_actual_satisfaction(year,month)
                            refresh_calendar_subscription_feeds([x for x in [rr.get("planned_backup"),current_actual,actual] if x])
                            st.success("Dublio perėmimas patvirtintas su perspėjimo ACK." if lang=="LT" else "Backup takeover confirmed with warning acknowledgement."); st.rerun()
                    else:
                        if st.button(tr("record_actual"),use_container_width=True,disabled=(bool(rr.get("completed_at")) or bool(actual and preview and not preview.get("ok")))):
                            db.set_actual_backup(rid,actual or None); persist_actual_satisfaction(year,month)
                            refresh_calendar_subscription_feeds([x for x in [rr.get("planned_backup"),current_actual,actual] if x])
                            st.success(tr("actual_saved")); st.rerun()
                with c2:
                    if actual and preview and preview.get("warnings") and preview.get("ok"):
                        if st.button("ATŠAUKTI" if lang=="LT" else "CANCEL",use_container_width=True,key=f"cancel_backup_override_{rid}"):
                            st.info("Nepatvirtinta — dublio perėmimas nepakeistas." if lang=="LT" else "Not confirmed — backup takeover unchanged.")
                    else:
                        if st.button(tr("clear_actual"),use_container_width=True,disabled=bool(rr.get("completed_at"))):
                            db.clear_actual_backup(rid); persist_actual_satisfaction(year,month)
                            refresh_calendar_subscription_feeds([x for x in [rr.get("planned_backup"),current_actual] if x])
                            st.success(tr("actual_cleared")); st.rerun()
                rr_fresh=next((x for x in db.list_backups(year,month) if int(x["id"])==int(rid)),rr)
                st.markdown(f"### {tr('backup_activation')}")
                effective_for_activation=str(rr_fresh.get("actual_backup") or rr_fresh.get("planned_backup") or "")
                activation_preview=(_preview_manual_backup_takeover(year,month,result,covered_slot,effective_for_activation,exclude_backup_id=rid) if effective_for_activation and covered_slot is not None else None)
                if rr_fresh.get("activated_at"):
                    st.success(f"{tr('backup_activated')}: {rr_fresh.get('activated_at')}")
                    if st.button(tr("undo_activation"),use_container_width=True):
                        db.clear_backup_activation(rid); st.success(tr("activation_undone")); st.rerun()
                else:
                    if activation_preview and activation_preview.get("blockers"):
                        st.error("Dublio aktyvuoti negalima, kol išlieka aukščiau parodytas absoliutus / teisinis blokatorius." if lang=="LT" else "Backup cannot be activated while an absolute/legal blocker remains.")
                    elif activation_preview and activation_preview.get("warnings"):
                        st.warning("Aktyvavus šį dublį perspėjimai taps realaus darbo pasekmėmis." if lang=="LT" else "Activating this backup will turn the warnings into actual-work consequences.")
                        activation_ack=st.checkbox("Patvirtinu, kad perimantis rezidentas sutiko su šiomis pasekmėmis." if lang=="LT" else "I confirm the covering resident agreed to these consequences.",key=f"backup_activation_ack_{rid}_{activation_preview.get('fingerprint')}")
                        if st.button("AKTYVUOTI IR PATVIRTINTI VIS TIEK" if lang=="LT" else "ACTIVATE AND CONFIRM ANYWAY",type="primary",use_container_width=True,disabled=not activation_ack):
                            oldmeta=_backup_override_note_decode(rr_fresh.get("note"))
                            oldmeta.update({"activation_ack_fingerprint":activation_preview.get("fingerprint"),"activation_ack_at":datetime.now(timezone.utc).isoformat(),"activation_warnings":activation_preview.get("warnings",[])})
                            db.set_backup_note(rid,_backup_override_note_encode(oldmeta,oldmeta.get("legacy_note","")))
                            db.activate_backup(rid)
                            rr_alert=next((x for x in db.list_backups(year,month) if int(x["id"])==int(rid)),rr_fresh)
                            ok_mail,detail_mail=send_backup_activation_email(year,month,result,rr_alert)
                            if ok_mail: st.success(tr("backup_email_sent"))
                            else: st.warning(f"{tr('backup_email_failed')} {detail_mail}")
                            st.rerun()
                    else:
                        if st.button(tr("activate_backup"),type="primary",use_container_width=True):
                            db.activate_backup(rid)
                            rr_alert=next((x for x in db.list_backups(year,month) if int(x["id"])==int(rid)),rr_fresh)
                            ok_mail,detail_mail=send_backup_activation_email(year,month,result,rr_alert)
                            if ok_mail: st.success(tr("backup_email_sent"))
                            else: st.warning(f"{tr('backup_email_failed')} {detail_mail}")
                            st.rerun()
                if rr_fresh.get("completed_at"):
                    st.success(f"{tr('completed_backup')}: {rr_fresh.get('completed_at')}")
                    if st.button(tr("undo_backup_completed"),use_container_width=True):
                        try:
                            db.undo_backup_credit(rid); st.success(tr("backup_completion_undone")); st.rerun()
                        except Exception as exc:
                            st.error(str(exc))
                else:
                    auto_type=("NIGHT" if covered_slot and covered_slot.block=="NIGHT" else covered_slot.block if covered_slot else "")
                    if auto_type=="AM": st.info(f"{tr('cover_credit_type')}: {tr('credit_am')}")
                    elif auto_type=="PM": st.info(f"{tr('cover_credit_type')}: {tr('credit_pm')}")
                    elif auto_type=="NIGHT": st.info(f"{tr('cover_credit_type')}: {tr('credit_night')}")
                    st.caption(tr("cover_credit_note"))
                    if st.button(tr("mark_backup_completed"),type="primary",use_container_width=True,disabled=auto_type not in ("AM","PM","NIGHT")):
                        try:
                            ev=db.complete_backup_cover(rid)
                            st.success(tr("backup_completed"))
                            if ev.get("coverer_effect")=="REST_EARNED": st.info(tr("cover_effect_rest"))
                        except Exception as exc:
                            st.error(str(exc))
                rest_all=db.all_rest_credit_balances()
                st.markdown(f"### {tr('credit_balances')}")
                all_ids=sorted(rest_all)
                st.dataframe(pd.DataFrame([{
                    tr("person"):i,
                    f"{tr('rest_credit_bank')} AM":rest_all.get(i,{}).get("AM",0),
                    f"{tr('rest_credit_bank')} PM":rest_all.get(i,{}).get("PM",0),
                    f"{tr('rest_credit_bank')} NIGHT":rest_all.get(i,{}).get("NIGHT",0),
                } for i in all_ids]),use_container_width=True,hide_index=True)
                st.dataframe(backup_table(year,month,result),use_container_width=True,hide_index=True)
pos+=1

# --- Swaps ---
with tabs[pos]:
    SWAP_META_PREFIX="V2555_SWAP_META:"
    def _swap_meta_decode(reason):
        raw=str(reason or "")
        if raw.startswith(SWAP_META_PREFIX):
            try:
                data=json.loads(raw[len(SWAP_META_PREFIX):])
                return data if isinstance(data,dict) else {}
            except Exception:
                return {}
        if raw.startswith("V2554_SWAP_META:"):
            try:
                old=json.loads(raw[len("V2554_SWAP_META:"):])
                return {"phase":old.get("phase","pending"),"impact_ack":{}}
            except Exception:
                pass
        if raw=="accepted_pending_senior_apply":
            return {"phase":"accepted_pending_senior_apply","impact_ack":{}}
        return {"phase":"pending","impact_ack":{}}
    def _swap_meta_encode(meta):
        return SWAP_META_PREFIX+json.dumps(meta,ensure_ascii=False,separators=(",",":"))
    def _impact_acks(meta):
        return {str(k):str(v) for k,v in ((meta or {}).get("impact_ack") or {}).items()}
    def _is_swap_slot_conflict(exc):
        msg=str(exc or "")
        return "SWAP_SLOT_ALREADY_PENDING" in msg or "BACKUP_SWAP_SLOT_ALREADY_PENDING" in msg
    def _impact_rows(stats,initials):
        return list((((stats or {}).get("global",{}) or {}).get("swap_warning_rows") or {}).get(initials,[]) or [])
    def _render_impact_table(stats,initials,title=None):
        rows=_impact_rows(stats,initials)
        if title:
            st.markdown(f"**{title}**")
        if not rows:
            st.caption("Papildomų swapo perspėjimų nėra." if lang=="LT" else "No additional swap warnings.")
            return
        df=pd.DataFrame([{
            ("Lygis" if lang=="LT" else "Level"):r.get("severity","ACK"),
            ("Kas keičiasi" if lang=="LT" else "Impact"):r.get("kind",""),
            ("Data / langas" if lang=="LT" else "Date / window"):r.get("date",""),
            ("Prieš" if lang=="LT" else "Before"):r.get("before",""),
            ("Po" if lang=="LT" else "After"):r.get("after",""),
            ("Paaiškinimas" if lang=="LT" else "Explanation"):r.get("explanation",""),
        } for r in rows])
        st.dataframe(df,use_container_width=True,hide_index=True)

    st.subheader(tr("swap_title")); st.write(tr("swap_note")); st.caption(tr("multiple_swap_help"))
    swap_flash=st.session_state.pop("_swap_response_flash",None)
    if swap_flash:
        level,msg=swap_flash
        if level=="success": st.success(msg)
        elif level=="warning": st.warning(msg)
        else: st.info(msg)
    refresh_col,_=st.columns([1,4])
    with refresh_col:
        if st.button("↻ ATNAUJINTI SWAP STATUSĄ" if lang=="LT" else "↻ REFRESH SWAP STATUS",key=f"swap_refresh_{year}_{month}",use_container_width=True):
            st.rerun()
    st.caption(
        "Swap statusas visada perskaitomas iš DB. Kito rezidento jau atidarytas langas gali rodyti seną būseną iki refresh/rerun."
        if lang=="LT" else
        "Swap status is always read from the database. Another resident's already-open page may show stale state until refresh/rerun."
    )
    currentp=db.load_schedule(year,month,"current")
    swap_perm=db.get_swap_permission_v2591(year,month) if currentp else {"allowed":False,"source":"not_open"}
    swap_create_allowed=bool(swap_perm.get("allowed",False))
    if currentp:
        src=str(swap_perm.get("source") or "")
        if swap_create_allowed and src=="window":
            dl=_parse_iso_dt(swap_perm.get("deadline")); txt=(dl.astimezone(ZoneInfo("Europe/Vilnius")).strftime("%Y-%m-%d %H:%M") if dl else "—")
            _workflow_card("APSIKEITIMŲ LANGAS ATIDARYTAS" if lang=="LT" else "SWAP WINDOW OPEN", (f"Naujus apsikeitimus galite kurti iki {txt}." if lang=="LT" else f"New swap requests may be created until {txt}."),"swap_open")
        elif swap_create_allowed and src=="late":
            exp=_parse_iso_dt(swap_perm.get("expires_at")); txt=(exp.astimezone(ZoneInfo("Europe/Vilnius")).strftime("%Y-%m-%d %H:%M") if exp else "—")
            _workflow_card("INDIVIDUALUS PAVĖLUOTAS LEIDIMAS" if lang=="LT" else "INDIVIDUAL LATE ACCESS", (f"Suteikta individuali prieiga iki {txt}; liko {swap_perm.get('remaining','—')} naujų prašymų." if lang=="LT" else f"Individual access is active until {txt}; {swap_perm.get('remaining','—')} new request(s) remain."),"swap_closed")
        else:
            state_name="FINAL" if src=="final" else ("TERMINAS PASIBAIGĖ" if src=="expired" else "NAUJI APSIKEITIMAI UŽDARYTI")
            _workflow_card(state_name if lang=="LT" else ("FINAL" if src=="final" else "NEW SWAPS CLOSED"), ("Naujų apsikeitimo prašymų kurti negalite. Jau egzistuojančius prašymus dar galima priimti / atmesti, kol operatorius užbaigs priežiūrą." if src!="final" and lang=="LT" else "Nauji apsikeitimai uždaryti." if lang=="LT" else "No new swap requests can be created. Existing requests may still be responded to until an operator completes oversight." if src!="final" else "The schedule is FINAL; new swaps are closed."),"final" if src=="final" else "expired")
    if not currentp: st.info(tr("not_published"))
    elif not resident_ok: st.error(tr("bad_pin"))
    else:
        st.markdown(
            "### ↔ NORMALŪS APSIKEITIMAI"
            if lang=="LT" else
            "### ↔ NORMAL SWAPS"
        )
        st.caption(
            "Tai dvišalis request → accept/reject → seniūnės apply srautas. ONE-WAY emergency rescue yra atskiras uždarytas blokas žemiau ir niekada nėra automatiškai atidaromas po REQUEST SWAP."
            if lang=="LT" else
            "This is the bilateral request → accept/reject → senior apply flow. ONE-WAY emergency rescue is a separate collapsed section below and is never opened automatically after REQUEST SWAP."
        )
        result=refresh_result_payload(currentp,year,month)
        slots={s.idx:s for s in make_slots(year,month)}
        mine=[sid for sid,w in result.assignments.items() if w==active_user]
        theirs=[sid for sid,w in result.assignments.items() if w!=active_user]
        def sl(sid):
            s=slots[sid]
            return f"#{sid} · {s.day:02d} {WEEKDAYS[lang][s.weekday]} · {s.department} · {block_label(s.block)} · {result.assignments[sid]}"
        if mine and theirs:
            a,b=st.columns(2)
            with a:
                sa=st.selectbox(
                    tr("my_assignment"),mine,format_func=sl,key="swap_a"
                )
            with b:
                sb=st.selectbox(
                    tr("their_assignment"),theirs,format_func=sl,key="swap_b"
                )

            target_person=result.assignments[sb]

            # V2.5.80: same structured colored selection UI as Emergency Rescue.
            st.markdown(
                "#### Apsikeitimas"
                if lang=="LT" else
                "#### Swap"
            )
            _render_swap_people_line(active_user,target_person,"↔")
            sv1,sv2=st.columns(2)
            with sv1:
                _render_shift_tile(
                    "MANO PAMAINA"
                    if lang=="LT" else
                    "MY CURRENT SHIFT",
                    active_user,slots.get(sa),PERSON_COLORS.get(active_user)
                )
            with sv2:
                _render_shift_tile(
                    "KITO REZIDENTO PAMAINA"
                    if lang=="LT" else
                    "OTHER RESIDENT'S SHIFT",
                    target_person,slots.get(sb),PERSON_COLORS.get(target_person)
                )

            swap_people=people_for_stored_result(result,year,month)
            preview_ok,preview_reason,preview_stats,preview_needed=preview_swap(
                year,month,swap_people,result,sa,sb,backup_assignments=db.list_backups(year,month)
            )
            my_fp=preview_needed.get(active_user) if preview_ok else None
            their_fp=preview_needed.get(target_person) if preview_ok else None
            my_ack=True
            if not preview_ok:
                block_rows=((preview_stats or {}).get("global",{}) or {}).get("swap_hard_block_rows") or []
                if block_rows:
                    _render_swap_hard_block(preview_stats,preview_reason)
                else:
                    st.error(tr("swap_preview_invalid").format(reason=preview_reason))
            else:
                _render_impact_table(preview_stats,active_user,("Tavo swapo pasekmės" if lang=="LT" else "Your swap consequences"))
                if my_fp:
                    st.warning(tr("swap_48_warning"))
                    st.caption(tr("swap_48_only_exception"))
                    my_ack=st.checkbox(tr("swap_48_ack"),key=f"swapimpact_proposer_{sa}_{sb}")
                if their_fp:
                    st.info(tr("swap_48_other"))
                    _render_impact_table(preview_stats,target_person,("Kito rezidento pasekmės" if lang=="LT" else "Other resident consequences"))
            if st.button(tr("request_swap"),type="primary",disabled=(not preview_ok or not my_ack or not swap_create_allowed)):
                meta={"phase":"pending","impact_ack":{}}
                if my_fp and my_ack:
                    meta["impact_ack"][active_user]=str(my_fp)
                try:
                    inserted=db.create_swap_request(
                        year,month,sa,sb,active_user,target_person,
                        reason=_swap_meta_encode(meta)
                    )
                    saved=(inserted[0] if inserted else {})
                    email_ok,email_detail=send_swap_request_email(year,month,saved) if saved else (False,"request row missing")
                    request_msg=(
                        f"SWAP REQUEST #{saved.get('id','—')} išsaugotas. "
                        + ("El. laiškas gavėjui išsiųstas." if email_ok else f"DB išsaugota, bet email nepavyko: {email_detail}")
                        if lang=="LT" else
                        f"SWAP REQUEST #{saved.get('id','—')} saved. "
                        + ("Email notification sent." if email_ok else f"DB saved, but email failed: {email_detail}")
                    )
                    if email_ok:
                        st.success(request_msg)
                    else:
                        st.warning(request_msg)
                    # Deliberately NO st.rerun(): the request list below re-reads DB
                    # in the same Streamlit pass. This prevents scroll restoration
                    # from landing the user in the separate Emergency Rescue block.
                except Exception as exc:
                    if "SWAP_WINDOW_CLOSED" in str(exc):
                        st.warning("Naujų apsikeitimų langas uždarytas. Jei pavėlavote, kreipkitės į tvarkaraščio operatorių dėl individualaus leidimo." if lang=="LT" else "New swaps are closed. If you are late, ask a schedule operator for individual late access.")
                    elif _is_swap_slot_conflict(exc):
                        st.warning(tr("swap_shift_busy"))
                    else:
                        st.error(tr("swap_preview_invalid").format(reason=("Nepavyko išsaugoti pasiūlymo." if lang=="LT" else "Could not save the offer.")))

        reqs=db.list_swap_requests(year,month,active_user)
        outgoing=[
            r for r in reqs
            if r["person_a"]==active_user
            and r["status"]=="pending"
            and _swap_meta_decode(r.get("reason")).get("kind") not in {"emergency_actual","emergency_rescue"}
        ]
        if outgoing:
            st.markdown(
                f"### {tr('my_outgoing_swaps')} · {len(outgoing)}"
            )
            for idx,r in enumerate(outgoing,start=1):
                with st.container(border=True):
                    _render_swap_request_card(r,idx,len(outgoing),slots,incoming=False)
                    if st.button(
                        tr("cancel_my_swap"),
                        key=f"cancel_swap_{r['id']}",
                        use_container_width=True
                    ):
                        try:
                            saved=db.cancel_swap_request(r["id"])
                            if saved.get("status")!="rejected":
                                raise RuntimeError(f"Unexpected saved status: {saved.get('status')}")
                            st.session_state["_swap_response_flash"]=(
                                "success",
                                ("Pasiūlymas atšauktas ir DB būsena patvirtinta: REJECTED."
                                 if lang=="LT" else
                                 "Offer cancelled; authoritative DB status confirmed: REJECTED.")
                            )
                            st.rerun()
                        except Exception as exc:
                            st.error((f"Nepavyko atšaukti swapo: {exc}" if lang=="LT" else f"Could not cancel swap: {exc}"))
        incoming=[
            r for r in reqs
            if r["person_b"]==active_user
            and r["status"]=="pending"
            and _swap_meta_decode(r.get("reason")).get("kind") not in {"emergency_actual","emergency_rescue"}
        ]
        st.markdown(f"### {tr('incoming')} · {len(incoming)}")
        if not incoming:
            st.caption("Naujų apsikeitimo prašymų nėra." if lang=="LT" else "No new swap requests.")
        for idx,r in enumerate(incoming,start=1):
            with st.container(border=True):
                _render_swap_request_card(r,idx,len(incoming),slots,incoming=True)
                fresh=refresh_result_payload(db.load_schedule(year,month,"current"),year,month)
                stale=(fresh.assignments.get(r["slot_a"])!=r["person_a"] or fresh.assignments.get(r["slot_b"])!=r["person_b"])
                meta=_swap_meta_decode(r.get("reason")); acks=_impact_acks(meta)
                pv_ok,pv_reason,pv_stats,pv_needed=(False,"stale",None,{}) if stale else preview_swap(
                    year,month,people_for_stored_result(fresh,year,month),fresh,r["slot_a"],r["slot_b"],backup_assignments=db.list_backups(year,month)
                )
                proposer_fp=pv_needed.get(r["person_a"]) if pv_ok else None
                proposer_missing=bool(proposer_fp and acks.get(r["person_a"])!=proposer_fp)
                target_fp=pv_needed.get(active_user) if pv_ok else None
                target_ack=True
                if stale:
                    st.error(
                        "Šis requestas paseno, nes ACTUAL grafikas nuo jo sukūrimo jau pasikeitė."
                        if lang=="LT" else
                        "This request is stale because ACTUAL changed after it was created."
                    )
                elif not pv_ok:
                    block_rows=((pv_stats or {}).get("global",{}) or {}).get("swap_hard_block_rows") or []
                    if block_rows:
                        _render_swap_hard_block(pv_stats,pv_reason)
                    else:
                        st.error(tr("swap_preview_invalid").format(reason=pv_reason))
                elif proposer_missing:
                    st.warning(tr("swap_48_reaccept"))
                else:
                    _render_impact_table(pv_stats,active_user,("Tavo swapo pasekmės" if lang=="LT" else "Your swap consequences"))
                    if target_fp:
                        st.warning(tr("swap_48_warning"))
                        st.caption(tr("swap_48_only_exception"))
                        target_ack=st.checkbox(tr("swap_48_ack"),key=f"swapimpact_target_{r['id']}")
                c1,c2=st.columns(2)
                with c1:
                    if st.button(tr("accept"),key=f"ac{r['id']}",use_container_width=True,disabled=(stale or not pv_ok or proposer_missing or not target_ack)):
                        if target_fp and target_ack:
                            meta.setdefault("impact_ack",{})[active_user]=str(target_fp)
                        meta["phase"]="accepted_pending_senior_apply"
                        try:
                            saved=db.respond_swap_request_v2578(
                                r["id"],"accept",_swap_meta_encode(meta)
                            )
                            if saved.get("status")!="approved":
                                raise RuntimeError(f"Unexpected saved status: {saved.get('status')}")
                            st.session_state["_swap_response_flash"]=(
                                "success",
                                ("SWAP PRIIMTAS — DB būsena APPROVED. Dabar laukia seniūnės galutinio pritaikymo."
                                 if lang=="LT" else
                                 "SWAP ACCEPTED — authoritative DB status APPROVED. It now awaits senior application.")
                            )
                            st.rerun()
                        except Exception as exc:
                            if _is_swap_slot_conflict(exc):
                                st.warning(tr("swap_shift_busy"))
                            else:
                                st.error((f"Nepavyko išsaugoti ACCEPT: {exc}" if lang=="LT" else f"Could not save ACCEPT: {exc}"))
                with c2:
                    if st.button(tr("reject"),key=f"rj{r['id']}",use_container_width=True):
                        try:
                            saved=db.respond_swap_request_v2578(r["id"],"reject","declined")
                            if saved.get("status")!="rejected":
                                raise RuntimeError(f"Unexpected saved status: {saved.get('status')}")
                            st.session_state["_swap_response_flash"]=(
                                "success",
                                ("SWAP ATMESTAS — DB būsena patvirtinta: REJECTED."
                                 if lang=="LT" else
                                 "SWAP REJECTED — authoritative DB status confirmed: REJECTED.")
                            )
                            st.rerun()
                        except Exception as exc:
                            st.error((f"Nepavyko išsaugoti REJECT: {exc}" if lang=="LT" else f"Could not save REJECT: {exc}"))
        st.divider(); st.markdown(f"### {tr('backup_swap_title')}"); st.caption(tr("backup_swap_help"))
        all_backups=[r for r in db.list_backups(year,month) if not r.get("activated_at") and not r.get("completed_at")]
        slot_map_b={s.idx:s for s in make_slots(year,month)}
        my_b=[r for r in all_backups if (r.get("actual_backup") or r.get("planned_backup"))==active_user]
        other_b=[r for r in all_backups if (r.get("actual_backup") or r.get("planned_backup"))!=active_user]
        def blabel(r):
            s=slot_map_b.get(int(r["covered_slot"])); who=r.get("actual_backup") or r.get("planned_backup")
            return f"{s.day:02d} {WEEKDAYS[lang][s.weekday]} · {block_label(s.block)} · {who}" if s else str(r["covered_slot"])
        if my_b and other_b:
            ba,bb=st.columns(2)
            with ba:
                my_br=st.selectbox(
                    tr("my_backup_duty"),my_b,format_func=blabel,key="backup_swap_my"
                )
            with bb:
                other_br=st.selectbox(
                    tr("their_backup_duty"),other_b,format_func=blabel,key="backup_swap_other"
                )

            target=other_br.get("actual_backup") or other_br.get("planned_backup")
            my_backup_person=my_br.get("actual_backup") or my_br.get("planned_backup") or active_user
            my_backup_slot=slot_map_b.get(int(my_br["covered_slot"]))
            other_backup_slot=slot_map_b.get(int(other_br["covered_slot"]))

            # V2.5.80: same colored structured visual language as normal swaps/rescue.
            st.markdown(
                "#### Dublių apsikeitimas"
                if lang=="LT" else
                "#### Backup swap"
            )
            _render_swap_people_line(my_backup_person,target,"↔")
            bv1,bv2=st.columns(2)
            with bv1:
                _render_shift_tile(
                    "MANO DUBLIO VIETA"
                    if lang=="LT" else
                    "MY BACKUP DUTY",
                    my_backup_person,my_backup_slot,PERSON_COLORS.get(my_backup_person)
                )
                covered_a=str(my_br.get("covered_person") or "")
                if covered_a:
                    st.markdown(
                        ("**DENGIAMAS REZIDENTAS:** " if lang=="LT" else "**RESIDENT COVERED:** ")
                        + badge(covered_a,include_name=True),
                        unsafe_allow_html=True,
                    )
            with bv2:
                _render_shift_tile(
                    "KITO REZIDENTO DUBLIO VIETA"
                    if lang=="LT" else
                    "OTHER RESIDENT'S BACKUP DUTY",
                    target,other_backup_slot,PERSON_COLORS.get(target)
                )
                covered_b=str(other_br.get("covered_person") or "")
                if covered_b:
                    st.markdown(
                        ("**DENGIAMAS REZIDENTAS:** " if lang=="LT" else "**RESIDENT COVERED:** ")
                        + badge(covered_b,include_name=True),
                        unsafe_allow_html=True,
                    )

            if st.button(tr("request_backup_swap"),key="request_backup_swap_btn",use_container_width=True,disabled=(not swap_create_allowed)):
                try:
                    inserted=db.create_backup_swap_request(
                        year,month,active_user,int(my_br["covered_slot"]),
                        target,int(other_br["covered_slot"])
                    )
                    saved=(inserted[0] if inserted else {})
                    email_ok,email_detail=send_backup_swap_request_email(year,month,saved) if saved else (False,"request row missing")
                    backup_msg=(
                        f"Dublio swap request #{saved.get('id','—')} išsaugotas. "
                        + ("Email gavėjui išsiųstas." if email_ok else f"DB išsaugota, bet email nepavyko: {email_detail}")
                        if lang=="LT" else
                        f"Backup swap request #{saved.get('id','—')} saved. "
                        + ("Email sent." if email_ok else f"DB saved, but email failed: {email_detail}")
                    )
                    if email_ok:
                        st.success(backup_msg)
                    else:
                        st.warning(backup_msg)
                    # Same no-rerun rule as normal swaps: keep the user in the swap UI.
                except Exception as exc:
                    if _is_swap_slot_conflict(exc): st.warning(tr("backup_swap_shift_busy"))
                    else: st.error(tr("backup_swap_invalid"))
        breqs=db.list_backup_swap_requests(year,month,None if is_seniune_account else active_user)
        backup_outgoing=[r for r in breqs if r.get("requester")==active_user and r.get("status")=="pending"]
        if backup_outgoing:
            st.markdown(f"#### {tr('my_outgoing_swaps')} · {tr('backup_swap_title')} · {len(backup_outgoing)}")
            for idx,r in enumerate(backup_outgoing,start=1):
                with st.container(border=True):
                    st.markdown(
                        f"**{'DUBLIO PASIŪLYMAS' if lang=='LT' else 'BACKUP OFFER'} {idx}/{len(backup_outgoing)} · DB #{r['id']}**"
                    )
                    _render_swap_people_line(r["requester"],r["target"],"→")
                    st.caption(
                        f"{blabel({'covered_slot':r['requester_slot'],'actual_backup':r['requester']})} "
                        f"↔ {blabel({'covered_slot':r['target_slot'],'actual_backup':r['target']})}"
                    )
                    if st.button(tr("cancel_my_swap"),key=f"cancel_backup_swap_{r['id']}",use_container_width=True):
                        try:
                            db.cancel_backup_swap_request(r["id"])
                            st.session_state["_swap_response_flash"]=(
                                "success",
                                ("Backup swapo pasiūlymas atšauktas."
                                 if lang=="LT" else
                                 "Backup swap offer cancelled.")
                            )
                            st.rerun()
                        except Exception as exc:
                            st.error((f"Nepavyko atšaukti backup swapo: {exc}" if lang=="LT" else f"Could not cancel backup swap: {exc}"))
        backup_incoming=[x for x in breqs if x["target"]==active_user and x["status"]=="pending" and not x.get("participant_accepted_at")]
        if backup_incoming:
            st.markdown(
                f"#### {'GAUTI DUBLIŲ REQUESTAI' if lang=='LT' else 'INCOMING BACKUP REQUESTS'} · {len(backup_incoming)}"
            )
        for idx,r in enumerate(backup_incoming,start=1):
            with st.container(border=True):
                st.markdown(
                    f"**{'GAUTAS DUBLIO REQUESTAS' if lang=='LT' else 'INCOMING BACKUP REQUEST'} {idx}/{len(backup_incoming)} · DB #{r['id']}**"
                )
                _render_swap_people_line(r["requester"],r["target"],"→")
                st.caption(
                    f"{blabel({'covered_slot':r['requester_slot'],'actual_backup':r['requester']})} "
                    f"↔ {blabel({'covered_slot':r['target_slot'],'actual_backup':r['target']})}"
                )
                bc1,bc2=st.columns(2)
                with bc1:
                    if st.button(tr("accept"),key=f"bac{r['id']}",use_container_width=True):
                        fresh=refresh_result_payload(db.load_schedule(year,month,"current"),year,month); smap_b={s.idx:s for s in make_slots(year,month)}
                        s_a=smap_b.get(int(r["requester_slot"])); s_b=smap_b.get(int(r["target_slot"])); people_now=people_for_stored_result(fresh,year,month)
                        elig_a=_eligible_backup_candidates(year,month,fresh,s_a,people_now) if s_a else []; elig_b=_eligible_backup_candidates(year,month,fresh,s_b,people_now) if s_b else []
                        if r["target"] not in elig_a or r["requester"] not in elig_b:
                            st.error(tr("backup_swap_invalid"))
                        else:
                            try:
                                db.accept_backup_swap_participant_v25111(r["id"])
                                st.session_state["_swap_response_flash"]=(
                                    "success",
                                    ("Abu rezidentai sutiko dėl dublio swapo. Jis DAR NEPRITAIKYTAS — laukia SP galutinio patvirtinimo."
                                     if lang=="LT" else
                                     "Both residents consented to the backup swap. It is NOT applied yet — awaiting SP final approval.")
                                )
                                st.rerun()
                            except Exception:
                                st.error(tr("backup_swap_invalid"))
                with bc2:
                    if st.button(tr("reject"),key=f"bar{r['id']}",use_container_width=True):
                        try:
                            db.reject_backup_swap_request(r["id"])
                            st.session_state["_swap_response_flash"]=(
                                "success",
                                ("Backup swapas atmestas." if lang=="LT" else "Backup swap rejected.")
                            )
                            st.rerun()
                        except Exception as exc:
                            st.error((f"Nepavyko atmesti backup swapo: {exc}" if lang=="LT" else f"Could not reject backup swap: {exc}"))
        if is_seniune_account:
            _sr_backup_waiting=[r for r in breqs if r.get("status")=="pending" and r.get("participant_accepted_at") and not r.get("senior_decision")]
            if _sr_backup_waiting:
                st.markdown("#### SP GALUTINIS DUBLIŲ SWAPŲ PATVIRTINIMAS" if lang=="LT" else "#### SP FINAL BACKUP-SWAP APPROVAL")
                st.caption("Abu rezidentai jau sutiko. ACTUAL dar nepakeistas. Tik tavo APPROVE pritaiko swapą." if lang=="LT" else "Both residents already consented. ACTUAL is still unchanged. Only your APPROVE applies the swap.")
                for r in _sr_backup_waiting:
                    with st.container(border=True):
                        st.markdown(f"**DB #{r['id']} · {r['requester']} ↔ {r['target']}**")
                        _render_swap_people_line(r["requester"],r["target"],"↔")
                        ac,dc=st.columns(2)
                        with ac:
                            if st.button("SP APPROVE + APPLY",key=f"sr_backup_approve_{r['id']}",use_container_width=True,type="primary"):
                                try:
                                    db.approve_backup_swap_by_sr_v25111(r["id"])
                                    persist_actual_satisfaction(year,month)
                                    refresh_calendar_subscription_feeds([r["requester"],r["target"]])
                                    st.session_state["_swap_response_flash"]=("success","SP patvirtino dublio swapą — ACTUAL atnaujintas." if lang=="LT" else "SP approved the backup swap — ACTUAL updated.")
                                    st.rerun()
                                except Exception as exc: st.error(str(exc))
                        with dc:
                            if st.button("SP DECLINE",key=f"sr_backup_decline_{r['id']}",use_container_width=True):
                                try:
                                    db.reject_backup_swap_by_sr_v25111(r["id"])
                                    st.session_state["_swap_response_flash"]=("success","SP atmetė dublio swapą." if lang=="LT" else "SP declined the backup swap.")
                                    st.rerun()
                                except Exception as exc: st.error(str(exc))

        if breqs:
            st.markdown("#### Dublių apsikeitimų istorija" if lang=="LT" else "#### Backup swap history")
            for idx,r in enumerate(breqs,start=1):
                with st.container(border=True):
                    applied=bool(r.get("status")=="accepted")
                    st.markdown(f"**DUBLIO SWAP #{idx} · DB #{r.get('id')} · {str(r.get('status','')).upper()}**")
                    _render_swap_people_line(str(r.get("requester") or ""),str(r.get("target") or ""),"↔")
                    st.caption(
                        ("Jei statusas ACCEPTED, DELETE kartu bandys saugiai grąžinti ankstesnius dublio turėtojus." if lang=="LT" else
                         "If status is ACCEPTED, DELETE will also safely restore the previous backup holders.")
                        if applied else
                        ("DELETE pašalins šį request/history įrašą; ACTUAL backup planas nuo jo dar nepakeistas." if lang=="LT" else
                         "DELETE removes this request/history row; it has not changed the ACTUAL backup plan.")
                    )
                    if _delete_confirm(f"backup_swap_{r['id']}",applied=applied):
                        try:
                            saved=db.delete_backup_swap_v2586(int(r["id"]))
                            st.session_state["_swap_response_flash"]=(
                                "success",
                                ("Dublio swapas ištrintas"+(" ir ankstesnis backup planas atkurtas." if saved.get("undone_actual") else ".")
                                 if lang=="LT" else
                                 "Backup swap deleted"+(" and previous backup holders restored." if saved.get("undone_actual") else "."))
                            )
                            st.rerun()
                        except Exception as exc:
                            st.error(("DELETE nepavyko: " if lang=="LT" else "DELETE failed: ")+str(exc))
        hist=db.list_swap_requests(year,month,None if is_seniune_account else active_user)
        if is_seniune_account:
            pending_apply=[]
            for r in hist:
                if r.get("status")!="approved": continue
                meta=_swap_meta_decode(r.get("reason"))
                if meta.get("phase")=="accepted_pending_senior_apply" or r.get("reason")=="accepted_pending_senior_apply":
                    pending_apply.append(r)
            for r in pending_apply:
                meta=_swap_meta_decode(r.get("reason")); acks=_impact_acks(meta)
                st.info(f"{r['person_a']} ↔ {r['person_b']} · #{r['slot_a']} ↔ #{r['slot_b']}")
                if acks:
                    st.caption(("Swapo pasekmių ACK: " if lang=="LT" else "Swap consequence ACK: ")+", ".join(sorted(acks)))
                _sra,_srd=st.columns(2)
                with _sra:
                    _sr_apply=st.button(tr("finalize_swap"),key=f"finalize_{r['id']}",use_container_width=True,type="primary")
                with _srd:
                    _sr_decline=st.button("SP — DECLINE",key=f"decline_{r['id']}",use_container_width=True)
                if _sr_decline:
                    try:
                        db.mark_normal_swap_sr_decision_v25111(r["id"],"declined")
                        st.session_state["_swap_response_flash"]=("success","SP atmetė swapą — ACTUAL nepakeistas." if lang=="LT" else "SP declined the swap — ACTUAL unchanged.")
                        st.rerun()
                    except Exception as exc: st.error(str(exc))
                if _sr_apply:
                    fresh=refresh_result_payload(db.load_schedule(year,month,"current"),year,month)
                    if fresh.assignments.get(r["slot_a"])!=r["person_a"] or fresh.assignments.get(r["slot_b"])!=r["person_b"]:
                        db.update_swap_request(r["id"],"rejected","stale"); st.error(tr("hard_reject")); st.rerun()
                    pv_ok,pv_reason,pv_stats,pv_needed=preview_swap(
                        year,month,people_for_stored_result(fresh,year,month),fresh,r["slot_a"],r["slot_b"],backup_assignments=db.list_backups(year,month)
                    )
                    if not pv_ok:
                        db.update_swap_request(r["id"],"rejected",pv_reason)
                        block_rows=((pv_stats or {}).get("global",{}) or {}).get("swap_hard_block_rows") or []
                        if block_rows:
                            _render_swap_hard_block(pv_stats,pv_reason)
                        else:
                            st.error(
                                ("Swapo pritaikyti nepavyko: " if lang=="LT" else "Could not apply swap: ")
                                + str(pv_reason)
                            )
                        st.stop()
                    missing=[who for who,fp in pv_needed.items() if acks.get(who)!=fp]
                    if missing:
                        who=missing[0]
                        db.update_swap_request(r["id"],"rejected",f"impact_ack_missing:{who}")
                        st.error(tr("swap_48_reaccept")); st.rerun()
                    ok,reason,_=attempt_swap(
                        year,month,people_for_stored_result(fresh,year,month),fresh,r["slot_a"],r["slot_b"],
                        backup_assignments=db.list_backups(year,month),acknowledged_fingerprints=acks
                    )
                    if ok:
                        # Bilateral voluntary swap changes ACTUAL only. Consequence ACKs are
                        # fingerprinted so a changed preview forces fresh resident consent.
                        if acks:
                            audit=list((fresh.stats.get("global",{}) or {}).get("swap_ack_audit") or [])
                            audit.append({"swap_id":int(r["id"]),"people":dict(acks),"applied_at":datetime.now(timezone.utc).isoformat()})
                            fresh.stats["global"]["swap_ack_audit"]=audit
                        db.save_current(year,month,serialize_result(fresh))
                        sync_backup_plan(year,month,fresh)
                        persist_actual_satisfaction(year,month)
                        refresh_calendar_subscription_feeds([r["person_a"],r["person_b"]])
                        meta["phase"]="applied"
                        db.update_swap_request(r["id"],"approved",_swap_meta_encode(meta))
                        try: db.mark_normal_swap_sr_decision_v25111(r["id"],"approved")
                        except Exception: pass
                        st.success(tr("swap_applied")); st.rerun()
                    else:
                        db.update_swap_request(r["id"],"rejected",reason)
                        st.error(
                            ("Swapo pritaikyti nepavyko: " if lang=="LT" else "Could not apply swap: ")
                            + str(reason)
                        )
                        st.stop()
        regular_hist=[r for r in hist if _swap_meta_decode(r.get("reason")).get("kind") not in {"emergency_actual","emergency_rescue"}]
        if regular_hist:
            smap={"pending":tr("pending"),"approved":tr("approved"),"rejected":tr("rejected_status")}
            st.markdown(f"### {tr('history')}")
            for idx,r in enumerate(regular_hist,start=1):
                meta=_swap_meta_decode(r.get("reason"))
                applied=bool(meta.get("phase")=="applied")
                with st.container(border=True):
                    st.markdown(f"**SWAP #{idx} · DB #{r.get('id')} · {smap.get(r.get('status'),r.get('status'))}**")
                    _render_swap_people_line(str(r.get("person_a") or ""),str(r.get("person_b") or ""),"↔")
                    sa=slots.get(int(r.get("slot_a") or -1)); sb=slots.get(int(r.get("slot_b") or -1))
                    hc1,hc2=st.columns(2)
                    with hc1: _render_shift_tile("A",str(r.get("person_a") or ""),sa,PERSON_COLORS.get(str(r.get("person_a") or "")))
                    with hc2: _render_shift_tile("B",str(r.get("person_b") or ""),sb,PERSON_COLORS.get(str(r.get("person_b") or "")))
                    st.caption(
                        ("Šis swapas jau pritaikytas ACTUAL. DELETE bandys atlikti UNDO tik jei abu slotai po to nebuvo pakeisti dar kartą." if lang=="LT" else
                         "This swap is already applied to ACTUAL. DELETE will UNDO it only if neither slot changed again afterwards.")
                        if applied else
                        ("Šis įrašas ACTUAL dar nepakeitė; DELETE pašalins request/history įrašą." if lang=="LT" else
                         "This row has not changed ACTUAL; DELETE removes the request/history row.")
                    )
                    if _delete_confirm(f"normal_swap_{r['id']}",applied=applied):
                        try:
                            saved=_delete_swap_row(r,year,month)
                            st.session_state["_swap_response_flash"]=(
                                "success",
                                ("Swapas ištrintas"+(" ir ACTUAL grąžintas į būseną prieš šį swapą." if saved.get("undone_actual") else ".")
                                 if lang=="LT" else
                                 "Swap deleted"+(" and ACTUAL restored to its pre-swap state." if saved.get("undone_actual") else "."))
                            )
                            st.rerun()
                        except Exception as exc:
                            st.error(("DELETE / UNDO nepavyko: " if lang=="LT" else "DELETE / UNDO failed: ")+str(exc))

        # V2.5.81 — always-visible EMERGENCY RESCUE operational panel.
        # The logged-in resident is pulled from their own lower-priority optional
        # post into an overlapping critical SPS post. The source becomes vacant;
        # SYSTEM fairness remains frozen.
        st.divider()
        with st.container(border=True):
            st.markdown(
                """
                <div style="display:flex;align-items:center;gap:16px;margin:2px 0 10px 0;">
                    <span style="font-size:3.1rem;line-height:1;">🚨</span>
                    <div>
                        <div style="font-size:1.55rem;font-weight:900;letter-spacing:.025em;">EMERGENCY RESCUE</div>
                        <div style="font-size:.88rem;opacity:.72;font-weight:650;">CURRENT LOCATION → MOVING TO → RESCUED PERSON</div>
                    </div>
                </div>
                """,
                unsafe_allow_html=True,
            )
            st.info(
                "Naudok, kai realiai dirbdamas savo poste esi skubiai perkeltas į svarbesnį SPS RO / SPS UG postą. "
                "CURRENT LOCATION lieka tuščias; RESCUED PERSON atleidžiamas nuo kritinio posto. "
                "Keičiamas ACTUAL grafikas ir iškart perskaičiuojama ACTUAL fairness statistika; SYSTEM publikavimo baseline lieka užšaldytas auditui. Jokio future catch-up / future catch-up nėra."
                if lang=="LT" else
                "Use this when, while working your assigned post, you are urgently moved into a more important SPS RO / SPS UG post. "
                "CURRENT LOCATION becomes vacant; the RESCUED PERSON is released from the critical post. "
                "ACTUAL changes and ACTUAL fairness statistics are recalculated immediately; the publication-time SYSTEM baseline stays frozen for audit. There is no future catch-up or future catch-up."
            )

            rescue_current=refresh_result_payload(
                db.load_schedule(year,month,"current"),year,month
            )
            rescue_slots={s.idx:s for s in make_slots(year,month)}

            # The constitutional workflow is self-recording: the resident who was
            # actually moved records the rescue from their own account.
            mover=active_user
            source_options=[]
            for sid,who in (rescue_current.assignments or {}).items():
                sl=rescue_slots.get(int(sid))
                if who!=mover or sl is None:
                    continue
                if not is_emergency_lower_priority_donor_slot(sl):
                    continue
                target_exists=any(
                    t.idx!=sl.idx
                    and is_emergency_critical_slot(t)
                    and t.day==sl.day
                    and t.block==sl.block
                    and rescue_current.assignments.get(t.idx)
                    and rescue_current.assignments.get(t.idx)!=mover
                    for t in rescue_slots.values()
                )
                if target_exists:
                    source_options.append(sl.idx)

            source_options=sorted(
                source_options,
                key=lambda sid:(
                    rescue_slots[sid].day,
                    {"AM":0,"PM":1,"FULL":2}.get(rescue_slots[sid].block,9),
                    rescue_slots[sid].department,
                    sid,
                )
            )

            if not source_options:
                st.info(
                    "Šiuo metu tavo ACTUAL grafike nėra tinkamos žemesnio prioriteto pamainos, kurią tuo pačiu laiku būtų galima perkelti į kritinį SPS RO / SPS UG postą."
                    if lang=="LT" else
                    "Your ACTUAL schedule currently has no eligible lower-priority assignment that can be moved at the same time into a critical SPS RO / SPS UG post."
                )
            else:
                source_sid=st.selectbox(
                    "CURRENT LOCATION — iš kur mane perkelia"
                    if lang=="LT" else
                    "CURRENT LOCATION — where I am being moved from",
                    source_options,
                    format_func=lambda sid:_swap_shift_text(rescue_slots[sid]),
                    key="emergency_rescue_source",
                )
                source_slot=rescue_slots[source_sid]

                target_options=[
                    t.idx for t in rescue_slots.values()
                    if t.idx!=source_sid
                    and is_emergency_critical_slot(t)
                    and t.day==source_slot.day
                    and t.block==source_slot.block
                    and rescue_current.assignments.get(t.idx)
                    and rescue_current.assignments.get(t.idx)!=mover
                ]
                target_options=sorted(
                    target_options,
                    key=lambda sid:(
                        rescue_slots[sid].department,
                        rescue_current.assignments.get(sid,""),
                        sid,
                    )
                )

                if not target_options:
                    st.info(
                        "Šiai CURRENT LOCATION pamainai nėra to paties laiko kritinio SPS targeto."
                        if lang=="LT" else
                        "There is no same-time critical SPS target for this CURRENT LOCATION."
                    )
                else:
                    target_sid=st.selectbox(
                        "MOVING TO — į kurį svarbesnį postą mane perkelia"
                        if lang=="LT" else
                        "MOVING TO — critical post I am moving to",
                        target_options,
                        format_func=lambda sid:(
                            f"{_swap_shift_text(rescue_slots[sid])} · "
                            f"RESCUED: {rescue_current.assignments.get(sid,'—')}"
                        ),
                        key="emergency_rescue_target",
                    )
                    target_slot=rescue_slots[target_sid]
                    rescued_person=str(rescue_current.assignments.get(target_sid) or "")

                    st.markdown("#### Emergency Rescue" if lang=="LT" else "#### Emergency Rescue")
                    _render_swap_people_line(mover,rescued_person,"→")
                    rc1,rc2=st.columns(2)
                    with rc1:
                        _render_shift_tile(
                            "CURRENT LOCATION — BUS PALIKTA TUŠČIA"
                            if lang=="LT" else
                            "CURRENT LOCATION — WILL BECOME VACANT",
                            mover,source_slot,PERSON_COLORS.get(mover)
                        )
                    with rc2:
                        _render_shift_tile(
                            "MOVING TO — KRITINIS POSTAS"
                            if lang=="LT" else
                            "MOVING TO — CRITICAL POST",
                            mover,target_slot,PERSON_COLORS.get(mover)
                        )
                    st.markdown(
                        (
                            '<div style="margin:12px 0;padding:12px 14px;border-radius:14px;'
                            'border:2px dashed #777;background:rgba(127,127,127,.08);">'
                            '<b>RESCUED PERSON:</b> '
                            + badge(rescued_person,include_name=True)
                            + '<br><span style="opacity:.82;">'
                            + html.escape(
                                "Šis žmogus atleidžiamas nuo pasirinkto kritinio posto ir NĖRA perkeliamas į tavo CURRENT LOCATION."
                                if lang=="LT" else
                                "This person is released from the selected critical post and is NOT moved to your CURRENT LOCATION."
                            )
                            + '</span></div>'
                        ),
                        unsafe_allow_html=True,
                    )

                    rescue_note=st.text_input(
                        "Trumpa operational pastaba"
                        if lang=="LT" else
                        "Short operational note",
                        key="emergency_rescue_note",
                    )
                    rescue_confirm=st.checkbox(
                        (
                            "PATVIRTINU: tai realiai įvykęs ONE-WAY emergency rescue. "
                            "Aš pereinu iš CURRENT LOCATION į MOVING TO; mano senas optional postas lieka tuščias; "
                            f"{rescued_person} yra rescued ir neina į mano seną vietą."
                        )
                        if lang=="LT" else
                        (
                            "I CONFIRM: this is an already-occurred ONE-WAY emergency rescue. "
                            "I move from CURRENT LOCATION to MOVING TO; my old optional post becomes vacant; "
                            f"{rescued_person} is rescued and does not move to my old post."
                        ),
                        key=f"emergency_rescue_confirm_{source_sid}_{target_sid}",
                    )

                    if st.button(
                        "🚨 ĮRAŠYTI ONE-WAY RESCUE Į ACTUAL"
                        if lang=="LT" else
                        "🚨 RECORD ONE-WAY RESCUE IN ACTUAL",
                        type="primary",
                        use_container_width=True,
                        disabled=not rescue_confirm,
                        key="emergency_rescue_apply",
                    ):
                        fresh_rescue=refresh_result_payload(
                            db.load_schedule(year,month,"current"),year,month
                        )
                        if (
                            fresh_rescue.assignments.get(source_sid)!=mover
                            or fresh_rescue.assignments.get(target_sid)!=rescued_person
                        ):
                            st.error(
                                "ACTUAL grafikas jau pasikeitė. Atnaujink puslapį ir pasirink CURRENT LOCATION / MOVING TO iš naujo."
                                if lang=="LT" else
                                "ACTUAL changed. Refresh and select CURRENT LOCATION / MOVING TO again."
                            )
                        elif (
                            source_slot.day!=target_slot.day
                            or source_slot.block!=target_slot.block
                            or not is_emergency_lower_priority_donor_slot(source_slot)
                            or not is_emergency_critical_slot(target_slot)
                        ):
                            st.error(
                                "Rescue neatitinka vienpusio same-time lower-priority → critical SPS modelio."
                                if lang=="LT" else
                                "Rescue no longer matches the same-time lower-priority → critical SPS model."
                            )
                        else:
                            repaired=apply_emergency_critical_transfer(
                                fresh_rescue.assignments,
                                target_slot,
                                mover,
                                source_slot=source_slot,
                            )
                            fresh_rescue.assignments=repaired
                            # Revalidate in-memory under ACTUAL operational rules so the
                            # stored payload carries refreshed stats + frozen workload credit.
                            fresh_rescue=revalidate_loaded_result(
                                year,month,people_for_stored_result(fresh_rescue,year,month),fresh_rescue,
                                backup_assignments=None,
                                validation_mode="emergency_rescue",
                            )
                            desired_backups,backup_errors=plan_backups(year,month,fresh_rescue)
                            if backup_errors:
                                st.error(
                                    ("Emergency Rescue negalimas, nes po perkėlimo nepavyksta sudaryti privalomo backup plano: " if lang=="LT" else
                                     "Emergency Rescue cannot be applied because required backup coverage cannot be rebuilt: ")
                                    + "; ".join(map(str,backup_errors[:3]))
                                )
                                st.stop()
                            # Final payload is validated against the NEW backup plan that
                            # will be committed in the same DB transaction.
                            fresh_rescue=revalidate_loaded_result(
                                year,month,people_for_stored_result(fresh_rescue,year,month),fresh_rescue,
                                backup_assignments=desired_backups,
                                validation_mode="emergency_rescue",
                            )
                            final_rescue_errors=list((fresh_rescue.stats or {}).get("global",{}).get("errors",[]) or [])
                            if final_rescue_errors:
                                st.error(
                                    ("Emergency Rescue negalimas dėl operational HARD taisyklės: " if lang=="LT" else
                                     "Emergency Rescue is blocked by an operational HARD rule: ")
                                    + str(final_rescue_errors[0])
                                )
                                st.stop()

                            meta={
                                "kind":"emergency_rescue",
                                "phase":"applied",
                                "mover":mover,
                                "rescued_person":rescued_person,
                                "source_slot":int(source_sid),
                                "target_slot":int(target_sid),
                                "source_department":source_slot.department,
                                "target_department":target_slot.department,
                                "day":int(source_slot.day),
                                "block":source_slot.block,
                                "source_vacated":True,
                                "bilateral_swap":False,
                                "fairness_neutral":True,
                                "workload_credit_neutral":True,
                                "workload_credit_source":"PUBLISHED_SYSTEM",
                                
                                "rescued_person_absence_outcome":"OUTSIDE_SCHEDULER_HR",
                                "recorded_by":active_user,
                                "recorded_at":datetime.now(timezone.utc).isoformat(),
                                "note":str(rescue_note or ""),
                            }
                            try:
                                saved=db.apply_emergency_rescue_atomic_v2585(
                                    year,month,source_sid,target_sid,mover,rescued_person,
                                    serialize_result(fresh_rescue),desired_backups,
                                    reason=_swap_meta_encode(meta),
                                )
                            except Exception as exc:
                                st.error(
                                    ("Emergency Rescue NEĮRAŠYTAS — transakcija atšaukta, ACTUAL grafikas ir backup planas nepakeisti. " if lang=="LT" else
                                     "Emergency Rescue NOT recorded — transaction rolled back; ACTUAL and backup plan were left unchanged. ")
                                    + str(exc)
                                )
                                st.stop()
                            refresh_calendar_subscription_feeds([mover,rescued_person])

                            st.session_state["_swap_response_flash"]=(
                                "success",
                                (
                                    f"ONE-WAY RESCUE įrašytas: {mover} "
                                    f"{source_slot.department} → {target_slot.department}; "
                                    f"RESCUED {rescued_person}. CURRENT source paliktas tuščias. Workload credit visiems nepakitęs."
                                )
                                if lang=="LT" else
                                (
                                    f"ONE-WAY RESCUE recorded: {mover} "
                                    f"{source_slot.department} → {target_slot.department}; "
                                    f"RESCUED {rescued_person}. Source post left vacant. Workload credit unchanged for everyone."
                                )
                            )
                            st.rerun()

            rescue_all=[
                r for r in db.list_swap_requests(
                    year,month,None if senior_mode else active_user
                )
                if _swap_meta_decode(r.get("reason")).get("kind") in {"emergency_rescue","emergency_actual"}
            ]
            if rescue_all:
                st.markdown(
                    "#### ONE-WAY rescue žurnalas"
                    if lang=="LT" else
                    "#### ONE-WAY rescue log"
                )
                for idx,r in enumerate(rescue_all,start=1):
                    meta=_swap_meta_decode(r.get("reason"))
                    with st.container(border=True):
                        if meta.get("kind")=="emergency_rescue":
                            mover_i=str(r.get("person_a") or "")
                            rescued_i=str(meta.get("rescued_person") or r.get("person_b") or "")
                            st.markdown(
                                f"**RESCUE #{idx} · DB #{r.get('id')}**"
                            )
                            _render_swap_people_line(mover_i,rescued_i,"→")
                            lc1,lc2=st.columns(2)
                            with lc1:
                                st.markdown(
                                    f"**CURRENT LOCATION**  \n{meta.get('source_department','—')} · "
                                    f"{meta.get('day','—')} · {meta.get('block','—')}"
                                )
                            with lc2:
                                st.markdown(
                                    f"**MOVING TO**  \n{meta.get('target_department','—')} · "
                                    f"{meta.get('day','—')} · {meta.get('block','—')}"
                                )
                            st.markdown(
                                "**RESCUED PERSON:** "+badge(rescued_i,include_name=True),
                                unsafe_allow_html=True,
                            )
                            st.caption(
                                (
                                    "Source postas paliktas tuščias; rescued person neperkeltas atgal. "
                                    + (f"Pastaba: {meta.get('note')}" if meta.get("note") else "")
                                )
                                if lang=="LT" else
                                (
                                    "Source post left vacant; rescued person was not moved back. "
                                    + (f"Note: {meta.get('note')}" if meta.get("note") else "")
                                )
                            )
                            st.caption(
                                "DELETE / UNDO atkurs mover į CURRENT LOCATION ir RESCUED PERSON į ankstesnį kritinį postą tik jei šie slotai po Rescue nebuvo pakeisti dar kartą."
                                if lang=="LT" else
                                "DELETE / UNDO restores the mover to CURRENT LOCATION and the RESCUED PERSON to the prior critical post only if those slots have not changed again."
                            )
                            if _delete_confirm(f"emergency_rescue_{r['id']}",applied=True):
                                try:
                                    saved=_delete_swap_row(r,year,month)
                                    st.session_state["_swap_response_flash"]=(
                                        "success",
                                        ("Emergency Rescue ištrintas ir ACTUAL saugiai atkurtas." if lang=="LT" else
                                         "Emergency Rescue deleted and ACTUAL safely restored.")
                                    )
                                    st.rerun()
                                except Exception as exc:
                                    st.error(("Emergency Rescue DELETE / UNDO nepavyko: " if lang=="LT" else "Emergency Rescue DELETE / UNDO failed: ")+str(exc))
                        else:
                            st.warning(
                                (
                                    f"LEGACY emergency_actual #{r.get('id')}: "
                                    f"{r.get('person_a')} ↔ {r.get('person_b')}. "
                                    "Tai senas bilateralinis įrašas iš ankstesnės, klaidingai pavadintos Emergency logikos."
                                )
                                if lang=="LT" else
                                (
                                    f"LEGACY emergency_actual #{r.get('id')}: "
                                    f"{r.get('person_a')} ↔ {r.get('person_b')}. "
                                    "This is a historical bilateral record from the old misnamed Emergency flow."
                                )
                            )

        # V2.5.13 — senior-only fairness-neutral post-publication repair workflow.
        if senior_mode:
            st.divider(); st.markdown(f"### {tr('repair_title')}"); st.caption(tr("repair_help"))
            fresh=refresh_result_payload(db.load_schedule(year,month,"current"),year,month)
            slots_by_id={s.idx:s for s in make_slots(year,month)}
            assigned_slots=[s for s in make_slots(year,month) if fresh.assignments.get(s.idx)]
            assigned_slots=sorted(assigned_slots,key=lambda x:(x.day,x.department,x.block,x.idx))
            def _repair_slot_label(sl):
                return f"{sl.day:02d} · {sl.department} · {block_label(sl.block)} · {fresh.assignments.get(sl.idx,'—')}"
            repair_rows=db.list_schedule_repairs(year,month)
            repair_load={}
            for rr in repair_rows:
                to_i=rr.get("to_person")
                if to_i: repair_load[to_i]=repair_load.get(to_i,0)+1
            if assigned_slots:
                chosen=st.selectbox(tr("repair_assignment"),assigned_slots,format_func=_repair_slot_label,key="repair_slot")
                from_person=fresh.assignments.get(chosen.idx)
                target_critical=is_emergency_critical_slot(chosen)
                candidate_rows=[]
                base_rh=int(fresh.stats.get("global",{}).get("resident_hard_total_losses",0) or 0)
                if target_critical:
                    # V2.5.56: critical sickness/absence rescue hierarchy. Keep SPS
                    # RO / SPS UG covered by moving a resident out of an overlapping
                    # lower-priority OPTIONAL post first. Only if no safe transfer
                    # exists do we expose a free-resident fallback.
                    shown_candidates=_critical_repair_candidate_rows(year,month,fresh,chosen,repair_load)
                else:
                    for prow in DEFAULT_PEOPLE:
                        cand=prow["initials"]
                        if cand==from_person: continue
                        ok,why,cstats=_repair_candidate_check(year,month,fresh,chosen.idx,cand)
                        if ok:
                            cg=(cstats or {}).get("global",{})
                            candidate_rows.append({
                                "initials":cand,
                                "source_slot":None,
                                "source_department":"",
                                "source_block":"",
                                "mode":"DIRECT",
                                "rh_total":int(cg.get("resident_hard_total_losses",0) or 0),
                                "rh_delta":int(cg.get("resident_hard_total_losses",0) or 0)-base_rh,
                                "rh_max":int(cg.get("resident_hard_max_loss_per_resident",0) or 0),
                                "rh_cum_spread":int(cg.get("resident_hard_cumulative_spread",0) or 0),
                                "critical_spread":int(cg.get("critical_worst_spread",0) or 0),
                                "noncritical_spread":int(cg.get("noncritical_worst_spread",0) or 0),
                            })
                    # Post-publication repairs follow the same RESIDENT-HARD
                    # constitution when this is not a critical pull-down repair.
                    strict_candidates=[r for r in candidate_rows if r["rh_delta"]<=0]
                    shown_candidates=strict_candidates or candidate_rows
                    shown_candidates=sorted(shown_candidates,key=lambda r:(
                        r["rh_total"],r["rh_max"],r["rh_cum_spread"],repair_load.get(r["initials"],0),r["initials"]
                    ))
                if shown_candidates:
                    c1,c2=st.columns(2)
                    with c1:
                        repl_row=st.selectbox(
                            tr("repair_replacement"),shown_candidates,
                            format_func=lambda r:(
                                f"{r['initials']} · "
                                + (f"PULL iš {r['source_department']} {block_label(r['source_block'])} · " if r.get('source_slot') is not None else ("BLOCK-FREE fallback · " if r.get('mode')=="FREE_FALLBACK" else ""))
                                + f"RH Δ {r['rh_delta']:+d} · repair {repair_load.get(r['initials'],0)}"
                            ),
                            key="repair_replacement"
                        )
                        repl=repl_row["initials"]
                        if target_critical and repl_row.get("source_slot") is not None:
                            st.success(
                                f"CRITICAL COVER: {repl} bus perkeltas iš neprivalomo posto "
                                f"{repl_row['source_department']} ({block_label(repl_row['source_block'])}) į "
                                f"{chosen.department} ({block_label(chosen.block)}). Donorinis optional postas lieka tuščias."
                            )
                        elif target_critical and repl_row.get("mode")=="FREE_FALLBACK":
                            st.warning(
                                "Nerasta saugaus tos pačios pamainos donorinio rezidento iš žemesnės hierarchijos optional posto; "
                                "todėl rodomas tame laiko bloke laisvo rezidento fallback."
                            )
                        if repl_row["rh_delta"]>0:
                            st.warning("RESIDENT HARD conflict: this candidate is invalid under V2.5.107 and must not be used for SYSTEM generation.")
                    with c2:
                        reason_options=[("sickness",tr("repair_reason_sickness")),("leave",tr("repair_reason_leave")),("approved_absence",tr("repair_reason_approved")),("force_majeure",tr("repair_reason_force"))]
                        reason=st.selectbox(tr("repair_reason"),reason_options,format_func=lambda x:x[1],key="repair_reason")[0]
                    note=st.text_input(tr("repair_note"),key="repair_note")
                    st.caption(tr("repair_load_help"))
                    if st.button(tr("apply_repair"),type="primary",key="apply_repair_btn"):
                        source_sid=repl_row.get("source_slot")
                        ok,why,stats=_repair_candidate_check(year,month,fresh,chosen.idx,repl,source_sid)
                        if not ok:
                            st.error(f"{tr('repair_invalid')}: {why}")
                        else:
                            old_person=fresh.assignments[chosen.idx]
                            source_sl=slots_by_id.get(int(source_sid)) if source_sid is not None else None
                            if target_critical:
                                fresh.assignments=apply_emergency_critical_transfer(
                                    fresh.assignments,chosen,repl,source_sl
                                )
                            else:
                                fresh.assignments[chosen.idx]=repl
                            fresh.stats=stats
                            auto_note=""
                            if source_sl is not None:
                                auto_note=(
                                    f"V2.5.56 CRITICAL PULL-DOWN: {repl} moved from optional "
                                    f"{source_sl.department} {source_sl.block} to mandatory {chosen.department} {chosen.block}; "
                                    f"source slot #{source_sl.idx} intentionally left unfilled."
                                )
                            final_note=(auto_note + (" | " + note if note else "")).strip(" |")
                            # ACTUAL changes; baseline_json and fairness_history are deliberately untouched.
                            db.apply_schedule_repair(year,month,serialize_result(fresh),chosen.idx,chosen.day,chosen.department,chosen.block,old_person,repl,reason,final_note)
                            sync_backup_plan(year,month,fresh)
                            persist_actual_satisfaction(year,month)
                            refresh_calendar_subscription_feeds([old_person,repl])
                            st.success(tr("repair_applied")); st.rerun()
                else:
                    st.warning(tr("repair_no_candidate"))
            if repair_rows:
                st.markdown(f"#### {tr('repair_history')}")
                rlabels={"sickness":tr("repair_reason_sickness"),"leave":tr("repair_reason_leave"),"approved_absence":tr("repair_reason_approved"),"force_majeure":tr("repair_reason_force")}
                rdf=pd.DataFrame([{tr("repair_date"):f"{int(r['day']):02d} · {r['department']} · {block_label(r['block'])}",tr("repair_from"):r['from_person'],tr("repair_to"):r['to_person'],tr("repair_reason"):rlabels.get(r.get('reason'),r.get('reason','')),tr("repair_note"):r.get('note') or "—",tr("status"):tr("repair_fairness_neutral")} for r in repair_rows])
                st.dataframe(rdf,use_container_width=True,hide_index=True)
                ldf=pd.DataFrame([{tr("person"):i,tr("repair_load"):n} for i,n in sorted(repair_load.items(),key=lambda kv:(-kv[1],kv[0]))])
                if not ldf.empty: st.dataframe(ldf,use_container_width=True,hide_index=True)
pos+=1

# --- Calendar ---
with tabs[pos]:
    st.subheader(tr("calendar_title")); currentp=db.load_schedule(year,month,"current")
    if not currentp: st.info(tr("not_published"))
    elif not resident_ok: st.error(tr("bad_pin"))
    else:
        result=refresh_result_payload(currentp,year,month)
        st.markdown(badge(active_user),unsafe_allow_html=True)
        if active_user in (SENIOR_INITIALS,WESTON_CREDITOR_INITIALS):
            try:
                _weston_calendar=db.weston_beer_stats_v25110(year,month)
                if active_user==SENIOR_INITIALS:
                    st.metric(("WESTON skola ŠR" if lang=="LT" else "WESTON debt to ŠR"),int(_weston_calendar.get("total_beers",0)),help=("Kiekvienas SP Generate/Rebuild paspaudimas = +1." if lang=="LT" else "Every SP Generate/Rebuild click = +1."))
                else:
                    st.metric(("WESTON, kuriuos SP tau skolinga" if lang=="LT" else "WESTONs SP owes you"),int(_weston_calendar.get("total_beers",0)),help=("Tas pats persistent skaičius, kurį SP mato kaip skolą." if lang=="LT" else "The same persistent total SP sees as debt."))
            except Exception:
                pass
        st.markdown("### Mano normalios darbo pamainos" if lang=="LT" else "### My normal work shifts")
        st.caption(
            "Tai vienintelis darbo grafiko sluoksnis. Teoriniai dubliai čia NĖRA skaičiuojami kaip darbas."
            if lang=="LT" else
            "This is the actual work-schedule layer. Theoretical backup duties are NOT counted as work here."
        )
        st.dataframe(personal_schedule_df(year,month,result,active_user),use_container_width=True,hide_index=True)
        st.markdown("### Teorinis dublių / pavadavimo sluoksnis" if lang=="LT" else "### Theoretical backup / standby layer")
        st.caption(
            "Dublis yra tik standby planas: jis nekeičia pageidavimų, darbo krūvio, poilsio, water-fill ar normalios pamainos statistikos. Tik pažymėtas COMPLETED realus pavadavimas tampa ACTUAL darbu."
            if lang=="LT" else
            "A backup is standby only: it does not change preferences, workload, rest, water-fill or normal-shift statistics. Only a COMPLETED real-life cover becomes ACTUAL work."
        )
        st.dataframe(backup_grid(year,month,result,active_user),use_container_width=True)
        ics_bytes=build_ics(year,month,result,active_user)
        st.download_button(tr("download_ics"),ics_bytes,file_name=f"{safe_filename(active_user)}_{year}_{month:02d}.ics",mime="text/calendar",type="primary")
        st.caption(tr("calendar_help"))
        try:
            feed_ics=build_calendar_subscription_ics(active_user)
            feed_url=db.publish_calendar_feed(active_user,feed_ics)
            st.markdown(f"### {tr('calendar_feed')}")
            st.warning(tr("calendar_feed_private"))
            st.code(feed_url,language=None)
            cga,cap,coth=st.columns(3)
            with cga:
                st.link_button(tr("calendar_google"),"https://calendar.google.com/calendar/u/0/r/settings/addbyurl",use_container_width=True)
                st.caption(tr("calendar_google_help"))
            with cap:
                apple_url="webcal://"+feed_url.split("://",1)[-1]
                st.markdown(f'<a href="{html.escape(apple_url,quote=True)}" style="display:block;text-align:center;padding:.55rem .75rem;border:1px solid rgba(128,128,128,.35);border-radius:.5rem;text-decoration:none;font-weight:600;">{html.escape(tr("calendar_apple"))}</a>',unsafe_allow_html=True)
                st.caption(tr("calendar_apple_help"))
            with coth:
                st.link_button(tr("calendar_other"),"https://outlook.live.com/calendar/0/addcalendar",use_container_width=True)
                st.caption(tr("calendar_other_help"))
                st.download_button(("ATSISIŲSTI .ics" if lang=="LT" else "DOWNLOAD .ics"),ics_bytes,file_name=f"{safe_filename(active_user)}_{year}_{month:02d}.ics",mime="text/calendar",use_container_width=True,key=f"ics_fallback_{year}_{month}_{active_user}")
        except Exception as exc:
            st.caption(("Kalendoriaus prenumeratos nuorodos dar nepavyko atnaujinti; .ics atsisiuntimas veikia." if lang=="LT" else "Could not refresh the subscription feed yet; .ics download still works.")+f" ({exc})")
pos+=1



# ===== V2.5.41 RESEARCH: AVAILABLE GPT + HUMAN vs MY ENGINE TOOL =====

_RESEARCH_PREF_COLUMNS = [
    "initials","name","unavailable","unavailable_am","unavailable_pm","vacation","justified_absence","long_duty",
    "soft_free","soft_free_am","soft_free_pm","preferred","preferred_am","preferred_pm",
    "spread_preference","holiday_preference","shift_length_preference","avoid_doubles","target_adjustment",
    "prior_weekend_count","prior_friday_count","prior_double_count","prior_weekday_day_count",
] + [f"prior_rotation__{cat}" for cat in ROTATION_CATEGORIES]

def _research_norm(x):
    s="" if x is None else str(x)
    s=unicodedata.normalize("NFKD",s).encode("ascii","ignore").decode("ascii").lower().strip()
    s=re.sub(r"\s+"," ",s)
    s=re.sub(r"[^a-z0-9]+","",s)
    return s

def _research_colmap(df):
    aliases={
        "initials":["initials","inic","inicialai","rezidentas","person","zmogus"],
        "name":["name","vardas","vardaspavarde"],
        "unavailable":["unavailable","negaliudirbti","hardoff","negaliu"],
        "unavailable_am":["unavailableam","negaliuryte","negaliuam"],
        "unavailable_pm":["unavailablepm","negaliupopiete","negaliupm"],
        "vacation":["vacation","atostogos","atostogauja","leave"],
        "justified_absence":["justifiedabsence","pateisinamasneatvykimas","liga","absence"],
        "long_duty":["longduty","24h","budėjimas24h","budejimas24h"],
        "soft_free":["softfree","noriulaisvos","pageidaujalaisvos","preferredoff"],
        "soft_free_am":["softfreeam","noriulaisvoryto"],
        "soft_free_pm":["softfreepm","noriulaisvospopietes"],
        "preferred":["preferred","pageidaujadirbti","pageidaujudirbti","pageidauju","noriudirbti","preferredwork"],
        "preferred_am":["preferredam","pageidaujadirbtiryte"],
        "preferred_pm":["preferredpm","pageidaujadirbtipopiete"],
        "weekday_preference":["weekdaypreference","darbdienupreference"],
        "weekend_preference":["weekendpreference","savaitgaliupreference"],
        "spread_preference":["spreadpreference","issklaidymas","koncentracija"],
        "holiday_preference":["holidaypreference","svenciupreference","sventes","publicholidays"],
        "shift_length_preference":["shiftlengthpreference","darbotrukme","pamainustrukme","workdaylength"],
        "avoid_doubles":["avoiddoubles","vengtidubliu","vengtidvigubu"],
        "target_adjustment":["targetadjustment","targetokorekcija"],
        "prior_weekend_count":["priorweekendcount","ankstesnisavaitgaliai"],
        "prior_friday_count":["priorfridaycount","ankstesnipenktadieniai"],
        "prior_double_count":["priordoublecount","ankstesnidubliai"],
        "prior_weekday_day_count":["priorweekdaydaycount","ankstesnesdarbdienos"],
    }
    actual={_research_norm(c):c for c in df.columns}
    out={}
    for key,opts in aliases.items():
        for opt in opts:
            n=_research_norm(opt)
            if n in actual:
                out[key]=actual[n]; break
    return out

def _research_days(v, ndays, y=None, m=None):
    """Parse day lists from real-world monthly wishes Excel cells.

    Supports single days, comma lists, Lithuanian/English ranges (22-24, 5 iki 14),
    and the phrase "Visi savaitgaliai šio mėnesio". Clock times are stripped before
    day parsing so e.g. "6 d 08:00-14:00" becomes day 6 rather than 6/8/14.
    """
    if v is None or (isinstance(v,float) and pd.isna(v)):
        return set()
    if isinstance(v,(int,float)) and not isinstance(v,bool):
        d=int(v)
        return {d} if 1 <= d <= ndays else set()

    raw=str(v).strip()
    ascii_text=unicodedata.normalize("NFKD",raw).encode("ascii","ignore").decode("ascii").lower()
    out=set()

    # Natural-language month shortcut used by the original August workbook.
    if y is not None and m is not None and (
        "visi savaitgaliai" in ascii_text
        or "all weekends" in ascii_text
        or "kiekviena savaitgali" in ascii_text
    ):
        out |= {
            d for d in range(1,ndays+1)
            if date(int(y),int(m),d).weekday() >= 5
        }

    # Remove clock expressions before interpreting day numbers.
    cleaned=re.sub(
        r"(?<!\d)\d{1,2}[:.]\d{2}\s*(?:-|–|—|iki|to)\s*\d{1,2}[:.]\d{2}(?!\d)",
        " ",ascii_text,flags=re.I
    )
    cleaned=re.sub(r"(?<!\d)\d{1,2}[:.]\d{2}(?!\d)"," ",cleaned)

    # Expand day ranges first, including Lithuanian "iki".
    spans=[]
    for mt in re.finditer(r"(?<!\d)([0-3]?\d)\s*(?:-|–|—|iki|to)\s*([0-3]?\d)(?!\d)",cleaned,re.I):
        a,b=int(mt.group(1)),int(mt.group(2))
        if 1 <= a <= ndays and 1 <= b <= ndays:
            lo,hi=sorted((a,b))
            out.update(range(lo,hi+1))
            spans.append(mt.span())

    # Blank matched spans so their endpoints are not handled twice (harmless but cleaner).
    chars=list(cleaned)
    for a,b in spans:
        for i in range(a,b):
            chars[i]=" "
    remainder="".join(chars)
    out |= {
        int(x) for x in re.findall(r"(?<!\d)([0-3]?\d)(?!\d)",remainder)
        if 1 <= int(x) <= ndays
    }
    return out

def _research_int(v, default=0):
    if v is None or (isinstance(v,float) and pd.isna(v)):
        return default
    try: return int(float(v))
    except Exception: return default

def _research_bool(v):
    if isinstance(v,bool): return v
    return _research_norm(v) in {"1","true","yes","taip","y","t"}

def research_preferences_template(y,m):
    rows=[]
    for p in DEFAULT_PEOPLE:
        rows.append({c:"" for c in _RESEARCH_PREF_COLUMNS})
        rows[-1]["initials"]=p["initials"]; rows[-1]["name"]=p["name"]
        rows[-1]["target_adjustment"]=p.get("target_adjustment",0)
    guide=pd.DataFrame([
        ["Date lists","Use day numbers separated by commas, e.g. 2, 5, 17. Leave blank = N/A / no preference."],
        ["HARD","unavailable / vacation / other justified absence / long duty are HARD."],
        ["SOFT","soft_free and preferred fields are active SOFT preferences; blank means N/A."],
        ["SOFT-3 month shape","spread_preference: -2..+2 (clustered ↔ dispersed); blank/0 = N/A. Broad weekday/weekend direction is deprecated and ignored in V2.5.52."],
        ["Prior fairness","prior_* columns are optional. Leave 0 if historical carry-in is unavailable."],
        ["Prior workplace exposure","prior_rotation__* columns are optional cumulative SYSTEM counts from earlier months. Leave blank/0 if unavailable."],
        ["Month",f"{y}-{m:02d}"],
    ],columns=["Field","Instruction"])
    bio=BytesIO()
    with pd.ExcelWriter(bio,engine="xlsxwriter") as w:
        pd.DataFrame(rows).to_excel(w,index=False,sheet_name="preferences")
        guide.to_excel(w,index=False,sheet_name="README")
    return bio.getvalue()

def research_schedule_template(y,m):
    rows=[]
    for s in make_slots(y,m):
        if s.blocked: continue
        rows.append({
            "slot_id":s.idx,
            "date":f"{y}-{m:02d}-{s.day:02d}",
            "department":s.department,
            "shift":s.block,
            "person":"",
        })
    guide=pd.DataFrame([
        ["person","Enter initials exactly as roster uses them, e.g. ŠR / MG"],
        ["Rows","Leave unfilled/unused optional slots blank. Do not delete slot_id if using this template."],
        ["Existing Excel","Long-format schedules are also accepted if columns equivalent to date/day, department, shift, person exist."],
        ["Month",f"{y}-{m:02d}"],
    ],columns=["Field","Instruction"])
    bio=BytesIO()
    with pd.ExcelWriter(bio,engine="xlsxwriter") as w:
        pd.DataFrame(rows).to_excel(w,index=False,sheet_name="schedule")
        guide.to_excel(w,index=False,sheet_name="README")
    return bio.getvalue()


_RESEARCH_LT_MONTHS={
    1:["sausis","sausio"],2:["vasaris","vasario"],3:["kovas","kovo"],4:["balandis","balandzio"],
    5:["geguze","geguzes"],6:["birzelis","birzelio"],7:["liepa","liepos"],8:["rugpjutis","rugpjucio"],
    9:["rugsejis","rugsejo"],10:["spalis","spalio"],11:["lapkritis","lapkricio"],12:["gruodis","gruodzio"],
}
_RESEARCH_EN_MONTHS={
    1:["january","jan"],2:["february","feb"],3:["march","mar"],4:["april","apr"],
    5:["may"],6:["june","jun"],7:["july","jul"],8:["august","aug"],
    9:["september","sep","sept"],10:["october","oct"],11:["november","nov"],12:["december","dec"],
}

def _research_file_bytes(uploaded):
    if hasattr(uploaded,"getvalue"):
        return uploaded.getvalue()
    if isinstance(uploaded,(bytes,bytearray)):
        return bytes(uploaded)
    pos=None
    try:
        pos=uploaded.tell()
    except Exception:
        pass
    raw=uploaded.read()
    if pos is not None:
        try: uploaded.seek(pos)
        except Exception: pass
    return raw


def _research_workbook_hash(uploaded):
    return hashlib.sha256(_research_file_bytes(uploaded)).hexdigest()


def _research_sheet_month_hint(sheet_name):
    n=unicodedata.normalize("NFKD",str(sheet_name or "")).encode("ascii","ignore").decode("ascii").lower()
    n=re.sub(r"[^a-z0-9]+"," ",n)
    found=set()
    for mm,toks in _RESEARCH_LT_MONTHS.items():
        if any(re.search(rf"\b{re.escape(t)}\w*\b",n) for t in toks):
            found.add(mm)
    for mm,toks in _RESEARCH_EN_MONTHS.items():
        if any(re.search(rf"\b{re.escape(t)}\w*\b",n) for t in toks):
            found.add(mm)
    # Numeric hints only when clearly month-like.
    for mm in range(1,13):
        if re.search(rf"(?:^|[\s_\-./])0?{mm}(?:$|[\s_\-./])",str(sheet_name or "")):
            found.add(mm)
    return sorted(found)


def _research_make_unique_headers(values):
    out=[]; seen={}
    for idx,v in enumerate(values):
        if v is None or (isinstance(v,float) and pd.isna(v)):
            base=f"__col_{idx+1}"
        else:
            base=str(v).strip() or f"__col_{idx+1}"
        count=seen.get(base,0)
        seen[base]=count+1
        out.append(base if count==0 else f"{base}__{count+1}")
    return out


def _research_table_candidate(raw,header_row,kind):
    if header_row<0 or header_row>=len(raw):
        return None
    headers=_research_make_unique_headers(raw.iloc[header_row].tolist())
    df=raw.iloc[header_row+1:].copy()
    df.columns=headers
    df=df.dropna(axis=0,how="all").dropna(axis=1,how="all")
    if df.empty:
        return None

    if kind=="preferences":
        cmap=_research_colmap(df)
        mapped=set(cmap)
        # prior_rotation columns are valid preference metadata too.
        prior_cols=[]
        for cat in ROTATION_CATEGORIES:
            target=_research_norm(f"prior_rotation__{cat}")
            for c in df.columns:
                if _research_norm(c)==target:
                    prior_cols.append(c); break
        # Original clinic wishes workbooks commonly identify residents by full name only.
        # Either initials OR a name column is therefore a valid identity field.
        if "initials" not in cmap and "name" not in cmap:
            return None
        nonidentity=mapped-{"initials","name"}
        score=10+len(nonidentity)*3+len(prior_cols)
        # A resident-directory/helper sheet with only initials/name is not enough.
        if not nonidentity and not prior_cols:
            return None
        return {"df":df,"cmap":cmap,"score":score,"prior_cols":prior_cols}

    if kind=="schedule":
        cmap=_research_schedule_colmap(df)
        if "person" not in cmap:
            return None
        has_locator=("slot_id" in cmap) or (
            ("date" in cmap or "day" in cmap)
            and "department" in cmap and "shift" in cmap
        )
        if not has_locator:
            return None
        score=10+len(cmap)*3
        return {"df":df,"cmap":cmap,"score":score}

    return None


def _research_scan_workbook(uploaded,kind,y,m):
    """Scan every worksheet and detect the best compatible table per sheet."""
    raw_bytes=_research_file_bytes(uploaded)
    wb_hash=hashlib.sha256(raw_bytes).hexdigest()
    xls=pd.ExcelFile(BytesIO(raw_bytes))
    tables=[]; audit=[]
    for sheet in xls.sheet_names:
        month_hints=_research_sheet_month_hint(sheet)
        if month_hints and int(m) not in month_hints:
            audit.append({
                "sheet":sheet,"status":"skipped_other_month",
                "reason":f"sheet name suggests month(s) {month_hints}, selected month is {m}",
            })
            continue
        try:
            raw=pd.read_excel(xls,sheet_name=sheet,header=None,dtype=object)
        except Exception as exc:
            audit.append({"sheet":sheet,"status":"read_error","reason":str(exc)})
            continue
        # Keep original row indices intact for exact Excel row provenance.
        raw=raw.dropna(axis=1,how="all")
        if raw.dropna(axis=0,how="all").empty:
            audit.append({"sheet":sheet,"status":"ignored_empty","reason":"empty worksheet"})
            continue

        candidates=[]
        max_header=min(len(raw)-1,60)
        for hr in range(max_header+1):
            cand=_research_table_candidate(raw,hr,kind)
            if cand is not None:
                candidates.append((cand["score"],-hr,hr,cand))
        if not candidates and kind=="schedule":
            grid=_research_grid_schedule_candidate(raw,y,m)
            if grid is not None:
                rec={
                    "sheet":sheet,
                    "header_row":grid["header_row"]+1,
                    "score":grid["score"],
                    "rows":len(grid["recognized_rows"]),
                    "mode":"grid",
                    "raw":raw,
                    "day_cols":grid["day_cols"],
                    "recognized_rows":grid["recognized_rows"],
                    "workbook_hash":wb_hash,
                }
                tables.append(rec)
                audit.append({
                    "sheet":sheet,"status":"used_grid_schedule",
                    "header_row":grid["header_row"]+1,
                    "rows":len(grid["recognized_rows"]),
                    "score":grid["score"],
                })
                continue

        if not candidates:
            audit.append({
                "sheet":sheet,"status":"ignored_unrecognized",
                "reason":f"no compatible {kind} table detected in first {max_header+1} rows",
            })
            continue
        _score,_neg_hr,hr,cand=max(candidates,key=lambda x:(x[0],x[1]))
        rec={
            "sheet":sheet,"header_row":hr+1,"score":cand["score"],
            "rows":len(cand["df"]),"df":cand["df"],"cmap":cand["cmap"],
            "workbook_hash":wb_hash,"mode":"long",
        }
        if "prior_cols" in cand:
            rec["prior_cols"]=cand["prior_cols"]
        tables.append(rec)
        audit.append({
            "sheet":sheet,"status":"used_long_table","header_row":hr+1,
            "rows":len(cand["df"]),"score":cand["score"],
        })
    return tables,{
        "kind":kind,
        "workbook_hash":wb_hash,
        "sheet_count":len(xls.sheet_names),
        "used_sheet_count":len(tables),
        "sheets":audit,
    }


def _research_nonempty(v):
    if v is None:
        return False
    if isinstance(v,float) and pd.isna(v):
        return False
    return bool(str(v).strip()) and str(v).strip().lower()!="nan"


def _research_scalar_conflict_key(field,value):
    if field=="name":
        return _research_norm(value)
    if field=="avoid_doubles":
        return str(bool(_research_bool(value)))
    if field in {
        "weekday_preference","weekend_preference","holiday_preference","spread_preference","target_adjustment",
        "prior_weekend_count","prior_friday_count","prior_double_count","prior_weekday_day_count",
    }:
        return str(_research_int(value,0))
    return str(value).strip()


def _research_explicit_shift(v):
    """Return AM/PM/FULL only when the cell explicitly contains shift/time information.

    Plain day numbers like "15 d" must never be mistaken for a 15:00 PM shift.
    """
    if not _research_nonempty(v):
        return ""
    raw=str(v).strip()
    n=_research_norm(raw)
    times=[(int(h),int(mm)) for h,mm in re.findall(r"(?<!\d)([0-2]?\d)[:.]([0-5]\d)(?!\d)",raw)]
    if len(times)>=2:
        start=times[0][0]+times[0][1]/60.0
        end=times[1][0]+times[1][1]/60.0
        if start <= 9 and end >= 16.5:
            return "FULL"
        if start <= 9 and end <= 14.5:
            return "AM"
        if start >= 13:
            return "PM"
    if re.search(r"(?:^|[^a-z])(am)(?:$|[^a-z])",raw,re.I) or "ryt" in n:
        return "AM"
    if re.search(r"(?:^|[^a-z])(pm)(?:$|[^a-z])",raw,re.I) or "popiet" in n or "vakar" in n:
        return "PM"
    if "visadiena" in n or "pilnadiena" in n or "full" in n:
        return "FULL"
    return ""


def _research_resolve_resident(raw_value, exact_initials, known_names):
    """Resolve initials or a full name to the canonical roster initials.

    Exact matches are preferred. A conservative fuzzy fallback handles harmless source
    spelling differences such as Stašinskas/Strašinskas without silently matching
    genuinely different people.
    """
    if not _research_nonempty(raw_value):
        return None, None
    raw_text=str(raw_value).strip()
    direct=exact_initials.get(raw_text.casefold()) or known_names.get(_research_norm(raw_text))
    if direct:
        return direct, None

    import difflib
    target=_research_norm(raw_text)
    ranked=sorted(
        ((difflib.SequenceMatcher(None,target,nm).ratio(),ini,nm) for nm,ini in known_names.items()),
        reverse=True
    )
    if ranked and ranked[0][0] >= 0.92 and (len(ranked)==1 or ranked[0][0]-ranked[1][0] >= 0.05):
        return ranked[0][1], f"fuzzy name match '{raw_text}' -> {ranked[0][1]} ({ranked[0][0]:.2f})"
    return None, None


def research_people_from_excel(uploaded,y,m,return_audit=False):
    tables,audit=_research_scan_workbook(uploaded,"preferences",y,m)
    if not tables:
        raise ValueError(
            "No compatible preferences/HARD table was found anywhere in the workbook. "
            "The importer scanned every worksheet. A compatible table may identify residents "
            "by initials OR by full name and should contain at least one preference/HARD column."
        )

    ndays=calendar.monthrange(y,m)[1]
    exact_initials={p["initials"].casefold():p["initials"] for p in DEFAULT_PEOPLE}
    exact_initials.update({"sk":"SŠ","sr":"SP"})  # V2.5.113 historical-import compatibility
    known_names={_research_norm(p["name"]):p["initials"] for p in DEFAULT_PEOPLE}
    set_fields={
        "unavailable","unavailable_am","unavailable_pm","vacation","justified_absence","long_duty",
        "soft_free","soft_free_am","soft_free_pm","preferred","preferred_am","preferred_pm",
    }
    scalar_fields={
        "name","weekday_preference","weekend_preference","holiday_preference","spread_preference",
        "avoid_doubles","target_adjustment","prior_weekend_count","prior_friday_count",
        "prior_double_count","prior_weekday_day_count",
    }

    collected={p["initials"]:{
        "sets":{k:set() for k in set_fields},
        "scalars":{},
        "prior_rotation":{cat:[] for cat in ROTATION_CATEGORIES},
        "sources":[],
    } for p in DEFAULT_PEOPLE}
    warnings=[]; conflicts=[]

    for table in tables:
        if table.get("mode")=="grid":
            raw=table["raw"]
            local_used=set()
            day_cols=table["day_cols"]
            recognized_rows=sorted(table["recognized_rows"],key=lambda x:x[0])

            for ridx,fam,descriptor in recognized_rows:
                explicit_block=_research_shift(descriptor)
                for col_idx,day in day_cols:
                    if col_idx>=raw.shape[1]:
                        continue
                    cell=raw.iloc[ridx,col_idx]
                    persons,unknown_parts=_research_cell_people(
                        cell,exact_initials,known_names
                    )
                    for unknown in unknown_parts:
                        warnings.append(
                            f"{table['sheet']} row {int(ridx)+1}, day {day}: "
                            f"unknown resident/cell value '{unknown}'."
                        )
                    if not persons:
                        continue

                    candidates=[
                        s for s in slots
                        if s.day==day
                        and not s.blocked
                        and _research_slot_family(s)==fam
                        and s.idx not in local_used
                    ]
                    if explicit_block:
                        candidates=[s for s in candidates if s.block==explicit_block]

                    for person in persons:
                        if not candidates:
                            warnings.append(
                                f"{table['sheet']} row {int(ridx)+1}, day {day}: "
                                f"no remaining slot for {fam}"
                                + (f" {explicit_block}" if explicit_block else "")
                                + f" while reading '{cell}'."
                            )
                            break
                        s=sorted(candidates,key=lambda x:x.idx)[0]
                        candidates=[x for x in candidates if x.idx!=s.idx]
                        local_used.add(s.idx)
                        sid=s.idx
                        src={
                            "workbook_hash":table["workbook_hash"],
                            "sheet":table["sheet"],
                            "source_row":int(ridx)+1,
                            "source_column":int(col_idx)+1,
                            "day":int(day),
                            "header_row":table["header_row"],
                            "mode":"grid",
                        }
                        if sid in assignments:
                            previous=assignments[sid]
                            prev_src=provenance[sid]
                            if previous==person:
                                duplicates.append({
                                    "slot_id":sid,"person":person,
                                    "first_source":prev_src,"duplicate_source":src,
                                })
                            else:
                                conflicts.append({
                                    "slot_id":sid,
                                    "existing_person":previous,
                                    "new_person":person,
                                    "existing_source":prev_src,
                                    "new_source":src,
                                })
                            continue
                        assignments[sid]=person
                        provenance[sid]=src
            continue

        df=table["df"]; cmap=table["cmap"]
        for ridx,r in df.iterrows():
            # Prefer initials when present, otherwise use the full-name column from the
            # original clinic worksheet.
            raw_identity=None
            if "initials" in cmap and _research_nonempty(r.get(cmap["initials"])):
                raw_identity=r.get(cmap["initials"])
            elif "name" in cmap and _research_nonempty(r.get(cmap["name"])):
                raw_identity=r.get(cmap["name"])
            if not _research_nonempty(raw_identity):
                continue
            raw_text=str(raw_identity).strip()
            ini,match_note=_research_resolve_resident(raw_identity,exact_initials,known_names)
            if not ini:
                warnings.append(
                    f"{table['sheet']} row {int(ridx)+1}: unknown resident '{raw_text}'."
                )
                continue
            if match_note:
                warnings.append(f"{table['sheet']} row {int(ridx)+1}: {match_note}.")
            src={
                "workbook_hash":table["workbook_hash"],
                "sheet":table["sheet"],
                "source_row":int(ridx)+1,
                "header_row":table["header_row"],
            }
            collected[ini]["sources"].append(src)

            for field in set_fields:
                if field not in cmap:
                    continue
                raw=r.get(cmap[field])
                if not _research_nonempty(raw):
                    continue
                days=_research_days(raw,ndays,y,m)
                target_field=field
                # Original sheets sometimes encode half-day wishes as a time range in
                # the ordinary column, e.g. "6 d 08:00-14:00". Route that date to AM/PM.
                if field in {"unavailable","soft_free","preferred"}:
                    shift=_research_explicit_shift(raw)
                    if shift=="AM":
                        target_field=f"{field}_am"
                    elif shift=="PM":
                        target_field=f"{field}_pm"
                collected[ini]["sets"][target_field] |= days

            for field in scalar_fields:
                if field not in cmap:
                    continue
                raw=r.get(cmap[field])
                if not _research_nonempty(raw):
                    continue
                norm_val=_research_scalar_conflict_key(field,raw)
                collected[ini]["scalars"].setdefault(field,[]).append(
                    {"value":raw,"norm":norm_val,"source":src}
                )

            for cat in ROTATION_CATEGORIES:
                target=_research_norm(f"prior_rotation__{cat}")
                matching=next((c for c in df.columns if _research_norm(c)==target),None)
                if matching is None:
                    continue
                raw=r.get(matching)
                if _research_nonempty(raw):
                    collected[ini]["prior_rotation"][cat].append(
                        {"value":raw,"source":src}
                    )

    # True scalar conflicts block official lock rather than silently overwriting.
    for ini,data in collected.items():
        for field,vals in data["scalars"].items():
            distinct={}
            for rec in vals:
                distinct.setdefault(rec["norm"],rec)
            if len(distinct)>1:
                conflicts.append({
                    "resident":ini,"field":field,
                    "values":[str(v["value"]) for v in distinct.values()],
                    "sources":[v["source"] for v in distinct.values()],
                })
        for cat,vals in data["prior_rotation"].items():
            parsed={_research_int(v["value"],0) for v in vals}
            if len(parsed)>1:
                conflicts.append({
                    "resident":ini,"field":f"prior_rotation__{cat}",
                    "values":sorted(parsed),
                    "sources":[v["source"] for v in vals],
                })

    audit["conflicts"]=conflicts
    audit["warnings"]=warnings
    if conflicts:
        preview="; ".join(
            f"{c['resident']} {c['field']}={c['values']}" for c in conflicts[:6]
        )
        raise ValueError(
            "Whole-workbook preference import found conflicting scalar values. "
            "Official lock is blocked until resolved: "+preview
        )

    people=[]
    for base in DEFAULT_PEOPLE:
        ini=base["initials"]; data=collected[ini]
        def scalar(field,default):
            vals=data["scalars"].get(field) or []
            return default if not vals else vals[0]["value"]
        people.append(Person(
            initials=ini,
            name=str(scalar("name",base["name"]) or base["name"]),
            unavailable=set(data["sets"]["unavailable"]),
            unavailable_am=set(data["sets"]["unavailable_am"]),
            unavailable_pm=set(data["sets"]["unavailable_pm"]),
            vacation=set(data["sets"].get("vacation",set())),
            justified_absence=set(data["sets"]["justified_absence"]),
            long_duty=set(data["sets"]["long_duty"]),
            soft_free=set(data["sets"]["soft_free"]),
            soft_free_am=set(data["sets"]["soft_free_am"]),
            soft_free_pm=set(data["sets"]["soft_free_pm"]),
            preferred=set(data["sets"]["preferred"]),
            preferred_am=set(data["sets"]["preferred_am"]),
            preferred_pm=set(data["sets"]["preferred_pm"]),
            weekday_preference=max(-2,min(2,_research_int(scalar("weekday_preference",0),0))),
            weekend_preference=max(-2,min(2,_research_int(scalar("weekend_preference",0),0))),
            holiday_preference=max(-1,min(1,_research_int(scalar("holiday_preference",0),0))),
            spread_preference=max(-2,min(2,_research_int(scalar("spread_preference",0),0))),
            shift_length_preference=max(0,min(3,_research_int(scalar("shift_length_preference",0),0))),
            avoid_doubles=_research_bool(scalar("avoid_doubles",False)),
            target_adjustment=_research_int(
                scalar("target_adjustment",base.get("target_adjustment",0)),
                base.get("target_adjustment",0)
            ),
            prior_weekend_count=max(0,_research_int(scalar("prior_weekend_count",0),0)),
            prior_friday_count=max(0,_research_int(scalar("prior_friday_count",0),0)),
            prior_double_count=max(0,_research_int(scalar("prior_double_count",0),0)),
            prior_weekday_day_count=max(0,_research_int(scalar("prior_weekday_day_count",0),0)),
            prior_rotation_counts={
                cat:max(0,_research_int(
                    (data["prior_rotation"][cat][0]["value"] if data["prior_rotation"][cat] else 0),0
                ))
                for cat in ROTATION_CATEGORIES
            },
        ))

    audit["resident_sources"]={
        ini:data["sources"] for ini,data in collected.items() if data["sources"]
    }
    return (people,audit,warnings) if return_audit else people



def _research_shift(v):
    """Normalize historical shift labels/times to AM / PM / FULL."""
    raw="" if v is None else str(v)
    n=_research_norm(raw)
    compact=re.sub(r"[^0-9]","",raw)

    if (
        "full" in n or "visadiena" in n or "pilnadiena" in n
        or "0817" in compact or "817" in compact
    ):
        return "FULL"
    if (
        "pm" in n or "popiet" in n or "vak" in n
        or "1420" in compact or "14002000" in compact
    ):
        return "PM"
    if (
        "am" in n or "ryt" in n
        or "0814" in compact or "814" in compact or "08001400" in compact
    ):
        return "AM"
    # Exact common clock starts.
    mt=re.search(r"(?<!\d)(\d{1,2})[:.]?(\d{2})?(?!\d)",raw)
    if mt:
        hour=int(mt.group(1))
        if hour>=13:
            return "PM"
        if hour<=9:
            return "AM"
    return ""


def _research_grid_day_value(v,y,m):
    if v is None or (isinstance(v,float) and pd.isna(v)):
        return None
    ndays=calendar.monthrange(y,m)[1]
    if isinstance(v,(pd.Timestamp,datetime,date)):
        if int(v.year)==int(y) and int(v.month)==int(m):
            return int(v.day)
        return None
    if isinstance(v,(int,np.integer)):
        d=int(v)
        return d if 1<=d<=ndays else None
    if isinstance(v,float) and float(v).is_integer():
        d=int(v)
        return d if 1<=d<=ndays else None
    s=str(v).strip()
    mt=re.search(r"(\d{4})[-/.](\d{1,2})[-/.](\d{1,2})",s)
    if mt:
        yy,mm,dd=map(int,mt.groups())
        return dd if yy==y and mm==m and 1<=dd<=ndays else None
    # Header cells like "1", "01", "1 d.", "01 Mon".
    mt=re.match(r"^\s*([0-3]?\d)(?:\s*(?:d\.?|diena|mon|tue|wed|thu|fri|sat|sun|pr|an|tr|kt|pn|st|sk))?\s*$",s,re.I)
    if mt:
        d=int(mt.group(1))
        return d if 1<=d<=ndays else None
    return None


def _research_grid_dept_family(v):
    n=_research_norm(v)
    if not n:
        return None
    if "centroro" in n:
        return "CENTRO RO"
    if "onkoro" in n or ("onko" in n and "ro" in n):
        return "Onko RO"
    if "centroug" in n:
        return "Centro UG"
    if "spsug" in n:
        return "SPS UG"
    if "spsro" in n:
        return "SPS RO"
    if "144" in n:
        return "ADC 144"
    if re.search(r"(?:^|adc)145",n) or n.startswith("145"):
        return "ADC 145"
    if "vaik" in n and "ug" in n:
        return "Vaikų UG"
    if "mamograf" in n:
        return "Mamografijos"
    return None


def _research_grid_schedule_candidate(raw,y,m):
    """Detect classic schedule matrix: department rows × calendar-day columns."""
    best=None
    max_header=min(len(raw)-1,60)
    for hr in range(max_header+1):
        day_cols=[]
        for col_idx,v in enumerate(raw.iloc[hr].tolist()):
            day=_research_grid_day_value(v,y,m)
            if day is not None:
                day_cols.append((col_idx,day))
        unique_days={d for _,d in day_cols}
        if len(unique_days)<3:
            continue

        recognized=[]
        for ridx in range(hr+1,len(raw)):
            vals=raw.iloc[ridx].tolist()
            descriptor=" | ".join(
                str(vals[c]).strip()
                for c in range(min([x[0] for x in day_cols]+[len(vals)]))
                if c<len(vals) and _research_nonempty(vals[c])
            )
            fam=_research_grid_dept_family(descriptor)
            if fam:
                recognized.append((ridx,fam,descriptor))
        if len(recognized)<2:
            continue
        score=20+len(unique_days)+len(recognized)*2
        cand={
            "score":score,"header_row":hr,
            "day_cols":day_cols,"recognized_rows":recognized,
            "raw":raw,
        }
        if best is None or score>best["score"]:
            best=cand
    return best


def _research_schedule_colmap(df):
    aliases={
        "slot_id":["slot_id","slotid","id"],
        "date":["date","data"],
        "day":["day","diena"],
        "department":["department","skyrius","vieta","padalinys"],
        "shift":["shift","pamaina","block","laikas"],
        "person":["person","zmogus","rezidentas","initials","inic"],
    }
    actual={_research_norm(c):c for c in df.columns}
    out={}
    for k,opts in aliases.items():
        for opt in opts:
            if _research_norm(opt) in actual:
                out[k]=actual[_research_norm(opt)]; break
    return out

def _research_day_from_value(v,y,m):
    if isinstance(v,(pd.Timestamp,datetime,date)):
        return int(v.day)
    s=str(v)
    mt=re.search(r"(\d{4})[-/.](\d{1,2})[-/.](\d{1,2})",s)
    if mt:
        yy,mm,dd=map(int,mt.groups())
        if yy!=y or mm!=m: raise ValueError(f"Date {s} is outside selected month {y}-{m:02d}.")
        return dd
    nums=re.findall(r"(?<!\d)([0-3]?\d)(?!\d)",s)
    if nums: return int(nums[-1])
    raise ValueError(f"Could not read date/day value: {v}")

def _research_dept_key(v):
    n=_research_norm(v)
    n=n.replace("kabinetas","").replace("kab","")
    return n

def _research_slot_family(slot):
    return _research_grid_dept_family(slot.department)


def _research_cell_people(v,exact_initials,known_names):
    if not _research_nonempty(v):
        return [],[]
    raw=str(v).strip()
    whole=exact_initials.get(raw.casefold()) or known_names.get(_research_norm(raw))
    if whole:
        return [whole],[]
    parts=[
        p.strip()
        for p in re.split(r"[\n,;/|]+",raw)
        if p and p.strip()
    ]
    found=[]; unknown=[]
    for part in parts:
        person=exact_initials.get(part.casefold()) or known_names.get(_research_norm(part))
        if person:
            found.append(person)
        else:
            unknown.append(part)
    return found,unknown


def research_assignments_from_excel(uploaded,y,m,return_audit=False):
    tables,audit=_research_scan_workbook(uploaded,"schedule",y,m)
    if not tables:
        raise ValueError(
            "No compatible schedule table was found anywhere in the workbook. "
            "The importer scanned every worksheet."
        )

    slots=make_slots(y,m)
    slot_by_id={s.idx:s for s in slots}
    assignments={}
    provenance={}
    warnings=[]
    duplicates=[]
    conflicts=[]
    exact_initials={p["initials"].casefold():p["initials"] for p in DEFAULT_PEOPLE}
    exact_initials.update({"sk":"SŠ","sr":"SP"})  # V2.5.113 historical-import compatibility
    known_names={_research_norm(p["name"]):p["initials"] for p in DEFAULT_PEOPLE}

    for table in tables:
        if table.get("mode")=="grid":
            raw=table["raw"]
            local_used=set()
            day_cols=table["day_cols"]
            recognized_rows=sorted(table["recognized_rows"],key=lambda x:x[0])

            for ridx,fam,descriptor in recognized_rows:
                explicit_block=_research_shift(descriptor)
                for col_idx,day in day_cols:
                    if col_idx>=raw.shape[1]:
                        continue
                    cell=raw.iloc[ridx,col_idx]
                    persons,unknown_parts=_research_cell_people(cell,exact_initials,known_names)
                    for unknown in unknown_parts:
                        warnings.append(
                            f"{table['sheet']} row {int(ridx)+1}, day {day}: "
                            f"unknown resident/cell value '{unknown}'."
                        )
                    if not persons:
                        continue

                    candidates=[
                        s for s in slots
                        if s.day==day
                        and not s.blocked
                        and _research_slot_family(s)==fam
                        and s.idx not in local_used
                    ]
                    if explicit_block:
                        candidates=[s for s in candidates if s.block==explicit_block]

                    for person in persons:
                        if not candidates:
                            warnings.append(
                                f"{table['sheet']} row {int(ridx)+1}, day {day}: "
                                f"no remaining slot for {fam}"
                                + (f" {explicit_block}" if explicit_block else "")
                                + f" while reading '{cell}'."
                            )
                            break
                        s=sorted(candidates,key=lambda x:x.idx)[0]
                        candidates=[x for x in candidates if x.idx!=s.idx]
                        local_used.add(s.idx)
                        sid=s.idx
                        src={
                            "workbook_hash":table["workbook_hash"],
                            "sheet":table["sheet"],
                            "source_row":int(ridx)+1,
                            "source_column":int(col_idx)+1,
                            "day":int(day),
                            "header_row":table["header_row"],
                            "mode":"grid",
                        }
                        if sid in assignments:
                            previous=assignments[sid]
                            prev_src=provenance[sid]
                            if previous==person:
                                duplicates.append({
                                    "slot_id":sid,"person":person,
                                    "first_source":prev_src,"duplicate_source":src,
                                })
                            else:
                                conflicts.append({
                                    "slot_id":sid,
                                    "existing_person":previous,
                                    "new_person":person,
                                    "existing_source":prev_src,
                                    "new_source":src,
                                })
                            continue
                        assignments[sid]=person
                        provenance[sid]=src
            continue

        df=table["df"]; cmap=table["cmap"]
        for ridx,r in df.iterrows():
            raw_person=r.get(cmap["person"])
            if not _research_nonempty(raw_person):
                continue
            raw_person_text=str(raw_person).strip()
            person=exact_initials.get(raw_person_text.casefold()) or known_names.get(_research_norm(raw_person_text))
            src={
                "workbook_hash":table["workbook_hash"],
                "sheet":table["sheet"],
                "source_row":int(ridx)+1,
                "header_row":table["header_row"],
            }
            if not person:
                warnings.append(
                    f"{table['sheet']} row {int(ridx)+1}: unknown resident '{raw_person_text}'."
                )
                continue

            sid=None
            if "slot_id" in cmap:
                raw=r.get(cmap["slot_id"])
                try:
                    candidate=int(float(raw))
                    if candidate in slot_by_id:
                        sid=candidate
                except Exception:
                    pass

            if sid is None:
                if "date" in cmap:
                    day=_research_day_from_value(r.get(cmap["date"]),y,m)
                elif "day" in cmap:
                    day=_research_day_from_value(r.get(cmap["day"]),y,m)
                else:
                    warnings.append(
                        f"{table['sheet']} row {int(ridx)+1}: no usable slot_id/date/day."
                    )
                    continue
                if day is None:
                    warnings.append(
                        f"{table['sheet']} row {int(ridx)+1}: date/day does not belong to {y}-{m:02d}."
                    )
                    continue
                if "department" not in cmap or "shift" not in cmap:
                    warnings.append(
                        f"{table['sheet']} row {int(ridx)+1}: schedule row needs department and shift."
                    )
                    continue
                dep=_research_dept_key(r.get(cmap["department"]))
                block=_research_shift(r.get(cmap["shift"]))
                candidates=[
                    s for s in slots
                    if s.day==day and s.block==block and not s.blocked
                ]
                exact=[s for s in candidates if _research_dept_key(s.department)==dep]
                if not exact and dep.startswith("centroro"):
                    exact=[s for s in candidates if _research_dept_key(s.department).startswith("centroro")]
                if not exact:
                    warnings.append(
                        f"{table['sheet']} row {int(ridx)+1}: no matching slot for day {day}, "
                        f"department '{r.get(cmap['department'])}', shift '{block}'."
                    )
                    continue

                # If a generic department label maps to several identical CENTRO rows,
                # choose the first still-unused slot. This preserves every assignment
                # instead of collapsing repeated rows onto one slot.
                unused=[s for s in sorted(exact,key=lambda s:s.idx) if s.idx not in assignments]
                sid=(unused[0] if unused else sorted(exact,key=lambda s:s.idx)[0]).idx

            if sid in assignments:
                previous=assignments[sid]
                prev_src=provenance[sid]
                if previous==person:
                    duplicates.append({
                        "slot_id":sid,"person":person,
                        "first_source":prev_src,"duplicate_source":src,
                    })
                    continue
                conflicts.append({
                    "slot_id":sid,
                    "existing_person":previous,
                    "new_person":person,
                    "existing_source":prev_src,
                    "new_source":src,
                })
                continue

            assignments[sid]=person
            provenance[sid]=src

    audit["warnings"]=warnings
    audit["duplicates"]=duplicates
    audit["conflicts"]=conflicts
    audit["assignment_count"]=len(assignments)
    audit["assignment_provenance"]={str(k):v for k,v in provenance.items()}

    if conflicts:
        preview="; ".join(
            f"slot {c['slot_id']}: {c['existing_person']} vs {c['new_person']} "
            f"({c['existing_source']['sheet']} / {c['new_source']['sheet']})"
            for c in conflicts[:6]
        )
        raise ValueError(
            "Whole-workbook schedule import found conflicting assignments for the same slot. "
            "Official lock is blocked until resolved: "+preview
        )

    if return_audit:
        return assignments,warnings,audit
    return assignments,warnings



def research_manual_result(y,m,people,assignments):
    targets=_research_targets(y,m,people)
    stats=validate_schedule(y,m,people,make_slots(y,m),assignments,targets)
    return SolveResult(ok=stats["global"]["hard_errors"]==0,
                       message="Imported AVAILABLE GPT + HUMAN schedule",assignments=assignments,targets=targets,stats=stats)


def _research_json_safe(value):
    """Convert engine/research structures into stable JSON-compatible primitives."""
    if value is None or isinstance(value,(str,bool,int)):
        return value
    if isinstance(value,float):
        if pd.isna(value):
            return None
        return float(value)
    if isinstance(value,(date,datetime)):
        return value.isoformat()
    if isinstance(value,set):
        return sorted(_research_json_safe(v) for v in value)
    if isinstance(value,(list,tuple)):
        return [_research_json_safe(v) for v in value]
    if isinstance(value,dict):
        return {str(k):_research_json_safe(v) for k,v in value.items()}
    try:
        if pd.isna(value):
            return None
    except Exception:
        pass
    try:
        return value.item()
    except Exception:
        return str(value)


def _research_people_snapshot(people):
    rows=[]
    for p in people:
        rows.append({
            "initials":p.initials,
            "name":p.name,
            "unavailable":sorted(p.unavailable),
            "unavailable_am":sorted(p.unavailable_am),
            "unavailable_pm":sorted(p.unavailable_pm),
            "vacation":sorted(p.vacation),
            "justified_absence":sorted(p.justified_absence),
            "long_duty":sorted(p.long_duty),
            "reserved_backup":[list(x) for x in sorted(p.reserved_backup)],
            "soft_free":sorted(p.soft_free),
            "soft_free_am":sorted(p.soft_free_am),
            "soft_free_pm":sorted(p.soft_free_pm),
            "preferred":sorted(p.preferred),
            "preferred_am":sorted(p.preferred_am),
            "preferred_pm":sorted(p.preferred_pm),
            "weekday_preference":p.weekday_preference,
            "weekend_preference":p.weekend_preference,
            "holiday_preference":p.holiday_preference,
            "spread_preference":p.spread_preference,
            "shift_length_preference":int(getattr(p,"shift_length_preference",0) or 0),
            "avoid_doubles":bool(p.avoid_doubles),
            "target_adjustment":p.target_adjustment,
            "prior_weekend_count":p.prior_weekend_count,
            "prior_friday_count":p.prior_friday_count,
            "prior_double_count":p.prior_double_count,
            "prior_weekday_day_count":p.prior_weekday_day_count,
            "prior_rotation_counts":dict(p.prior_rotation_counts or {}),
            "prior_consecutive_weekend_streak":int(getattr(p,"prior_consecutive_weekend_streak",0) or 0),
            "prior_last_day_onko":bool(getattr(p,"prior_last_day_onko",False)),
            "prior_resident_hard_loss_count":int(getattr(p,"prior_resident_hard_loss_count",0) or 0),
            "request_items":_research_json_safe(list(getattr(p,"request_items",[]) or [])),
            "rest_credit_am_to_use":int(getattr(p,"rest_credit_am_to_use",0) or 0),
            "rest_credit_pm_to_use":int(getattr(p,"rest_credit_pm_to_use",0) or 0),
            "note":str(getattr(p,"note","") or ""),
        })
    return _research_json_safe(rows)


def _research_people_from_snapshot(rows):
    people=[]
    for r in rows or []:
        people.append(Person(
            initials=str(r.get("initials","")),
            name=str(r.get("name","")),
            unavailable=set(r.get("unavailable") or []),
            unavailable_am=set(r.get("unavailable_am") or []),
            unavailable_pm=set(r.get("unavailable_pm") or []),
            vacation=set(r.get("vacation") or []),
            justified_absence=set(r.get("justified_absence") or []),
            long_duty=set(r.get("long_duty") or []),
            reserved_backup={tuple(x) for x in (r.get("reserved_backup") or [])},
            soft_free=set(r.get("soft_free") or []),
            soft_free_am=set(r.get("soft_free_am") or []),
            soft_free_pm=set(r.get("soft_free_pm") or []),
            preferred=set(r.get("preferred") or []),
            preferred_am=set(r.get("preferred_am") or []),
            preferred_pm=set(r.get("preferred_pm") or []),
            weekday_preference=int(r.get("weekday_preference") or 0),
            weekend_preference=int(r.get("weekend_preference") or 0),
            holiday_preference=max(-1,min(1,int(r.get("holiday_preference") or 0))),
            spread_preference=int(r.get("spread_preference") or 0),
            shift_length_preference=max(0,min(3,int(r.get("shift_length_preference") or 0))),
            avoid_doubles=bool(r.get("avoid_doubles",False)),
            target_adjustment=int(r.get("target_adjustment") or 0),
            prior_weekend_count=int(r.get("prior_weekend_count") or 0),
            prior_friday_count=int(r.get("prior_friday_count") or 0),
            prior_double_count=int(r.get("prior_double_count") or 0),
            prior_weekday_day_count=int(r.get("prior_weekday_day_count") or 0),
            prior_rotation_counts={str(k):int(v) for k,v in (r.get("prior_rotation_counts") or {}).items()},
            prior_consecutive_weekend_streak=int(r.get("prior_consecutive_weekend_streak") or 0),
            prior_last_day_onko=bool(r.get("prior_last_day_onko",False)),
            prior_resident_hard_loss_count=int(r.get("prior_resident_hard_loss_count") or 0),
            request_items=list(r.get("request_items") or []),
            rest_credit_am_to_use=int(r.get("rest_credit_am_to_use") or 0),
            rest_credit_pm_to_use=int(r.get("rest_credit_pm_to_use") or 0),
            note=str(r.get("note") or ""),
        ))
    return people


def _research_hash(value):
    raw=json.dumps(
        _research_json_safe(value),
        ensure_ascii=False,sort_keys=True,separators=(",",":")
    ).encode("utf-8")
    return hashlib.sha256(raw).hexdigest()


def _research_input_hash(y,m,people):
    # Rule Profile / engine version are stored separately and server-locked per case.
    return _research_hash({
        "cycle_year":int(y),
        "cycle_month":int(m),
        "people":_research_people_snapshot(people),
    })


def _research_schedule_hash(assignments):
    return _research_hash({str(k):v for k,v in sorted((assignments or {}).items())})


def _research_locked_comparator(case,people):
    assignments={int(k):v for k,v in (case.get("comparator_assignments") or {}).items()}
    return research_manual_result(int(case["cycle_year"]),int(case["cycle_month"]),people,assignments)


def _research_result_from_run(run,people,y,m):
    assignments={int(k):v for k,v in (run.get("assignments") or {}).items()}
    frozen_stats=run.get("raw_metrics") or {}
    if not isinstance(frozen_stats,dict) or "global" not in frozen_stats:
        # Old/partial data safety fallback; research runs created by V2.5.41 store full stats.
        frozen_stats=validate_schedule(y,m,people,make_slots(y,m),assignments,_research_targets(y,m,people))
    return SolveResult(
        ok=bool(run.get("success")),
        message=f"Frozen research run {run.get('run_no')}",
        assignments=assignments,
        targets=_research_targets(y,m,people),
        stats=frozen_stats,
        objective_value=None,
    )


def _research_metric_rows(comparator,engine_result,engine_label="MY ENGINE — RUN 1"):
    cg=comparator.stats.get("global",{})
    eg=engine_result.stats.get("global",{})
    defs=[
        ("HARD errors","hard_errors",False),
        ("Monthly fairness %","monthly_fairness_score",True),
        ("Cumulative fairness %","cumulative_fairness_score",True),
        ("Preference fulfilment %","mean_preference_score",True),
        ("Monthly workplace imbalance","rotation_monthly_imbalance",False),
        ("Cumulative workplace imbalance","rotation_cumulative_imbalance",False),
        ("Weekend spread","weekend_monthly_spread",False),
        ("Friday spread","friday_monthly_spread",False),
        ("Double spread","double_monthly_spread",False),
        ("Weekday-day spread","weekday_day_monthly_spread",False),
        ("Average distinct workplaces","mean_distinct_rotations",True),
        ("Real gap count","optional_gap_count",False),
        ("Gap workplace spread","optional_gap_category_spread",False),
    ]
    rows=[]
    for label,key,higher in defs:
        a=cg.get(key); b=eg.get(key)
        delta=None
        if isinstance(a,(int,float)) and isinstance(b,(int,float)):
            delta=round(float(b)-float(a),2)
        better=""
        if delta is not None and delta!=0:
            better=engine_label if (delta>0)==higher else "AVAILABLE GPT + HUMAN"
        elif delta==0:
            better="Tie"
        rows.append({
            "Metric":label,
            "AVAILABLE GPT + HUMAN":a,
            engine_label:b,
            f"Δ {engine_label} − GPT+Human":delta,
            "Better":better,
        })
    return pd.DataFrame(rows)



def _research_post_matrix_df(result,people,include_summary=True):
    """Resident × workplace assignment-count matrix for paper-ready export."""
    pdata=(result.stats or {}).get("people",{})
    rows=[]
    for p in people:
        d=pdata.get(p.initials,{})
        rc=d.get("rotation_counts") or {}
        row={
            "Initials":p.initials,
            "Name":p.name,
        }
        for cat in ROTATION_CATEGORIES:
            row[cat]=int(rc.get(cat,0) or 0)
        row["Distinct workplaces"]=int(d.get("distinct_rotations",0) or 0)
        row["Total assignments"]=int(d.get("assignments",0) or 0)
        row["Workload"]=d.get("workload")
        rows.append(row)

    df=pd.DataFrame(rows)
    if not include_summary or df.empty:
        return df

    summary=[]
    for label,func in [
        ("MAX",lambda s:int(s.max())),
        ("MIN",lambda s:int(s.min())),
        ("SPREAD (MAX-MIN)",lambda s:int(s.max()-s.min())),
    ]:
        r={"Initials":label,"Name":""}
        for cat in ROTATION_CATEGORIES:
            r[cat]=func(df[cat])
        r["Distinct workplaces"]=""
        r["Total assignments"]=""
        r["Workload"]=""
        summary.append(r)
    return pd.concat([df,pd.DataFrame(summary)],ignore_index=True)


def _research_post_stats(result,people):
    """Per-workplace supply and resident spread."""
    matrix=_research_post_matrix_df(result,people,include_summary=False)
    rows={}
    for cat in ROTATION_CATEGORIES:
        vals=[int(v) for v in matrix[cat].tolist()] if not matrix.empty else []
        rows[cat]={
            "total":sum(vals),
            "mean":round(sum(vals)/len(vals),2) if vals else 0.0,
            "min":min(vals) if vals else 0,
            "max":max(vals) if vals else 0,
            "spread":(max(vals)-min(vals)) if vals else 0,
        }
    return rows


def _research_post_spread_comparison_df(
    comparator,people,run1_result=None,best_result=None,best_label="Best"
):
    arms=[
        ("GPT+Human",comparator),
        ("Run 1",run1_result),
        (best_label,best_result),
    ]
    arm_stats={
        label:_research_post_stats(result,people)
        for label,result in arms if result is not None
    }
    rows=[]
    for cat in ROTATION_CATEGORIES:
        row={"Workplace":cat}
        for label,_result in arms:
            stt=(arm_stats.get(label) or {}).get(cat)
            if stt is None:
                row[f"{label} total"]=None
                row[f"{label} min"]=None
                row[f"{label} max"]=None
                row[f"{label} spread"]=None
            else:
                row[f"{label} total"]=stt["total"]
                row[f"{label} min"]=stt["min"]
                row[f"{label} max"]=stt["max"]
                row[f"{label} spread"]=stt["spread"]

        gh=(arm_stats.get("GPT+Human") or {}).get(cat)
        r1=(arm_stats.get("Run 1") or {}).get(cat)
        rb=(arm_stats.get(best_label) or {}).get(cat)
        row["Δ Run1 spread vs GPT+Human"]=(
            None if gh is None or r1 is None else int(r1["spread"]-gh["spread"])
        )
        row[f"Δ {best_label} spread vs GPT+Human"]=(
            None if gh is None or rb is None else int(rb["spread"]-gh["spread"])
        )
        rows.append(row)

    # Overall row = sum of the nine independent post spreads, matching the
    # engine's monthly workplace imbalance concept.
    total_row={"Workplace":"TOTAL POST IMBALANCE (Σ spreads)"}
    for label,_result in arms:
        stats=arm_stats.get(label)
        total_row[f"{label} total"]=None
        total_row[f"{label} min"]=None
        total_row[f"{label} max"]=None
        total_row[f"{label} spread"]=(
            None if stats is None else sum(int(stats[c]["spread"]) for c in ROTATION_CATEGORIES)
        )
    gh_total=total_row.get("GPT+Human spread")
    r1_total=total_row.get("Run 1 spread")
    best_total=total_row.get(f"{best_label} spread")
    total_row["Δ Run1 spread vs GPT+Human"]=(
        None if gh_total is None or r1_total is None else int(r1_total-gh_total)
    )
    total_row[f"Δ {best_label} spread vs GPT+Human"]=(
        None if gh_total is None or best_total is None else int(best_total-gh_total)
    )
    rows.append(total_row)
    return pd.DataFrame(rows)


def _research_run_quality_key(run):
    """Transparent secondary best-of-run selector using the V2.5.49 constitution.

    Order is lexicographic: ABSOLUTE-HARD validity -> mandatory zero RESIDENT-HARD
    violations -> post equality ->
    MAX-MIN SOFT satisfaction -> mean SOFT satisfaction -> monthly fairness ->
    diversity. Overall request satisfaction is a final descriptive tie-breaker.
    """
    hard=run.get("hard_errors")
    hard=9999 if hard is None else int(hard)
    g=((run.get("full_stats") or {}).get("global",{}) if isinstance(run.get("full_stats"),dict) else {})
    rh_total=float(g.get("resident_hard_total_losses") if g.get("resident_hard_total_losses") is not None else 1e9)
    rh_max=float(g.get("resident_hard_max_loss_per_resident") if g.get("resident_hard_max_loss_per_resident") is not None else 1e9)
    rh_cum_spread=float(g.get("resident_hard_cumulative_spread") if g.get("resident_hard_cumulative_spread") is not None else 1e9)
    worst_post=float(g.get("worst_monthly_post_spread") if g.get("worst_monthly_post_spread") is not None else 1e9)
    imb=run.get("monthly_workplace_imbalance")
    imb=1e9 if imb is None else float(imb)
    soft_min=g.get("min_soft_preference_score")
    soft_min=-1e9 if soft_min is None else float(soft_min)
    soft_mean=g.get("mean_soft_preference_score")
    soft_mean=-1e9 if soft_mean is None else float(soft_mean)
    fair=run.get("monthly_fairness")
    fair=-1e9 if fair is None else float(fair)
    div=run.get("mean_distinct_workplaces")
    div=-1e9 if div is None else float(div)
    overall=g.get("mean_preference_score")
    overall=-1e9 if overall is None else float(overall)
    return (
        hard,rh_total,rh_max,rh_cum_spread,worst_post,imb,
        -soft_min,-soft_mean,-fair,-div,-overall,int(run.get("run_no") or 999)
    )


def _research_run_log_df(runs):
    rows=[]
    total=0.0
    for r in runs:
        total+=float(r.get("elapsed_seconds") or 0)
        rg=((r.get("full_stats") or {}).get("global",{}) if isinstance(r.get("full_stats"),dict) else {})
        rows.append({
            "Run":int(r.get("run_no") or 0),
            "Role":"FIRST SHOT (PRIMARY)" if int(r.get("run_no") or 0)==1 else "IMPROVEMENT ATTEMPT",
            "Success":bool(r.get("success")),
            "Time, s":round(float(r.get("elapsed_seconds") or 0),2),
            "Cumulative time, s":round(total,2),
            "Engine version":r.get("app_version"),
            "Rule Profile":r.get("rule_profile_version"),
            "Solver stage":r.get("solver_stage"),
            "ABSOLUTE HARD errors":r.get("hard_errors"),
            "RESIDENT HARD losses":rg.get("resident_hard_total_losses"),
            "RESIDENT HARD max/person":rg.get("resident_hard_max_loss_per_resident"),
            "RESIDENT HARD cumulative spread":rg.get("resident_hard_cumulative_spread"),
            "Monthly fairness %":r.get("monthly_fairness"),
            "Overall request satisfaction %":rg.get("mean_preference_score"),
            "Worst SOFT satisfaction %":rg.get("min_soft_preference_score"),
            "SOFT satisfaction spread, pp":rg.get("soft_preference_score_spread"),
            "Worst post spread":rg.get("worst_monthly_post_spread"),
            "SYSTEM-HARD worst-post lock":rg.get("post_system_hard_worst_spread_lock"),
            "SYSTEM-HARD 9-post total spread lock":rg.get("post_system_hard_total_spread_lock"),
            "Distinct workplace spread":rg.get("distinct_rotation_spread"),
            "Post structural water-fill gate":rg.get("post_spread_quality_gate_passed"),
            "Monthly workplace imbalance":r.get("monthly_workplace_imbalance"),
            "Weekend spread":r.get("weekend_spread"),
            "Friday spread":r.get("friday_spread"),
            "Friday 0-1 gate":("PASS" if (r.get("friday_spread") is not None and int(r.get("friday_spread") or 0)<=1) else "LEGACY / INVALID"),
            "Double spread":r.get("double_spread"),
            "Weekday-day spread":r.get("weekday_day_spread"),
            "Distinct workplaces":r.get("mean_distinct_workplaces"),
            "Gap count":r.get("gap_count"),
            "Gap workplace spread":r.get("gap_category_spread"),
        })
    return pd.DataFrame(rows)


def _research_result_wish_totals(result):
    """Comparable request totals from one validated result."""
    out={"active":0,"honored":0,"missed":0,"hard":0,"hard_missed":0,"exact_soft":0,"exact_soft_missed":0}
    for initials,d in ((result.stats or {}).get("people",{}) or {}).items():
        for r in (d.get("request_detail_rows") or []):
            if not r.get("included_in_score"):
                continue
            out["active"]+=1
            ok=bool(r.get("fulfilled"))
            out["honored"]+=int(ok); out["missed"]+=int(not ok)
            if r.get("kind")=="resident_hard":
                out["hard"]+=1; out["hard_missed"]+=int(not ok)
            if r.get("kind") in ("soft_free","preferred"):
                out["exact_soft"]+=1; out["exact_soft_missed"]+=int(not ok)
    out["percent"]=(None if out["active"]==0 else round(100.0*out["honored"]/out["active"],1))
    return out


def _research_wish_summary_comparison_df(comparator,engine_result,people):
    """Per-resident wish fulfilment using the SAME frozen input snapshot."""
    rows=[]
    cp=(comparator.stats or {}).get("people",{}) or {}
    ep=(engine_result.stats or {}).get("people",{}) or {}
    for p in people:
        c=cp.get(p.initials,{}) or {}; e=ep.get(p.initials,{}) or {}
        cdet=[r for r in (c.get("request_detail_rows") or []) if r.get("included_in_score")]
        edet=[r for r in (e.get("request_detail_rows") or []) if r.get("included_in_score")]
        active=max(len(cdet),len(edet))
        c_ok=sum(1 for r in cdet if r.get("fulfilled")); e_ok=sum(1 for r in edet if r.get("fulfilled"))
        rows.append({
            "Resident":p.initials,
            "Name":p.name,
            "Active wishes":active,
            "GPT+Human honored":c_ok,
            "GPT+Human missed":max(0,len(cdet)-c_ok),
            "GPT+Human satisfaction %":(None if not cdet else round(100*c_ok/len(cdet),1)),
            "GPT+Human cannot-work violations":int(c.get("resident_hard_losses",0) or 0),
            "Engine honored":e_ok,
            "Engine missed":max(0,len(edet)-e_ok),
            "Engine satisfaction %":(None if not edet else round(100*e_ok/len(edet),1)),
            "Engine cannot-work violations":int(e.get("resident_hard_losses",0) or 0),
        })
    return pd.DataFrame(rows)


def _research_wish_request_comparison_df(comparator,engine_result,people):
    """Request-by-request paired outcomes; request_id is stable under one frozen snapshot."""
    rows=[]
    cp=(comparator.stats or {}).get("people",{}) or {}
    ep=(engine_result.stats or {}).get("people",{}) or {}
    for p in people:
        crows={str(r.get("request_id")):r for r in (cp.get(p.initials,{}).get("request_detail_rows") or []) if r.get("included_in_score")}
        erows={str(r.get("request_id")):r for r in (ep.get(p.initials,{}).get("request_detail_rows") or []) if r.get("included_in_score")}
        keys=list(dict.fromkeys(list(crows)+list(erows)))
        for key in keys:
            c=crows.get(key) or {}; e=erows.get(key) or {}
            base=c or e
            c_ok=(None if not c else bool(c.get("fulfilled")))
            e_ok=(None if not e else bool(e.get("fulfilled")))
            if c_ok is True and e_ok is True: verdict="Both honored"
            elif c_ok is False and e_ok is True: verdict="ENGINE only"
            elif c_ok is True and e_ok is False: verdict="GPT+HUMAN only"
            elif c_ok is False and e_ok is False: verdict="Both missed"
            else: verdict="Audit mismatch"
            rows.append({
                "Resident":p.initials,
                "Name":p.name,
                "Request ID":key,
                "Priority":base.get("priority"),
                "Request":base.get("type"),
                "Date":base.get("date") or "—",
                "Block":base.get("block") or "—",
                "Requested value":base.get("requested_value"),
                "GPT+Human":("HONORED" if c_ok is True else "MISSED" if c_ok is False else "MISSING"),
                "Engine Run 1":("HONORED" if e_ok is True else "MISSED" if e_ok is False else "MISSING"),
                "Outcome":verdict,
            })
    return pd.DataFrame(rows)


def research_locked_comparison_xlsx(y,m,case,people,comparator,runs,questionnaires=None):
    bio=BytesIO()
    valid_runs=[r for r in runs if bool(r.get("success")) and int(r.get("hard_errors") or 0)==0]
    best=min(valid_runs,key=_research_run_quality_key) if valid_runs else None
    first=next((r for r in runs if int(r.get("run_no") or 0)==1),None)

    with pd.ExcelWriter(bio,engine="xlsxwriter") as w:
        # Method / audit trail.
        method_rows=[
            ["Study comparison","AVAILABLE GPT + HUMAN vs MY ENGINE TOOL"],
            ["Cycle",f"{y}-{m:02d}"],
            ["Primary engine endpoint","Immutable FIRST SHOT — Run 1"],
            ["Secondary endpoint","BEST-OF-5 after up to 4 additional improvement attempts"],
            ["Best-of-5 selector","ABSOLUTE-HARD-valid → zero RESIDENT-HARD violations → lower post imbalance → max-min SOFT → monthly fairness → diversity"],
            ["Max engine runs",case.get("engine_max_runs",5)],
            ["Locked input hash",case.get("input_hash")],
            ["Wish comparison rule","Both schedules are validated against the exact same frozen input snapshot"],
            ["Locked GPT+Human schedule hash",case.get("comparator_schedule_hash")],
            ["Engine version at lock",case.get("app_version_at_lock")],
            ["Rule Profile at lock",case.get("rule_profile_version_at_lock")],
            ["GPT+Human iteration/rework count",case.get("gpt_human_iterations")],
            ["GPT+Human estimated total time, min",case.get("gpt_human_minutes")],
            ["GPT+Human method note",case.get("method_note")],
            ["Questionnaires stored",len(questionnaires or [])],
            ["Questionnaire respondents",", ".join(sorted(str(q.get("respondent_initials")) for q in (questionnaires or [])))],
            ["Case locked at",case.get("created_at")],
            ["Critical integrity rule","Run 1 and all later run records are immutable; no cherry-picking replaces Run 1."],
        ]
        pd.DataFrame(method_rows,columns=["Field","Value"]).to_excel(w,index=False,sheet_name="method")

        _research_run_log_df(runs).to_excel(w,index=False,sheet_name="run_log")
        snapshot_excel=[]
        for row in _research_people_snapshot(people):
            snapshot_excel.append({
                k:(json.dumps(v,ensure_ascii=False,sort_keys=True) if isinstance(v,(list,dict)) else v)
                for k,v in row.items()
            })
        pd.DataFrame(snapshot_excel).to_excel(w,index=False,sheet_name="input_snapshot")
        schedule_list_df(y,m,comparator).to_excel(w,index=False,sheet_name="gpt_human_schedule")

        # Comparator vs first shot / best.
        if first and first.get("success"):
            first_result=_research_result_from_run(first,people,y,m)
            _research_metric_rows(comparator,first_result,"MY ENGINE — RUN 1").to_excel(
                w,index=False,sheet_name="comparison_run1"
            )
            schedule_list_df(y,m,first_result).to_excel(w,index=False,sheet_name="engine_run1")
            _research_wish_summary_comparison_df(comparator,first_result,people).to_excel(
                w,index=False,sheet_name="wish_summary_run1"
            )
            _research_wish_request_comparison_df(comparator,first_result,people).to_excel(
                w,index=False,sheet_name="wish_request_compare"
            )
        if best:
            best_result=_research_result_from_run(best,people,y,m)
            _research_metric_rows(comparator,best_result,f"MY ENGINE — BEST RUN {best['run_no']}").to_excel(
                w,index=False,sheet_name="comparison_best"
            )
            schedule_list_df(y,m,best_result).to_excel(w,index=False,sheet_name="engine_best")

        # Every frozen engine run.
        for r in runs:
            if not r.get("assignments"):
                continue
            rr=_research_result_from_run(r,people,y,m)
            schedule_list_df(y,m,rr).to_excel(
                w,index=False,sheet_name=f"run_{int(r['run_no'])}_schedule"
            )

        # Per-resident comparator vs first and best.
        comp_people=comparator.stats.get("people",{})
        run1_result=_research_result_from_run(first,people,y,m) if first and first.get("success") else None
        best_result=_research_result_from_run(best,people,y,m) if best else None
        rows=[]
        for p in people:
            c=comp_people.get(p.initials,{})
            r1=(run1_result.stats.get("people",{}).get(p.initials,{}) if run1_result else {})
            rb=(best_result.stats.get("people",{}).get(p.initials,{}) if best_result else {})
            rows.append({
                "initials":p.initials,"name":p.name,
                "GPT+Human workload":c.get("workload"),
                "Run1 workload":r1.get("workload"),
                "Best workload":rb.get("workload"),
                "GPT+Human weekends":c.get("weekend_assignments"),
                "Run1 weekends":r1.get("weekend_assignments"),
                "Best weekends":rb.get("weekend_assignments"),
                "GPT+Human Fridays":c.get("friday_assignments"),
                "Run1 Fridays":r1.get("friday_assignments"),
                "Best Fridays":rb.get("friday_assignments"),
                "GPT+Human doubles":c.get("doubles"),
                "Run1 doubles":r1.get("doubles"),
                "Best doubles":rb.get("doubles"),
                "GPT+Human workplaces":c.get("distinct_rotations"),
                "Run1 workplaces":r1.get("distinct_rotations"),
                "Best workplaces":rb.get("distinct_rotations"),
                "GPT+Human pref %":c.get("preference_score"),
                "Run1 pref %":r1.get("preference_score"),
                "Best pref %":rb.get("preference_score"),
            })
        pd.DataFrame(rows).to_excel(w,index=False,sheet_name="per_resident")

        # Paper-ready resident × workplace matrices.
        _research_post_matrix_df(comparator,people).to_excel(
            w,index=False,sheet_name="post_matrix_gpt_human"
        )
        if run1_result is not None:
            _research_post_matrix_df(run1_result,people).to_excel(
                w,index=False,sheet_name="post_matrix_run1"
            )
        if best_result is not None:
            _research_post_matrix_df(best_result,people).to_excel(
                w,index=False,sheet_name="post_matrix_best"
            )

        spread_compare=_research_post_spread_comparison_df(
            comparator,people,
            run1_result=run1_result,
            best_result=best_result,
            best_label=("Best" if best is None else f"Best Run {best['run_no']}")
        )
        spread_compare.to_excel(
            w,index=False,sheet_name="post_spread_compare"
        )

        # Whole-workbook import audit / provenance.
        import_audit_rows=[]
        for item in case.get("import_warnings") or []:
            if not isinstance(item,dict) or item.get("type") not in ("WHOLE_WORKBOOK_IMPORT_AUDIT_V2544","LIVE_APP_SAME_INPUT_AUDIT_V25108"):
                continue
            if item.get("type")=="LIVE_APP_SAME_INPUT_AUDIT_V25108":
                pa=item.get("preferences") or {}
                import_audit_rows.append({
                    "Kind":"preferences",
                    "Workbook SHA-256":pa.get("snapshot_hash",""),
                    "Sheet":"LIVE APP DB",
                    "Status":"frozen_live_snapshot",
                    "Header row":"",
                    "Rows":pa.get("resident_count",""),
                    "Reason":item.get("selected_cycle","")+" · SAME snapshot used for both schedules",
                })
            for kind in (("schedule",) if item.get("type")=="LIVE_APP_SAME_INPUT_AUDIT_V25108" else ("preferences","schedule")):
                audit=item.get(kind) or {}
                wb_hash=audit.get("workbook_hash")
                for sh in audit.get("sheets",[]) or []:
                    import_audit_rows.append({
                        "Kind":kind,
                        "Workbook SHA-256":wb_hash,
                        "Sheet":sh.get("sheet"),
                        "Status":sh.get("status"),
                        "Header row":sh.get("header_row"),
                        "Rows":sh.get("rows"),
                        "Reason":sh.get("reason",""),
                    })
                if kind=="schedule":
                    for sid,src in (audit.get("assignment_provenance") or {}).items():
                        import_audit_rows.append({
                            "Kind":"schedule_assignment_provenance",
                            "Workbook SHA-256":src.get("workbook_hash"),
                            "Sheet":src.get("sheet"),
                            "Status":f"slot_id={sid}",
                            "Header row":src.get("header_row"),
                            "Rows":src.get("source_row"),
                            "Reason":(
                                f"source_col={src.get('source_column','')} "
                                f"day={src.get('day','')} mode={src.get('mode','long')}"
                            ),
                        })
            import_audit_rows.append({
                "Kind":"case",
                "Workbook SHA-256":"",
                "Sheet":"",
                "Status":"same_workbook" if item.get("same_workbook") else "separate_workbooks",
                "Header row":"",
                "Rows":"",
                "Reason":item.get("selected_cycle",""),
            })
        pd.DataFrame(import_audit_rows or [{
            "Kind":"","Workbook SHA-256":"","Sheet":"","Status":"No whole-workbook audit stored",
            "Header row":"","Rows":"","Reason":""
        }]).to_excel(w,index=False,sheet_name="import_audit")

        # Retrospective questionnaire audit trail.
        q_rows=[]
        for q in questionnaires or []:
            payload=q.get("parser_payload") or {}
            extracted=str(q.get("extracted_text") or "")
            if len(extracted)>30000:
                extracted=extracted[:30000]+"… [TRUNCATED IN XLSX; DB COPY RETAINED]"
            q_rows.append({
                "Respondent":q.get("respondent_initials"),
                "File":q.get("file_name"),
                "File SHA-256":q.get("file_hash"),
                "MIME":q.get("mime_type"),
                "Bytes":q.get("file_size"),
                "Parsed iterations":q.get("parsed_iterations"),
                "Parsed total time, min":q.get("parsed_minutes"),
                "Iteration evidence":payload.get("iteration_evidence"),
                "Time evidence":payload.get("time_evidence"),
                "Locked at":q.get("created_at"),
                "Extracted text":extracted,
            })
        pd.DataFrame(q_rows or [{
            "Respondent":"","File":"","File SHA-256":"","MIME":"","Bytes":"",
            "Parsed iterations":"","Parsed total time, min":"",
            "Iteration evidence":"","Time evidence":"","Locked at":"",
            "Extracted text":"No questionnaires stored",
        }]).to_excel(w,index=False,sheet_name="questionnaires")

        # Frozen errors/findings.
        error_rows=[]
        for e in comparator.stats.get("global",{}).get("errors",[]) or []:
            error_rows.append({"Source":"AVAILABLE GPT + HUMAN","Finding":str(e)})
        for r in runs:
            raw=r.get("raw_metrics") or {}
            for e in (raw.get("global",{}).get("errors",[]) if isinstance(raw,dict) else []) or []:
                error_rows.append({"Source":f"ENGINE RUN {r.get('run_no')}","Finding":str(e)})
        for e in case.get("import_warnings") or []:
            error_rows.append({"Source":"IMPORT","Finding":str(e)})
        pd.DataFrame(error_rows or [{"Source":"","Finding":"No findings"}]).to_excel(
            w,index=False,sheet_name="errors_findings"
        )

    return bio.getvalue()


def _research_parse_optional_int(v):
    s=str(v or "").strip()
    if not s:
        return None
    return max(1,int(float(s)))


def _research_parse_optional_float(v):
    s=str(v or "").strip()
    if not s:
        return None
    return max(0.0,float(s))



def _questionnaire_extract_text(uploaded):
    """Extract text locally from research questionnaire uploads."""
    if uploaded is None:
        return {"ok":False,"text":"","error":"No file","file_name":"","file_hash":"","file_size":0,"mime_type":""}

    raw=uploaded.getvalue()
    name=str(getattr(uploaded,"name","questionnaire"))
    suffix=Path(name).suffix.lower()
    mime=str(getattr(uploaded,"type","") or "")
    file_hash=hashlib.sha256(raw).hexdigest()
    text=""
    error=""

    try:
        if suffix==".pdf":
            reader=PdfReader(BytesIO(raw))
            parts=[]
            for pageno,page in enumerate(reader.pages,1):
                page_text=page.extract_text() or ""
                if page_text.strip():
                    parts.append(f"[PAGE {pageno}]\n{page_text}")
            text="\n\n".join(parts)
            if not text.strip():
                error="PDF text extraction returned no text. If this is a scanned/image-only questionnaire, enter the values manually."
        elif suffix==".docx":
            doc=Document(BytesIO(raw))
            parts=[]
            for p in doc.paragraphs:
                if p.text.strip():
                    parts.append(p.text)
            for table in doc.tables:
                for row in table.rows:
                    vals=[cell.text.strip() for cell in row.cells]
                    if any(vals):
                        parts.append(" | ".join(vals))
            text="\n".join(parts)
        elif suffix in {".xlsx",".xls"}:
            xls=pd.ExcelFile(BytesIO(raw))
            parts=[]
            for sheet in xls.sheet_names:
                df=pd.read_excel(xls,sheet_name=sheet,header=None)
                parts.append(f"[SHEET {sheet}]")
                for _,row in df.iterrows():
                    vals=[
                        str(v).strip()
                        for v in row.tolist()
                        if v is not None and not (isinstance(v,float) and pd.isna(v))
                    ]
                    if vals:
                        parts.append(" | ".join(vals))
            text="\n".join(parts)
        elif suffix==".csv":
            try:
                text=raw.decode("utf-8")
            except UnicodeDecodeError:
                text=raw.decode("latin-1",errors="replace")
        elif suffix in {".txt",".md"}:
            try:
                text=raw.decode("utf-8")
            except UnicodeDecodeError:
                text=raw.decode("latin-1",errors="replace")
        else:
            error=f"Unsupported questionnaire type: {suffix or 'unknown'}"
    except Exception as exc:
        error=f"{type(exc).__name__}: {exc}"

    cleaned=re.sub(r"\r\n?","\n",text or "")
    cleaned=re.sub(r"[ \t]+"," ",cleaned)
    cleaned=re.sub(r"\n{3,}","\n\n",cleaned).strip()

    return {
        "ok":bool(cleaned),
        "text":cleaned,
        "error":error,
        "file_name":name,
        "file_hash":file_hash,
        "file_size":len(raw),
        "mime_type":mime,
    }


def _questionnaire_context(text,start,end,radius=130):
    lo=max(0,start-radius); hi=min(len(text),end+radius)
    return re.sub(r"\s+"," ",text[lo:hi]).strip()


def _questionnaire_parse_process_metrics(text):
    """Heuristic extraction of GPT-human iteration count and total time.

    It intentionally returns evidence/confidence rather than silently deciding.
    Final research metadata remains human-confirmed.
    """
    raw=str(text or "")
    norm=unicodedata.normalize("NFKD",raw).encode("ascii","ignore").decode("ascii").lower()

    iteration_candidates=[]
    iter_patterns=[
        r"(?:perdarym\w*|iteracij\w*|bandym\w*|generavim\w*|prompt\w*|atsakym\w*|iterations?|attempts?|regenerations?|revisions?)[^0-9]{0,45}(\d{1,3})",
        r"(\d{1,3})[^a-z0-9]{0,15}(?:kart\w*|iteracij\w*|bandym\w*|perdarym\w*|iterations?|attempts?|times?|revisions?)",
        r"(?:kiek\s+kart\w*|how\s+many\s+(?:times|iterations|attempts|revisions))[^0-9]{0,35}(\d{1,3})",
    ]
    for pat in iter_patterns:
        for m in re.finditer(pat,norm,flags=re.I):
            try:
                value=int(m.group(1))
            except Exception:
                continue
            if 1 <= value <= 100:
                iteration_candidates.append({
                    "value":value,
                    "context":_questionnaire_context(raw,m.start(),m.end()),
                    "pattern":pat,
                })

    time_candidates=[]
    time_patterns=[
        (r"(\d+(?:[.,]\d+)?)\s*(?:val(?:\.|and\w*)?|hours?|hrs?|h)\b",60.0,"hours"),
        (r"(\d+(?:[.,]\d+)?)\s*(?:min(?:\.|uc\w*)?|minutes?|mins?)\b",1.0,"minutes"),
        (r"(?:uztruk\w*|laik\w*|trukm\w*|duration|time)[^0-9]{0,45}(\d+(?:[.,]\d+)?)\s*(?:min(?:\.|uc\w*)?|minutes?|mins?)",1.0,"minutes"),
        (r"(?:uztruk\w*|laik\w*|trukm\w*|duration|time)[^0-9]{0,45}(\d+(?:[.,]\d+)?)\s*(?:val(?:\.|and\w*)?|hours?|hrs?|h)\b",60.0,"hours"),
    ]
    for pat,mult,unit in time_patterns:
        for m in re.finditer(pat,norm,flags=re.I):
            try:
                value=float(m.group(1).replace(",", "."))*mult
            except Exception:
                continue
            if 0 < value <= 24*60:
                time_candidates.append({
                    "minutes":round(value,2),
                    "context":_questionnaire_context(raw,m.start(),m.end()),
                    "unit":unit,
                    "pattern":pat,
                })

    # Prefer candidates whose local context explicitly mentions GPT/schedule/rework.
    def iter_rank(c):
        ctx=unicodedata.normalize("NFKD",c["context"]).encode("ascii","ignore").decode("ascii").lower()
        score=sum(k in ctx for k in ["gpt","grafik","tvarkar","iter","perdary","bandym","prompt"])
        return (-score,c["value"])

    def time_rank(c):
        ctx=unicodedata.normalize("NFKD",c["context"]).encode("ascii","ignore").decode("ascii").lower()
        score=sum(k in ctx for k in ["gpt","grafik","tvarkar","laik","truk","uztruk","schedule"])
        return (-score,c["minutes"])

    iteration_candidates=sorted(iteration_candidates,key=iter_rank)
    time_candidates=sorted(time_candidates,key=time_rank)

    best_iter=iteration_candidates[0] if iteration_candidates else None
    best_time=time_candidates[0] if time_candidates else None

    return {
        "iterations":None if best_iter is None else int(best_iter["value"]),
        "minutes":None if best_time is None else float(best_time["minutes"]),
        "iteration_evidence":None if best_iter is None else best_iter["context"],
        "time_evidence":None if best_time is None else best_time["context"],
        "iteration_candidates":iteration_candidates[:10],
        "time_candidates":time_candidates[:10],
    }


def _questionnaire_consensus(parsed_by_respondent):
    iter_vals=[
        int(v.get("iterations"))
        for v in parsed_by_respondent.values()
        if v and v.get("iterations") is not None
    ]
    time_vals=[
        float(v.get("minutes"))
        for v in parsed_by_respondent.values()
        if v and v.get("minutes") is not None
    ]
    return {
        "iterations":None if not iter_vals else int(round(median(iter_vals))),
        "minutes":None if not time_vals else round(float(median(time_vals)),1),
        "n_iterations":len(iter_vals),
        "n_minutes":len(time_vals),
    }


def _questionnaire_render_result(initials,extract,parsed):
    if not extract:
        return
    if extract.get("error"):
        st.warning(f"{initials}: {extract['error']}")
    if extract.get("ok"):
        st.success(
            f"{initials}: perskaityta {extract.get('file_name')} · "
            f"SHA-256 {extract.get('file_hash','')[:12]}…"
        )
        q1,q2=st.columns(2)
        q1.metric(
            "Rastos iteracijos",
            "Nerasta" if parsed.get("iterations") is None else parsed.get("iterations")
        )
        q2.metric(
            "Rastas bendras laikas",
            "Nerasta" if parsed.get("minutes") is None else f"{parsed.get('minutes'):g} min"
        )
        if parsed.get("iteration_evidence") or parsed.get("time_evidence"):
            with st.expander(f"{initials} — ką parseris rado",expanded=False):
                if parsed.get("iteration_evidence"):
                    st.markdown("**Iterations evidence**")
                    st.caption(parsed["iteration_evidence"])
                if parsed.get("time_evidence"):
                    st.markdown("**Time evidence**")
                    st.caption(parsed["time_evidence"])
                st.caption(
                    "Tai automatinė teksto interpretacija, ne galutinis research outcome. "
                    "Prieš lock skaičius patvirtink / pataisyk laukeliuose žemiau."
                )



def _research_targets(y,m,people):
    """Research shadow uses the exact same workload-target calculation as Sudarymas."""
    return calculate_targets(y,m,people)


def _research_shadow_result_from_run(run,people,y,m):
    assignments={int(k):v for k,v in (run.get("assignments") or {}).items()}
    frozen_stats=run.get("full_stats") or {}
    if not isinstance(frozen_stats,dict) or "global" not in frozen_stats:
        frozen_stats=validate_schedule(
            y,m,people,make_slots(y,m),assignments,_research_targets(y,m,people)
        )
    proof=_friday_waterfill_proof(frozen_stats)
    frozen_ok=bool(run.get("success")) and bool(proof.get("passed"))
    msg=f"Frozen research shadow run {run.get('run_no')}"
    if bool(run.get("success")) and not proof.get("passed"):
        msg+=(
            f" — LEGACY FRIDAY WATER-FILL INVALID: total {proof['total']} requires "
            f"{proof['floor']}-{proof['ceil']} each, observed spread {proof['spread']}"
        )
    return SolveResult(
        ok=frozen_ok,
        message=msg,
        assignments=assignments,
        targets=_research_targets(y,m,people),
        stats=frozen_stats,
        objective_value=None,
    )


def _research_shadow_run_log_df(runs):
    rows=[]
    total=0.0
    for r in runs:
        total+=float(r.get("elapsed_seconds") or 0)
        rows.append({
            "Run":int(r.get("run_no") or 0),
            "Role":"FIRST SHOT (PRIMARY)" if int(r.get("run_no") or 0)==1 else "IMPROVEMENT ATTEMPT",
            "Success":bool(r.get("success")),
            "Time, s":round(float(r.get("elapsed_seconds") or 0),2),
            "Cumulative time, s":round(total,2),
            "Engine version":r.get("app_version"),
            "Rule Profile":r.get("rule_profile_version"),
            "Solver stage":r.get("solver_stage"),
            "HARD errors":r.get("hard_errors"),
            "Monthly fairness %":r.get("monthly_fairness"),
            "Preference %":r.get("preference_mean"),
            "Monthly workplace imbalance":r.get("monthly_workplace_imbalance"),
            "Weekend spread":r.get("weekend_spread"),
            "Friday spread":r.get("friday_spread"),
            "Double spread":r.get("double_spread"),
            "Weekday-day spread":r.get("weekday_day_spread"),
            "Distinct workplaces":r.get("mean_distinct_workplaces"),
            "Gap count":r.get("gap_count"),
            "Gap workplace spread":r.get("gap_category_spread"),
        })
    return pd.DataFrame(rows)


def _research_people_stats_df(result):
    """Full per-resident research stats without reading operational backup tables."""
    rows=[]
    for initials,d in (result.stats or {}).get("people",{}).items():
        row={
            "Initials":initials,
            "Name":d.get("name",""),
            "Target":d.get("target"),
            "Workload":d.get("workload"),
            "Assignments":d.get("assignments"),
            "Weekday assignments":d.get("weekday_assignments"),
            "Distinct weekdays":d.get("weekday_days"),
            "Weekend assignments":d.get("weekend_assignments"),
            "Prior weekends":d.get("prior_weekend_count"),
            "Cumulative weekends":d.get("cumulative_weekend_count"),
            "Friday assignments (frozen SYSTEM run)":d.get("friday_assignments"),
            "Doubles":d.get("doubles"),
            "Max consecutive days":d.get("max_consecutive_days"),
            "Max rolling-7 hours":d.get("max_rolling7_hours"),
            "Max calendar-week hours":d.get("max_calendar_week_hours"),
            "Free days":d.get("fully_free_days"),
            "Consecutive double pairs":d.get("consecutive_double_pairs"),
            "Worked day after two doubles":d.get("worked_after_two_doubles"),
            "Distinct workplaces":d.get("distinct_rotations"),
            "RESIDENT HARD requested":d.get("resident_hard_requested"),
            "RESIDENT HARD honored":d.get("resident_hard_honored"),
            "RESIDENT HARD losses":d.get("resident_hard_losses"),
            "RESIDENT HARD fulfilment %":d.get("resident_hard_score"),
            "Prior RESIDENT HARD losses":d.get("prior_resident_hard_loss_count"),
            "Cumulative RESIDENT HARD losses":d.get("cumulative_resident_hard_losses"),
            "Exact SOFT requests":d.get("exact_preference_requests"),
            "Exact SOFT honored":d.get("exact_preference_honored"),
            "Exact SOFT fulfilment %":d.get("exact_preference_score"),
            "SOFT preference %":d.get("soft_preference_score"),
            "Overall resident-request satisfaction %":d.get("overall_request_score",d.get("preference_score")),
        }
        rc=d.get("rotation_counts") or {}
        for cat in ROTATION_CATEGORIES:
            row[cat]=int(rc.get(cat,0) or 0)
        rows.append(row)
    return pd.DataFrame(rows)


def _research_global_stats_df(result):
    g=(result.stats or {}).get("global",{})
    rows=[]
    for key in sorted(g):
        value=g.get(key)
        if isinstance(value,(dict,list,tuple,set)):
            value=json.dumps(_research_json_safe(value),ensure_ascii=False,sort_keys=True)
        rows.append({"Metric":key,"Value":value})
    return pd.DataFrame(rows)


def _research_hard_errors_df(result):
    errors=list(((result.stats or {}).get("global",{}) or {}).get("errors") or [])
    return pd.DataFrame([{"#":i+1,"HARD / solver detail":e} for i,e in enumerate(errors)])


def _research_feasibility_precheck_df(y,m,people):
    """Two generous individual upper bounds: ABSOLUTE HARD vs strict RESIDENT HARD.

    The ABSOLUTE-only bound answers whether even relaxing every personal
    `Negaliu dirbti` request could ever reach the target. The strict bound shows
    whether some RESIDENT-HARD relaxation may be necessary before group
    competition and other fairness rules are considered.
    """
    targets=_research_targets(y,m,people)
    slots=make_slots(y,m)
    rows=[]

    def upper_for(p,strict_resident_hard):
        upper2=0
        for d in range(1,calendar.monthrange(y,m)[1]+1):
            ds=[]
            for sl in slots:
                if sl.day!=d or sl.blocked:
                    continue
                blocked=(
                    hard_unavailable_for_block(p,d,sl.block)
                    if strict_resident_hard else
                    absolute_unavailable_for_block(p,d,sl.block)
                )
                if not blocked:
                    ds.append(sl)
            best=0
            for a in ds:
                best=max(best,int(a.workload2))
            for ai,a in enumerate(ds):
                for b in ds[ai+1:]:
                    if blocks_overlap(a.block,b.block):
                        continue
                    best=max(best,int(a.workload2)+int(b.workload2))
            upper2+=best
        return upper2/2.0

    for p in people:
        target=float(targets.get(p.initials,0))
        strict_upper=upper_for(p,True)
        absolute_upper=upper_for(p,False)
        rows.append({
            "Initials":p.initials,
            "Name":p.name,
            "Target":target,
            "Strict zero-RESIDENT-HARD-loss max (generous)":strict_upper,
            "ABSOLUTE-HARD-only max after Resident-HARD relaxation (generous)":absolute_upper,
            "Resident-HARD relaxation may be required":bool(target>strict_upper+1e-9 and target<=absolute_upper+1e-9),
            "Clear ABSOLUTE infeasibility":bool(target>absolute_upper+1e-9),
        })
    return pd.DataFrame(rows)


def _research_input_preferences_df(people, year=None, month=None):
    """Human-readable audit of scored requests plus ignored legacy signals.

    Holiday preference is active only for months that actually contain an official
    public holiday, so a standing account setting does not create a fake scored
    request in a holiday-free research month.
    """
    rows=[]
    for p in people:
        exact_soft=(
            len(p.soft_free)+len(p.soft_free_am)+len(p.soft_free_pm)
            +len(p.preferred)+len(p.preferred_am)+len(p.preferred_pm)
        )
        holiday_active = bool(
            int(getattr(p,"holiday_preference",0) or 0)
            and year is not None and month is not None
            and public_holiday_days_in_month(int(year),int(month))
        )
        directional=sum([
            1 if p.spread_preference else 0,
            1 if holiday_active else 0,
            1 if p.avoid_doubles else 0,
        ])
        ignored_legacy_directional=sum([
            1 if p.weekday_preference else 0,
            1 if p.weekend_preference else 0,
        ])
        resident_hard=(len(p.unavailable)+len(p.unavailable_am)+len(p.unavailable_pm))
        absolute_hard=(len(p.vacation)+len(p.justified_absence)+len(p.long_duty))
        choice_items=sum(
            1 for x in (p.request_items or [])
            if x.get("included_in_score") and x.get("kind") in ("backup_claim","rest_credit")
        )
        active_overall=bool(resident_hard or exact_soft or directional or choice_items)
        rows.append({
            "Initials":p.initials,
            "Name":p.name,
            "RESIDENT HARD request units":resident_hard,
            "ABSOLUTE HARD audit units":absolute_hard,
            "Exact SOFT requests":exact_soft,
            "Scored directional SOFT settings":directional,
            "Holiday setting":({1:"Prefer work",0:"Neutral",-1:"Prefer rest"}.get(int(getattr(p,"holiday_preference",0) or 0),"Neutral")),
            "Holiday preference active this month":holiday_active,
            "Ignored legacy weekday/weekend signals":ignored_legacy_directional,
            "Other structured resident choices":choice_items,
            "Overall request % status":("ACTIVE" if active_overall else "N/A — no scored request submitted"),
            "SOFT % status":("ACTIVE" if exact_soft or directional else "N/A — no SOFT submitted"),
            "Preferred work":", ".join(map(str,sorted(p.preferred))) or "—",
            "Preferred AM":", ".join(map(str,sorted(p.preferred_am))) or "—",
            "Preferred PM":", ".join(map(str,sorted(p.preferred_pm))) or "—",
            "Soft free":", ".join(map(str,sorted(p.soft_free))) or "—",
            "RESIDENT HARD unavailable":", ".join(map(str,sorted(p.unavailable))) or "—",
            "ABSOLUTE vacation":", ".join(map(str,sorted(p.vacation))) or "—",
            "ABSOLUTE other justified absence":", ".join(map(str,sorted(p.justified_absence))) or "—",
        })
    return pd.DataFrame(rows)


def _research_write_schedule_grid(writer,sheet_name,y,m,result):
    """Write the same Sudarymas grid to research XLSX and apply resident colors."""
    grid=schedule_grid(y,m,result)
    grid.to_excel(writer,sheet_name=sheet_name,index=True)
    ws=writer.sheets[sheet_name]
    ws.freeze_panes(1,1)
    ws.set_column(0,0,28)
    ws.set_column(1,len(grid.columns),8)
    wb=writer.book
    resident_formats={}
    for initials,color in PERSON_COLORS.items():
        resident_formats[initials]=wb.add_format({
            "bg_color":color,
            "font_color":contrast_text(color),
            "bold":True,
            "align":"center",
            "valign":"vcenter",
            "border":1,
        })
    block_fmt=wb.add_format({"bg_color":"#D9D9D9","font_color":"#555555","bold":True,"align":"center","border":1})
    for r_idx,row in enumerate(grid.itertuples(index=False),start=1):
        for c_idx,value in enumerate(row,start=1):
            if value in resident_formats:
                ws.write(r_idx,c_idx,value,resident_formats[value])
            elif value=="BLOCK":
                ws.write(r_idx,c_idx,value,block_fmt)


def research_shadow_xlsx(y,m,case,people,runs):
    bio=BytesIO()
    valid=[
        r for r in runs
        if bool(r.get("success"))
        and int(r.get("hard_errors") or 0)==0
        and bool(r.get("assignments"))
    ]
    best=min(valid,key=_research_run_quality_key) if valid else None
    first=next((r for r in runs if int(r.get("run_no") or 0)==1),None)

    with pd.ExcelWriter(bio,engine="xlsxwriter") as w:
        pd.DataFrame([
            ["Research object","MY ENGINE SHADOW / FAKE GENERATOR"],
            ["Cycle",f"{y}-{m:02d}"],
            ["Operational status","RESEARCH ONLY — unconfirmed; never published; never changes the operational Seniūnė schedule"],
            ["Profile","ŠR resident profile / researcher access"],
            ["Input source","Uploaded wishes / HARD workbook"],
            ["Input SHA-256",case.get("input_hash")],
            ["Engine version at lock",case.get("app_version_at_lock")],
            ["Rule Profile at lock",case.get("rule_profile_version_at_lock")],
            ["Primary endpoint","Frozen Run 1 / FIRST SHOT"],
            ["Reset policy","Researcher may delete/reset the active frozen experiment; reset event is audit-logged"],
            ["Created at",case.get("created_at")],
        ],columns=["Field","Value"]).to_excel(w,index=False,sheet_name="method")

        _research_shadow_run_log_df(runs).to_excel(w,index=False,sheet_name="run_log")
        _research_feasibility_precheck_df(y,m,people).to_excel(w,index=False,sheet_name="feasibility_precheck")

        snapshot_excel=[]
        for row in _research_people_snapshot(people):
            snapshot_excel.append({
                k:(json.dumps(v,ensure_ascii=False,sort_keys=True) if isinstance(v,(list,dict)) else v)
                for k,v in row.items()
            })
        pd.DataFrame(snapshot_excel).to_excel(w,index=False,sheet_name="input_snapshot")

        audit=case.get("import_audit") or {}
        pd.DataFrame([{
            "Workbook SHA-256":audit.get("workbook_hash"),
            "Sheet":r.get("sheet"),
            "Status":r.get("status"),
            "Header row":r.get("header_row"),
            "Rows":r.get("rows"),
            "Reason":r.get("reason",""),
        } for r in audit.get("sheets",[]) or []] or [{
            "Workbook SHA-256":audit.get("workbook_hash",""),
            "Sheet":"","Status":"No audit rows","Header row":"","Rows":"","Reason":""
        }]).to_excel(w,index=False,sheet_name="input_import_audit")

        # Every run gets its own schedule + all statistics. Failed runs get a non-empty
        # diagnostic sheet, so research downloads are never misleadingly blank.
        for r in runs:
            n=int(r.get("run_no") or 0)
            rr=_research_shadow_result_from_run(r,people,y,m)
            if r.get("assignments"):
                _research_write_schedule_grid(w,f"r{n}_grid",y,m,rr)
                schedule_list_df(y,m,rr).to_excel(w,index=False,sheet_name=f"r{n}_schedule")
                _research_people_stats_df(rr).to_excel(w,index=False,sheet_name=f"r{n}_people_stats")
                _research_post_matrix_df(rr,people).to_excel(w,index=False,sheet_name=f"r{n}_post_matrix")
                _research_global_stats_df(rr).to_excel(w,index=False,sheet_name=f"r{n}_global_stats")
                _research_hard_errors_df(rr).to_excel(w,index=False,sheet_name=f"r{n}_hard_errors")
            else:
                g=(r.get("full_stats") or {}).get("global",{}) if isinstance(r.get("full_stats"),dict) else {}
                errors=g.get("errors") or []
                pd.DataFrame([
                    {"Field":"Run","Value":n},
                    {"Field":"Success","Value":bool(r.get("success"))},
                    {"Field":"Solver stage","Value":r.get("solver_stage") or g.get("solve_stage")},
                    {"Field":"HARD errors","Value":r.get("hard_errors")},
                    {"Field":"Elapsed seconds","Value":r.get("elapsed_seconds")},
                    {"Field":"Failure / diagnostic","Value":" | ".join(map(str,errors)) if errors else "No assignments were returned by the engine."},
                ]).to_excel(w,index=False,sheet_name=f"r{n}_failure")

        if first and first.get("assignments"):
            r1=_research_shadow_result_from_run(first,people,y,m)
            _research_write_schedule_grid(w,"PRIMARY_grid",y,m,r1)
            _research_people_stats_df(r1).to_excel(w,index=False,sheet_name="PRIMARY_people")
            _research_global_stats_df(r1).to_excel(w,index=False,sheet_name="PRIMARY_metrics")

        if best and best.get("assignments"):
            rb=_research_shadow_result_from_run(best,people,y,m)
            _research_write_schedule_grid(w,"BEST_grid",y,m,rb)
            _research_people_stats_df(rb).to_excel(w,index=False,sheet_name="BEST_people")
            _research_global_stats_df(rb).to_excel(w,index=False,sheet_name="BEST_metrics")

        # Make ordinary table sheets readable.
        for name,ws in w.sheets.items():
            if name.endswith("_grid") or name in ("PRIMARY_grid","BEST_grid"):
                continue
            ws.freeze_panes(1,0)
            ws.set_column(0,0,28)
            ws.set_column(1,30,20)

    return bio.getvalue()


def _render_research_shadow_result(rr,run_no,y,m,people,primary=False):
    g=(rr.stats or {}).get("global",{})
    label=("PRIMARY — MOCK FIRST SHOT" if primary else f"MOCK RUN {run_no}")
    st.markdown(f"### {label}")
    st.info(
        "RESEARCH SHADOW / UNCONFIRMED — tai yra Sudarymas juodraščio ekvivalentas. "
        "Nėra Publish / Confirm / Authenticate mygtuko ir šis grafikas nepatenka į operacinį grafiką."
    )
    m1,m2,m3,m4=st.columns(4)
    m1.metric("HARD errors *",g.get("hard_errors","—"))
    m2.metric("Monthly fairness",("—" if g.get("monthly_fairness_score") is None else f"{g.get('monthly_fairness_score')}%"))
    m3.metric("Worst post spread",g.get("worst_monthly_post_spread","—"))
    m4.metric("Post structural water-fill","PASS" if g.get("post_spread_quality_gate_passed") else "FAIL")
    q1,q2,q3,q4=st.columns(4)
    q1.metric("Active SOFT residents",g.get("active_preference_residents",0))
    q2.metric("Worst preference %",g.get("min_preference_score") if g.get("min_preference_score") is not None else "N/A")
    q3.metric("Mean preference %",g.get("mean_preference_score") if g.get("mean_preference_score") is not None else "N/A")
    q4.metric("Preference spread, pp",g.get("preference_score_spread") if g.get("preference_score_spread") is not None else "N/A")
    p1,p2,p3,p4=st.columns(4)
    p1.metric("SYSTEM-HARD worst post lock",g.get("post_system_hard_worst_spread_lock","—"))
    p2.metric("SYSTEM-HARD total 9-post spread",g.get("post_system_hard_total_spread_lock","—"))
    p3.metric("Distinct workplace spread",g.get("distinct_rotation_spread","—"))
    p4.metric("Post-stage proof",("OPTIMAL" if g.get("post_system_hard_stage_optimal") else "BEST FOUND"))
    if g.get("generation_quality_issues"):
        st.warning("Generation quality diagnostics: "+"; ".join(map(str,g.get("generation_quality_issues") or [])))

    # The actual visible Sudarymas-style grid — same helper and permanent resident colors.
    st.dataframe(
        style_schedule(schedule_grid(y,m,rr)),
        use_container_width=True,
        height=620,
    )

    friday_proof=_friday_waterfill_proof(rr.stats)
    if not friday_proof.get("passed"):
        counts=friday_proof.get("counts") or {}
        st.error(
            "FRIDAY WATER-FILL INVALID šiame frozen run: "
            f"{friday_proof['total']} Friday assignments / {friday_proof['n']} rezidentų → "
            f"teisingas entitlement {friday_proof['floor']}-{friday_proof['ceil']} kiekvienam; "
            f"šiame run observed {min(counts.values()) if counts else 0}-{max(counts.values()) if counts else 0} "
            f"(spread {friday_proof['spread']}). Tai legacy frozen rezultatas; naujas V2.5.86 run su tokiu spread negali būti pažymėtas validžiu."
        )
    else:
        st.success(
            f"Friday SYSTEM water-fill PASS: {friday_proof['total']} assignments / {friday_proof['n']} residents → "
            f"{friday_proof['floor']}-{friday_proof['ceil']} each, raw spread {friday_proof['spread']}."
        )

    t1,t2,t3,t4=st.tabs(["Resident stats","Post matrix","Global metrics","HARD / diagnostics"])
    with t1:
        st.dataframe(_research_people_stats_df(rr),use_container_width=True,hide_index=True,height=520)
    with t2:
        st.dataframe(_research_post_matrix_df(rr,people),use_container_width=True,hide_index=True,height=520)
    with t3:
        st.dataframe(_research_global_stats_df(rr),use_container_width=True,hide_index=True,height=520)
    with t4:
        render_hard_error_explainer(g,lang,key_suffix=f"shadow_{y}_{m}_{run_no}")
        hdf=_research_hard_errors_df(rr)
        if not hdf.empty:
            st.dataframe(hdf,use_container_width=True,hide_index=True)


def render_research_shadow_generator():
    st.subheader("FAKE GENERATOR — RESEARCH SHADOW")
    st.success("ŠR RESEARCHER ACCESS — embedded in the same single ŠR account window. No profile switching is required.")
    st.error(
        "RESEARCH ONLY. Šis generatorius NIEKADA nepublikuoja grafiko, "
        "nekeičia realaus operacinio Seniūnės grafiko, backupų, fairness_history ar operacinių duomenų."
    )
    st.caption(
        "GENERATE MOCK SCHEDULE naudoja tą patį solve_schedule engine kaip normalus Seniūnės Sudarymas. "
        "Skirtumas: rezultatas saugomas tik research-shadow lentelėse ir neturi Publish/Confirm/Authenticate veiksmo."
    )

    st.markdown(f"### {month_label(year,month)}")
    case=db.get_research_shadow_case_v2545(year,month)

    if not case:
        st.markdown("### 1. UPLOAD WISHES / HARD CONSTRAINTS")
        wishes_file=st.file_uploader(
            "Pageidavimai / HARD taisyklės (.xlsx / .xls)",
            type=["xlsx","xls"],
            key=f"research_shadow_wishes_{year}_{month}"
        )
        st.caption(
            "Galima kelti originalų mėnesio Excel. Importeris skanuoja VISUS worksheet'us, "
            "sujungia suderinamas pageidavimų/HARD lenteles ir saugo sheet/row provenance."
        )

        preflight_ok=False
        people=None
        audit=None
        warnings=[]
        if wishes_file is not None:
            try:
                people,audit,warnings=research_people_from_excel(
                    wishes_file,year,month,return_audit=True
                )
                st.markdown("### WHOLE-WORKBOOK WISHES PREFLIGHT")
                a1,a2,a3=st.columns(3)
                a1.metric("Workbook sheets",audit.get("sheet_count"))
                a2.metric("Prefs/HARD sheets used",audit.get("used_sheet_count"))
                a3.metric("Warnings",len(warnings or []))
                scan_df=pd.DataFrame([{
                    "Sheet":r.get("sheet"),
                    "Status":r.get("status"),
                    "Header row":r.get("header_row"),
                    "Rows":r.get("rows"),
                    "Reason":r.get("reason",""),
                } for r in audit.get("sheets",[])])
                st.dataframe(scan_df,use_container_width=True,hide_index=True)
                if warnings:
                    with st.expander(f"Import warnings ({len(warnings)})"):
                        for x in warnings:
                            st.write("• "+str(x))

                pref_input_df=_research_input_preferences_df(people,year,month)
                active_requests=int((pref_input_df["Overall request % status"]=="ACTIVE").sum())
                resident_hard_units=int(pref_input_df["RESIDENT HARD request units"].sum())
                exact_soft=int(pref_input_df["Exact SOFT requests"].sum())
                st.markdown("### IMPORTED RESIDENT INPUT AUDIT")
                p1,p2,p3,p4=st.columns(4)
                p1.metric("Residents in roster",len(pref_input_df))
                p2.metric("Residents with scored requests",active_requests)
                p3.metric("RESIDENT HARD units",resident_hard_units)
                p4.metric("Exact SOFT units",exact_soft)
                st.dataframe(pref_input_df,use_container_width=True,hide_index=True,height=520)
                st.caption(
                    "V2.5.49: `Negaliu dirbti` yra RESIDENT HARD ir įeina į bendrą rezidento prašymų išpildymą. "
                    "Liga / atostogos / teisės-poilsio sauga yra ABSOLUTE HARD: jos audituojamos atskirai ir į procento vardiklį neįtraukiamos, nes jų negalima aukoti. "
                    "SOFT fairness lieka MAX-MIN: tarp aktyvių SOFT rezidentų 85/85/85 yra geriau už 100/100/55."
                )

                feas=_research_feasibility_precheck_df(year,month,people)
                impossible=feas[feas["Clear ABSOLUTE infeasibility"]==True]
                with st.expander("ABSOLUTE / RESIDENT HARD feasibility pre-check",expanded=not impossible.empty):
                    st.dataframe(feas,use_container_width=True,hide_index=True)
                    st.caption(
                        "Pirmas upper bound saugo visus RESIDENT HARD; antras leidžia tik minimaliai aukoti `Negaliu dirbti`, bet vis tiek niekada nelaužo ABSOLUTE HARD. "
                        "Tik kai Target viršija net ABSOLUTE-HARD-only upper bound, tai aiškus individualus matematinis neįmanomumas."
                    )
                if not impossible.empty:
                    st.warning(
                        "Pre-check found at least one resident whose target exceeds even a generous maximum allowed by the uploaded HARD dates. "
                        "You may still run the engine so the FIRST SHOT failure is recorded, or correct the source file before generating."
                    )
                preflight_ok=True
                st.success("WHOLE-WORKBOOK WISHES PREFLIGHT — PASSED")
            except Exception as exc:
                st.error("WHOLE-WORKBOOK WISHES PREFLIGHT — BLOCKED")
                st.error(str(exc))

        if st.button(
            "GENERATE MOCK SCHEDULE — FIRST SHOT",
            type="primary",
            use_container_width=True,
            disabled=(not preflight_ok),
            key=f"research_shadow_run1_{year}_{month}"
        ):
            snapshot=_research_people_snapshot(people)
            input_hash=_research_input_hash(year,month,people)
            case=db.create_research_shadow_case_v2545(
                year,month,input_hash,snapshot,_research_json_safe(audit),
                APP_VERSION,ACTIVE_RULE_PROFILE_VERSION
            )

            t0=perf_counter()
            engine_message=""
            try:
                result=solve_schedule(year,month,people,time_limit=180.0)
                elapsed=perf_counter()-t0
                full_stats=_research_json_safe(result.stats or {})
                run_ok=result.ok
                assignments=result.assignments
                engine_message=str(result.message or "")
                if not run_ok:
                    full_stats.setdefault("global",{})
                    full_stats["global"].setdefault("errors",[])
                    if engine_message and engine_message not in full_stats["global"]["errors"]:
                        full_stats["global"]["errors"].append(engine_message)
                    full_stats["global"].setdefault("solve_stage","NO_VALID_FIRST_SHOT")
            except Exception as solve_exc:
                elapsed=perf_counter()-t0
                engine_message=f"ENGINE EXCEPTION: {type(solve_exc).__name__}: {solve_exc}"
                full_stats={
                    "global":{
                        "hard_errors":None,
                        "errors":[engine_message],
                        "solve_stage":"ENGINE_EXCEPTION",
                    },
                    "people":{}
                }
                run_ok=False
                assignments={}

            db.record_research_shadow_run_v2545(
                case["id"],input_hash,elapsed,run_ok,
                APP_VERSION,ACTIVE_RULE_PROFILE_VERSION,
                _research_json_safe(full_stats),assignments
            )
            if run_ok:
                st.success("MOCK FIRST SHOT generated and frozen in research storage. It remains UNCONFIRMED / NOT PUBLISHED.")
            else:
                st.error("FIRST SHOT returned no valid schedule. The failure and diagnostics were frozen for research; you can delete/reset this experiment below and try again.")
            st.rerun()
        return

    people=_research_people_from_snapshot(case.get("input_snapshot") or [])
    runs=db.list_research_shadow_runs_v2545(case["id"])
    st.success(
        f"LOCKED SHADOW INPUT {year}-{month:02d} · "
        f"SHA-256 {str(case.get('input_hash'))[:16]}…"
    )
    st.caption(
        f"Frozen engine: {case.get('app_version_at_lock')} · "
        f"Rule Profile v{case.get('rule_profile_version_at_lock')} · "
        f"created {str(case.get('created_at') or '')[:19]}"
    )

    # Researcher-only reset: intentionally available in ŠR resident profile so
    # operational Seniūnė role changes never remove ŠR research control.
    with st.expander("DELETE / RESET FROZEN MOCK EXPERIMENT",expanded=False):
        st.warning(
            "This deletes the active research-shadow case and all its frozen runs for this month, then lets you upload wishes and generate a new FIRST SHOT. "
            "It does NOT delete or change the operational Seniūnė draft/published schedule, backups, swaps or fairness_history. A reset audit tombstone is kept."
        )
        reset_token=f"DELETE SHADOW {year}-{month:02d}"
        typed=st.text_input(
            f"Type exactly: {reset_token}",
            key=f"research_shadow_delete_confirm_{case['id']}"
        )
        reason=st.text_input(
            "Reset note (optional)",
            key=f"research_shadow_delete_reason_{case['id']}",
            placeholder="e.g. testing input corrected / regenerate experiment"
        )
        if st.button(
            "DELETE FROZEN FIRST SHOT / ALL SHADOW RUNS",
            type="primary",
            use_container_width=True,
            disabled=(typed.strip()!=reset_token),
            key=f"research_shadow_delete_{case['id']}"
        ):
            db.reset_research_shadow_case_v2546(year,month,reason or "researcher requested reset")
            st.session_state.pop(f"research_shadow_wishes_{year}_{month}",None)
            st.success("Research shadow experiment deleted/reset. You can now generate a completely new mock FIRST SHOT from the same or corrected wishes file.")
            st.rerun()

    st.markdown("### Shadow run log")
    if runs:
        st.dataframe(_research_shadow_run_log_df(runs),use_container_width=True,hide_index=True)

    version_ok=(
        str(case.get("app_version_at_lock"))==APP_VERSION
        and int(case.get("rule_profile_version_at_lock") or 0)==int(ACTIVE_RULE_PROFILE_VERSION)
    )
    if not version_ok:
        st.error(
            "FROZEN ENGINE MISMATCH — šio mėnesio shadow testas turi būti tęsiamas tik su "
            f"{case.get('app_version_at_lock')} / Rule Profile v{case.get('rule_profile_version_at_lock')}."
        )

    next_no=len(runs)+1
    if next_no<=5:
        if st.button(
            f"PERTIKRINTI / GERINTI MOCK GRAFIKĄ — RUN {next_no}/5",
            use_container_width=True,
            disabled=not version_ok,
            key=f"research_shadow_next_{case['id']}_{next_no}"
        ):
            locked_hash=_research_input_hash(year,month,people)
            if locked_hash!=case.get("input_hash"):
                st.error("Locked input hash validation failed. Run aborted.")
                return
            t0=perf_counter()
            try:
                result=solve_schedule(year,month,people,time_limit=180.0)
                elapsed=perf_counter()-t0
                full_stats=_research_json_safe(result.stats or {})
                run_ok=result.ok
                assignments=result.assignments
                engine_message=str(result.message or "")
                if not run_ok:
                    full_stats.setdefault("global",{})
                    full_stats["global"].setdefault("errors",[])
                    if engine_message and engine_message not in full_stats["global"]["errors"]:
                        full_stats["global"]["errors"].append(engine_message)
                    full_stats["global"].setdefault("solve_stage","NO_VALID_SCHEDULE")
            except Exception as solve_exc:
                elapsed=perf_counter()-t0
                full_stats={
                    "global":{
                        "hard_errors":None,
                        "errors":[f"ENGINE EXCEPTION: {type(solve_exc).__name__}: {solve_exc}"],
                        "solve_stage":"ENGINE_EXCEPTION",
                    },
                    "people":{}
                }
                run_ok=False
                assignments={}
            saved=db.record_research_shadow_run_v2545(
                case["id"],locked_hash,elapsed,run_ok,
                APP_VERSION,ACTIVE_RULE_PROFILE_VERSION,
                _research_json_safe(full_stats),assignments
            )
            if run_ok:
                st.success(f"MOCK RUN {saved.get('run_no',next_no)} frozen. FIRST SHOT remains the primary endpoint.")
            else:
                st.warning(f"RUN {saved.get('run_no',next_no)} did not return a valid schedule; diagnostics were saved.")
            st.rerun()
    else:
        st.success("5/5 shadow runs complete. Reset the experiment if you intentionally want to start over.")

    first=next((r for r in runs if int(r.get("run_no") or 0)==1),None)
    valid=[
        r for r in runs
        if bool(r.get("success")) and int(r.get("hard_errors") or 0)==0 and bool(r.get("assignments"))
    ]
    best=min(valid,key=_research_run_quality_key) if valid else None

    # Always show the frozen FIRST SHOT diagnostics. If it has a schedule, render
    # the actual Sudarymas-style colored grid immediately in this window.
    if first:
        if first.get("assignments"):
            rr=_research_shadow_result_from_run(first,people,year,month)
            _render_research_shadow_result(rr,1,year,month,people,primary=True)
        else:
            st.markdown("### PRIMARY — MOCK FIRST SHOT")
            st.error("The engine returned no valid schedule for Run 1, so there is no grid to display for this frozen attempt.")
            fg=(first.get("full_stats") or {}).get("global",{}) if isinstance(first.get("full_stats"),dict) else {}
            diag=fg.get("errors") or []
            if diag:
                st.dataframe(pd.DataFrame({"FIRST SHOT diagnostic":diag}),use_container_width=True,hide_index=True)
            feas=_research_feasibility_precheck_df(year,month,people)
            st.dataframe(feas,use_container_width=True,hide_index=True)
            st.caption(
                "If a target exceeds the generous HARD-only maximum, the same operational Sudarymas engine is mathematically unable to create a valid schedule from that frozen input. "
                "Use DELETE / RESET above after correcting the input or target semantics."
            )

    if best and (not first or int(best.get("run_no") or 0)!=1):
        st.divider()
        rb=_research_shadow_result_from_run(best,people,year,month)
        _render_research_shadow_result(rb,int(best.get("run_no") or 0),year,month,people,primary=False)
        st.caption(
            f"Secondary best-of-{len(runs)} = Run {best.get('run_no')} · "
            f"fairness {best.get('monthly_fairness')}%. FIRST SHOT remains the primary research endpoint."
        )

    inspectable=[r for r in runs if r.get("assignments")]
    if len(inspectable)>1:
        st.divider()
        chosen_no=st.selectbox(
            "VIEW ANY FROZEN MOCK RUN",
            [int(r.get("run_no") or 0) for r in inspectable],
            format_func=lambda n:(
                f"Run {n} — FIRST SHOT" if n==1 else f"Run {n} — improvement attempt"
            ),
            key=f"research_shadow_view_run_{case['id']}"
        )
        chosen=next(r for r in inspectable if int(r.get("run_no") or 0)==int(chosen_no))
        # Avoid duplicating Run 1 / current best unless the researcher explicitly wants
        # to inspect it here; this selector is intentionally an experiment workbench.
        with st.expander(f"Inspect frozen Run {chosen_no}",expanded=False):
            chosen_result=_research_shadow_result_from_run(chosen,people,year,month)
            _render_research_shadow_result(
                chosen_result,int(chosen_no),year,month,people,primary=(int(chosen_no)==1)
            )

    st.download_button(
        "DOWNLOAD COMPLETE RESEARCH SHADOW DATASET (.xlsx)",
        research_shadow_xlsx(year,month,case,people,runs),
        file_name=f"my_engine_shadow_{year}_{month:02d}.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        type="primary",
        use_container_width=True
    )
    st.info(
        "Research export includes the schedule grid/list for every successful run, per-resident stats, post matrices, full global metrics, HARD diagnostics, run timings, input snapshot/provenance and feasibility pre-check."
    )

def render_available_gpt_vs_engine_research():
    st.subheader("AVAILABLE GPT + HUMAN vs MY ENGINE TOOL")
    st.success("ŠR RESEARCHER-ONLY · OFFICIAL RUN LOCK")
    st.caption(
        "Comparatorius = bendrinis / nespecializuotas GPT + seniūnė, kuri iteravo rankiniu būdu, "
        "rašė papildymus ir vizualiai sprendė, kada atsakymas jau pakankamai geras. "
        "MY ENGINE = ši specializuota schedulerio versija su formalizuotais HARD, fairness ir audit metrics."
    )
    st.info(
        "PRIMARY = FIRST SHOT (Run 1) ir jis DB užrakinamas visam laikui. "
        "SECONDARY = iki 4 papildomų bandymų, iš viso max 5. Run 1 niekada nepakeičiamas geresniu vėlesniu rezultatu."
    )

    c1,c2=st.columns(2)
    ry=int(c1.number_input("Comparison year",min_value=2026,max_value=2100,value=int(year),step=1,key="research_lock_year"))
    rm=int(c2.selectbox("Comparison month",options=list(range(1,13)),index=int(month)-1,key="research_lock_month"))

    case=db.get_research_scheduler_case_v2541(ry,rm)
    all_cases=db.list_research_scheduler_cases_v2541()

    # Engine freeze audit across already locked months.
    other_versions=sorted({
        (str(c.get("app_version_at_lock")),int(c.get("rule_profile_version_at_lock") or 0))
        for c in all_cases
    })
    if other_versions:
        st.caption("Already locked research engine(s): "+", ".join(
            f"{v} / Rule Profile v{rp}" for v,rp in other_versions
        ))

    if not case:
        t1,t2=st.columns(2)
        t1.download_button(
            "DOWNLOAD HISTORICAL INPUT TEMPLATE (.xlsx)",
            research_preferences_template(ry,rm),
            file_name=f"historical_inputs_{ry}_{rm:02d}.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            use_container_width=True
        )
        t2.download_button(
            "DOWNLOAD AVAILABLE GPT + HUMAN SCHEDULE TEMPLATE (.xlsx)",
            research_schedule_template(ry,rm),
            file_name=f"available_gpt_human_schedule_{ry}_{rm:02d}.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            use_container_width=True
        )

        st.warning(
            "Naudok tik realiai žinotus inputus. 2026-09 palyginimui rekomenduojama imti LIVE APP WISHES: "
            "užrakinimo momentu sistema nukopijuos dabartinius rugsėjo pageidavimus/HARD ir tą PATĮ frozen snapshot naudos abiem grafikams."
        )
        input_source=st.radio(
            "Wishes / HARD input source",
            ["LIVE APP WISHES / HARD — same selected month", "UPLOAD HISTORICAL INPUT WORKBOOK"],
            horizontal=True,key=f"research_input_source_{ry}_{rm}"
        )
        use_live_inputs=input_source.startswith("LIVE APP")
        u1,u2=st.columns(2)
        if use_live_inputs:
            u1.success("1. LIVE APP WISHES selected — no wishes Excel needed")
            pref_file=None
        else:
            pref_file=u1.file_uploader(
                "1. Historical preferences / HARD constraints",
                type=["xlsx","xls"],key=f"research_lock_pref_{ry}_{rm}"
            )
        sched_file=u2.file_uploader(
            "2. REAL / AVAILABLE GPT + HUMAN final schedule",
            type=["xlsx","xls"],key=f"research_lock_sched_{ry}_{rm}"
        )

        # V2.5.108 preflight. Live DB wishes OR whole-workbook historical wishes are frozen before comparison.
        import_preflight_ok=False
        preflight_people=None
        preflight_assignments=None
        preflight_warnings=[]
        pref_audit=None
        sched_audit=None

        if use_live_inputs or pref_file is not None or sched_file is not None:
            st.markdown("### SAME-INPUT COMPARISON PREFLIGHT")
            st.caption(
                "Svarbiausia taisyklė: REAL / GPT+HUMAN ir MY ENGINE vertinami pagal IDENTIŠKĄ užšaldytą pageidavimų snapshot. "
                "Schedule Excel importer still scans every worksheet and preserves provenance."
            )

        try:
            if use_live_inputs:
                preflight_people=load_people(ry,rm)
                pref_audit={
                    "source_type":"LIVE_APP_DB",
                    "selected_cycle":f"{ry}-{rm:02d}",
                    "snapshot_hash":_research_input_hash(ry,rm,preflight_people),
                    "resident_count":len(preflight_people),
                    "used_sheet_count":0,
                    "sheet_count":0,
                    "sheets":[],
                }
                st.success(
                    f"LIVE APP WISHES loaded: {len(preflight_people)} residents · snapshot {_research_input_hash(ry,rm,preflight_people)[:12]}…"
                )
                st.dataframe(_research_input_preferences_df(preflight_people,ry,rm),use_container_width=True,hide_index=True,height=420)
                st.info(
                    "These exact wishes/HARD will be frozen on LOCK and applied to BOTH schedules. "
                    "The hand-made schedule does not get its own easier/different wish set."
                )
            elif pref_file is not None:
                preflight_people,pref_audit,pref_warn=research_people_from_excel(
                    pref_file,ry,rm,return_audit=True
                )
                preflight_warnings.extend(pref_warn or [])
            if sched_file is not None:
                preflight_assignments,sched_warn,sched_audit=research_assignments_from_excel(
                    sched_file,ry,rm,return_audit=True
                )
                preflight_warnings.extend(sched_warn or [])

            if preflight_people is not None and sched_file is not None and preflight_assignments is not None:
                import_preflight_ok=True
                same_workbook=(
                    (not use_live_inputs) and pref_audit and sched_audit
                    and pref_audit.get("workbook_hash")==sched_audit.get("workbook_hash")
                )
                if same_workbook:
                    st.success(
                        "Tas pats Excel workbook įkeltas į abu laukus — gerai. "
                        "Preferences/HARD ir schedule importeriai jį skenuoja ATSKIRAI ir ima informaciją iš skirtingų atpažintų sheets."
                    )

                p1,p2,p3,p4=st.columns(4)
                p1.metric("Wish input",("LIVE APP DB" if use_live_inputs else "Excel workbook"))
                p2.metric("Residents frozen",len(preflight_people) if preflight_people is not None else "—")
                p3.metric("Schedule sheets used",sched_audit.get("used_sheet_count") if sched_audit else "—")
                p4.metric("Schedule assignments read",sched_audit.get("assignment_count") if sched_audit else "—")

                def _audit_sheet_df(audit):
                    if not audit:
                        return pd.DataFrame()
                    return pd.DataFrame([{
                        "Sheet":r.get("sheet"),
                        "Status":r.get("status"),
                        "Header row":r.get("header_row"),
                        "Rows":r.get("rows"),
                        "Reason":r.get("reason",""),
                    } for r in audit.get("sheets",[])])

                left_a,right_a=st.columns(2)
                with left_a:
                    st.markdown("**Wishes / HARD source audit**")
                    if use_live_inputs:
                        st.dataframe(pd.DataFrame([{
                            "Source":"LIVE APP DB",
                            "Cycle":f"{ry}-{rm:02d}",
                            "Snapshot hash":pref_audit.get("snapshot_hash"),
                            "Residents":pref_audit.get("resident_count"),
                        }]),use_container_width=True,hide_index=True)
                    else:
                        st.dataframe(_audit_sheet_df(pref_audit),use_container_width=True,hide_index=True)
                with right_a:
                    st.markdown("**Schedule worksheet scan**")
                    st.dataframe(_audit_sheet_df(sched_audit),use_container_width=True,hide_index=True)

                if sched_audit.get("duplicates"):
                    st.info(
                        f"Schedule import found {len(sched_audit['duplicates'])} exact duplicate rows across sheets; "
                        "they were audited and not double-counted."
                    )
                if preflight_warnings:
                    with st.expander(f"Import warnings ({len(preflight_warnings)})",expanded=False):
                        for w in preflight_warnings:
                            st.write("• "+str(w))

                st.success(
                    "SAME-INPUT PREFLIGHT — PASSED. Official lock will freeze exactly these wishes/HARD and use them for BOTH schedules."
                )
        except Exception as import_exc:
            import_preflight_ok=False
            st.error("SAME-INPUT PREFLIGHT — BLOCKED")
            st.error(str(import_exc))

        st.markdown("### GPT+human proceso klausimynai")
        st.caption(
            "Gali įkelti MG ir VL to mėnesio retrospektyvius klausimynus. "
            "Sistema perskaitys tekstą ir pasiūlys iterations / time reikšmes, bet jos nėra aklai užrakinamos — žemiau visada matai ir patvirtini galutinį skaičių."
        )
        q1,q2=st.columns(2)
        gm_q=q1.file_uploader(
            "MG klausimynas",
            type=["pdf","docx","xlsx","xls","csv","txt","md"],
            key=f"gm_questionnaire_{ry}_{rm}"
        )
        lv_q=q2.file_uploader(
            "VL klausimynas",
            type=["pdf","docx","xlsx","xls","csv","txt","md"],
            key=f"lv_questionnaire_{ry}_{rm}"
        )

        questionnaire_payloads={}
        parsed_by_respondent={}
        for initials,uploaded in [("MG",gm_q),("VL",lv_q)]:
            if uploaded is None:
                continue
            extract=_questionnaire_extract_text(uploaded)
            parsed=_questionnaire_parse_process_metrics(extract.get("text","")) if extract.get("ok") else {
                "iterations":None,"minutes":None,
                "iteration_evidence":None,"time_evidence":None,
                "iteration_candidates":[],"time_candidates":[]
            }
            questionnaire_payloads[initials]=(extract,parsed)
            parsed_by_respondent[initials]=parsed
            _questionnaire_render_result(initials,extract,parsed)

        consensus=_questionnaire_consensus(parsed_by_respondent)
        if parsed_by_respondent:
            st.info(
                "Questionnaire consensus suggestion: "
                + ("iterations = N/A" if consensus["iterations"] is None else f"iterations ≈ {consensus['iterations']}")
                + " · "
                + ("time = N/A" if consensus["minutes"] is None else f"time ≈ {consensus['minutes']:g} min")
                + ". Galutiniai research metadata laukai lieka redaguojami."
            )

        iter_key=f"gpt_iter_{ry}_{rm}"
        minutes_key=f"gpt_minutes_{ry}_{rm}"
        if consensus["iterations"] is not None and not str(st.session_state.get(iter_key,"")).strip():
            st.session_state[iter_key]=str(consensus["iterations"])
        if consensus["minutes"] is not None and not str(st.session_state.get(minutes_key,"")).strip():
            st.session_state[minutes_key]=str(consensus["minutes"])

        m1,m2=st.columns(2)
        gpt_iter_text=m1.text_input(
            "GPT+human perdarymų / iteracijų skaičius (patvirtintas)",
            placeholder="pvz. 6",key=iter_key
        )
        gpt_minutes_text=m2.text_input(
            "Apytikslis bendras GPT+human laikas, min. (patvirtintas)",
            placeholder="pvz. 45",key=minutes_key
        )
        method_note=st.text_area(
            "GPT+human proceso pastaba (nebūtina)",
            placeholder="Pvz. seniūnė kelis kartus taisė promptą ir vizualiai vertino rezultatą; formalios fairness statistikos neturėjo.",
            key=f"gpt_method_{ry}_{rm}"
        )

        version_mismatch=bool(all_cases) and any(
            str(c.get("app_version_at_lock"))!=APP_VERSION
            or int(c.get("rule_profile_version_at_lock") or 0)!=int(ACTIVE_RULE_PROFILE_VERSION)
            for c in all_cases
        )
        if version_mismatch:
            st.error(
                "ENGINE FREEZE MISMATCH: jau yra užrakintas kitas tyrimo mėnuo su kita engine arba Rule Profile versija. "
                "Official August/September comparison turi būti atliekamas ta pačia frozen versija."
            )

        confirm=st.checkbox(
            "PATVIRTINU: tai oficialus šio mėnesio tyrimo input snapshot; po užrakinimo jo nekeisiu.",
            key=f"confirm_research_case_{ry}_{rm}"
        )
        if st.button(
            "LOCK CASE + RUN 1 (FIRST SHOT)",
            type="primary",use_container_width=True,
            disabled=(not confirm or not sched_file or version_mismatch or not import_preflight_ok),
            key=f"lock_run1_{ry}_{rm}"
        ):
            try:
                if not import_preflight_ok or preflight_people is None or preflight_assignments is None:
                    raise ValueError("WHOLE_WORKBOOK_PREFLIGHT_REQUIRED")
                people=preflight_people
                assignments=preflight_assignments
                warnings=list(preflight_warnings or [])
                warnings.append(_research_json_safe({
                    "type":("LIVE_APP_SAME_INPUT_AUDIT_V25108" if use_live_inputs else "WHOLE_WORKBOOK_IMPORT_AUDIT_V2544"),
                    "selected_cycle":f"{ry}-{rm:02d}",
                    "input_source":("LIVE_APP_DB" if use_live_inputs else "UPLOADED_WORKBOOK"),
                    "preferences":pref_audit,
                    "schedule":sched_audit,
                    "same_workbook":(
                        (not use_live_inputs) and bool(pref_audit and sched_audit)
                        and pref_audit.get("workbook_hash")==sched_audit.get("workbook_hash")
                    ),
                    "same_frozen_snapshot_for_both_schedules":True,
                }))
                snapshot=_research_people_snapshot(people)
                input_hash=_research_input_hash(ry,rm,people)
                schedule_hash=_research_schedule_hash(assignments)
                gpt_iters=_research_parse_optional_int(gpt_iter_text)
                gpt_minutes=_research_parse_optional_float(gpt_minutes_text)

                case=db.create_research_scheduler_case_v2541(
                    ry,rm,input_hash,schedule_hash,snapshot,
                    {str(k):v for k,v in assignments.items()},warnings,
                    gpt_iters,gpt_minutes,method_note,
                    APP_VERSION,ACTIVE_RULE_PROFILE_VERSION,
                )

                questionnaire_save_warnings=[]
                for initials,(extract,parsed) in questionnaire_payloads.items():
                    if not extract.get("ok"):
                        questionnaire_save_warnings.append(f"{initials}: questionnaire text was not extracted; not saved.")
                        continue
                    try:
                        db.save_research_scheduler_questionnaire_v2542(case["id"],initials,extract,parsed)
                    except Exception as qexc:
                        questionnaire_save_warnings.append(f"{initials}: {qexc}")

                t0=perf_counter()
                try:
                    algorithm=solve_schedule(ry,rm,people,time_limit=180.0)
                    elapsed=perf_counter()-t0
                    full_stats=_research_json_safe(algorithm.stats or {})
                    gg=full_stats.get("global",{}) if isinstance(full_stats,dict) else {}
                    run_assignments=algorithm.assignments
                    run_ok=algorithm.ok
                except Exception as solve_exc:
                    elapsed=perf_counter()-t0
                    full_stats={"global":{
                        "hard_errors":None,
                        "errors":[f"ENGINE EXCEPTION: {type(solve_exc).__name__}: {solve_exc}"],
                        "solve_stage":"ENGINE_EXCEPTION",
                    },"people":{}}
                    gg=full_stats["global"]
                    run_assignments={}
                    run_ok=False

                db.record_research_scheduler_run_v2541(
                    case["id"],input_hash,elapsed,run_ok,
                    APP_VERSION,ACTIVE_RULE_PROFILE_VERSION,gg,
                    run_assignments,raw_metrics=_research_json_safe(full_stats),
                )
                if questionnaire_save_warnings:
                    st.warning("Questionnaire audit: "+" | ".join(questionnaire_save_warnings))
                st.success("CASE LOCKED. RUN 1 / FIRST SHOT išsaugotas immutable DB įrašu.")
                st.rerun()
            except Exception as e:
                msg=str(e)
                if "RESEARCH_CASE_LOCKED_INPUT_MISMATCH" in msg:
                    st.error("Šis mėnuo jau užrakintas su kitu input/schedule snapshot. Official case negalima perrašyti.")
                else:
                    st.exception(e)
            return

        st.caption("Kol nepaspaudei LOCK CASE + RUN 1, gali saugiai keisti/taisytis importo failus.")
        return

    # ===== LOCKED CASE =====
    people=_research_people_from_snapshot(case.get("input_snapshot") or [])
    comparator=_research_locked_comparator(case,people)
    runs=db.list_research_scheduler_runs_v2541(case["id"])

    st.success(
        f"LOCKED CASE {ry}-{rm:02d} · input {str(case.get('input_hash'))[:12]}… · "
        f"GPT+Human schedule {str(case.get('comparator_schedule_hash'))[:12]}…"
    )
    st.caption(
        f"Frozen engine: {case.get('app_version_at_lock')} · Rule Profile v{case.get('rule_profile_version_at_lock')} · "
        f"locked {str(case.get('created_at') or '')[:19]}"
    )

    questionnaires=db.list_research_scheduler_questionnaires_v2542(case["id"])
    q_by={q.get("respondent_initials"):q for q in questionnaires}

    st.markdown("### MG / VL retrospective questionnaires")
    if questionnaires:
        st.dataframe(pd.DataFrame([{
            "Respondent":q.get("respondent_initials"),
            "File":q.get("file_name"),
            "SHA-256":str(q.get("file_hash") or "")[:16]+"…",
            "Parsed iterations":q.get("parsed_iterations"),
            "Parsed total time, min":q.get("parsed_minutes"),
            "Locked at":str(q.get("created_at") or "")[:19],
        } for q in questionnaires]),use_container_width=True,hide_index=True)

    for initials in ("MG","VL"):
        if initials in q_by:
            q=q_by[initials]
            payload=q.get("parser_payload") or {}
            with st.expander(f"{initials} questionnaire audit",expanded=False):
                st.caption(
                    f"{q.get('file_name')} · SHA-256 {q.get('file_hash')} · "
                    f"{q.get('file_size')} bytes"
                )
                st.write(
                    f"Parser: iterations={q.get('parsed_iterations')} · "
                    f"minutes={q.get('parsed_minutes')}"
                )
                if payload.get("iteration_evidence"):
                    st.markdown("**Iteration evidence**")
                    st.caption(str(payload.get("iteration_evidence")))
                if payload.get("time_evidence"):
                    st.markdown("**Time evidence**")
                    st.caption(str(payload.get("time_evidence")))
                if q.get("extracted_text"):
                    st.markdown("**Extracted questionnaire text**")
                    st.text_area(
                        f"{initials} extracted text",
                        value=str(q.get("extracted_text")),
                        height=220,disabled=True,
                        key=f"locked_questionnaire_text_{q.get('id')}",
                        label_visibility="collapsed"
                    )
        else:
            uploaded=st.file_uploader(
                f"{initials} klausimynas — papildyti locked case",
                type=["pdf","docx","xlsx","xls","csv","txt","md"],
                key=f"locked_questionnaire_upload_{case['id']}_{initials}"
            )
            if uploaded is not None:
                extract=_questionnaire_extract_text(uploaded)
                parsed=_questionnaire_parse_process_metrics(extract.get("text","")) if extract.get("ok") else {
                    "iterations":None,"minutes":None,
                    "iteration_evidence":None,"time_evidence":None,
                    "iteration_candidates":[],"time_candidates":[]
                }
                _questionnaire_render_result(initials,extract,parsed)
                if extract.get("ok") and st.button(
                    f"LOCK {initials} QUESTIONNAIRE TO CASE",
                    key=f"lock_questionnaire_{case['id']}_{initials}",
                    use_container_width=True
                ):
                    try:
                        db.save_research_scheduler_questionnaire_v2542(
                            case["id"],initials,extract,parsed
                        )
                        st.success(f"{initials} questionnaire immutable išsaugotas.")
                        st.rerun()
                    except Exception as qexc:
                        if "QUESTIONNAIRE_ALREADY_LOCKED" in str(qexc):
                            st.error(f"{initials} questionnaire jau užrakintas kitu failu.")
                        else:
                            st.error(str(qexc))

    if questionnaires:
        stored_parsed={
            q.get("respondent_initials"):{
                "iterations":q.get("parsed_iterations"),
                "minutes":q.get("parsed_minutes"),
            }
            for q in questionnaires
        }
        q_consensus=_questionnaire_consensus(stored_parsed)
        st.info(
            "Stored questionnaire consensus suggestion: "
            + ("iterations = N/A" if q_consensus["iterations"] is None else f"iterations ≈ {q_consensus['iterations']}")
            + " · "
            + ("time = N/A" if q_consensus["minutes"] is None else f"time ≈ {q_consensus['minutes']:g} min")
        )
        if (q_consensus["iterations"] is not None or q_consensus["minutes"] is not None):
            if st.button(
                "APPLY QUESTIONNAIRE CONSENSUS TO GPT+HUMAN METADATA",
                key=f"apply_q_consensus_{case['id']}",
                use_container_width=True
            ):
                db.update_research_scheduler_process_v2541(
                    case["id"],
                    q_consensus["iterations"] if q_consensus["iterations"] is not None else case.get("gpt_human_iterations"),
                    q_consensus["minutes"] if q_consensus["minutes"] is not None else case.get("gpt_human_minutes"),
                    str(case.get("method_note") or "")
                )
                st.success("Consensus perkeltas į GPT+human proceso metadata; questionnaire audit įrašai liko atskiri.")
                st.rerun()

    # Allow later enrichment ONLY of human-process metadata, never inputs or schedules.
    with st.expander("Papildyti GPT+human proceso duomenis (inputų/grafiko tai nekeičia)",expanded=False):
        meta1,meta2=st.columns(2)
        it_txt=meta1.text_input(
            "Perdarymų / iteracijų skaičius",
            value="" if case.get("gpt_human_iterations") is None else str(case.get("gpt_human_iterations")),
            key=f"update_iters_{case['id']}"
        )
        min_txt=meta2.text_input(
            "Apytikslis bendras laikas, min.",
            value="" if case.get("gpt_human_minutes") is None else str(case.get("gpt_human_minutes")),
            key=f"update_minutes_{case['id']}"
        )
        note_txt=st.text_area(
            "Proceso pastaba",
            value=str(case.get("method_note") or ""),
            key=f"update_method_{case['id']}"
        )
        if st.button("IŠSAUGOTI TIK PROCESO METADATA",key=f"save_process_{case['id']}"):
            try:
                db.update_research_scheduler_process_v2541(
                    case["id"],_research_parse_optional_int(it_txt),
                    _research_parse_optional_float(min_txt),note_txt
                )
                st.success("Proceso metadata atnaujinta; locked input/schedule nepakeisti.")
                st.rerun()
            except Exception as e:
                st.error(str(e))

    # Frozen run log.
    st.markdown("### Engine run log")
    if runs:
        logdf=_research_run_log_df(runs)
        st.dataframe(logdf,use_container_width=True,hide_index=True)
        st.caption(
            f"Engine clicks recorded: {len(runs)}/5 · total engine solve time: "
            f"{sum(float(r.get('elapsed_seconds') or 0) for r in runs):.2f} s."
        )
    else:
        st.error("Case exists but Run 1 is missing. Do not recreate the case; run the next engine attempt below.")

    # Engine version is frozen server-side by the case; block future run if local build changed.
    version_ok=(
        str(case.get("app_version_at_lock"))==APP_VERSION
        and int(case.get("rule_profile_version_at_lock") or 0)==int(ACTIVE_RULE_PROFILE_VERSION)
    )
    if not version_ok:
        st.error(
            "FROZEN ENGINE MISMATCH — šis case turi būti tęsiamas tik su "
            f"{case.get('app_version_at_lock')} / Rule Profile v{case.get('rule_profile_version_at_lock')}. "
            "Naujo run su kita versija nedaryk."
        )

    next_no=len(runs)+1
    if next_no<=5:
        role="FIRST SHOT" if next_no==1 else f"IMPROVEMENT ATTEMPT {next_no-1}"
        if st.button(
            f"RUN ENGINE #{next_no}/5 — {role}",
            type="primary" if next_no==1 else "secondary",
            use_container_width=True,disabled=not version_ok,
            key=f"research_next_run_{case['id']}_{next_no}"
        ):
            try:
                # Reconstruct locked input directly from DB; no re-upload/cherry-picking.
                locked_hash=_research_input_hash(ry,rm,people)
                if locked_hash!=case.get("input_hash"):
                    st.error("Locked input hash validation failed. Run aborted.")
                    return
                t0=perf_counter()
                try:
                    algorithm=solve_schedule(ry,rm,people,time_limit=180.0)
                    elapsed=perf_counter()-t0
                    full_stats=_research_json_safe(algorithm.stats or {})
                    gg=full_stats.get("global",{}) if isinstance(full_stats,dict) else {}
                    run_assignments=algorithm.assignments
                    run_ok=algorithm.ok
                except Exception as solve_exc:
                    elapsed=perf_counter()-t0
                    full_stats={"global":{
                        "hard_errors":None,
                        "errors":[f"ENGINE EXCEPTION: {type(solve_exc).__name__}: {solve_exc}"],
                        "solve_stage":"ENGINE_EXCEPTION",
                    },"people":{}}
                    gg=full_stats["global"]
                    run_assignments={}
                    run_ok=False

                saved=db.record_research_scheduler_run_v2541(
                    case["id"],locked_hash,elapsed,run_ok,
                    APP_VERSION,ACTIVE_RULE_PROFILE_VERSION,gg,
                    run_assignments,raw_metrics=_research_json_safe(full_stats),
                )
                st.success(f"RUN {saved.get('run_no',next_no)} immutable įrašytas.")
                st.rerun()
            except Exception as e:
                if "RESEARCH_RUN_LIMIT_REACHED" in str(e):
                    st.error("Max 5 engine runs jau pasiekti.")
                else:
                    st.exception(e)
    else:
        st.success("5/5 engine runs užbaigti. Tyrimo generation loop užrakintas.")

    # Primary comparison = immutable Run 1.
    first=next((r for r in runs if int(r.get("run_no") or 0)==1),None)
    valid_runs=[r for r in runs if bool(r.get("success")) and int(r.get("hard_errors") or 0)==0]
    best=min(valid_runs,key=_research_run_quality_key) if valid_runs else None

    cg=comparator.stats.get("global",{})
    st.markdown("### Comparator baseline — AVAILABLE GPT + HUMAN")
    b1,b2,b3,b4=st.columns(4)
    b1.metric("HARD errors *",cg.get("hard_errors"))
    b2.metric("Monthly fairness",cg.get("monthly_fairness_score"))
    b3.metric("Preference %",cg.get("mean_preference_score") if cg.get("mean_preference_score") is not None else "N/A")
    b4.metric("Workplace imbalance",cg.get("rotation_monthly_imbalance"))
    if cg.get("errors"):
        with st.expander(f"GPT+Human HARD findings ({len(cg.get('errors',[]))})",expanded=False):
            for e in cg.get("errors",[]):
                st.write("• "+_hard_error_explanation(e,lang))

    run1_result=(
        _research_result_from_run(first,people,ry,rm)
        if first and first.get("success") else None
    )
    best_result=(
        _research_result_from_run(best,people,ry,rm)
        if best else None
    )

    if first:
        if run1_result is not None:
            st.markdown("### PRIMARY — AVAILABLE GPT + HUMAN vs ENGINE FIRST SHOT")
            st.dataframe(
                _research_metric_rows(comparator,run1_result,"MY ENGINE — RUN 1"),
                use_container_width=True,hide_index=True
            )
            left,right=st.columns(2)
            with left:
                st.markdown("**AVAILABLE GPT + HUMAN — REAL USED SCHEDULE**")
                st.dataframe(schedule_list_df(ry,rm,comparator),use_container_width=True,hide_index=True,height=480)
            with right:
                st.markdown("**MY ENGINE — IMMUTABLE RUN 1 / FIRST SHOT**")
                st.dataframe(schedule_list_df(ry,rm,run1_result),use_container_width=True,hide_index=True,height=480)

            st.markdown("### SAME FROZEN WISHES — DIRECT OUTCOME COMPARISON")
            st.caption(
                "This is the apples-to-apples analysis: BOTH schedules are revalidated against the exact same frozen resident wishes/HARD snapshot. "
                "No wishes are inferred from the hand-made schedule."
            )
            _ct=_research_result_wish_totals(comparator); _et=_research_result_wish_totals(run1_result)
            w1,w2,w3,w4=st.columns(4)
            w1.metric("Active frozen wishes",_ct["active"])
            w2.metric("GPT+Human honored",f"{_ct['honored']}/{_ct['active']}" if _ct['active'] else "N/A")
            w3.metric("Engine honored",f"{_et['honored']}/{_et['active']}" if _et['active'] else "N/A")
            w4.metric("Cannot-work violations GPT / Engine",f"{_ct['hard_missed']} / {_et['hard_missed']}")
            st.dataframe(
                _research_wish_summary_comparison_df(comparator,run1_result,people),
                use_container_width=True,hide_index=True,height=520
            )
            _wishcmp=_research_wish_request_comparison_df(comparator,run1_result,people)
            _diff=_wishcmp[_wishcmp["Outcome"]!="Both honored"] if not _wishcmp.empty else _wishcmp
            if _diff.empty:
                st.success("ALL FROZEN WISHES WERE HONORED BY BOTH SCHEDULES.")
            else:
                st.markdown("**Missed / different wish outcomes**")
                st.dataframe(_diff,use_container_width=True,hide_index=True,height=520)
            with st.expander("Show every frozen wish request",expanded=False):
                st.dataframe(_wishcmp,use_container_width=True,hide_index=True,height=620)
        else:
            st.error(
                f"PRIMARY FIRST SHOT FAILED after {float(first.get('elapsed_seconds') or 0):.2f}s. "
                "Tai yra validus primary outcome ir negali būti pakeistas vėlesniu run."
            )

    if best_result is not None:
        st.markdown(f"### SECONDARY — BEST-OF-{len(runs)} = RUN {best['run_no']}")
        st.caption(
            "Best selector: HARD-valid → highest monthly fairness → lower workplace imbalance → "
            "higher preference fulfilment → higher diversity. FIRST SHOT lieka primary ir nėra pakeičiamas."
        )
        st.dataframe(
            _research_metric_rows(comparator,best_result,f"MY ENGINE — BEST RUN {best['run_no']}"),
            use_container_width=True,hide_index=True
        )

    st.markdown("### POST / WORKPLACE DISTRIBUTION")
    st.caption(
        "Čia matosi ne tik composite fairness %. Kiekvienam postui rodoma assignment supply, "
        "mažiausias ir didžiausias rezidento skaičius bei MAX−MIN spread. "
        "Neigiamas Δ spread reiškia, kad MY ENGINE paskirstė tą postą tolygiau už AVAILABLE GPT + HUMAN."
    )
    post_compare=_research_post_spread_comparison_df(
        comparator,people,
        run1_result=run1_result,
        best_result=best_result,
        best_label=("Best" if best is None else f"Best Run {best['run_no']}")
    )
    st.dataframe(post_compare,use_container_width=True,hide_index=True)

    with st.expander("Resident × post matrices",expanded=False):
        st.markdown("**AVAILABLE GPT + HUMAN**")
        st.dataframe(
            _research_post_matrix_df(comparator,people),
            use_container_width=True,hide_index=True
        )
        if run1_result is not None:
            st.markdown("**MY ENGINE — RUN 1 / FIRST SHOT**")
            st.dataframe(
                _research_post_matrix_df(run1_result,people),
                use_container_width=True,hide_index=True
            )
        if best_result is not None:
            st.markdown(f"**MY ENGINE — BEST RUN {best['run_no']}**")
            st.dataframe(
                _research_post_matrix_df(best_result,people),
                use_container_width=True,hide_index=True
            )

    warnings=case.get("import_warnings") or []
    if warnings:
        with st.expander(f"Import warnings ({len(warnings)})"):
            for w in warnings:
                st.write("• "+str(w))

    report=research_locked_comparison_xlsx(ry,rm,case,people,comparator,runs,questionnaires)
    st.download_button(
        "DOWNLOAD LOCKED RESEARCH DATASET (.xlsx)",
        report,
        file_name=f"available_gpt_human_vs_engine_{ry}_{rm:02d}.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        type="primary",use_container_width=True
    )
    st.caption(
        "Official interpretation: AVAILABLE GPT + HUMAN = real iterative general-GPT-assisted human workflow; "
        "ENGINE RUN 1 = primary first-shot performance; RUN 2–5 = secondary improvement/click/time efficiency."
    )


# --- ŠR-only Sudarymas: isolated research shadow / fake generator ---
if active_user==RESEARCHER_INITIALS and research_shadow_tab_index is not None:
    with tabs[research_shadow_tab_index]:
        render_research_shadow_generator()

# --- ŠR-only locked research comparison: general GPT+human vs specialized engine ---
if active_user==RESEARCHER_INITIALS and advanced_mode:
    with tabs[0]:
        render_available_gpt_vs_engine_research()

# --- Research ---
with tabs[pos]:
    st.subheader(tr("research_title"))
    st.caption(tr("research_privacy"))
    st.markdown(f"### {tr('research_study_plan')}")
    st.info(f"{tr('research_study_period')}  \n{tr('research_primary_outcomes')}")

    # Resident survey: exactly three planned checkpoints across the study.
    st.markdown(f"### {tr('research_survey')}")
    st.caption(tr("research_resident_note"))
    cp_codes=[c for c,_,_,_ in RESIDENT_RESEARCH_CHECKPOINTS]
    # Prefer baseline before launch, then month-3, then month-6. Users can still open any checkpoint for QA/testing.
    completed={}
    for c,ph,yy,mm in RESIDENT_RESEARCH_CHECKPOINTS:
        completed[c]=bool(db.get_my_research_survey(ph,yy,mm))
    default_cp=next((c for c in cp_codes if not completed[c]),cp_codes[-1])
    cp=st.selectbox(tr("research_checkpoint"),cp_codes,index=cp_codes.index(default_cp),format_func=research_checkpoint_label,key="research_checkpoint_select")
    phase,sy,sm=research_checkpoint_for_storage(cp)
    status_cols=st.columns(3)
    current_period=(year,month)
    activation={"baseline":(2026,9),"month3":(2026,12),"month6":(2027,3)}
    for i,c in enumerate(cp_codes):
        if completed[c]:
            status="✓ " + tr("research_checkpoint_done")
        elif current_period < activation[c]:
            status=tr("research_checkpoint_locked")
        else:
            status=tr("research_checkpoint_pending")
        status_cols[i].metric(research_checkpoint_label(c),status)

    existing=db.get_my_research_survey(phase,sy,sm) or {}
    old_answers=existing.get("answers") or {}
    old_free=existing.get("free_text") or {}
    st.caption(tr("research_likert_help"))
    item_labels=RESEARCH_ITEMS_LT if lang=="LT" else RESEARCH_ITEMS
    with st.form(f"research_form_{active_user}_{cp}"):
        answers={}
        for key,label in item_labels.items():
            default=int(old_answers.get(key,3)) if str(old_answers.get(key,"")).isdigit() else 3
            answers[key]=st.slider(label,1,5,default,1,key=f"rq_{cp}_{key}")
        answers["stress"]=st.slider(tr("research_stress"),0,10,int(old_answers.get("stress",5) or 5),1,key=f"rq_{cp}_stress")
        answers["change_count"]=st.selectbox(tr("research_changes"),[0,1,2,3,4],index=min(int(old_answers.get("change_count",0) or 0),4),format_func=lambda x:"4+" if x==4 else str(x),key=f"rq_{cp}_changes")
        contact_opts=["never","rarely","sometimes","often","very_often"]
        contact_labels={"LT":["Niekada","Retai","Kartais","Dažnai","Labai dažnai"],"EN":["Never","Rarely","Sometimes","Often","Very often"]}[lang]
        old_contact_code=int(old_answers.get("contact_frequency_code",0) or 0)
        old_contact_code=max(0,min(old_contact_code,4))
        contact=st.selectbox(tr("research_contact"),contact_opts,index=old_contact_code,format_func=lambda x:contact_labels[contact_opts.index(x)],key=f"rq_{cp}_contact")
        answers["contact_frequency_code"]=contact_opts.index(contact)
        if phase=="followup":
            for key,label_key in [("easy","research_easy"),("mobile","research_mobile"),("actual","research_actual"),("system_actual","research_system_actual"),("continue","research_continue")]:
                answers[key]=st.slider(tr(label_key),1,5,int(old_answers.get(key,3) or 3),1,key=f"rq_{cp}_{key}")
        problem=st.text_area(tr("research_problem"),value=old_free.get("problem","") or "",key=f"rq_{cp}_problem")
        improve=st.text_area(tr("research_improve"),value=old_free.get("improve","") or "",key=f"rq_{cp}_improve")
        submitted=st.form_submit_button(tr("research_submit"),type="primary")
        if submitted:
            db.submit_research_survey(phase,sy,sm,answers,{"problem":problem,"improve":improve,"checkpoint":cp})
            st.success(tr("research_saved"))

    # Current Seniūnė: workflow-burden checkpoints, separate from the resident survey.
    if active_user==SENIOR_INITIALS:
        st.divider(); st.markdown(f"### {tr('research_scheduler_section')}")
        st.caption(tr("research_scheduler_intro"))
        study_options=[f"{yy:04d}-{mm:02d}" for yy,mm in STUDY_MONTHS]
        current_key=f"{year:04d}-{month:02d}"
        default_idx=study_options.index(current_key) if current_key in study_options else 0
        month_key=st.selectbox(tr("research_scheduler_month"),study_options,index=default_idx,format_func=lambda x:study_month_label(int(x[:4]),int(x[5:])),key="senior_research_month")
        gy,gm=map(int,month_key.split("-"))
        checkpoint=st.radio(tr("research_scheduler_checkpoint"),["post_creation","post_month"],format_func=lambda x:tr("research_after_creation") if x=="post_creation" else tr("research_after_month"),horizontal=True,key="senior_research_checkpoint")
        old=db.get_my_scheduler_research_checkpoint(gy,gm,checkpoint) or {}
        oa=old.get("answers") or {}; of=old.get("free_text") or {}
        if checkpoint=="post_creation":
            methods=["tool","excel","shadow"]
            method_labels={"tool":tr("research_method_tool"),"excel":tr("research_method_excel"),"shadow":tr("research_method_shadow")}
            with st.form(f"senior_creation_{gy}_{gm}"):
                method=st.selectbox(tr("research_workflow_method"),methods,index=methods.index(oa.get("method","shadow") if oa.get("method","shadow") in methods else "shadow"),format_func=lambda x:method_labels[x])
                c1,c2=st.columns(2)
                total_minutes=c1.number_input(tr("research_total_minutes"),0,2000,int(oa.get("total_minutes",0) or 0),5)
                corrections=c2.number_input(tr("research_corrections"),0,200,int(oa.get("corrections",0) or 0),1)
                c3,c4=st.columns(2)
                contacts=c3.number_input(tr("research_resident_contacts"),0,500,int(oa.get("resident_contacts",0) or 0),1)
                comm_minutes=c4.number_input(tr("research_communication_minutes"),0,2000,int(oa.get("communication_minutes",0) or 0),5)
                stress=st.slider(tr("research_scheduler_stress"),0,10,int(oa.get("stress",5) or 5),1)
                q1,q2,q3=st.columns(3)
                fairness=q1.slider(tr("research_fairness_confidence"),1,5,int(oa.get("fairness_confidence",3) or 3),1)
                hard=q2.slider(tr("research_hard_confidence"),1,5,int(oa.get("hard_confidence",3) or 3),1)
                sat=q3.slider(tr("research_scheduler_satisfaction"),1,5,int(oa.get("satisfaction",3) or 3),1)
                excel_minutes=tool_minutes=excel_corr=tool_corr=0
                if method=="shadow":
                    a,b=st.columns(2); excel_minutes=a.number_input(tr("research_excel_minutes"),0,2000,int(oa.get("excel_minutes",0) or 0),5); tool_minutes=b.number_input(tr("research_tool_minutes"),0,2000,int(oa.get("tool_minutes",0) or 0),5)
                    a,b=st.columns(2); excel_corr=a.number_input(tr("research_excel_corrections"),0,200,int(oa.get("excel_corrections",0) or 0),1); tool_corr=b.number_input(tr("research_tool_corrections"),0,200,int(oa.get("tool_corrections",0) or 0),1)
                notes=st.text_area(tr("research_scheduler_notes"),value=of.get("notes","") or "")
                if st.form_submit_button(tr("research_submit"),type="primary"):
                    db.submit_scheduler_research_checkpoint(gy,gm,checkpoint,{"method":method,"total_minutes":total_minutes,"corrections":corrections,"resident_contacts":contacts,"communication_minutes":comm_minutes,"stress":stress,"fairness_confidence":fairness,"hard_confidence":hard,"satisfaction":sat,"excel_minutes":excel_minutes,"tool_minutes":tool_minutes,"excel_corrections":excel_corr,"tool_corrections":tool_corr},{"notes":notes})
                    st.success(tr("research_scheduler_saved"))
        else:
            with st.form(f"senior_month_{gy}_{gm}"):
                c1,c2=st.columns(2)
                post_minutes=c1.number_input(tr("research_post_minutes"),0,3000,int(oa.get("post_minutes",0) or 0),5)
                interventions=c2.number_input(tr("research_post_interventions"),0,500,int(oa.get("interventions",0) or 0),1)
                contacts=st.number_input(tr("research_post_contacts"),0,1000,int(oa.get("post_contacts",0) or 0),1)
                stress=st.slider(tr("research_stress"),0,10,int(oa.get("stress",5) or 5),1,key=f"senior_post_stress_{gy}_{gm}")
                q1,q2=st.columns(2)
                actual_conf=q1.slider(tr("research_actual_confidence"),1,5,int(oa.get("actual_confidence",3) or 3),1)
                sat=q2.slider(tr("research_scheduler_satisfaction"),1,5,int(oa.get("satisfaction",3) or 3),1,key=f"senior_post_sat_{gy}_{gm}")
                use_opts=["yes","unsure","no"]; use_labels={"yes":tr("research_yes"),"unsure":tr("research_unsure"),"no":tr("research_no")}
                old_use=oa.get("use_next","unsure") if oa.get("use_next","unsure") in use_opts else "unsure"
                use_next=st.selectbox(tr("research_use_next"),use_opts,index=use_opts.index(old_use),format_func=lambda x:use_labels[x])
                notes=st.text_area(tr("research_scheduler_notes"),value=of.get("notes","") or "",key=f"senior_post_notes_{gy}_{gm}")
                if st.form_submit_button(tr("research_submit"),type="primary"):
                    db.submit_scheduler_research_checkpoint(gy,gm,checkpoint,{"post_minutes":post_minutes,"interventions":interventions,"post_contacts":contacts,"stress":stress,"actual_confidence":actual_conf,"satisfaction":sat,"use_next":use_next},{"notes":notes})
                    st.success(tr("research_scheduler_saved"))
        # Show automatic operational counts for the same month to reduce manual counting burden.
        cp_current=db.load_schedule(gy,gm,"current"); cp_base=db.load_schedule(gy,gm,"baseline")
        if cp_current:
            cr=deserialize_result(cp_current); br=deserialize_result(cp_base or cp_current); gg=br.stats.get("global",{})
            try: changed=len(observer_assignment_changes_df(gy,gm,br,cr))
            except Exception: changed=0
            nsw=len(db.list_swap_requests(gy,gm,None)); bsw=len(db.list_backup_swap_requests(gy,gm,None)); covers=sum(1 for r in db.list_backups(gy,gm) if r.get("completed_at"))
            m1,m2,m3,m4=st.columns(4); m1.metric(tr("research_changed_assignments"),changed); m2.metric(tr("research_normal_swaps"),nsw); m3.metric(tr("research_backup_swaps"),bsw); m4.metric(tr("research_completed_covers"),covers)

    # Aggregate dashboard: current Seniūnė gets group-only results; ŠR gets full research QA / exports.
    if active_user in (RESEARCHER_INITIALS,SENIOR_INITIALS) and advanced_mode:
        st.divider(); st.markdown(f"### {tr('research_dashboard')}")
        st.info(tr("research_rs_note") if active_user==RESEARCHER_INITIALS else tr("research_gm_note"))
        counts=db.research_checkpoint_counts(); summary=db.research_checkpoint_summary()
        count_map={(int(r.get("cycle_year")),int(r.get("cycle_month"))):int(r.get("response_count",0)) for r in counts}
        c1,c2,c3=st.columns(3)
        c1.metric(tr("research_baseline_checkpoint"),f"{count_map.get((2026,9),0)}/{RESEARCH_EXPECTED_RESIDENTS}")
        c2.metric(tr("research_month3_checkpoint"),f"{count_map.get((2026,12),0)}/{RESEARCH_EXPECTED_RESIDENTS}")
        c3.metric(tr("research_month6_checkpoint"),f"{count_map.get((2027,3),0)}/{RESEARCH_EXPECTED_RESIDENTS}")
        if summary:
            sdf=pd.DataFrame(summary)
            label_map={(2026,9):("Pradinis vertinimas" if lang=="LT" else "Baseline"),(2026,12):("3 mėn." if lang=="LT" else "Month 3"),(2027,3):("6 mėn." if lang=="LT" else "Month 6")}
            sdf["Checkpoint"]=[label_map.get((int(y),int(m)),f"{y}-{int(m):02d}") for y,m in zip(sdf["cycle_year"],sdf["cycle_month"])]
            piv=sdf.pivot(index="question_key",columns="Checkpoint",values="mean_score").reset_index()
            st.markdown(f"### {tr('research_survey_results')}"); st.dataframe(piv,use_container_width=True,hide_index=True)
        else:
            st.caption(tr("research_no_data"))

        # Monthly objective metrics across the full six-month prospective period.
        monthly=[]
        for yy,mm in STUDY_MONTHS:
            curp=db.load_schedule(yy,mm,"current"); basep=db.load_schedule(yy,mm,"baseline")
            row={("Laikotarpis" if lang=="LT" else "Period"):f"{yy}-{mm:02d}",( "Paskelbtas" if lang=="LT" else "Published"):bool(curp)}
            if curp:
                # V2.5.49 retrospective satisfaction: SYSTEM stays frozen at publication,
                # while ACTUAL is recalculated against the SAME frozen ORIGINAL request set.
                # This lets the six-month study later measure whether resident-led swaps
                # improved or worsened realized request satisfaction without rewriting
                # algorithmic fairness history.
                cr=refresh_result_payload(curp,yy,mm,use_actual_backups=True) or deserialize_result(curp)
                br=refresh_result_payload(basep or curp,yy,mm,use_actual_backups=False) or deserialize_result(basep or curp)
                gg=br.stats.get("global",{})
                ag=cr.stats.get("global",{})
                published_sat=gg.get("mean_preference_score")
                actual_sat=ag.get("mean_preference_score")
                sat_delta=(round(float(actual_sat)-float(published_sat),1) if published_sat is not None and actual_sat is not None else None)
                sys_live=calculate_live_fairness_snapshot(yy,mm,br.assignments,people_initials=[p["initials"] for p in DEFAULT_PEOPLE],backup_assignments=[])["global"]
                act_live=calculate_live_fairness_snapshot(yy,mm,cr.assignments,people_initials=[p["initials"] for p in DEFAULT_PEOPLE],backup_assignments=db.list_backups(yy,mm))["global"]
                row.update({
                    ("Privalomų taisyklių klaidos" if lang=="LT" else "HARD"):gg.get("hard_errors"),
                    ("SYSTEM mėnesio fairness" if lang=="LT" else "SYSTEM monthly fairness"):sys_live.get("monthly_fairness_score"),
                    ("ACTUAL mėnesio fairness" if lang=="LT" else "ACTUAL monthly fairness"):act_live.get("monthly_fairness_score"),
                    ("ACTUAL−SYSTEM fairness, p.p." if lang=="LT" else "ACTUAL−SYSTEM fairness, pp"):round(float(act_live.get("monthly_fairness_score",0))-float(sys_live.get("monthly_fairness_score",0)),1),
                    ("SYSTEM postų imbalance" if lang=="LT" else "SYSTEM post imbalance"):sys_live.get("rotation_monthly_imbalance"),
                    ("ACTUAL postų imbalance" if lang=="LT" else "ACTUAL post imbalance"):act_live.get("rotation_monthly_imbalance"),
                    ("SYSTEM pageidavimų išpildymas %" if lang=="LT" else "SYSTEM request satisfaction %"):published_sat,
                    ("ACTUAL pageidavimų išpildymas %" if lang=="LT" else "ACTUAL request satisfaction %"):actual_sat,
                    ("Pokytis po apsikeitimų, proc. p." if lang=="LT" else "Change after swaps, pp"):sat_delta,
                    ("SYSTEM RESIDENT HARD pažeidimai" if lang=="LT" else "SYSTEM RESIDENT HARD violations"):gg.get("resident_hard_total_losses",0),
                    ("ACTUAL RESIDENT HARD pažeidimai" if lang=="LT" else "ACTUAL RESIDENT HARD violations"):ag.get("resident_hard_total_losses",0),
                })
                try: row[("Sistemos→faktinio grafiko pakeitimai" if lang=="LT" else "SYSTEM→ACTUAL changes")]=len(observer_assignment_changes_df(yy,mm,br,cr))
                except Exception: row[("Sistemos→faktinio grafiko pakeitimai" if lang=="LT" else "SYSTEM→ACTUAL changes")]=0
                row["Normal swaps"]=len(db.list_swap_requests(yy,mm,None)); row["Backup swaps"]=len(db.list_backup_swap_requests(yy,mm,None)); row["Completed covers"]=sum(1 for r in db.list_backups(yy,mm) if r.get("completed_at"))
            monthly.append(row)
        monthly_df=pd.DataFrame(monthly)
        st.markdown(f"### {tr('research_monthly_table')}"); st.dataframe(monthly_df,use_container_width=True,hide_index=True)

        gm_rows=db.research_scheduler_dashboard()
        if gm_rows:
            st.markdown(f"### {tr('research_scheduler_status')}")
            st.dataframe(pd.DataFrame(gm_rows),use_container_width=True,hide_index=True)
        gen_rows=db.research_generation_dashboard()
        if gen_rows:
            st.markdown(f"### {tr('research_generation_telemetry')}")
            st.dataframe(pd.DataFrame(gen_rows),use_container_width=True,hide_index=True)

        if active_user=="ŠR":
            st.caption(tr("research_researcher_only"))
            raw=db.research_survey_deidentified()
            if raw:
                rows=[]
                for r in raw:
                    flat={"code":r.get("response_code"),"phase":r.get("phase"),"year":r.get("cycle_year"),"month":r.get("cycle_month"),"submitted_at":r.get("submitted_at")}
                    flat.update(r.get("answers") or {})
                    for k,v in (r.get("free_text") or {}).items(): flat[f"comment_{k}"]=v
                    rows.append(flat)
                raw_df=pd.DataFrame(rows)
                st.markdown(f"### {tr('research_deidentified')}"); st.dataframe(raw_df,use_container_width=True,hide_index=True)
                st.download_button(tr("research_download_surveys"),raw_df.to_csv(index=False).encode("utf-8-sig"),file_name="research_surveys_deidentified.csv",mime="text/csv")
            st.download_button(tr("research_download_monthly"),monthly_df.to_csv(index=False).encode("utf-8-sig"),file_name="research_monthly_metrics.csv",mime="text/csv")
            comments=db.research_comments_v2510()
            if comments:
                st.markdown(f"### {tr('research_comments')}"); st.dataframe(pd.DataFrame(comments),use_container_width=True,hide_index=True)
            obs=db.research_observer_dashboard()
            if obs:
                st.markdown(f"### {tr('research_observer_tab')}"); st.dataframe(pd.DataFrame(obs),use_container_width=True,hide_index=True)
pos+=1

# --- Proof ---
if advanced_mode:
    with tabs[pos]:
        st.subheader(tr("proof_title")); st.write(tr("proof_intro")); currentp=db.load_schedule(year,month,"current"); basep=db.load_schedule(year,month,"baseline")
        if not currentp:
            st.info(tr("not_published"))
        else:
            current=refresh_result_payload(currentp,year,month); base=refresh_result_payload(basep or currentp,year,month,use_actual_backups=False); pref=db.get_preference(year,month,active_user) or {}; slots=make_slots(year,month)
            def worked_days(res): return {s.day for s in slots if res.assignments.get(s.idx)==active_user}
            cdays=worked_days(current)
            hard=set(pref.get("unavailable",set()))
            hard_am=set(pref.get("unavailable_am",set()))
            hard_pm=set(pref.get("unavailable_pm",set()))
            soft=set(pref.get("soft_free",set())); soft_am=set(pref.get("soft_free_am",set())); soft_pm=set(pref.get("soft_free_pm",set()))
            wanted=set(pref.get("preferred",set())); wanted_am=set(pref.get("preferred_am",set())); wanted_pm=set(pref.get("preferred_pm",set()))

            current_person_slots=[s for s in slots if current.assignments.get(s.idx)==active_user]
            hard_bad=[]
            for d in sorted(hard):
                if any(s.day==d for s in current_person_slots):
                    hard_bad.append(f"{d} · {tr('full_day')}")
            for d in sorted(hard_am):
                if any(s.day==d and blocks_overlap(s.block,"AM") for s in current_person_slots):
                    hard_bad.append(f"{d} · {tr('morning')}")
            for d in sorted(hard_pm):
                if any(s.day==d and blocks_overlap(s.block,"PM") for s in current_person_slots):
                    hard_bad.append(f"{d} · {tr('afternoon')}")
            hard_total=len(hard)+len(hard_am)+len(hard_pm)
            def req_label(day,block):
                label=tr("full_day") if block=="FULL" else tr("morning") if block=="AM" else tr("afternoon")
                return f"{day} · {label}"
            def has_overlap(day,block):
                if block=="FULL": return any(s.day==day for s in current_person_slots)
                return any(s.day==day and blocks_overlap(s.block,block) for s in current_person_slots)
            soft_requests=[(d,"FULL") for d in soft]+[(d,"AM") for d in soft_am]+[(d,"PM") for d in soft_pm]
            pref_requests=[(d,"FULL") for d in wanted]+[(d,"AM") for d in wanted_am]+[(d,"PM") for d in wanted_pm]
            soft_miss=[req_label(d,b) for d,b in sorted(soft_requests) if has_overlap(d,b)]
            pref_miss=[req_label(d,b) for d,b in sorted(pref_requests) if not has_overlap(d,b)]
            bd=base.stats["people"].get(active_user,{}); cd=current.stats["people"].get(active_user,{})
            fair=base.stats["global"].get("fairness_score"); rr=balance_ratio(cd.get("preference_score"),fair)

            # Top visual summary.
            a,b,c,d=st.columns(4)
            a.metric(tr("hard_errors"),len(hard_bad))
            b.metric(tr("workload_ok"),f"{cd.get('workload')} / {cd.get('target')}")
            c.metric(tr("preference_score"),tr("not_applicable") if cd.get("preference_score") is None else f"{cd.get('preference_score')}%")
            d.metric(tr("balance_ratio"),tr("not_applicable") if rr is None else f"{rr:.2f}")

            comps=cd.get("preference_components",{})
            soft_problem=bool(soft_miss or pref_miss or any(v<80 for v in comps.values()))
            if hard_bad:
                st.error(tr("proof_hard_issue"))
            elif soft_problem:
                st.warning(tr("proof_soft_issues"))
            else:
                st.success(tr("proof_all_good"))

            if advanced_mode:
                rows=[]
                hard_score=100.0 if not hard_bad else max(0.0,100.0*(hard_total-len(hard_bad))/max(1,hard_total))
                rows.append({tr("criterion"):tr("hard_ok"),tr("result"):component_status(hard_score),tr("score"):f"{hard_score:.0f}%",tr("explanation"):("—" if not hard_bad else f"{tr('missed_dates')}: {', '.join(map(str,hard_bad))}")})
                if soft_requests:
                    score=round(100*(len(soft_requests)-len(soft_miss))/len(soft_requests),1)
                    rows.append({tr("criterion"):tr("soft_off_ok"),tr("result"):component_status(score),tr("score"):f"{score}%",tr("explanation"):("—" if not soft_miss else f"{tr('missed_dates')}: {', '.join(map(str,soft_miss))}")})
                if pref_requests:
                    score=round(100*(len(pref_requests)-len(pref_miss))/len(pref_requests),1)
                    rows.append({tr("criterion"):tr("preferred_ok"),tr("result"):component_status(score),tr("score"):f"{score}%",tr("explanation"):("—" if not pref_miss else f"{tr('missed_dates')}: {', '.join(map(str,pref_miss))}")})
                wl_credit=float(cd.get("workload_credit",cd.get("workload",0)) or 0)
                wl_target=float(cd.get("target",0) or 0)
                wl_ok=abs(wl_credit-wl_target)<1e-9
                rows.append({tr("criterion"):tr("workload_ok"),tr("result"):tr("matches") if wl_ok else tr("mismatch"),tr("score"):"100%" if wl_ok else "0%",tr("explanation"):(f"{wl_credit:g} / {wl_target:g} · SYSTEM krūvio kreditas užšaldytas publikavimo metu" if lang=="LT" else f"{wl_credit:g} / {wl_target:g} · SYSTEM workload credit frozen at publication")})
                labels={"weekday_preference":tr("weekday_pref"),"weekend_preference":tr("weekend_pref"),"spread_preference":tr("spread_pref"),"avoid_doubles":tr("avoid_double_shifts")}
                for key,label in labels.items():
                    if key in comps:
                        val=float(comps[key]); rows.append({tr("criterion"):label,tr("result"):component_status(val),tr("score"):f"{val:.1f}%",tr("explanation"):"—"})
                st.dataframe(pd.DataFrame(rows),use_container_width=True,hide_index=True)

                # Native progress visuals for active soft components.
                active_progress=[]
                if soft_requests: active_progress.append((tr("soft_off_ok"),100*(len(soft_requests)-len(soft_miss))/len(soft_requests)))
                if pref_requests: active_progress.append((tr("preferred_ok"),100*(len(pref_requests)-len(pref_miss))/len(pref_requests)))
                for key,label in labels.items():
                    if key in comps: active_progress.append((label,float(comps[key])))
                if active_progress:
                    st.divider()
                    for label,val in active_progress:
                        st.write(f"{label}: {val:.1f}%")
                        st.progress(max(0.0,min(1.0,val/100.0)))

                st.divider(); x,y,z=st.columns(3)
                x.metric(tr("baseline"),tr("not_applicable") if bd.get("preference_score") is None else f"{bd.get('preference_score')}%")
                y.metric(tr("current"),tr("not_applicable") if cd.get("preference_score") is None else f"{cd.get('preference_score')}%")
                z.metric(tr("balance_ratio"),tr("not_applicable") if rr is None else f"{rr:.2f}")
                if hard_bad or soft_problem: st.info(tr("swap_suggestion"))
            else:
                if hard_bad or soft_problem:
                    st.info(tr("swap_suggestion"))
                st.caption(
                    ("Detali kriterijų lentelė ir komponentų progresas rodomi Išplėstiniame režime."
                     if lang=="LT" else
                     "Detailed criterion tables and component progress are shown in Advanced mode.")
                )
    pos+=1

# --- Senior usability / audit guide ---
if senior_mode:
    with tabs[pos]:
        st.subheader(tr("senior_guide"))
        if lang=="LT":
            st.info(
                "Tikslas nėra aklai pasitikėti algoritmu. Tikslas — vietoje viso grafiko konstravimo ranka gauti "
                "stiprų juodraštį su įrodymais ir seniūnei palikti trumpą, kryptingą išimčių auditą."
            )
            st.markdown("### 5 minučių audito protokolas")
            audit_rows=[
                {"Žingsnis":"1","Kur žiūrėti":"HARD / diagnostics","Ką patikrinti":"0 TRUE ABSOLUTE HARD klaidų; jei RESIDENT HARD prarastas — turi būti aiškiai nurodyta kas, kada ir kodėl.","Jei blogai":"Nepublikuoti."},
                {"Žingsnis":"2","Kur žiūrėti":"Post matrix","Ką patikrinti":"SPS UG ir nesavanoriško SPS RO / šeštadienių / sekmadienių fairness spread 0–1; RAW savaitgalių spread gali būti didesnis, jei žmogus pats aiškiai pageidavo tas pamainas. Kiti postai pagal einamojo mėnesio struktūrinį water-fill guardrail.","Jei blogai":"Nepublikuoti arba aiškiai patikrinti, ar nukrypimas matematiškai neišvengiamas."},
                {"Žingsnis":"3","Kur žiūrėti":"Resident stats","Ką patikrinti":"Nėra vieno žmogaus su neproporcingu savaitiniu krūviu; generatoriaus max rolling-7 ir doubles/recovery rodikliai logiški.","Jei blogai":"Regeneruoti / taisyti prieš publikavimą."},
                {"Žingsnis":"4","Kur žiūrėti":"Išplėstinis / Patikra","Ką patikrinti":"Pasirinkti 3–5 rezidentus, ypač mažiausio ir didžiausio pageidavimų išpildymo, ir ranka patikrinti 1–2 konkrečius įrankio teiginius kiekvienam.","Jei blogai":"Jei teiginys nesutampa su SYSTEM grafiku, laikyti tai metrikos / programos klaida ir nepublikuoti, kol ištaisyta."},
                {"Žingsnis":"5","Kur žiūrėti":"Grafikas + Proof","Ką patikrinti":"Coverage, nepaaiškintos skylės, akivaizdūs overlap'ai ir ar galutinis grafikas atitinka tai, ką rodo suvestinės.","Jei blogai":"Nepublikuoti."},
            ]
            st.dataframe(pd.DataFrame(audit_rows),use_container_width=True,hide_index=True)
            st.markdown("### Įrankio teiginių patikra — paprastai, be žargono")
            st.info(
                "**Teiginys** = vienas konkretus sakinys, kurį įrankis sako apie grafiką. Pvz.: "
                "„ŠR 18 d. PM prašė laisvos, bet SYSTEM grafike paskirtas SPS UG PM, todėl pageidavimas neįvykdytas.“ "
                "Patikra reiškia tik viena: atsidaryti 18 d. ir pažiūrėti, ar tas SPS UG PM tikrai yra. "
                "Nereikia iš naujo perskaičiuoti viso mėnesio."
            )
            st.markdown(
                "**Ką tikrinti:** 1) ko žmogus prašė; 2) ką SYSTEM grafikas jam realiai paskyrė; "
                "3) ar iš to logiškai seka „įvykdyta / neįvykdyta“; 4) ar Post Matrix skaičiai sutampa su grafiku."
            )

            # V2.5.60: interactive senior verification card. Load its own schedule
            # state so this tab is safe even when other tabs did not define base/current.
            _guide_currentp=db.load_schedule(year,month,"current")
            _guide_basep=db.load_schedule(year,month,"baseline")
            if _guide_currentp:
                _guide_system=refresh_result_payload(_guide_basep or _guide_currentp,year,month,use_actual_backups=False)
                _guide_people=[p["initials"] for p in DEFAULT_PEOPLE]
                _guide_person=st.selectbox("Pasirink rezidentą teiginių patikrai",_guide_people,key=f"senior_verify_person_{year}_{month}")
                _guide_pd=(_guide_system.stats.get("people",{}).get(_guide_person,{}) or {})
                _gh,_gs,_go,_gw=st.columns(4)
                _gh.metric("RESIDENT HARD",tr("not_applicable") if _guide_pd.get("resident_hard_score") is None else f"{_guide_pd.get('resident_hard_score')}%")
                _gs.metric("SOFT",tr("not_applicable") if _guide_pd.get("soft_preference_score") is None else f"{_guide_pd.get('soft_preference_score')}%")
                _go.metric("Visi prašymai",tr("not_applicable") if _guide_pd.get("overall_request_score") is None else f"{_guide_pd.get('overall_request_score')}%")
                _gw.metric("SPS UG",int((_guide_pd.get("rotation_counts") or {}).get("SPS UG",0) or 0))
                _guide_rows=(_guide_pd.get("resident_hard_conflicts") or [])+(_guide_pd.get("soft_request_misses") or [])
                if _guide_rows:
                    st.markdown("#### Ką konkrečiai įrankis teigia apie šį rezidentą")
                    st.dataframe(request_details_df(_guide_rows,_guide_person),use_container_width=True,hide_index=True)
                else:
                    st.success("Šiam rezidentui struktūruotų neįvykdytų RESIDENT HARD / SOFT prašymų nėra.")
                with st.expander("Rodyti ir įvykdytus prašymus — kontroliniam spot-check",expanded=False):
                    _guide_honored=_guide_pd.get("honored_request_details") or []
                    if _guide_honored:
                        st.dataframe(request_details_df(_guide_honored,_guide_person),use_container_width=True,hide_index=True)
                    else:
                        st.caption("Nėra į score įtrauktų struktūruotų prašymų.")
                st.caption("Postų skaičius tikrink Post Matrix / SYSTEM grafike. Pvz. jei rodoma SPS UG = 2, turi rasti lygiai du SYSTEM SPS UG paskyrimus šiam rezidentui.")
            else:
                st.caption("Interaktyvi teiginių patikros kortelė atsiras, kai šiam mėnesiui bus paskelbtas grafikas.")

            st.success(
                "Principas: **ne skaityti visą grafiką nuo nulio, o tikrinti konkrečius įrankio teiginius ir išimtis.** "
                "Jei sakinys nesutampa su pačiu SYSTEM grafiku, tai ne „nuomonės skirtumas“, o aiški metrikos / programos klaida."
            )
            st.markdown("### Kada grafiko NEPUBLIKUOTI")
            st.dataframe(pd.DataFrame([
                {"Raudona vėliava":"TRUE ABSOLUTE HARD klaida","Veiksmas":"Blokuoti publikavimą."},
                {"Raudona vėliava":"Trūksta privalomo SPS RO / SPS UG / savaitgalio coverage","Veiksmas":"Blokuoti publikavimą."},
                {"Raudona vėliava":"Kritinis nesavanoriško SPS RO / SPS UG / savaitgalių fairness spread >1 be aiškios diagnostikos","Veiksmas":"Regeneruoti / tirti."},
                {"Raudona vėliava":"Bet koks RESIDENT HARD pažeidimas SYSTEM juodraštyje","Veiksmas":"Regeneruoti / tirti."},
                {"Raudona vėliava":"Teiginys apie pageidavimą / postą / krūvį nesutampa su pačiu SYSTEM grafiku","Veiksmas":"Laikyti metrikos klaida; nepublikuoti, kol išspręsta."},
                {"Raudona vėliava":"Importuoti pageidavimai akivaizdžiai nepilni / rodo N/A, nors žmogus pateikė prašymus","Veiksmas":"Taisyti importą prieš solve."},
            ]),use_container_width=True,hide_index=True)
            st.markdown("### Kas po publikavimo yra normalu")
            st.markdown(
                "- **Swapai** keičia ACTUAL grafiką; SYSTEM fairness baseline lieka užšaldytas.\n"
                "- **Liga / neatvykimas** gali perkelti jau dirbantį žmogų iš optional posto į SPS; SYSTEM baseline dėl to nesikeičia, tačiau ACTUAL postų/fairness statistika perskaičiuojama pagal realų darbą.\n"
                "- **SOFT neįvykdymas** savaime nėra klaida, jei aukštesnio rango taisyklės ir horizontalus water-filling paaiškina rezultatą.\n"
                "- **Ordinary posto spread iki guardrail** gali būti sąmoningas kompromisas su FUTURE CATCH-UP kompensacija kitais mėnesiais."
            )
            st.markdown("### Seniūnės darbo eiga — nuo nulio iki mėnesio uždarymo")
            workflow=pd.DataFrame([
                {"Etapas":"1. Paruošti mėnesį","Seniūnės veiksmas":"Patikrinti aktyvų taisyklių profilį, šventes, etatus/targetus ir administracinius uždarymus.","Įrankio darbas":"Sukuria aktualią slotų ir taisyklių erdvę."},
                {"Etapas":"2. Surinkti pageidavimus","Seniūnės veiksmas":"Stebėti, kas nepateikė; peržiūrėti tik neaiškius / konfliktinius įrašus.","Įrankio darbas":"Normalizuoja ABSOLUTE / RESIDENT HARD / SOFT ir ilgalaikius nustatymus."},
                {"Etapas":"3. Sugeneruoti","Seniūnės veiksmas":"Spausti Generate; nebandyti iš anksto ranka konstruoti viso grafiko.","Įrankio darbas":"Sprendžia coverage, critical water-fill, Resident-HARD, krūvį, postus ir SOFT."},
                {"Etapas":"4. Audituoti","Seniūnės veiksmas":"5 min. red-flag + konkrečių teiginių spot-check; tikrinti išimtis, ne kiekvieną langelį.","Įrankio darbas":"Rodo diagnostics, post matrix, resident stats, satisfaction ledger ir misses."},
                {"Etapas":"5. Koreguoti","Seniūnės veiksmas":"Tik jei auditas randa realią klaidą arba administracinę išimtį.","Įrankio darbas":"Regenerate / manual correction su pakartotine validacija."},
                {"Etapas":"6. Publikuoti","Seniūnės veiksmas":"Publikuoti tik po Proof patikros.","Įrankio darbas":"Užšaldo SYSTEM fairness baseline ir išsiunčia grafiką."},
                {"Etapas":"7. Eksploatuoti","Seniūnės veiksmas":"Neperimti kiekvieno swapo į asmenines žinutes; naudoti platformos swap/repair srautą.","Įrankio darbas":"ACTUAL grafikas keičiasi, SYSTEM lieka tyrimo/fairness baseline."},
                {"Etapas":"8. Uždaryti mėnesį","Seniūnės veiksmas":"Peržiūrėti SYSTEM vs ACTUAL ir eksportuoti tyrimo duomenis.","Įrankio darbas":"Palieka audit trail, satisfaction, post/debt ir pakeitimų istoriją."},
            ])
            st.dataframe(workflow,use_container_width=True,hide_index=True)
            with st.expander("Pilnas seniūnės naudojimo ir audito vadovas"):
                st.markdown(SENIOR_GUIDE_LT)
        else:
            st.info(
                "The aim is not blind trust in an algorithm. The aim is to replace full manual construction with a strong draft, "
                "explicit evidence and a short exception-focused human audit."
            )
            st.markdown("### Five-minute audit protocol")
            audit_rows=[
                {"Step":"1","Where":"HARD / diagnostics","Verify":"Zero TRUE ABSOLUTE HARD errors and zero Resident-HARD / Unavailable violations.","If failed":"Do not publish."},
                {"Step":"2","Where":"Post matrix","Verify":"SPS RO, SPS UG and weekend spread 0–1; other posts remain within the current-month structural water-fill guardrails.","If failed":"Do not publish unless the deviation is explicitly proven unavoidable."},
                {"Step":"3","Where":"Resident stats","Verify":"No disproportionate weekly load; rolling-7 and double/recovery metrics are plausible.","If failed":"Regenerate / correct before publication."},
                {"Step":"4","Where":"Advanced / Proof","Verify":"Spot-check 3–5 residents, including the lowest and highest satisfaction, against the actual grid.","If failed":"Treat the tool statement as a metric defect until corrected."},
                {"Step":"5","Where":"Schedule + Proof","Verify":"Coverage, gaps, overlaps and that summary claims match the schedule.","If failed":"Do not publish."},
            ]
            st.dataframe(pd.DataFrame(audit_rows),use_container_width=True,hide_index=True)
            st.success("Principle: verify concrete tool statements against the grid instead of rebuilding the entire month by hand.")
            with st.expander("Full senior usability and audit guide"):
                st.markdown(SENIOR_GUIDE_EN)
    pos+=1

def explanatory_manual_only(content: str) -> str:
    """Hide obsolete version-locked operational sections.

    Operational rules are rendered live from the active Rule Profile above.
    The manual below is explanatory context only and must not duplicate mutable
    engine parameters.
    """
    locked_prefixes=(
        "## Dubliai — V2.5.2 LOCKED",
        "## Savaitgalio dublių savitarna — V2.5.3 LOCKED",
        "## LOCKED V2.5.28",
        "## LOCKED V2.5.29",
        "## LOCKED V2.5.32",
    )
    lines=str(content or "").splitlines()
    out=[]
    skipping=False
    for line in lines:
        if any(line.startswith(p) for p in locked_prefixes):
            skipping=True
            continue
        if skipping and line.startswith("## "):
            skipping=False
        if not skipping:
            out.append(line)
    return "\n".join(out).strip()


# --- Rules ---
with tabs[pos]:
    st.subheader(tr("rules_title"))
    st.info(
        ("LOCKED V2.5.40 — PADENGIMAS IR SKYLĖS. Mandatory administraciniai postai lieka SPS RO d.d., SPS UG ir savaitgalio SPS RO budėjimai. "
         "Kiti neuždaryti postai gali likti neužpildyti tik tiek, kiek matematiškai reikia tiksliam mėnesio targetui išlaikyti. "
         "Gap dienos pirmiausia tolygiai išdėstomos per visą mėnesį; vieną kalendorinę dieną gali būti daugiausia 1 reali skylė, įskaitant leidžiamą Onko išimtį. "
         "Konkretų optional postą tą dieną parenka solveris, o gap'ų pasiskirstymo tarp postų grupių spread negali viršyti 2. "
         "Aiškios uždarymo išimtys, pvz. Mamografijos PM penktadieniais, nėra skylės ir lieka uždarytos."
         if lang=="LT" else
         "LOCKED V2.5.40 — COVERAGE AND GAPS. Mandatory administrative posts remain SPS RO d.d., SPS UG and weekend SPS RO duty. "
         "Other open posts may remain unfilled only as mathematically required by the exact monthly target. "
         "Gap dates are first spread evenly across the month; at most one real gap is allowed per calendar day, including the Onko exception. "
         "The solver chooses the concrete optional row on each gap date, while workplace gap-count spread may not exceed 2. Explicit closures are not gaps.")
    )

    if lang=="LT":
        st.success(
            "V2.5.65 — PIRMA DARBO DIENOS, TADA DARBO VIETOS. Pirmiausia parenkama, kuriomis dienomis ir kuriuo metu kiekvienas rezidentas dirba, kuo labiau saugant „Negaliu dirbti“, poilsį ir kitus pageidavimus. Tik tada parenkamos konkrečios darbo vietos. Kai tam tikro posto per mėnesį yra nedaug, bet jų užtenka bent po vieną kiekvienam, sistema pirmiausia stengiasi kiekvienam duoti bent vieną kartą, o tik tada skiria antrą. Todėl, pavyzdžiui, 22 Onko / Centro UG / Vaikų UG vietos 16 rezidentų normaliai turi pasiskirstyti po 1–2, o ne 0–2. SPS UG ir nesavanoriškas SPS RO / šeštadienių / sekmadienių krūvis laikomas kuo lygesnis; aiškiai savanoriškai „Pageidauju dirbti“ pasirinkti savaitgaliai gali RAW skaičių padidinti. Konkreti SPS data nėra užrakinama žmogui, jei tą patį bendrą kiekį galima išlaikyti kitu paskyrimu ir geriau įvykdyti jo pageidavimą. Patvirtintos atostogos yra privalomos nedarbo dienos ir proporcingai sumažina to žmogaus mėnesio darbo tikslą. "
            "V2.5.74 — VISŲ POSTŲ STRUKTŪRINIS WATER-FILL. Prieš atiduodama SYSTEM grafiką sistema pirmiausia užrakina darbo datas/blokus, tada VISUS ne-Onko postų labelius sprendžia kartu. SPS UG ir kiekvienas įprastas ne-Onko postas pirmiausia bandomi 0–1 koridoriuje; savaitgaliams ir savaitgalio SPS RO 0–1 taikomas likusiam nesavanoriškam krūviui, o aiškiai savanoriškos „Pageidauju dirbti“ pamainos gali RAW spread padidinti. Platesnis 0–2 ar 0–3 koridorius leidžiamas tik po matematinio įrodymo, kad siauresnis koridorius neįmanomas. Timeout nėra įrodymas. Tokiu būdu 1-vs-3 negali likti, jei validus dviejų ar kelių rezidentų postų perkeitimas gali padaryti 2-vs-2 nepakeičiant darbo datos/bloko. "
            "Generatorius lieka konservatyvus: ~40 val./7 d. tikslas, ≤48 val./7 d., ≤6 darbo dienos/7 d. ir recovery po dvigubų. "
            "Po publikavimo bilateral voluntary swapas tikrina ABSOLUTE / operacinius ir darbo-laiko blokatorius, o naują 12 h double, >40/>48 h krūvį, "
            "6 dienų seką, post-double recovery ar savo RESIDENT HARD override parodo pasekmių lentelėje ir prašo paveikto rezidento ACK. "
            "ACK niekada neapeina >12 h/d., <11 h poilsio, >6 darbo dienų/7 d., >60 h/7 d., pateisinamo neatvykimo, overlap/coverage ar privalomo poilsio po 24 h budėjimo."
        )
        st.markdown("### Kaip derinama darbo vietų lygybė ir asmeniniai pageidavimai")
        st.info(
            "Sistema saugo ne konkretaus žmogaus SPS datą, o bendrą mėnesio paskirstymo lygumą. "
            "Pavyzdžiui, jei rezidentas nori 18 d. PM laisvos, jo SPS UG pamaina gali būti perkelta kitam tinkamam rezidentui arba jam parinkta kita SPS UG data, "
            "jeigu galutinis kiekis tarp žmonių išlieka kuo lygesnis ir nepažeidžiamos svarbesnės taisyklės. "
            "Todėl vien faktas, kad SPS yra aukšto prioriteto kategorija, nereiškia, kad konkretus SOFT pageidavimas turi būti atmestas."
        )
        st.markdown("### Supaprastinta vertikali prioritetų lentelė")
        st.dataframe(pd.DataFrame([
            {"Rangas":"1. TRUE ABSOLUTE HARD","Kas įeina":"Sauga, patvirtinta liga/atostogos, fizinis neįmanomumas, coverage; generatoriui ≤48 val./7 d. ir bent 1 laisva diena/7 d.","Kaip sprendžiama":"Generuojant 100%. Po publikavimo tik 48h riba gali būti savanoriškai viršyta normaliu bilateral swapu su aiškiu asmens sutikimu; kiti HARD lieka"},
            {"Rangas":"2. CRITICAL STRUCTURAL","Kas įeina":"SPS RO + SPS UG + šeštadieniai + sekmadieniai + penktadieniai","Kaip sprendžiama":"SPS UG water-fillinamas 0–1; šeštadienių, sekmadienių ir savaitgalio SPS RO 0–1 taikomas nesavanoriškam krūviui. Aiškiai pageidautos savaitgalio pamainos gali RAW spread padidinti. Vengiama clustering."},
            {"Rangas":"3. RESIDENT HARD","Kas įeina":"Negaliu dirbti — data / AM / PM / recurring","Kaip sprendžiama":"0 pažeidimų privaloma; jei tokio SYSTEM grafiko nėra, juodraštis negrąžinamas"},
            {"Rangas":"4. WEEKLY LOAD + RECOVERY","Kas įeina":"Valandos per slenkančias 7 d., kalendorinės savaitės, 12 val. dvigubų dienų seka","Kaip sprendžiama":"Taikosi į ~40 val./7 d.; lygina savaitinį krūvį. Po 1 double vengia kito double; po 2 doubles kita diena PM arba laisva, preferuojama laisva"},
            {"Rangas":"5. ŠVENČIŲ WATER-FILL","Kas įeina":"Oficialios Lietuvos švenčių dienos ir ilgalaikis pasirinkimas: noriu dirbti / neutralu / noriu ilsėtis","Kaip sprendžiama":"Pirmiausia norintys dirbti, tada neutralūs, o norintys ilsėtis — tik kai reikia. Kiekvienoje grupėje 1 visiems → 2 visiems; žiūrima ankstesnė SYSTEM švenčių našta ir mėnesio krūvis."},
            {"Rangas":"6. Kitas struktūrinis krūvis","Kas įeina":"Dvigubų pamainų bendras skaičius ir kitas consecutive/fatigue","Kaip sprendžiama":"Lyginama grupėje nebloginant aukštesnių užraktų; penktadieniai jau užrakinti aukščiau raw 0–1"},
            {"Rangas":"7. OTHER POST CORE","Kas įeina":"CENTRO RO, Onko RO, Centro UG, ADC 144, ADC 145, Vaikų UG, Mamografijos","Kaip sprendžiama":"Struktūrinis water-filling; target 0–1. Mamografija pildoma paskutinė; rezidentams su 0 Onko RO pirmiausia stengiamasi duoti bent vieną likusią Mamografijos ekspoziciją. 0–2 / 0–3 tik solveriui įrodžius, kad siauresnis koridorius neįmanomas"},
            {"Rangas":"8. SOFT-1","Kas įeina":"Noriu laisvos; struktūruotas recovery / vengti dublių","Kaip sprendžiama":"Horizontalus water-filling: bendras sluoksnis visiems prieš papildomus vieno žmogaus prašymus"},
            {"Rangas":"9. SOFT-2","Kas įeina":"Pageidauju dirbti konkrečią datą / AM / PM","Kaip sprendžiama":"Horizontalus water-filling"},
            {"Rangas":"10. SOFT-3","Kas įeina":"Išsklaidymas / koncentracija","Kaip sprendžiama":"Tik po aukštesnių rangų"},
            {"Rangas":"11. CURRENT-MONTH POST OPTIMAL","Kas įeina":"Likęs einamojo mėnesio ordinary-post spread","Kaip sprendžiama":"SOFT rezultato neblogina; kuo labiau grąžina šio mėnesio spread į 0–1. Jokio future catch-up nėra."},
        ]),use_container_width=True,hide_index=True)

        st.markdown("### Švenčių dienų paskirstymo protokolas")
        st.dataframe(pd.DataFrame([
            {"Žingsnis":"1. Šventės atpažinimas","Veikimas":"Sistema automatiškai atpažįsta oficialias Lietuvos DK švenčių dienas. Darbo dieną sutampanti šventė naudoja ne darbo dienos SPS RO budėjimo modelį, o įprasti outpatient postai tą dieną uždaromi / nepriskiriami; aktyvūs lieka SPS RO budėjimo AM/PM slotai."},
            {"Žingsnis":"2. Noriu dirbti","Veikimas":"Jei keli rezidentai nustatymuose pažymėjo, kad linkę dirbti per šventes, šventinės pamainos pirmiausia skiriamos jiems, bet water-fill'inamos: 1 kiekvienam prieš 2 tam pačiam."},
            {"Žingsnis":"3. Neutralu","Veikimas":"Kai norinčių neužtenka arba nėra, naudojami neutralūs. Tarp lygiaverčių kandidatų prioritetą gauna turintis mažesnę ankstesnę SYSTEM švenčių naštą ir mažesnį einamojo mėnesio krūvį."},
            {"Žingsnis":"4. Noriu ilsėtis","Veikimas":"Rezidentai, pasirinkę poilsį per šventes, naudojami tik kai aukštesnių grupių nepakanka dėl coverage / HARD. Ir jų neišvengiama našta water-fill'inama kuo lygiau."},
            {"Žingsnis":"5. Istorija","Veikimas":"SYSTEM ir ACTUAL švenčių darbo istorija saugoma auditui. Kito mėnesio generatorius jos nenaudoja kompensaciniam catch-up; kiekvienas mėnuo pradeda nuo naujo water-fill baseline."},
        ]),use_container_width=True,hide_index=True)
        st.caption("Švenčių pasirinkimas yra normalizuotas SOFT signalas, o ne teisė visada gauti arba visada išvengti šventės. Jis veikia tik aukštesnių ABSOLUTE / critical SPS / RESIDENT HARD / recovery užraktų viduje. Taip išlaikomas ir norų tenkinimas, ir grupės fairness.")

        st.markdown("### Neplanuotas neatvykimas: kritinių SPS postų gelbėjimo hierarchija")
        st.dataframe(pd.DataFrame([
            {"Situacija":"Suserga / neatvyksta žmogus iš SPS RO arba SPS UG","Veiksmas":"Kritinis postas PALIEKAMAS padengtas. Pirmiausia ieškomas tos pačios dienos ir persidengiančio bloko rezidentas, jau dirbantis žemesnės hierarchijos NEPRIVALOMAME poste.","Kas nutinka donoriniam postui":"Rezidentas perkeliamas į SPS; optional donorinis postas gali likti tuščias."},
            {"Situacija":"Yra keli tinkami donorai","Veiksmas":"Neleisti naujo RESIDENT HARD konflikto SYSTEM generavimo metu. Donoro parinkimas neturi bandyti atkurti water-fill; po operacinio pakeitimo ACTUAL spread tiesiog perskaičiuojamas ir parodomas.","Kas nutinka donoriniam postui":"Aukštesnio prioriteto mandatory coverage laimi prieš optional coverage."},
            {"Situacija":"Nėra saugaus donorinio rezidento iš optional posto","Veiksmas":"Tik tada rodomas tame bloke laisvo rezidento fallback, jei jis ABSOLUTE-safe ir nekuria overlap / mandatory coverage problemos.","Kas nutinka donoriniam postui":"Nėra priverstinio critical posto aukojimo."},
            {"Situacija":"Neatvykstama iš paprasto optional posto","Veiksmas":"Šis postas nėra aukščiau SPS RO / SPS UG; kritinio SPS rezidento iš jo traukti negalima.","Kas nutinka donoriniam postui":"Post-publication ACTUAL grafike optional gap gali būti toleruojamas; SYSTEM fairness lieka frozen."},
        ]),use_container_width=True,hide_index=True)
        st.caption("Principas: liga / force majeure pirmiausia perstato jau suplanuotą tos pačios pamainos pajėgumą į privalomą SPS. SYSTEM publikavimo baseline nekinta, bet ACTUAL postų ekspozicija ir fairness perskaičiuojami pagal realią situaciją. Istorija auditinė — jokio ateities catch-up.")

        st.markdown("### Savaitinio krūvio ir savanoriško swapo protokolas")
        st.dataframe(pd.DataFrame([
            {"Taisyklė":"GENERATORIUS — ~40 val./7 d.","Veikimas":"Savaitinis krūvis water-fill'inamas tarp rezidentų; tai planavimo tikslas.","Statusas":"STRUCTURAL TARGET"},
            {"Taisyklė":"GENERATORIUS — ≤48 val./7 d.","Veikimas":"Sistema pati nekuria >48 h rolling-7 krūvio.","Statusas":"GENERATION HARD"},
            {"Taisyklė":"VOLUNTARY SWAP — >48 val./7 d.","Veikimas":"Vien 48 h viršijimas swapo neblokuoja. Paveiktas rezidentas mato pasekmių lentelę ir turi aiškiai patvirtinti.","Statusas":"ACK / WARNING"},
            {"Taisyklė":"VOLUNTARY SWAP — ≤60 val./7 d.","Veikimas":">60 h per bet kurias 7 paeiliui einančias dienas atmetama.","Statusas":"BLOCK"},
            {"Taisyklė":"≤12 val. per darbo dieną","Veikimas":"AM+PM = 12 h galima; >12 h atmetama. Nauja 12 h diena prieš sutikimą aiškiai parodoma.","Statusas":"BLOCK + ACK ties 12 h"},
            {"Taisyklė":"≥11 val. nepertraukiamo paros poilsio","Veikimas":"Jei tarp darbo dienų / pamainų po swapo lieka <11 h, swapas atmetamas.","Statusas":"BLOCK"},
            {"Taisyklė":"Po 6 darbo dienų — poilsis","Veikimas":"Negalima >6 darbo dienų per 7 paeiliui einančias dienas. 6 dienų seka leidžiama ir rodoma kaip perspėjimas.","Statusas":"7-a diena = BLOCK"},
            {"Taisyklė":"Recovery po doubles","Veikimas":"Generatorius po dviejų doubles kitą dieną riboja. Savanoriškame swape tai tampa ACK perspėjimu, jei 11 h / 12 h / 6 d. / 60 h ribos išlaikytos.","Statusas":"GENERATION HARD → SWAP ACK"},
            {"Taisyklė":"RESIDENT HARD per savanorišką swapą","Veikimas":"Jei rezidentas pats priima darbą per savo ankstesnį „Negaliu dirbti“, sistema rodo override ir prašo jo ACK; ORIGINAL pageidavimas istorijoje lieka.","Statusas":"SWAP ACK"},
        ]),use_container_width=True,hide_index=True)

        st.markdown("### Kas swapą BLOKUOJA ir kas tik PERSPĖJA")
        st.dataframe(pd.DataFrame([
            {"Tipas":"BLOKUOJA","Pavyzdžiai":"ABSOLUTE HARD / pateisinamas neatvykimas; overlap; >12 h/d.; <11 h poilsio; >6 darbo dienų/7 d.; >60 h/7 d.; 24 h post-duty rest; neįmanomas backup/coverage; mėnesio target ≠ tikslus; Onko 1/3/5","ACK":"Negali apeiti"},
            {"Tipas":"PERSPĖJA + ACK","Pavyzdžiai":"Nauja 12 h double; >40 ar >48 h/7 d.; 6 darbo dienų seka; consecutive doubles; darbas po 2 doubles; consecutive Onko; savo RESIDENT HARD override","ACK":"Kiekvienas paveiktas rezidentas patvirtina atskirai"},
            {"Tipas":"NEBLOKUOJA VOLUNTARY SWAP","Pavyzdžiai":"SYSTEM post spread, weekend/double fairness, SOFT satisfaction","ACK":"SYSTEM baseline frozen; ACTUAL perskaičiuojamas"},
        ]),use_container_width=True,hide_index=True)
        st.caption("ACK nėra teisinė išimtis: jis tik patvirtina rezidentui parodytas pasekmes. Darbo laiko režimo ir apskaitinio laikotarpio teisinį taikymą galutinai nustato darbdavys.")
        st.info("V2.5.66 — vienas rezidentas gali turėti kelis laukiančius apsikeitimus, jei jie liečia skirtingas pamainas. Ta pati konkreti pamaina vienu metu gali būti tik viename aktyviame pasiūlyme. Ta pati taisyklė taikoma dublių apsikeitimams. Savo dar nepriimtą pasiūlymą galima atšaukti. Jau pritaikytas ar atmestas pasiūlymas pamainos neberezervuoja.")
        st.info("V2.5.67 — mėnesio darbo krūvio targetas yra ABSOLIUTUS: 28 reiškia tiksliai 28.0, 26 reiškia tiksliai 26.0. Onko diena = 1.5 pamainos, todėl Onko skiriamas poromis (0, 2, 4...) ir mėnesio skirtumas tarp rezidentų negali viršyti 2. Kas šį mėnesį gauna mažiau Onko, turi catch-up prioritetą kitais mėnesiais pagal publikuotą istoriją.")
        st.info("V2.5.68 — Onko RO atsigavimo taisyklė yra ABSOLIUTI: tas pats rezidentas negali būti Onko RO dvi kalendorines dienas iš eilės. Jei dirbo Onko paskutinę ankstesnio mėnesio dieną, naujo mėnesio 1 d. Onko jam taip pat blokuojamas. Taisyklė negali būti paaukota dėl postų lygybės ar SOFT pageidavimų.")
        st.info("V2.5.73 — ONKO PORŲ ABSOLIUTI TAISYKLĖ: kiekvieno rezidento Onko skaičius SYSTEM ir ACTUAL grafike turi būti tik 0, 2, 4, 6... Kadangi viena Onko diena = 1.5 pamainos, nelyginis 1/3/5 sukurtų 0.5 krūvio trupmeną ir yra BLOKUOJAMAS net savanoriškame swape. Jei aktyvių mėnesio Onko dienų skaičius nelyginis, viena Onko diena paliekama neužpildyta, kad bendras užpildytų Onko skaičius būtų lyginis. Consecutive Onko po publikavimo gali likti tik ACK išimtis; parity ir tikslus mėnesio targetas — niekada.")
        st.info("V2.5.74 — VISŲ POSTŲ STRUCTURAL WATER-FILL: SYSTEM generavime, kai datos ir AM/PM blokai jau parinkti, visi ne-Onko postų labeliai sprendžiami kartu. Kiekvienam postui pirmiausia bandomas floor/ceil pasiskirstymas raw spread 0–1. Pvz., 38 Mamografijos vietos / 16 rezidentų → 10 rezidentų po 2 ir 6 rezidentai po 3; 1-vs-3 negali likti, jei egzistuoja validus postų perkeitimas ar kelių žmonių ciklas. Po publikavimo savanoriški ACTUAL swapai gali išbalansuoti postų ekspoziciją — fairness / UG / Mamografijos kiekiai swapo NEBLOKUOJA; SYSTEM fairness lieka užšaldytas.")
        st.info("V2.5.77 — PENKTADIENIŲ STRUCTURAL WATER-FILL: SYSTEM grafike visi penktadienio priskyrimai skaičiuojami kaip struktūrinė našta ir turi būti paskirstyti floor/ceil principu, raw max−min ≤1. Pvz., jei yra 72 penktadienio priskyrimai / 16 rezidentų, matematinis water-fill yra 8 rezidentai po 4 ir 8 po 5. Pageidautas penktadienis vis tiek skaičiuojamas kaip penktadienio ekspozicija. Phase 1 subalansuoja penktadienio darbo blokus, Phase 2 ant tų blokų kartu perbalansuoja visus ne-Onko postus. Po publikavimo abipusis ACTUAL swapas gali išbalansuoti penktadienius; SYSTEM baseline dėl to nesikeičia.")

        st.markdown("### Emergency — jau įvykusio pakeitimo registravimas")
        st.dataframe(pd.DataFrame([
            {"Situacija":"Skubiai realybėje sukeisti du rezidentai / jų pamainos","Ką darome":"Seniūnė arba vienas iš dalyvavusių rezidentų Apsikeitimai → Emergency lange įrašo abi buvusias pamainas.","Kas pasikeičia":"ACTUAL grafikas ir pakeitimų žurnalas."},
            {"Situacija":"Pakeitimą įrašo seniūnė","Ką darome":"Abiem dalyvavusiems rezidentams žurnale rodoma 🔔, kol jie pažymi, kad įrašą matė / jis teisingas.","Kas pasikeičia":"Tik peržiūros patvirtinimas; grafikas jau rodo faktą."},
            {"Situacija":"Pakeitimą įrašo pats rezidentas","Ką darome":"Jo peržiūra pažymima iš karto; kitam dalyviui paliekamas 🔔 patvirtinimas.","Kas pasikeičia":"ACTUAL + audito istorija."},
            {"Situacija":"Mėnuo uždaromas / eksportuojamas galutinis grafikas","Ką darome":"Naudojamas ACTUAL grafikas su visais įrašytais emergency pakeitimais.","Kas pasikeičia":"Galutinė faktinė istorija tiksli."},
            {"Situacija":"Fairness / mokomojo paskirstymo vertinimas","Ką darome":"Emergency įrašas jo neperskaičiuoja.","Kas pasikeičia":"Nieko: SYSTEM publikavimo bazė, postų istorija ir ateities kompensacijos lieka frozen."},
        ]),use_container_width=True,hide_index=True)
        st.caption("Jei apsikeitimas dar tik planuojamas, naudokite įprastą savanoriško apsikeitimo srautą. Emergency poskyris skirtas realiai jau įvykusiam / tą pačią dieną aiškiai sutartam faktui užregistruoti, kad galutinis ACTUAL grafikas atitiktų realybę.")
        st.info("Teisinis orientyras pagal VDI: bent 11 val. nepertraukiamo paros poilsio, bent 35 val. nepertraukiamo poilsio per 7 paeiliui einančias dienas ir ne daugiau kaip 6 darbo dienos per 7. 48 val. yra svarbus vidutinio darbo laiko slenkstis; atskirais režimais taikomi papildomi 52/60 val. limitai. Todėl >48 visada rodoma su tiksliu skaičiumi ir aiškiu perspėjimu, o tikros absoliučios ribos lieka blokuojamos.")

        st.markdown("### Savanoriško dublio perėmimo išimtis")
        st.dataframe(pd.DataFrame([
            {"Situacija":"Po dublio būtų 12 val. darbo diena","Ką mato žmogus":"Dabar X val. → po dublio 12 val.","Veiksmas":"Perspėjimas. Galima ATŠAUKTI arba PATVIRTINTI VIS TIEK, jei kitos absoliučios ribos išlaikytos."},
            {"Situacija":"Po dublio 7 dienų krūvis viršija ~40 / 48 val.","Ką mato žmogus":"Tikslus skaičius ir konkrečios 7 dienos, pvz. 54 val. (12–18 d.)","Veiksmas":"Savanoriškas ACK. 48 val. nėra tyliai ignoruojama — žmogus aiškiai mato pasekmę."},
            {"Situacija":"Rezidentas pats perima dublį per savo RESIDENT HARD laiką","Ką mato žmogus":"Kuri data / blokas buvo pažymėtas „Negaliu dirbti“","Veiksmas":"Galima tik aiškiai savanoriškai patvirtinus; ORIGINAL pageidavimas istorijoje lieka."},
            {"Situacija":"Po dublio būtų >12 val./d., aktyvios 7 d. ribos viršijimas, <11 val. paros poilsio arba >6 darbo dienos/7 d.","Ką mato žmogus":"Tikslus apskaičiuotas pažeidimas ir riba skliausteliuose","Veiksmas":"BLOKUOJAMA — šių ribų manual ACK neapeina."},
            {"Situacija":"Pateisinamas neatvykimas / privalomas post-duty poilsis / persidengianti pamaina","Ką mato žmogus":"Konkreti priežastis","Veiksmas":"BLOKUOJAMA."},
        ]),use_container_width=True,hide_index=True)
        st.caption("Principas paprastas: jei dublis tik pablogina planavimo komfortą, žmogus gali sąmoningai sutikti. Jei atsiranda tikras saugos / teisinis blokatorius, sistema jo nevadina „override“ ir neleidžia patvirtinti.")

        st.markdown("### Kaip lyginamas darbas skirtingose pozicijose")
        st.dataframe(pd.DataFrame([
            {"Grupė":"KRITINĖ","Pozicijos":"SPS RO, SPS UG, šeštadieniai, sekmadieniai","Taisyklė":"Kiekviena kategorija balansuojama atskirai. Šeštadieniai ir sekmadieniai turi atskirą water-fill 0–1, nes jų atlygio / naštos pobūdis skiriasi."},
            {"Grupė":"ONKO RO — SPECIALI HARD","Pozicijos":"Onko RO 08:00–17:00","Taisyklė":"1 diena = 1.5 pamainos, todėl SYSTEM skiriama lyginėmis poromis (0/2/4...), mėnesio skirtumas ≤2. Tas pats rezidentas NEGALI būti Onko dvi kalendorines dienas iš eilės, įskaitant mėnesio ribą."},
            {"Grupė":"KITI POSTAI","Pozicijos":"CENTRO RO, Centro UG, ADC 144, ADC 145, Vaikų UG, Mamografijos","Taisyklė":"Jei mėnesio vietų pakanka, pirmiausia kiekvienas turi gauti bent vieną galimybę. Toliau paskirstymas lyginamas kuo labiau; didesnis skirtumas leidžiamas tik kai lygesnis variantas neįmanomas arba būtinas svarbesnei taisyklei."},
            {"Grupė":"FUTURE CATCH-UP","Pozicijos":"Kiekvienas postas × rezidentas","Taisyklė":"Jei šį mėnesį žmogus konkrečioje darbo vietoje gavo mažiau nei kiti, kitą mėnesį sistema jam teikia pirmenybę pasivyti. Jei gavo daugiau, papildomas paskyrimas pirmiau siūlomas kitiems."},
            {"Grupė":"TEMPORAL SPACING","Pozicijos":"Ypač savaitgaliai, taip pat SPS RO/SPS UG","Taisyklė":"Vienodi skaičiai dar nereiškia vienodo nuovargio: tarp lygiaverčių variantų vengiami 2–3 savaitgaliai iš eilės ir bereikalingas SPS suspaudimas."},
        ]),use_container_width=True,hide_index=True)

        st.markdown("### Kokie SOFT priimami")
        st.dataframe(pd.DataFrame([
            {"Būsena":"PRIIMAMA","Pavyzdys":"Noriu laisvos konkrečią datą / AM / PM","Kodėl":"Aiškus asmeninio laiko poreikis; SOFT-1"},
            {"Būsena":"PRIIMAMA","Pavyzdys":"Pageidauju dirbti konkrečią datą / AM / PM","Kodėl":"Aiškus teigiamas darbo pageidavimas; SOFT-2"},
            {"Būsena":"PRIIMAMA","Pavyzdys":"Vengti dublių / labiau išsklaidytas ar koncentruotas mėnuo","Kodėl":"Standartizuotas recovery / schedule-shape signalas"},
            {"Būsena":"NEPRIIMAMA kaip SOFT","Pavyzdys":"Nedėk manęs į Mamografiją / tik SPS RO / kuo daugiau Centro RO","Kodėl":"Pageidavimų sistema nėra darbo vietų pasirinkimo meniu; visi turi gauti kuo lygesnes mokymosi galimybes."},
            {"Būsena":"NEPRIIMAMA kaip bendras SOFT","Pavyzdys":"Nenoriu savaitgalių / noriu mažiau darbo dienų","Kodėl":"Gali permesti kritinę naštą kitiems. Jei reikia konkrečios datos — rinktis ją; jei realiai negalite — RESIDENT HARD"},
        ]),use_container_width=True,hide_index=True)
        st.caption("Principas: sistema tenkina realius poreikius, bet kartu saugo lygybę visai grupei. Vien tai, kad žmogus pateikė daugiau pageidavimų, nesuteikia jam didesnės galios už kitus.")
    else:
        st.success(
            "V2.5.62 — EMERGENCY ACTUAL SWAP LOG + VOLUNTARY BACKUP OVERRIDE. V2.5.58 engine rules are preserved; a dedicated Senior guide adds a five-minute concrete-statement verification audit, red flags and end-to-end monthly workflow. "
            "Generation remains conservative (~40h target, ≤48h/7d, ≤6 workdays/7d and post-double recovery). "
            "After publication, bilateral voluntary swaps are blocked only by ABSOLUTE/operational and labour-time guardrails; new 12h doubles, >40/>48h load, six-day streaks, post-double recovery patterns and self-overridden Resident-HARD requests are shown in a consequence table and require acknowledgement."
        )
        st.markdown("### Unplanned absence: critical SPS rescue hierarchy")
        st.dataframe(pd.DataFrame([
            {"Situation":"Resident absent from SPS RO or SPS UG","Action":"Keep the critical post covered. First pull a resident already working the same day / overlapping block in a lower-priority NON-MANDATORY post.","Donor post":"Move the resident to SPS; the optional source post may remain empty."},
            {"Situation":"Several safe donors exist","Action":"Reject new Resident-HARD conflicts unless a later ACTUAL voluntary action explicitly changes the resident's own request. Do not force a donor choice to restore water-fill; after the operational move, ACTUAL exposure/fairness is simply recalculated and reported.","Donor post":"Mandatory critical coverage outranks optional coverage."},
            {"Situation":"No safe optional-post donor exists","Action":"Only then use a resident free in that block as fallback if ABSOLUTE-safe and overlap/coverage-valid.","Donor post":"Never sacrifice another critical SPS post."},
        ]),use_container_width=True,hide_index=True)

        st.markdown("### Simplified vertical-priority table")
        st.dataframe(pd.DataFrame([
            {"Rank":"1. TRUE ABSOLUTE HARD","Includes":"Safety/rest, approved absence, physical impossibility, coverage; generation <=48h/rolling7 and >=1 free day/7d","Method":"100% during generation. Post-publication voluntary swaps use consequence + ACK warnings, but ABSOLUTE/operational and labour-time blockers remain hard"},
            {"Rank":"2. CRITICAL STRUCTURAL","Includes":"SPS RO + SPS UG + Saturday + Sunday + Fridays","Method":"Saturday and Sunday water-fill independently; raw spread 0–1; reduce temporal clustering"},
            {"Rank":"3. RESIDENT HARD","Includes":"Unavailable date / AM / PM / recurring","Method":"Zero violations are mandatory in SYSTEM generation; if impossible, no draft is returned"},
            {"Rank":"4. WEEKLY LOAD + RECOVERY","Includes":"Rolling-7 hours, calendar-week load, double-shift sequences","Method":"Aim ~40h/7d; equalize weekly load; after 2 consecutive doubles next day PM-only or off, preferring off"},
            {"Rank":"5. OTHER STRUCTURAL","Includes":"Total doubles and other consecutive/fatigue","Method":"Balance without worsening higher locks; Fridays are already structurally locked at raw 0–1"},
            {"Rank":"6. OTHER POST CORE","Includes":"CENTRO RO, Centro UG, ADC 144/145, Paediatric UG, Mammography; Onko has its own special HARD structure","Method":"Ordinary non-Onko posts: structural floor/ceil water-fill with target raw spread <=1 before SOFT; <=2/<=3 only after the tighter corridor is proven infeasible. Onko: exact-workload even pairs, monthly spread <=2, never consecutive calendar days."},
            {"Rank":"7–9. SOFT","Includes":"SOFT-1 time/recovery; SOFT-2 exact desired work; SOFT-3 month shape","Method":"Vertical rank + horizontal resident water-fill"},
            {"Rank":"10. CURRENT-MONTH POST OPTIMAL","Includes":"Residual ordinary-post spread in this month","Method":"Improve toward 0–1 without worsening locked SOFT; no longitudinal catch-up"},
        ]),use_container_width=True,hide_index=True)
    # RULES = ENGINE: this summary is rendered from the active engine profile,
    # never from a separate hard-coded policy copy.
    enabled_backup=[]
    if bool(rule_value("backup_sps_ro")): enabled_backup.append("SPS RO")
    if bool(rule_value("backup_sps_ug")): enabled_backup.append("SPS UG")
    if bool(rule_value("backup_centro120_am")): enabled_backup.append("Centro 120 AM")
    if bool(rule_value("backup_onko_ro")): enabled_backup.append("Onko RO")
    if bool(rule_value("backup_centro_ro_best_effort")): enabled_backup.append("CENTRO RO best-effort")

    if lang=="LT":
        st.info(
            f"AKTYVUS TAISYKLIŲ PROFILIS v{ACTIVE_RULE_PROFILE_VERSION} — TAISYKLĖS = ENGINE. "
            f"Dublio padengimas: {', '.join(enabled_backup)}. "
            f"Target formulė: darbo dienos × {float(rule_value('target_daily_hours')):g} / {float(rule_value('target_shift_hours')):g}. "
            f"Min. poilsis: {float(rule_value('min_rest_hours')):g} val.; max. darbo valandų/d.: {float(rule_value('max_hours_per_day')):g}; "
            f"max. darbo dienų per 7 d.: {min(int(rule_value('max_workdays_rolling7')), int(FATIGUE_MAX_WORKDAYS_ROLLING7))}; "
            f"GENERATION HARD max. valandų per 7 d.: {min(float(rule_value('max_hours_rolling7')), float(FATIGUE_ROLLING7_HARD_CEILING_HOURS)):g}; voluntary swap >48 = ACK, absoliutus guardrail ≤{float(SWAP_ABSOLUTE_MAX_HOURS_ROLLING7):g}; "
            f"planavimo tikslas ~{float(WEEKLY_LOAD_SOFT_TARGET_HOURS):g} val./7 d. "
            f"Mėnesio krūvio targetas: TIKSLUS HARD (leidžiamas nuokrypis 0.0). "
            f"Onko: 1.5 pamainos, tik lyginės poros (0/2/4...), mėnesio skirtumas ≤2; jokio istorinio catch-up, niekada dvi kalendorines dienas iš eilės tam pačiam rezidentui; "
            f"savaitgalio unikalumo taisyklė: {'TAIP' if rule_value('weekend_unique_required') else 'NE'}. "
            f"Struktūrinis guardrail: SPS RO / SPS UG / ŠEŠTADIENIAI / SEKMADIENIAI / PENKTADIENIAI raw 0–1; Onko ≤2 poromis; VISI kiti postai pirmiausia raw 0–1. 0–2/0–3 leidžiama tik įrodžius, kad siauresnis variantas neįmanomas. Jokio future future catch-up nėra. "
            f"Kiti pagrindiniai burden spread baseline +{int(rule_value('general_guardrail_tolerance'))}. "
            f"Pageidavimų pateikimo terminas: ankstesnio mėnesio {int(rule_value('deadline_day'))} d."
        )
    else:
        st.info(
            f"ACTIVE RULE PROFILE v{ACTIVE_RULE_PROFILE_VERSION} — RULES = ENGINE. "
            f"Backup scope: {', '.join(enabled_backup)}. "
            f"Target formula: weekdays × {float(rule_value('target_daily_hours')):g} / {float(rule_value('target_shift_hours')):g}. "
            f"Minimum rest: {float(rule_value('min_rest_hours')):g}h; max hours/day: {float(rule_value('max_hours_per_day')):g}; "
            f"max workdays/7d: {min(int(rule_value('max_workdays_rolling7')), int(FATIGUE_MAX_WORKDAYS_ROLLING7))}; "
            f"GENERATION HARD max hours/7d: {min(float(rule_value('max_hours_rolling7')), float(FATIGUE_ROLLING7_HARD_CEILING_HOURS)):g}; voluntary swap >48 = ACK, absolute guardrail ≤{float(SWAP_ABSOLUTE_MAX_HOURS_ROLLING7):g}; "
            f"planning target ~{float(WEEKLY_LOAD_SOFT_TARGET_HOURS):g}h/7d. "
            f"Monthly workload target: EXACT HARD (allowed deviation 0.0). "
            f"Onko: 1.5 shift units, even pairs only (0/2/4...), monthly spread ≤2 with no historical catch-up, never on consecutive calendar days for the same resident; "
            f"weekend uniqueness: {'YES' if rule_value('weekend_unique_required') else 'NO'}. "
            f"Structural guardrail: SPS RO / SPS UG / SATURDAYS / SUNDAYS / FRIDAYS raw 0–1; Onko ≤2 in even pairs; ALL other posts first target raw 0–1. 0–2/0–3 is allowed only after the tighter corridor is proven infeasible. No future future catch-up exists. "
            f"Other main burden spreads baseline +{int(rule_value('general_guardrail_tolerance'))}. "
            f"Preference deadline: day {int(rule_value('deadline_day'))} of the preceding month."
        )

    if lang=="LT":
        workflow_rows=[
            {"Etapas":"0. Request pre-check","Sistema":"Užšaldo ORIGINAL request ledger ir pašalina nepriimamus/gaming SOFT signalus.","Vertina":"RESIDENT HARD, tikslias SOFT datas, recovery ir month-shape; generic weekday/weekend pattern ir postų vengimas neįeina.","Principas":"Pageidavimų skaičius nesuteikia daugiau balsų."},
            {"Etapas":"1. TRUE ABSOLUTE HARD","Sistema":"Randa tik saugų/fiziškai įmanomą grafiką; generuojant taiko ≤48h/7d ir Onko recovery guard.","Vertina":"Poilsį, valandas, patvirtintą neatvykimą, coverage/overlap, tikslų mėnesio krūvį ir lygines Onko poras.","Principas":"Tikslus mėnesio targetas ir Onko 0/2/4/... yra HARD SYSTEM ir ACTUAL. Consecutive Onko gali būti tik savanoriško swapo ACK pasekmė; parity niekada neapeinama."},
            {"Etapas":"2. Kritinių darbų lygybė","Sistema":"Kartu lygina SPS RO, SPS UG, šeštadienius, sekmadienius ir penktadienius tarp rezidentų.","Vertina":"Šeštadienis ir sekmadienis yra atskiros naštos / atlygio klasės.","Principas":"Neutralus SYSTEM baseline pradeda nuo 0–1 water-fill kiekvienai klasei; po publikavimo savanoriški swapai gali ACTUAL balansą pakeisti."},
            {"Etapas":"3. RESIDENT HARD","Sistema":"Užrakina `Negaliu dirbti` kaip privalomą 0-pažeidimų SYSTEM apribojimą prieš bet kokį fairness ar SOFT optimizavimą.","Vertina":"Whole-day, AM/PM ir recurring RESIDENT HARD.","Principas":"0 pažeidimų privaloma; jei safety/coverage/target su tuo nesuderinami, SYSTEM juodraštis negrąžinamas."},
            {"Etapas":"4. Critical spacing","Sistema":"Nejudindama kritinių count spreadų, išdėsto juos laike.","Vertina":"Consecutive weekends ir SPS dienų clustering, įskaitant ankstesnio mėnesio weekend tail.","Principas":"Vengti 2–3 savaitgalių iš eilės ir bereikalingo streso suspaudimo."},
            {"Etapas":"5. WEEKLY LOAD + RECOVERY","Sistema":"Water-fill'ina savaitinį valandų krūvį ir užrakina recovery frontier.","Vertina":"Rolling-7 valandas, kalendorinių savaičių spreadą, consecutive 12h doubles.","Principas":"~40 val./7 d. tikslas; ≤48 HARD; po 2 doubles kita diena tik PM arba laisva, laisva preferinama."},
            {"Etapas":"6. ŠVENČIŲ WATER-FILL","Sistema":"Šventines pamainas skirsto einamajame mėnesyje, atsižvelgdamas į aktyvų švenčių pageidavimą ir aukštesnius užraktus.","Vertina":"Tik einamojo mėnesio holiday burden ir mėnesio krūvį.","Principas":"Kur įmanoma 1 visiems prieš 2; ankstesni mėnesiai catch-up nesukuria."},
            {"Etapas":"7. Kitas burden fairness","Sistema":"Balansuoja bendrą dublių skaičių ir kitą consecutive/fatigue.","Vertina":"Likusią struktūrinę naštą.","Principas":"Penktadieniai jau HARD water-fillinti raw 0–1 ir čia nebeatlaisvinami."},
            {"Etapas":"8. Kitų darbo vietų lygybė","Sistema":"Patikrina likusių darbo vietų paskirstymą tarp rezidentų.","Vertina":"Kiek skiriasi daugiausiai ir mažiausiai konkrečią darbo vietą gavę rezidentai.","Principas":"Siekiama 0–1; įprastai leidžiama iki 2; 3 tik jei 2 tikrai neįmanoma dėl svarbesnių taisyklių."},
            {"Etapas":"9. SOFT-1 → SOFT-2 → SOFT-3","Sistema":"Kiekvieną rangą water-fill'ina horizontaliai ir užrakina.","Vertina":"Asmeninį laiką/recovery → tikslias darbo datas → month shape.","Principas":"2,2,3,4 pirmiausia 2,2,2,2; tik tada extras."},
            {"Etapas":"10. CURRENT-MONTH POST OPTIMAL","Sistema":"SOFT neblogindama grąžina šio mėnesio ordinary post spread kuo arčiau 0–1.","Vertina":"Tik einamojo mėnesio resident × post exposure.","Principas":"Istorija stebima, bet nekuria skolos ir nekeičia kito mėnesio paskyrimų."},
            {"Etapas":"11. ACTUAL + swaps/repairs","Sistema":"Po swap/repair perskaičiuoja ACTUAL grafiką, satisfaction ir live fairness. Kritinio SPS neatvykimo atveju pirmiausia perkelia žmogų iš tos pačios pamainos optional posto.","Vertina":"Mandatory SPS coverage, realią postų ekspoziciją, ACTUAL spread, konkrečius misses ir saugą.","Principas":"SYSTEM baseline lieka užšaldytas auditui; ACTUAL fairness seka realybę. Post-publication water-fill gali būti pralaužtas, bet jokio future catch-up nesukuria."},
        ]
        if advanced_mode:
            st.markdown("### Generatorius: workflow")
            st.dataframe(pd.DataFrame(workflow_rows),use_container_width=True,hide_index=True)
            st.caption("Kai žmogus neturi SOFT, jis neįtraukiamas į SOFT max-min; tačiau jo RESIDENT HARD vis tiek pilnai dalyvauja aukštesnio prioriteto request-fairness sluoksnyje.")
        else:
            st.caption("Pilną generatoriaus workflow lentelę gali matyti Išplėstiniame režime.")
    else:
        st.info(
            "V2.5.65 — WORKDAYS FIRST, THEN WORKPLACES. The engine first chooses when each resident works while protecting Resident-HARD, recovery and personal requests. It then assigns workplaces. For ordinary 1.0-unit sparse posts such as Centro UG / pediatric US, the system gives everyone a first exposure before avoidable second exposures where mathematically feasible. Onko is now governed by the later V2.5.67–68 exact-workload pair + recovery rules, not by first-exposure 1–2 logic. SPS RO, SPS UG and weekends remain as equal as possible. A specific SPS date is not locked to one resident if the same monthly amount can be preserved with another placement that honors the resident's request. Approved vacation is an absolute no-work period and proportionally lowers that resident's monthly workload target. "
            "V2.5.63 FAIRNESS FAILSAFE. A SYSTEM draft is returned only after the solver verifies an acceptably even distribution: SPS RO, SPS UG and weekends normally differ by no more than one assignment between residents; other main workplaces normally differ by no more than two. A timeout is not treated as proof that a wider imbalance is necessary. Concrete SPS dates remain flexible, so personal requests may still be honored whenever the same overall equality can be preserved."
        )
        st.info("V2.5.66 — a resident may have several pending swaps when they involve different shifts. The same concrete shift may be in only one active future offer at a time; the same rule applies to backup swaps. A requester may cancel their own still-pending offer. Applied/rejected offers release the shift.")
        st.info("V2.5.67 — the calculated monthly workload target is ABSOLUTE: 28 means exactly 28.0 and 26 means exactly 26.0. One Onko day = 1.5 shift units, so Onko is assigned in pairs (0, 2, 4...) with a monthly resident spread no greater than 2. Residents with fewer Onko exposures receive catch-up priority in later months using published history.")
        st.info("V2.5.68 — Onko RO recovery is ABSOLUTE: the same resident may not work Onko RO on two consecutive calendar days. If the resident worked Onko on the last day of the previous published month, day 1 of the new month is also blocked for Onko. Fairness or SOFT preferences may not override this rule.")
        st.info("V2.5.73 — ONKO PAIRING ABSOLUTE: every resident's Onko count must be 0, 2, 4, 6... in both SYSTEM and ACTUAL. Because one Onko day equals 1.5 workload units, odd 1/3/5 would create a half-unit monthly workload and is blocked even in a voluntary swap. If the number of active monthly Onko days is odd, one Onko day remains unfilled so the filled total is even. Consecutive Onko may remain a post-publication ACK exception; parity and exact monthly workload never are.")
        st.info("V2.5.74 — ALL-POST STRUCTURAL WATER-FILL: in SYSTEM generation, after dates and AM/PM blocks are frozen, all non-Onko post labels are solved jointly. Every post first targets its floor/ceil distribution with raw spread 0–1. Example: 38 Mammography slots / 16 residents → ten residents receive 2 and six receive 3; a 1-vs-3 pattern cannot remain when a valid post exchange or multi-person cycle can equalize it. After publication, voluntary ACTUAL swaps may unbalance exposure — post fairness / US / Mammography counts do NOT block a mutually accepted swap; SYSTEM fairness remains frozen.")
        workflow_rows=[
            {"Stage":"0. Request pre-check","System":"Freezes ORIGINAL request ledger and removes non-whitelisted/gaming SOFT signals.","Evaluates":"Resident-HARD, exact SOFT dates, recovery/month-shape; generic weekday/weekend patterns and station avoidance are excluded.","Principle":"More raw requests do not buy more priority."},
            {"Stage":"1. TRUE ABSOLUTE HARD","System":"Finds only safe/physically feasible schedules; generation applies <=48h/7d and the Onko recovery guard.","Evaluates":"Rest, hours, approved absence, coverage/overlap, exact monthly workload and even Onko pairing.","Principle":"Exact monthly workload and Onko 0/2/4/... are HARD in SYSTEM and ACTUAL. Consecutive Onko may be accepted only as a voluntary-swap ACK consequence; parity can never be overridden."},
            {"Stage":"2. CRITICAL WATER-FILL","System":"Co-optimizes SPS RO, SPS UG and all weekend exposure.","Evaluates":"Raw max-min in the three critical categories.","Principle":"0–1; first unit for everyone before second; ordinary SOFT cannot widen to 2."},
            {"Stage":"3. RESIDENT HARD","System":"Locks every Unavailable block as a mandatory zero-violation SYSTEM constraint before fairness or SOFT optimization.","Evaluates":"Whole-day, AM/PM and recurring Resident-HARD.","Principle":"Zero violations are mandatory; if safety/coverage/target cannot coexist with them, no SYSTEM draft is returned."},
            {"Stage":"4. Critical spacing","System":"Places equivalent critical counts more evenly in time.","Evaluates":"Consecutive weekends and SPS clustering, including prior-month weekend tail.","Principle":"Avoid concentrated fatigue."},
            {"Stage":"5. WEEKLY LOAD + RECOVERY","System":"Water-fills weekly hours and locks the recovery frontier.","Evaluates":"Rolling-7 hours, calendar-week spread, consecutive 12h doubles.","Principle":"Aim ~40h/7d; <=48h HARD; after 2 doubles next day PM-only or off, preferring off."},
            {"Stage":"6. HOLIDAY WATER-FILL","System":"Allocates public-holiday duty within the current month while respecting active holiday preferences and higher locks.","Evaluates":"Current-month holiday burden and monthly load only.","Principle":"One unit for everyone before seconds where feasible; prior months do not create catch-up."},
            {"Stage":"7. Other burden","System":"Balances total doubles and other consecutive/fatigue burden.","Evaluates":"Remaining structural burden.","Principle":"Fridays are already HARD water-filled at raw 0–1 and are not relaxed here."},
            {"Stage":"8. OTHER POST CORE","System":"Tests the all-post structural floor/ceil water-fill corridor.","Evaluates":"Seven noncritical post spreads.","Principle":"Target raw 0–1. <=2 and then <=3 are tried only after the tighter corridor is proven infeasible inside higher locks."},
            {"Stage":"9. SOFT-1 → SOFT-2 → SOFT-3","System":"Water-fills residents horizontally and locks each vertical rank.","Evaluates":"Time/recovery → exact desired work → month shape.","Principle":"Common entitlement layers before extras."},
            {"Stage":"10. CURRENT-MONTH POST OPTIMAL","System":"Without worsening SOFT, returns this month’s ordinary posts toward 0–1.","Evaluates":"Current-month resident × post exposure only.","Principle":"History is monitored but creates no debt and never steers a future month."},
            {"Stage":"11. ACTUAL + swaps/repairs","System":"Recomputes ACTUAL schedule, satisfaction and live fairness. For critical SPS absence, first transfers a resident from a same-block optional post.","Evaluates":"Mandatory SPS coverage, real post exposure, ACTUAL spreads, misses and safety.","Principle":"SYSTEM publication baseline stays frozen for audit; ACTUAL fairness follows reality. Water-fill may be broken post-publication and no future catch-up is created."},
        ]
        if advanced_mode:
            st.markdown("### Generator workflow")
            st.dataframe(pd.DataFrame(workflow_rows),use_container_width=True,hide_index=True)
        else:
            st.caption("The full generator workflow table is available in Advanced mode.")
        st.caption("Residents with no SOFT request are excluded from SOFT max-min, but any Resident-HARD requests still participate in the higher-priority burden-equity layer.")
    if senior_mode and advanced_mode:
        st.divider()
        st.markdown("### RULE CHANGE RESCUE" if lang=="EN" else "### TAISYKLIŲ KEITIMO RESCUE")
        st.caption(
            ("Čia keičiamos tik iš anksto saugiai sukonfigūruotos administracinės taisyklės. "
             "Kiekvienas pakeitimas sukuria naują versiją; ankstesnė versija lieka istorijoje ir gali būti grąžinta vienu veiksmu."
             if lang=="LT" else
             "Only pre-defined safe administrative rules are editable here. Every change creates a new version; the previous version remains in history and can be restored in one action.")
        )

        active_cfg=dict(ACTIVE_RULES)
        with st.expander("Keisti aktyvias taisykles" if lang=="LT" else "Edit active rules", expanded=False):
            with st.form(f"rule_profile_editor_v2534_{ACTIVE_RULE_PROFILE_VERSION}"):
                profile_name=st.text_input(
                    "Naujos versijos pavadinimas" if lang=="LT" else "New version name",
                    value=f"Rule profile v{ACTIVE_RULE_PROFILE_VERSION+1}"
                )
                profile_note=st.text_area(
                    "Kodėl keičiama?" if lang=="LT" else "Reason for change",
                    placeholder=("Pvz. nuo spalio SPS UG dubliams taikoma kita taisyklė." if lang=="LT" else "Example: SPS UG backup policy changes from October.")
                )

                st.markdown("#### Operacinės / HARD taisyklės" if lang=="LT" else "#### Operational / HARD rules")
                r1,r2,r3=st.columns(3)
                deadline_v=r1.number_input("Deadline day",1,28,int(active_cfg["deadline_day"]),1)
                rest_v=r2.number_input("Min rest, h",0.0,24.0,float(active_cfg["min_rest_hours"]),1.0)
                maxday_v=r3.number_input("Max hours/day",1.0,24.0,float(active_cfg["max_hours_per_day"]),1.0)
                r4,r5,r6=st.columns(3)
                maxdays7_v=r4.number_input("Max workdays / 7d",1,int(FATIGUE_MAX_WORKDAYS_ROLLING7),min(int(active_cfg["max_workdays_rolling7"]),int(FATIGUE_MAX_WORKDAYS_ROLLING7)),1)
                maxhours7_v=r5.number_input("Generation max hours / 7d",1.0,float(FATIGUE_ROLLING7_HARD_CEILING_HOURS),min(float(active_cfg["max_hours_rolling7"]),float(FATIGUE_ROLLING7_HARD_CEILING_HOURS)),1.0)
                maxassign_v=r6.number_input("Max assignments/day",1,4,int(active_cfg["max_assignments_per_day"]),1)
                swap_cap_v=st.number_input(
                    "Voluntary swap hard cap / 7d",
                    48.0,float(SWAP_ABSOLUTE_MAX_HOURS_ROLLING7),
                    min(float(active_cfg.get("swap_max_hours_rolling7",SWAP_ABSOLUTE_MAX_HOURS_ROLLING7)),float(SWAP_ABSOLUTE_MAX_HOURS_ROLLING7)),1.0,
                    help=("Tai nėra 'legalizavimo' mygtukas: 48 h lieka perspėjimo slenkstis; šis laukas tik nustato absoliutų techninį bloką pagal klinikos patvirtintą darbo laiko režimą." if lang=="LT" else "This does not legalize extra hours: 48h remains a warning threshold; this field only sets the absolute technical blocker for the employer-approved work-time regime.")
                )

                t1,t2=st.columns(2)
                target_daily_v=t1.number_input("Target norm h/day",1.0,24.0,float(active_cfg["target_daily_hours"]),0.1)
                target_shift_v=t2.number_input("Target shift h",1.0,24.0,float(active_cfg["target_shift_hours"]),0.5)

                st.markdown("#### Struktūrinės taisyklės" if lang=="LT" else "#### Structural rules")
                s1,s2=st.columns(2)
                onko_v=s1.toggle("Onko pairs + recovery — HARD",value=True,disabled=True,help=("V2.5.68: Onko = 1.5 pamainos, todėl skiriamas lyginiu skaičiumi; be to, tam pačiam rezidentui Onko negalima dvi kalendorines dienas iš eilės." if lang=="LT" else "V2.5.68: Onko = 1.5 shift units, so counts are even; additionally, the same resident cannot work Onko on consecutive calendar days."))
                weekend_unique_v=s2.toggle("Weekend uniqueness required",value=bool(active_cfg["weekend_unique_required"]))
                weekend_cap_v=st.number_input("Weekend max assignments/resident",1,4,int(active_cfg["weekend_max_assignments_per_resident"]),1)

                st.markdown("#### Dublių apimtis" if lang=="LT" else "#### Backup scope")
                b1,b2=st.columns(2)
                backup_sps_ro_v=b1.toggle("SPS RO — visos dienos / blokai",value=bool(active_cfg["backup_sps_ro"]))
                backup_sps_ug_v=b2.toggle("SPS UG — visos dienos / blokai",value=bool(active_cfg["backup_sps_ug"]))
                backup_weekends_v=False  # compatibility-only; generic weekend scope retired in V2.5.98
                b4,b5,b6=st.columns(3)
                backup_centro120_v=b4.toggle("Centro 120 AM",value=bool(active_cfg.get("backup_centro120_am",True)))
                backup_onko_v=b5.toggle("Onko RO",value=bool(active_cfg.get("backup_onko_ro",True)))
                backup_centro_ro_v=b6.toggle("CENTRO RO best-effort",value=bool(active_cfg.get("backup_centro_ro_best_effort",True)))

                st.markdown("#### Fairness guardrails")
                g1,g2=st.columns(2)
                post_tol_v=g1.number_input("Legacy post tolerance (V2.5.53 constitutional gates are fixed)",0,5,int(active_cfg["post_guardrail_tolerance"]),1,disabled=True)
                general_tol_v=g2.number_input("Other spread tolerance",0,5,int(active_cfg["general_guardrail_tolerance"]),1)
                st.caption("V2.5.74 fixed generation gates: SPS RO / SPS UG / weekends / FRIDAYS raw 0–1; every ordinary non-Onko post targets raw 0–1, widening only after proven infeasibility; generation ≤48 known hours and ≤6 worked days per rolling 7 days. Voluntary-swap hard cap is separately configurable 48–60h and should match the employer-approved legal work-time regime.")

                with st.expander("Optimizerio svoriai — keisti tik sąmoningai" if lang=="LT" else "Optimizer weights — change deliberately"):
                    w1,w2,w3=st.columns(3)
                    post_weight_v=w1.number_input("Monthly post weight",0.0,100000.0,float(active_cfg["monthly_post_spread_weight"]),50.0)
                    catchup_weight_v=w2.number_input("Legacy catch-up weight — DISABLED in V2.5.96",0.0,100000.0,0.0,1.0,disabled=True,help="Kept only for Rule Profile schema compatibility. Historical fairness is audit-only and never steers future generation.")
                    active_reward_v=w3.number_input("Active SOFT reward",0.0,100000.0,float(active_cfg["active_date_reward"]),10.0)

                candidate_cfg={
                    "deadline_day":deadline_v,
                    "target_daily_hours":target_daily_v,
                    "target_shift_hours":target_shift_v,
                    "max_assignments_per_day":maxassign_v,
                    "max_hours_per_day":maxday_v,
                    "min_rest_hours":rest_v,
                    "max_workdays_rolling7":maxdays7_v,
                    "max_hours_rolling7":maxhours7_v,
                    "swap_max_hours_rolling7":swap_cap_v,
                    "onko_even_required":True,
                    "weekend_unique_required":weekend_unique_v,
                    "weekend_max_assignments_per_resident":weekend_cap_v,
                    "backup_weekends":backup_weekends_v,
                    "backup_sps_ro":backup_sps_ro_v,
                    "backup_sps_ug":backup_sps_ug_v,
                    "backup_centro120_am":backup_centro120_v,
                    "backup_onko_ro":backup_onko_v,
                    "backup_centro_ro_best_effort":backup_centro_ro_v,
                    "post_guardrail_tolerance":post_tol_v,
                    "general_guardrail_tolerance":general_tol_v,
                    "monthly_post_spread_weight":post_weight_v,
                    "cumulative_post_catchup_weight":catchup_weight_v,
                    "active_date_reward":active_reward_v,
                }
                validate_btn=st.form_submit_button("VALIDUOTI PAKEITIMUS" if lang=="LT" else "VALIDATE CHANGES")
                activate_btn=st.form_submit_button("SUKURTI IR AKTYVUOTI NAUJĄ VERSIJĄ" if lang=="LT" else "CREATE & ACTIVATE NEW VERSION",type="primary")

            normalized_cfg,rule_errors=validate_rule_profile(candidate_cfg)
            if validate_btn:
                if rule_errors:
                    st.error(("Taisyklių profilis netinkamas: " if lang=="LT" else "Invalid rule profile: ")+"; ".join(rule_errors))
                else:
                    changes={k:(active_cfg.get(k),normalized_cfg.get(k)) for k in normalized_cfg if active_cfg.get(k)!=normalized_cfg.get(k)}
                    if changes:
                        st.success("VALIDATION — PASSED")
                        st.dataframe(pd.DataFrame([
                            {"Rule":k,"Current":a,"Proposed":b} for k,(a,b) in changes.items()
                        ]),use_container_width=True,hide_index=True)
                    else:
                        st.info("Nėra pakeitimų." if lang=="LT" else "No changes.")

            if activate_btn:
                if rule_errors:
                    st.error(("NEAKTYVUOTA: " if lang=="LT" else "NOT ACTIVATED: ")+"; ".join(rule_errors))
                else:
                    changes={k:(active_cfg.get(k),normalized_cfg.get(k)) for k in normalized_cfg if active_cfg.get(k)!=normalized_cfg.get(k)}
                    if not changes:
                        st.info("Nėra pakeitimų." if lang=="LT" else "No changes.")
                    else:
                        try:
                            created=db.create_and_activate_rule_profile(profile_name,normalized_cfg,profile_note)
                            st.success(
                                ("Naujas taisyklių profilis aktyvuotas. " if lang=="LT" else "New Rule Profile activated. ")
                                + f"v{created.get('version_no','?')}"
                            )
                            st.rerun()
                        except Exception as e:
                            st.error(("Aktyvuoti nepavyko: " if lang=="LT" else "Activation failed: ")+str(e))

        profiles=db.list_rule_profiles(20)
        if profiles:
            with st.expander("Versijų istorija / rollback" if lang=="LT" else "Version history / rollback", expanded=False):
                hist=pd.DataFrame([{
                    "Version":r.get("version_no"),
                    "Name":r.get("name"),
                    "Active":bool(r.get("is_active")),
                    "Created":str(r.get("created_at") or "")[:19],
                    "Note":r.get("note") or "",
                } for r in profiles])
                st.dataframe(hist,use_container_width=True,hide_index=True)
                older=[r for r in profiles if not r.get("is_active")]
                if older:
                    selected_profile=st.selectbox(
                        "Grąžinti versiją" if lang=="LT" else "Restore version",
                        older,
                        format_func=lambda r:f"v{r.get('version_no')} — {r.get('name')}",
                        key="rollback_rule_profile_v2534"
                    )
                    confirm_rule_rollback=st.checkbox(
                        "Patvirtinu rollback" if lang=="LT" else "Confirm rollback",
                        key="confirm_rule_rollback_v2534"
                    )
                    if st.button(
                        "ROLLBACK Į PASIRINKTĄ VERSIJĄ" if lang=="LT" else "ROLL BACK TO SELECTED VERSION",
                        disabled=not confirm_rule_rollback,
                        use_container_width=True
                    ):
                        try:
                            restored=db.activate_rule_profile(int(selected_profile["id"]))
                            st.success(("Aktyvuota " if lang=="LT" else "Activated ")+f"v{restored.get('version_no','?')}")
                            st.rerun()
                        except Exception as e:
                            st.error(("Rollback nepavyko: " if lang=="LT" else "Rollback failed: ")+str(e))

    content=db.get_manual(lang)
    explanatory=explanatory_manual_only(content)
    st.caption(
        ("Žemiau — tik paaiškinimai ir kontekstas. Keičiamos operacinės taisyklės negali būti aprašomos ranka: jos keičiamos tik per aktyvų Rule Profile aukščiau."
         if lang=="LT" else
         "Below is explanatory context only. Mutable operational rules must not be defined manually; they are changed only through the active Rule Profile above.")
    )
    mode=tr("read")
    if senior_mode: mode=st.radio("",[tr("read"),tr("edit")],horizontal=True,label_visibility="collapsed")
    if mode==tr("read"):
        st.markdown(explanatory)
    else:
        st.warning(
            "Šis tekstas nėra engine taisyklių šaltinis. HARD, backup, target, guardrail ir kiti keičiami parametrai turi būti keičiami tik per RULE CHANGE RESCUE."
            if lang=="LT" else
            "This text is not the engine rule source. HARD, backup, target, guardrail and other mutable parameters must be changed only through RULE CHANGE RESCUE."
        )
        edited=st.text_area(("Paaiškinimai / pastabos" if lang=="LT" else "Explanatory notes"),value=explanatory,height=760,label_visibility="collapsed")
        if st.button(tr("save_rules"),type="primary"):
            db.save_manual(lang,edited)
            st.success(tr("rules_saved"))
            st.rerun()
